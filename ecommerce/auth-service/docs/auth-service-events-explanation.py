"""
Script tao file docx mo ta chi tiet he thong Events trong auth-service.
Yeu cau: python-docx
Chay: python3 auth-service-events-explanation.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    """To mau nen cho cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc, code_text, language='java'):
    """Them khoi code voi font monospace + nen xam nhe."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    # Nen xam nhe
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    p_pr.append(shd)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Them bang co header to dam, nen xanh nhat."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], '4472C4')
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ''
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Zebra striping
            if r_idx % 2 == 0:
                set_cell_background(row_cells[c_idx], 'F2F6FC')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_note(doc, text, color='FFF4CE'):
    """Khung 'Luu y' co nen vang nhat."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run('Luu y: ' + text)
    run.font.size = Pt(10)
    run.italic = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    p_pr.append(shd)


# =====================================================
# BAT DAU TAO DOCUMENT
# =====================================================

doc = Document()

# Cau hinh font mac dinh
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Cau hinh le trang
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ----- TITLE PAGE -----
title = doc.add_heading('HỆ THỐNG EVENTS TRONG AUTH-SERVICE', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Tài liệu giải thích chi tiết kiến trúc Domain Events và Integration Events')
run.italic = True
run.font.size = Pt(13)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Dự án: ecommerce/auth-service  |  Ngôn ngữ: Java + Spring Boot')
run.font.size = Pt(11)

doc.add_paragraph()

# =====================================================
# CHUONG 1: TONG QUAN
# =====================================================
add_heading(doc, '1. Tổng quan kiến trúc Events', level=1)

add_paragraph(doc,
    'Auth-service áp dụng kiến trúc Domain-Driven Design (DDD) kết hợp với Hexagonal '
    'Architecture. Trong kiến trúc này, "Event" (sự kiện) đóng vai trò trung tâm để '
    'kết nối các thành phần trong hệ thống mà không tạo ra phụ thuộc trực tiếp '
    '(loose coupling). Có hai loại event hoàn toàn khác nhau cùng tồn tại trong dự án.')

add_heading(doc, '1.1. Hai loại Event', level=2)

add_table(doc,
    headers=['Tiêu chí', 'Domain Event', 'Integration Event'],
    rows=[
        ['Phạm vi sử dụng',
         'Trong cùng một JVM (cùng service)',
         'Liên service (qua message broker Kafka)'],
        ['Mục đích',
         'Thông báo trạng thái thay đổi của Aggregate cho các thành phần nội bộ',
         'Thông báo cho các service khác biết về sự kiện đã xảy ra'],
        ['Định dạng',
         'Class Java strongly-typed (có kiểu chặt chẽ)',
         'Cấu trúc phẳng dạng JSON, versioned'],
        ['Schema thay đổi',
         'Tự do refactor vì chỉ ảnh hưởng nội bộ',
         'Phải tăng version (.v1 sang .v2) để tương thích ngược'],
        ['Lifetime',
         'Ngắn (vài mili-giây trong process)',
         'Lâu dài (lưu trong Kafka topic theo retention)'],
        ['Nơi định nghĩa',
         'Tầng domain',
         'Tầng application'],
    ],
    col_widths=[3.5, 6.0, 6.0])

doc.add_paragraph()
add_heading(doc, '1.2. Lý do tách hai loại Event', level=2)

reasons = [
    ('Tôn trọng ranh giới Bounded Context: Domain layer chỉ chứa logic nghiệp vụ thuần tuý, '
     'không được biết về Kafka, JSON, hay framework cụ thể nào. Nếu trộn lẫn, khi đổi message '
     'broker (Kafka sang RabbitMQ) sẽ phải sửa cả tầng domain.'),
    ('Linh hoạt khi tiến hoá: Domain event đổi tên class, refactor field bao nhiêu cũng được '
     'vì chỉ ảnh hưởng nội bộ service. Trong khi đó, Integration event phải ổn định vì có hàng '
     'chục service khác đang consume, không thể bắt họ deploy đồng thời.'),
    ('Lọc thông tin nhạy cảm: Một số domain event có chứa dữ liệu nội bộ (ví dụ chi tiết về '
     'mật khẩu đã hash, thông tin debug). Tầng mapper sẽ lọc bớt, chỉ giữ lại field cần thiết '
     'trước khi đẩy ra ngoài qua Kafka.'),
    ('Phù hợp với pattern "Eventually Consistent": Các service consume Kafka chấp nhận độ trễ '
     'vài giây, không cần đồng bộ tức thì như component nội bộ.'),
]
for r in reasons:
    add_bullet(doc, r)

# =====================================================
# CHUONG 2: DOMAIN EVENTS
# =====================================================
doc.add_page_break()
add_heading(doc, '2. Domain Events - Sự kiện nội bộ', level=1)

add_heading(doc, '2.1. Base class DomainEvent', level=2)

add_paragraph(doc, 'Đường dẫn: domain/events/DomainEvent.java', italic=True)

add_paragraph(doc,
    'DomainEvent là lớp trừu tượng (abstract class) làm nền tảng cho tất cả các domain event '
    'cụ thể. Việc dùng abstract class thay vì interface nhằm chia sẻ logic chung như '
    'tự động sinh eventId và occurredOn.')

add_code_block(doc,
'''public abstract class DomainEvent {
    private final UUID eventId = UUID.randomUUID();
    private final Instant occurredOn = Instant.now();

    public final UUID eventId() { return eventId; }
    public final Instant occurredOn() { return occurredOn; }
    public String eventType() { return getClass().getSimpleName(); }
    public abstract String aggregateId();
}''')

add_heading(doc, 'Giải thích từng thành phần', level=3)

add_table(doc,
    headers=['Thành phần', 'Loại', 'Vai trò'],
    rows=[
        ['eventId', 'UUID auto-gen',
         'Định danh duy nhất cho mỗi instance event. Dùng để dedupe phía consumer khi gặp '
         'duplicate delivery (Kafka chỉ đảm bảo at-least-once).'],
        ['occurredOn', 'Instant auto-gen',
         'Thời điểm event được tạo ra (wall-clock time). Dùng để consumer biết thứ tự '
         'tuyệt đối các event và tính độ trễ.'],
        ['eventType()', 'Method mặc định',
         'Trả về tên đơn giản của class (ví dụ "UserCreatedEvent"). Dùng cho mục đích log '
         'và debug. Subclass có thể override nếu cần.'],
        ['aggregateId()', 'Method abstract',
         'Bắt buộc subclass phải implement. Trả về id của Aggregate đã raise event. '
         'Quan trọng vì sẽ dùng làm Kafka partition key sau này, đảm bảo các event của '
         'cùng một user/role đi vào cùng một partition (giữ thứ tự).'],
        ['toString()', 'Method override',
         'Tự format gọn gàng cho log. Không in payload chi tiết để tránh leak PII '
         '(Personally Identifiable Information) vào log aggregator.'],
    ],
    col_widths=[3.0, 3.0, 9.5])

doc.add_paragraph()
add_heading(doc, 'Lý do thiết kế', level=3)

design_reasons = [
    ('Tại sao là abstract class chứ không phải interface? Vì cần có field instance '
     '(eventId, occurredOn) được khởi tạo tự động trong constructor. Interface ở Java không '
     'thể có instance field.'),
    ('Tại sao field là "final"? Vì event đại diện cho "sự thật đã xảy ra trong quá khứ" '
     '(past tense). Không thể sửa được. Đảm bảo immutability.'),
    ('Tại sao tách "eventType" thành method mà không phải field? Vì có thể subclass muốn '
     'trả về tên khác (ví dụ format chuẩn ngành). Việc dùng "getClass().getSimpleName()" '
     'làm mặc định giúp 90 phần trăm trường hợp không cần code thêm.'),
    ('Tại sao "aggregateId()" là abstract? Vì mỗi loại event biết rõ aggregate mình thuộc '
     'về là gì (UserCreatedEvent thì là userId, PermissionGrantedEvent thì là roleId). '
     'Bắt buộc subclass override giúp compiler kiểm tra ngay tại compile time.'),
]
for r in design_reasons:
    add_bullet(doc, r)

# ----- 2.2 DomainEventPublisher -----
add_heading(doc, '2.2. Interface DomainEventPublisher', level=2)
add_paragraph(doc, 'Đường dẫn: domain/events/DomainEventPublisher.java', italic=True)

add_paragraph(doc,
    'Đây là Outbound Port theo thuật ngữ Hexagonal Architecture. Domain layer khai báo '
    '"tôi cần phát event ra ngoài" mà không quan tâm đến cơ chế cụ thể là gì.')

add_code_block(doc,
'''public interface DomainEventPublisher {
    void publish(DomainEvent event);

    default void publishAll(Iterable<? extends DomainEvent> events) {
        events.forEach(this::publish);
    }
}''')

add_paragraph(doc, 'Giải thích các phương thức:', bold=True)
add_bullet(doc,
    'publish(DomainEvent event): Phát một event ra event bus. Mọi implementation cụ thể '
    'sẽ tự quyết định gửi đi đâu (Spring bus, log, message broker).')
add_bullet(doc,
    'publishAll(Iterable events): Phát một loạt event cùng lúc. Là default method với '
    'implementation đơn giản là forEach. Implementation cụ thể có thể override để batch '
    'gửi cho hiệu năng tốt hơn.')

add_heading(doc, 'Lý do thiết kế', level=3)
add_bullet(doc,
    'Tại sao đặt interface ở tầng domain nhưng implement ở tầng infrastructure? Đây là '
    'Dependency Inversion Principle. Tầng domain định nghĩa "cần cái gì", tầng '
    'infrastructure cung cấp "cách làm". Khi đổi framework, chỉ cần viết implementation '
    'mới mà không động đến aggregate.')
add_bullet(doc,
    'Tại sao có "publishAll" default method? Để tránh aggregate phải tự loop. Aggregate '
    'thường có nhiều event tích luỹ trong buffer, gọi publishAll gọn hơn forEach.')

# ----- 2.3 AggregateRoot -----
add_heading(doc, '2.3. Base class AggregateRoot', level=2)
add_paragraph(doc, 'Đường dẫn: domain/shared/AggregateRoot.java', italic=True)

add_paragraph(doc,
    'AggregateRoot là class cha cho tất cả các aggregate trong domain (User, RefreshToken, '
    'Role). Nó cung cấp cơ chế tích luỹ event vào buffer nội bộ trước khi infrastructure '
    'rút ra để publish.')

add_code_block(doc,
'''public abstract class AggregateRoot {
    private final transient List<DomainEvent> domainEvents = new ArrayList<>();

    protected final void registerEvent(DomainEvent event) {
        if (event == null) return;
        this.domainEvents.add(event);
    }

    public final List<DomainEvent> pullDomainEvents() {
        if (domainEvents.isEmpty()) return List.of();
        List<DomainEvent> snapshot = List.copyOf(domainEvents);
        domainEvents.clear();
        return snapshot;
    }

    public final List<DomainEvent> peekDomainEvents() {
        return Collections.unmodifiableList(domainEvents);
    }

    protected final void clearDomainEvents() {
        domainEvents.clear();
    }
}''')

add_heading(doc, 'Giải thích từng phương thức', level=3)

methods_aggregate = [
    ('registerEvent(event)',
     'protected final',
     'Được gọi BÊN TRONG các phương thức nghiệp vụ của aggregate (như User.register, '
     'User.changeEmail). Đẩy event vào buffer nội bộ. Tại sao protected? Vì chỉ subclass '
     '(các aggregate cụ thể) mới được phép raise event, code bên ngoài không được phép.'),
    ('pullDomainEvents()',
     'public final',
     'Được gọi BÊN NGOÀI bởi repository sau khi save thành công. Phương thức này có hai '
     'hành động: (1) tạo snapshot bất biến của buffer, (2) xoá sạch buffer gốc, (3) trả về '
     'snapshot. Vì có hai hành động (lấy và xoá) nên tên là "pull" chứ không phải "get".'),
    ('peekDomainEvents()',
     'public final',
     'Trả về view chỉ-đọc của buffer mà KHÔNG xoá. Mục đích cho unit test hoặc debug, '
     'để kiểm tra event nào đang chờ mà không làm thay đổi state.'),
    ('clearDomainEvents()',
     'protected final',
     'Xoá hẳn buffer mà không trả về gì. Dùng cho trường hợp đặc biệt như rollback handling, '
     'khi muốn huỷ bỏ event đã raise.'),
]
add_table(doc,
    headers=['Phương thức', 'Modifier', 'Vai trò chi tiết'],
    rows=[(m[0], m[1], m[2]) for m in methods_aggregate],
    col_widths=[3.0, 2.5, 10.0])

doc.add_paragraph()
add_heading(doc, 'Lý do thiết kế quan trọng', level=3)

add_bullet(doc,
    'Tại sao field "domainEvents" có modifier "transient"? Vì khi serialize aggregate ra DB '
    '(qua JPA), không muốn serialize cả list event. Event chỉ là trạng thái tạm thời '
    'in-memory.')
add_bullet(doc,
    'Tại sao "pullDomainEvents" lại copy ra "snapshot" trước khi clear? Vì sau khi clear, '
    'reference đến buffer cũ sẽ trỏ về list rỗng. Nếu không copy, caller nhận được list '
    'rỗng. Việc dùng "List.copyOf" tạo ra immutable copy, an toàn tuyệt đối.')
add_bullet(doc,
    'Tại sao "pullDomainEvents" có ngữ nghĩa destructive (vừa lấy vừa xoá)? Để đảm bảo '
    'mỗi event chỉ được publish ĐÚNG MỘT LẦN. Nếu là getter thông thường, dễ bị bug '
    '"publish duplicate" khi vô tình gọi hai lần.')
add_bullet(doc,
    'Tại sao trả về "List.of()" khi rỗng thay vì "new ArrayList<>()"? Vì "List.of()" là '
    'singleton bất biến do JVM tạo sẵn, không tốn memory allocation. Tối ưu nhỏ.')

# ----- 2.4 User events -----
doc.add_page_break()
add_heading(doc, '2.4. Các Domain Event cụ thể', level=2)
add_paragraph(doc,
    'Auth-service có tổng cộng 12 domain event, chia thành 3 nhóm theo aggregate.')

add_heading(doc, 'Nhóm 1: User events (7 event)', level=3)
add_table(doc,
    headers=['Event', 'Khi nào raise', 'Field chính', 'Aggregate Id'],
    rows=[
        ['UserCreatedEvent',
         'User.register() tạo user mới', 'userId, username, email', 'userId'],
        ['UserLoggedInEvent',
         'User.recordLogin() khi đăng nhập thành công', 'userId, username', 'userId'],
        ['EmailChangedEvent',
         'User.changeEmail() khi đổi email', 'userId, oldEmail, newEmail', 'userId'],
        ['PasswordChangedEvent',
         'User.changePassword() khi đổi mật khẩu', 'userId, username', 'userId'],
        ['AccountStatusChangedEvent',
         'User.transitionStatus() khi active/suspend/delete', 'userId, username, transition', 'userId'],
        ['RoleAssignedEvent',
         'User.assignRole() khi gán role mới', 'userId, roleId', 'userId'],
        ['RoleRemovedEvent',
         'User.removeRole() khi bỏ role', 'userId, roleId', 'userId'],
    ],
    col_widths=[3.5, 4.5, 4.0, 2.0])

doc.add_paragraph()
add_heading(doc, 'Nhóm 2: RefreshToken events (3 event)', level=3)
add_table(doc,
    headers=['Event', 'Khi nào raise', 'Field chính', 'Aggregate Id'],
    rows=[
        ['TokenCreatedEvent',
         'Tạo refresh token mới (lúc login)', 'userId, family, generation', 'tokenId'],
        ['TokenRotatedEvent',
         'Refresh token được xoay (rotation security)', 'userId, family, oldTokenId, newGeneration', 'tokenId'],
        ['TokenRevokedEvent',
         'Token bị thu hồi (logout, security breach)', 'userId, family, reason', 'tokenId'],
    ],
    col_widths=[3.5, 4.5, 4.0, 2.0])

doc.add_paragraph()
add_heading(doc, 'Nhóm 3: Role events (2 event)', level=3)
add_table(doc,
    headers=['Event', 'Khi nào raise', 'Field chính', 'Aggregate Id'],
    rows=[
        ['PermissionGrantedEvent',
         'Cấp permission cho role', 'roleId, permissionId', 'roleId'],
        ['PermissionRevokedEvent',
         'Thu hồi permission của role', 'roleId, permissionId', 'roleId'],
    ],
    col_widths=[3.5, 4.5, 4.0, 2.0])

doc.add_paragraph()
add_heading(doc, 'Pattern chung của Domain Event', level=3)

add_paragraph(doc,
    'Tất cả domain event đều tuân theo cùng một pattern POJO immutable. Ví dụ UserCreatedEvent:')

add_code_block(doc,
'''public final class UserCreatedEvent extends DomainEvent {
    private final String userId;
    private final String username;
    private final String email;

    public UserCreatedEvent(String userId, String username, String email) {
        this.userId = userId;
        this.username = username;
        this.email = email;
    }

    @Override public String aggregateId() { return userId; }

    public String userId()   { return userId; }
    public String username() { return username; }
    public String email()    { return email; }
}''')

add_paragraph(doc, 'Các đặc điểm bắt buộc:', bold=True)
add_bullet(doc,
    'Class phải là "final" để không ai kế thừa được. Event không nên có hierarchy phức tạp.')
add_bullet(doc,
    'Tất cả field phải là "private final" để đảm bảo immutability.')
add_bullet(doc,
    'Không có setter. Chỉ có getter style record (tên method = tên field, không "get" '
    'prefix).')
add_bullet(doc,
    'Chỉ chứa primitive type, String, UUID, Instant. KHÔNG được chứa reference đến '
    'entity khác (như User user, Role role) vì sẽ gây memory leak và serialization vấn đề.')
add_bullet(doc,
    'Phải override "aggregateId()" để trả về id của aggregate đã raise event.')

# =====================================================
# CHUONG 3: INTEGRATION EVENTS
# =====================================================
doc.add_page_break()
add_heading(doc, '3. Integration Events - Sự kiện liên service', level=1)

add_heading(doc, '3.1. Interface IntegrationEvent', level=2)
add_paragraph(doc, 'Đường dẫn: application/events/IntegrationEvent.java', italic=True)

add_paragraph(doc,
    'IntegrationEvent là contract cho mọi event sẽ được publish ra Kafka, cho các service '
    'khác tiêu thụ. Đây là tầng "ngoại giao" của hệ thống - phải ổn định, versioned, và '
    'tương thích ngược.')

add_code_block(doc,
'''public interface IntegrationEvent {
    UUID eventId();
    String eventType();
    Instant occurredOn();
    String aggregateId();
    int schemaVersion();
}''')

add_heading(doc, 'Giải thích từng phương thức', level=3)
add_table(doc,
    headers=['Phương thức', 'Vai trò'],
    rows=[
        ['eventId()',
         'Định danh duy nhất, được COPY từ DomainEvent gốc. Quan trọng vì consumer dùng '
         'để dedupe khi Kafka delivery duplicate (Kafka chỉ đảm bảo at-least-once).'],
        ['eventType()',
         'KHÁC với DomainEvent.eventType(). Đây là chuỗi format chuẩn ngành như '
         '"auth.user.created.v1", "auth.user.email_changed.v1". Consumer (có thể viết bằng '
         'Python, Go, Node) dựa vào chuỗi này để biết loại event chứ không quan tâm tên '
         'class Java.'],
        ['occurredOn()',
         'Thời điểm event gốc đã raise. Copy từ DomainEvent.'],
        ['aggregateId()',
         'Id của aggregate. Quan trọng vì Kafka sẽ dùng làm "key" để partition. '
         'Các event của cùng một user sẽ vào cùng partition, giữ đúng thứ tự.'],
        ['schemaVersion()',
         'Số nguyên đánh version cho payload. Khi có breaking change (đổi tên field, '
         'xoá field, đổi kiểu), bump version từ 1 lên 2. Consumer biết schema để parse '
         'cho đúng.'],
    ],
    col_widths=[3.5, 12.0])

doc.add_paragraph()
add_note(doc,
    'Tại sao eventType là chuỗi mà không phải class name? Vì consumer là service khác, '
    'thường viết bằng ngôn ngữ khác, không có khái niệm "class Java". Format '
    'service.aggregate.action.version là convention quốc tế giúp consumer parse được '
    'mà không cần share code.')

add_heading(doc, '3.2. Hai implementation cụ thể', level=2)

add_paragraph(doc,
    'Auth-service có hai loại Integration Event, được implement dưới dạng Java record '
    '(immutable, terse syntax).')

add_heading(doc, 'UserLifecycleIntegrationEvent', level=3)
add_paragraph(doc, 'Đường dẫn: application/events/UserLifecycleIntegrationEvent.java', italic=True)

add_code_block(doc,
'''public record UserLifecycleIntegrationEvent(
    UUID eventId,
    String eventType,
    Instant occurredOn,
    String aggregateId,
    int schemaVersion,
    Map<String, Object> payload
) implements IntegrationEvent {}''')

add_paragraph(doc, 'Dùng cho các event lifecycle của User và Role:')
for e in ['auth.user.created.v1', 'auth.user.email_changed.v1', 'auth.user.password_changed.v1',
          'auth.user.status_changed.v1', 'auth.user.role_assigned.v1', 'auth.user.role_removed.v1',
          'auth.role.permission_granted.v1', 'auth.role.permission_revoked.v1']:
    add_bullet(doc, e)

add_heading(doc, 'SessionIntegrationEvent', level=3)
add_paragraph(doc, 'Đường dẫn: application/events/SessionIntegrationEvent.java', italic=True)

add_code_block(doc,
'''public record SessionIntegrationEvent(
    UUID eventId,
    String eventType,
    Instant occurredOn,
    String aggregateId,
    int schemaVersion,
    Map<String, Object> payload
) implements IntegrationEvent {}''')

add_paragraph(doc, 'Dùng cho các event liên quan session và token:')
for e in ['auth.user.logged_in.v1', 'auth.token.created.v1',
          'auth.token.rotated.v1', 'auth.token.revoked.v1']:
    add_bullet(doc, e)

doc.add_paragraph()
add_heading(doc, 'Câu hỏi: Tại sao hai record cấu trúc giống hệt nhau lại tách ra?', level=3)

add_paragraph(doc,
    'Mặc dù về mặt kỹ thuật hai record có signature giống hệt, nhưng chúng đại diện cho '
    'hai khái niệm nghiệp vụ khác nhau và đi vào hai Kafka topic khác nhau:')

add_table(doc,
    headers=['Topic', 'Loại event', 'Consumer điển hình'],
    rows=[
        ['auth.user.events.v1',
         'UserLifecycleIntegrationEvent',
         'profile-service, mailing-service, billing-service - cần biết user đổi gì'],
        ['auth.session.events.v1',
         'SessionIntegrationEvent',
         'audit-service, fraud-detection, security-monitoring - cần theo dõi session'],
    ],
    col_widths=[5.0, 5.5, 6.0])

doc.add_paragraph()
add_bullet(doc,
    'Tách topic giúp consumer subscribe đúng cái mình cần, không bị spam event không '
    'liên quan.')
add_bullet(doc,
    'Tách record giúp dễ tiến hoá: nếu sau này muốn thêm field "ipAddress" cho session '
    'event (security purpose), chỉ cần bump SessionIntegrationEvent lên v2, không động '
    'đến UserLifecycleIntegrationEvent.')

# ----- 3.3 KafkaTopics -----
add_heading(doc, '3.3. KafkaTopics - Hằng số topic name', level=2)
add_paragraph(doc, 'Đường dẫn: application/events/KafkaTopics.java', italic=True)

add_code_block(doc,
'''public final class KafkaTopics {
    private KafkaTopics() {}

    public static final String AUTH_USER_EVENTS = "auth.user.events.v1";
    public static final String AUTH_SESSION_EVENTS = "auth.session.events.v1";
}''')

add_paragraph(doc, 'Lý do thiết kế:', bold=True)
add_bullet(doc,
    'Tập trung tên topic vào một file để tránh hardcode rải rác trong code.')
add_bullet(doc,
    'Constructor private để không ai instantiate được (chỉ là namespace chứa hằng số).')
add_bullet(doc,
    'Suffix ".v1" trong tên topic: Khi có breaking schema change, sẽ tạo topic mới '
    '".v2" chạy song song. Consumer cũ tiếp tục đọc v1 trong khi consumer mới đọc v2. '
    'Sau khi tất cả đã migrate, mới retire v1.')

# =====================================================
# CHUONG 4: PUBLISHER
# =====================================================
doc.add_page_break()
add_heading(doc, '4. Publishers - Bộ phát event', level=1)

add_heading(doc, '4.1. SpringDomainEventPublisher', level=2)
add_paragraph(doc, 'Đường dẫn: infrastructure/events/SpringDomainEventPublisher.java', italic=True)

add_paragraph(doc,
    'Đây là implementation của interface DomainEventPublisher (Outbound Port). Vì project '
    'dùng Spring Boot, implementation sẽ delegate sang ApplicationEventPublisher của Spring.')

add_code_block(doc,
'''@Component
public class SpringDomainEventPublisher implements DomainEventPublisher {
    private static final Logger log = LoggerFactory.getLogger(SpringDomainEventPublisher.class);
    private final ApplicationEventPublisher delegate;

    public SpringDomainEventPublisher(ApplicationEventPublisher delegate) {
        this.delegate = delegate;
    }

    @Override
    public void publish(DomainEvent event) {
        if (event == null) return;
        log.debug("Publishing {} [eventId={}, aggregateId={}]",
                event.eventType(), event.eventId(), event.aggregateId());
        delegate.publishEvent(event);
    }
}''')

add_heading(doc, 'Giải thích chi tiết', level=3)

add_bullet(doc,
    'Annotation @Component: Đánh dấu class là Spring bean, sẽ được Spring container quản lý '
    'và inject vào nơi cần.')
add_bullet(doc,
    'Field "delegate" kiểu ApplicationEventPublisher: Đây là event bus có sẵn của Spring, '
    'cho phép publish event nội bộ trong cùng JVM. Spring tự inject thông qua constructor.')
add_bullet(doc,
    'Phương thức publish: Đầu tiên check null để phòng thủ, sau đó log debug (chỉ log '
    'metadata, KHÔNG log payload để tránh leak PII vào hệ thống log aggregator như ELK).')
add_bullet(doc,
    'Cuối cùng delegate.publishEvent(event) sẽ broadcast event tới mọi @EventListener và '
    '@TransactionalEventListener trong cùng ApplicationContext.')

add_heading(doc, 'Lý do thiết kế', level=3)
add_bullet(doc,
    'Tại sao không gọi thẳng ApplicationEventPublisher trong aggregate? Vì aggregate ở tầng '
    'domain không được biết về Spring. Việc tách ra DomainEventPublisher interface '
    '(ở domain) và SpringDomainEventPublisher (ở infrastructure) là Dependency Inversion. '
    'Domain định nghĩa "cần gì", infrastructure cung cấp "cách làm".')
add_bullet(doc,
    'Tại sao log eventId chứ không log toàn bộ event? Vì event có thể chứa email, username, '
    'thông tin nhạy cảm. Chỉ log eventId và aggregateId đủ để correlation trong khi vẫn '
    'tuân thủ GDPR/privacy.')

# ----- 4.2 KafkaIntegrationEventPublisher -----
add_heading(doc, '4.2. KafkaIntegrationEventPublisher', level=2)
add_paragraph(doc,
    'Đường dẫn: infrastructure/messaging/KafkaIntegrationEventPublisher.java', italic=True)

add_paragraph(doc,
    'Class này đóng vai trò cầu nối giữa Spring event bus và Kafka. Nó LẮNG NGHE các '
    'domain event được publish bởi SpringDomainEventPublisher, rồi chuyển đổi và gửi '
    'sang Kafka.')

add_code_block(doc,
'''@Component
@RequiredArgsConstructor
public class KafkaIntegrationEventPublisher {
    private static final Logger log = LoggerFactory.getLogger(KafkaIntegrationEventPublisher.class);
    private final KafkaTemplate<String, Object> kafkaTemplate;
    private final DomainEventToIntegrationEventMapper mapper;

    @TransactionalEventListener(phase = TransactionPhase.AFTER_COMMIT)
    public void onDomainEvent(DomainEvent event) {
        var routed = mapper.map(event);
        if (routed == null) return;

        try {
            kafkaTemplate.send(routed.topic(), routed.key(), routed.event());
            log.debug("Published {} to {} (key={})",
                    routed.event().eventType(), routed.topic(), routed.key());
        } catch (Exception e) {
            log.warn("Failed to publish {} to {}: {}",
                    routed.event().eventType(), routed.topic(), e.getMessage());
        }
    }
}''')

add_heading(doc, 'Giải thích chi tiết', level=3)

add_bullet(doc,
    '@TransactionalEventListener với phase = AFTER_COMMIT: Đây là điểm quan trọng NHẤT. '
    'Listener này chỉ chạy SAU KHI database transaction đã commit thành công. Nếu '
    'transaction rollback thì listener KHÔNG chạy, đảm bảo Kafka không nhận event của '
    'transaction đã thất bại.')
add_bullet(doc,
    'mapper.map(event): Chuyển đổi DomainEvent thành cấu trúc Routed (topic + key + event). '
    'Nếu mapper trả về null nghĩa là event này không cần publish ra ngoài (event nội bộ '
    'thuần tuý), thoát luôn.')
add_bullet(doc,
    'kafkaTemplate.send(topic, key, value): Gửi event lên Kafka broker. Argument "key" là '
    'aggregateId, đảm bảo các event cùng aggregate đi vào cùng partition (giữ thứ tự).')
add_bullet(doc,
    'try-catch và chỉ log warn: Best-effort. Nếu Kafka chết tạm thời, event sẽ mất. '
    'Không retry tự động. Đây là trade-off để giữ code đơn giản.')

add_heading(doc, 'Lý do thiết kế và trade-off', level=3)

add_paragraph(doc,
    'Vì sao chỉ log warn mà không retry hay throw exception?')

add_bullet(doc,
    'Throw exception sẽ làm crash transaction handler, ảnh hưởng các listener khác '
    '(ví dụ RolePermissionCacheInvalidator) - không mong muốn.')
add_bullet(doc,
    'Retry tự động đòi hỏi backoff strategy, retry queue, deduplication phức tạp. Hiện tại '
    'business chấp nhận được "best-effort delivery" cho integration event.')
add_bullet(doc,
    'Nếu nghiệp vụ yêu cầu at-least-once delivery cứng, phải nâng cấp lên pattern '
    'Transactional Outbox: lưu event vào DB cùng transaction, một background worker '
    'đọc và publish lên Kafka với retry. Comment trong code đã ghi chú điều này.')

# =====================================================
# CHUONG 5: MAPPER
# =====================================================
doc.add_page_break()
add_heading(doc, '5. DomainEventToIntegrationEventMapper - Cầu nối hai thế giới', level=1)
add_paragraph(doc,
    'Đường dẫn: infrastructure/messaging/DomainEventToIntegrationEventMapper.java', italic=True)

add_paragraph(doc,
    'Đây là class CHỦ CHỐT, là điểm duy nhất trong hệ thống mà Domain Event và Integration '
    'Event "gặp nhau". Mọi quy tắc chuyển đổi đều tập trung tại đây.')

add_heading(doc, '5.1. Tổng quan class', level=2)

add_code_block(doc,
'''@Component
public class DomainEventToIntegrationEventMapper {
    private static final String FIELD_USERNAME = "username";
    private static final String FIELD_FAMILY = "family";
    private static final String FIELD_USER_ID = "userId";
    private static final int SCHEMA_V1 = 1;

    public record Routed(String topic, String key, IntegrationEvent event) {}

    public Routed map(DomainEvent event) {
        return switch (event) {
            case UserCreatedEvent e -> lifecycle(e, "auth.user.created.v1",
                    Map.of(FIELD_USERNAME, e.username(), "email", e.email()));
            // ... các case khác
            default -> null;
        };
    }
}''')

add_heading(doc, 'Giải thích các thành phần', level=3)

add_table(doc,
    headers=['Thành phần', 'Vai trò'],
    rows=[
        ['FIELD_USERNAME, FIELD_FAMILY, FIELD_USER_ID',
         'Hằng số String được dùng nhiều lần. Tránh "magic string" rải rác, dễ gõ sai. '
         'Sửa một chỗ áp dụng mọi nơi.'],
        ['SCHEMA_V1 = 1',
         'Phiên bản schema hiện tại. Khi đổi format payload, sẽ bump lên SCHEMA_V2 = 2.'],
        ['record Routed',
         'Inner record chứa 3 thông tin cần để gửi Kafka: topic đích, key (aggregateId '
         'cho partitioning), và payload (IntegrationEvent). Đóng gói gọn vào một object '
         'thay vì trả về 3 giá trị riêng lẻ.'],
        ['Phương thức map(event)',
         'Phương thức chính. Dùng pattern matching trên kiểu domain event để chuyển đổi '
         'thành Routed phù hợp. Switch expression của Java 17+ đảm bảo exhaustive check '
         'tại compile-time.'],
        ['Phương thức private lifecycle(...) và session(...)',
         'Hai helper riêng cho hai loại topic. lifecycle build UserLifecycleIntegrationEvent '
         'và gửi vào topic AUTH_USER_EVENTS. session build SessionIntegrationEvent và gửi '
         'vào topic AUTH_SESSION_EVENTS.'],
        ['Phương thức private fields(...)',
         'Helper tạo Map<String, Object> từ danh sách cặp key-value. Skip các value null '
         'để payload Kafka gọn gàng hơn.'],
    ],
    col_widths=[4.5, 11.0])

add_heading(doc, '5.2. Chi tiết phương thức map()', level=2)

add_paragraph(doc,
    'Phương thức "map" sử dụng switch expression với pattern matching của Java 17+. '
    'Mỗi case xử lý một loại domain event.')

add_paragraph(doc, 'Ví dụ chuyển đổi UserCreatedEvent:', bold=True)

add_code_block(doc,
'''case UserCreatedEvent e -> lifecycle(
        e, "auth.user.created.v1",
        Map.of(FIELD_USERNAME, e.username(), "email", e.email()));''')

add_paragraph(doc, 'Phân tích từng phần:')
add_bullet(doc,
    '"case UserCreatedEvent e": Pattern matching - nếu event là instanceof UserCreatedEvent, '
    'gán vào biến local "e" với kiểu UserCreatedEvent (không cần cast).')
add_bullet(doc,
    '"-> lifecycle(...)": Gọi helper lifecycle với 3 tham số.')
add_bullet(doc,
    'Tham số 1 "e": event gốc, để extract eventId, occurredOn, aggregateId.')
add_bullet(doc,
    'Tham số 2 "auth.user.created.v1": eventType chuỗi gửi cho consumer ngoài.')
add_bullet(doc,
    'Tham số 3 Map: payload chỉ chứa username và email - đủ thông tin cho consumer, '
    'không leak password hay metadata nội bộ.')

add_heading(doc, '5.3. Quy tắc routing', level=2)

add_paragraph(doc,
    'Mapper quyết định mỗi domain event sẽ đi vào topic nào. Đây là quy tắc tổng hợp:')

add_table(doc,
    headers=['Domain Event', 'Integration Event Type', 'Kafka Topic'],
    rows=[
        ['UserCreatedEvent', 'auth.user.created.v1', 'auth.user.events.v1'],
        ['EmailChangedEvent', 'auth.user.email_changed.v1', 'auth.user.events.v1'],
        ['PasswordChangedEvent', 'auth.user.password_changed.v1', 'auth.user.events.v1'],
        ['AccountStatusChangedEvent', 'auth.user.status_changed.v1', 'auth.user.events.v1'],
        ['RoleAssignedEvent', 'auth.user.role_assigned.v1', 'auth.user.events.v1'],
        ['RoleRemovedEvent', 'auth.user.role_removed.v1', 'auth.user.events.v1'],
        ['UserLoggedInEvent', 'auth.user.logged_in.v1', 'auth.session.events.v1'],
        ['TokenCreatedEvent', 'auth.token.created.v1', 'auth.session.events.v1'],
        ['TokenRotatedEvent', 'auth.token.rotated.v1', 'auth.session.events.v1'],
        ['TokenRevokedEvent', 'auth.token.revoked.v1', 'auth.session.events.v1'],
        ['PermissionGrantedEvent', 'auth.role.permission_granted.v1', 'auth.user.events.v1'],
        ['PermissionRevokedEvent', 'auth.role.permission_revoked.v1', 'auth.user.events.v1'],
    ],
    col_widths=[4.5, 6.0, 5.0])

doc.add_paragraph()
add_heading(doc, '5.4. Helper "fields(...)" và lý do tồn tại', level=2)

add_code_block(doc,
'''private static Map<String, Object> fields(Object... keyValuePairs) {
    Map<String, Object> map = new LinkedHashMap<>();
    for (int i = 0; i < keyValuePairs.length; i += 2) {
        String key = (String) keyValuePairs[i];
        Object val = i + 1 < keyValuePairs.length ? keyValuePairs[i + 1] : null;
        if (val != null) map.put(key, val);
    }
    return map;
}''')

add_paragraph(doc, 'Giải thích tại sao cần helper này thay vì dùng Map.of() có sẵn:', bold=True)

add_bullet(doc,
    'Map.of(...) ném NullPointerException ngay lập tức nếu có value null. Trong khi event '
    'có thể có field optional (ví dụ TokenRotatedEvent.family có thể null cho token đầu '
    'tiên). Cần một helper tolerant với null.')
add_bullet(doc,
    'Map.of(...) không đảm bảo thứ tự key (dùng hash map). LinkedHashMap giữ nguyên thứ tự '
    'put, làm output JSON ổn định, dễ debug.')
add_bullet(doc,
    'Logic "if (val != null) map.put(key, val)" giúp tự động bỏ qua field null, payload '
    'Kafka gọn gàng (không có "field": null).')

add_heading(doc, 'Phân tích vòng for chi tiết', level=3)
add_paragraph(doc,
    'Varargs Object... gom các tham số thành mảng phẳng. Ví dụ '
    'fields("a", 1, "b", 2) tạo mảng ["a", 1, "b", 2]. Vòng for nhảy bước 2 (i += 2), '
    'mỗi vòng xử lý một cặp key-value. Index chẵn là key, index lẻ là value.')

add_paragraph(doc, 'Check "i + 1 < keyValuePairs.length" để phòng thủ trường hợp truyền '
    'số tham số lẻ (lập trình viên gõ thiếu), tránh ArrayIndexOutOfBoundsException.')

# =====================================================
# CHUONG 6: LISTENERS
# =====================================================
doc.add_page_break()
add_heading(doc, '6. Listeners - Các bộ lắng nghe event', level=1)

add_paragraph(doc,
    'Ngoài KafkaIntegrationEventPublisher, hệ thống còn có một listener nữa lắng nghe '
    'domain event để thực hiện công việc khác.')

add_heading(doc, '6.1. RolePermissionCacheInvalidator', level=2)
add_paragraph(doc,
    'Đường dẫn: infrastructure/messaging/RolePermissionCacheInvalidator.java', italic=True)

add_paragraph(doc,
    'Class này giữ Redis cache (role tới permissions) đồng bộ với database. Khi có thay '
    'đổi permission của role, cache phải được refresh ngay để tránh stale data.')

add_code_block(doc,
'''@Component
@RequiredArgsConstructor
@Slf4j
public class RolePermissionCacheInvalidator {
    private final RoleRepository roleRepository;
    private final PermissionRepository permissionRepository;
    private final RolePermissionCachePort cache;

    @TransactionalEventListener(phase = TransactionPhase.AFTER_COMMIT)
    public void onPermissionGranted(PermissionGrantedEvent event) { refresh(event); }

    @TransactionalEventListener(phase = TransactionPhase.AFTER_COMMIT)
    public void onPermissionRevoked(PermissionRevokedEvent event) { refresh(event); }

    private void refresh(DomainEvent event) {
        String roleIdRaw = event.aggregateId();
        try {
            Role role = roleRepository.findById(RoleId.of(roleIdRaw)).orElse(null);
            if (role == null) {
                log.debug("Role {} no longer exists; leaving cache as-is", roleIdRaw);
                return;
            }
            Set<String> permissionNames = role.getPermissions().isEmpty()
                    ? Set.of()
                    : permissionRepository.findByIdIn(role.getPermissions()).stream()
                            .map(p -> p.getName().value())
                            .collect(Collectors.toSet());
            cache.put(role.getName().value(), permissionNames);
            log.info("Role-permission cache refreshed for role={} ({} perms)",
                    role.getName().value(), permissionNames.size());
        } catch (Exception e) {
            log.warn("Failed to refresh role-permission cache for {}; evicting instead. Error: {}",
                    roleIdRaw, e.getMessage());
            cache.evictAll();
        }
    }
}''')

add_heading(doc, 'Giải thích chi tiết', level=3)

add_bullet(doc,
    'Hai listener method được tách riêng cho PermissionGrantedEvent và PermissionRevokedEvent. '
    'Cả hai đều gọi cùng helper "refresh(event)" vì logic xử lý giống nhau.')
add_bullet(doc,
    'Logic refresh: (1) Load lại Role từ DB theo aggregateId, (2) Lấy danh sách permission '
    'names hiện tại, (3) Ghi đè cache bằng cache.put. Ghi đè (replace) an toàn hơn '
    'partial update vì atomic.')
add_bullet(doc,
    'Trường hợp role không tồn tại (đã bị xoá): không làm gì, để cache tự expire qua TTL.')
add_bullet(doc,
    'Trường hợp lỗi (Redis down, DB lag): fallback "cache.evictAll()" để xoá toàn bộ cache. '
    'Lần read sau sẽ tự reload từ DB, đảm bảo không có stale data.')

add_heading(doc, 'Lý do thiết kế', level=3)

add_bullet(doc,
    'Tại sao listener này nằm song song với KafkaIntegrationEventPublisher mà không xếp '
    'chuỗi? Vì cả hai đều có nhu cầu phản ứng với cùng một event nhưng làm việc khác nhau '
    '(một gửi Kafka, một refresh Redis). Spring event bus đảm bảo cả hai đều nhận được '
    'cùng một event.')
add_bullet(doc,
    'Tại sao dùng AFTER_COMMIT giống Kafka publisher? Cùng lý do - chỉ refresh cache khi '
    'DB đã commit chắc chắn, tránh cache chứa dữ liệu của transaction đã rollback.')
add_bullet(doc,
    'Tại sao re-fetch toàn bộ Role và permissions thay vì update theo delta? Vì code đơn '
    'giản hơn, ít bug hơn. Chi phí một query thêm chấp nhận được vì permission change '
    'không xảy ra quá thường xuyên.')

add_note(doc,
    'Listener này CHỈ invalidate cache trong JVM hiện tại. Nếu auth-service chạy nhiều '
    'instance, các instance khác sẽ tự refresh khi nhận Kafka event hoặc đợi TTL hết. '
    'Đây là trade-off chấp nhận được vì permission cache không yêu cầu strong consistency.')

# =====================================================
# CHUONG 7: FLOW TONG QUAN
# =====================================================
doc.add_page_break()
add_heading(doc, '7. Flow tổng quan một use case', level=1)

add_paragraph(doc,
    'Để hiểu các thành phần phối hợp với nhau như thế nào, mình trace toàn bộ flow một '
    'use case cụ thể: User đổi email từ old@mail.com sang new@mail.com.')

add_heading(doc, 'Bước 1: Use case khởi động', level=2)
add_paragraph(doc,
    'ChangeEmailUseCase nhận request từ controller, validate input, mở transaction, '
    'load user từ database qua UserRepository.findById().')

add_heading(doc, 'Bước 2: Aggregate xử lý nghiệp vụ', level=2)
add_code_block(doc,
'''user.changeEmail(newEmail);
// Bên trong User.changeEmail():
//   - Validate email format
//   - Check email chưa tồn tại
//   - this.email = newEmail
//   - this.updatedAt = Instant.now()
//   - registerEvent(new EmailChangedEvent(userId, oldEmail, newEmail))''')

add_paragraph(doc,
    'Sau bước này, buffer "domainEvents" của user chứa một EmailChangedEvent. Event vẫn '
    'chưa được publish ra đâu cả, chỉ tích luỹ in-memory.')

add_heading(doc, 'Bước 3: Repository save và drain events', level=2)
add_code_block(doc,
'''public User save(User user) {
    UserJpaEntity saved = jpaRepository.save(mapper.toJpaEntity(user, ...));
    user.pullDomainEvents().forEach(eventPublisher::publish);
    return mapper.toDomain(saved);
}''')

add_paragraph(doc,
    'jpaRepository.save() ghi vào DB (vẫn trong transaction, chưa commit). Sau khi save '
    'thành công, gọi pullDomainEvents() để rút sạch buffer, mỗi event được pass vào '
    'eventPublisher.publish().')

add_heading(doc, 'Bước 4: SpringDomainEventPublisher broadcast', level=2)
add_paragraph(doc,
    'publish() gọi delegate.publishEvent(event) - đây là ApplicationEventPublisher của '
    'Spring. Event được đẩy vào event bus nội bộ, mọi @EventListener và '
    '@TransactionalEventListener đã đăng ký đều "biết" về event này.')

add_paragraph(doc,
    'Tuy nhiên các @TransactionalEventListener (AFTER_COMMIT) chưa chạy ngay - chúng đợi '
    'transaction commit thành công.')

add_heading(doc, 'Bước 5: Transaction commit', level=2)
add_paragraph(doc,
    'Use case kết thúc, Spring TransactionInterceptor commit transaction. DB lúc này thực '
    'sự lưu thay đổi vĩnh viễn.')

add_heading(doc, 'Bước 6: Các listener AFTER_COMMIT chạy', level=2)
add_paragraph(doc,
    'Spring kích hoạt các @TransactionalEventListener đã đăng ký với phase AFTER_COMMIT. '
    'Có hai listener nhận EmailChangedEvent:')

add_bullet(doc,
    'KafkaIntegrationEventPublisher.onDomainEvent(): map sang '
    'UserLifecycleIntegrationEvent với eventType "auth.user.email_changed.v1", payload '
    '{"oldEmail": "old@mail.com", "newEmail": "new@mail.com"}, gửi lên Kafka topic '
    '"auth.user.events.v1" với key là userId.')
add_bullet(doc,
    'RolePermissionCacheInvalidator: KHÔNG xử lý event này vì chỉ subscribe '
    'PermissionGrantedEvent và PermissionRevokedEvent.')

add_heading(doc, 'Bước 7: Consumer service xử lý', level=2)
add_paragraph(doc,
    'Các service khác (profile-service, mailing-service, audit-service) đã subscribe '
    'topic "auth.user.events.v1" sẽ nhận được event và xử lý theo logic riêng:')

add_bullet(doc,
    'profile-service: cập nhật field email trong bảng profile cục bộ (cache).')
add_bullet(doc,
    'mailing-service: gửi email xác nhận đến địa chỉ cũ (security notification).')
add_bullet(doc,
    'audit-service: ghi log thay đổi vào bảng audit_history.')

# =====================================================
# CHUONG 8: BANG TOM TAT
# =====================================================
doc.add_page_break()
add_heading(doc, '8. Bảng tổng kết toàn bộ file', level=1)

add_table(doc,
    headers=['File', 'Tầng', 'Vai trò'],
    rows=[
        ['DomainEvent.java',
         'domain/events',
         'Abstract base class cho mọi domain event nội bộ. Cung cấp eventId, occurredOn, '
         'eventType, aggregateId.'],
        ['DomainEventPublisher.java',
         'domain/events',
         'Interface (Outbound Port) khai báo nhu cầu "phát event". Domain không biết '
         'implementation.'],
        ['AggregateRoot.java',
         'domain/shared',
         'Base class cho mọi aggregate. Cung cấp buffer chứa event và phương thức '
         'registerEvent, pullDomainEvents.'],
        ['UserCreatedEvent, EmailChangedEvent...',
         'domain/models/user/events',
         '7 event cụ thể cho aggregate User.'],
        ['TokenCreatedEvent, TokenRotatedEvent, TokenRevokedEvent',
         'domain/models/refreshtoken/events',
         '3 event cho aggregate RefreshToken.'],
        ['PermissionGrantedEvent, PermissionRevokedEvent',
         'domain/models/role/events',
         '2 event cho aggregate Role.'],
        ['IntegrationEvent.java',
         'application/events',
         'Interface contract cho event gửi qua Kafka. Yêu cầu eventId, eventType, '
         'occurredOn, aggregateId, schemaVersion.'],
        ['UserLifecycleIntegrationEvent.java',
         'application/events',
         'Record implement IntegrationEvent, dùng cho event lifecycle của User và Role.'],
        ['SessionIntegrationEvent.java',
         'application/events',
         'Record implement IntegrationEvent, dùng cho event session và token.'],
        ['KafkaTopics.java',
         'application/events',
         'Hằng số tên Kafka topic, version trong tên (.v1).'],
        ['SpringDomainEventPublisher.java',
         'infrastructure/events',
         'Implement DomainEventPublisher bằng Spring ApplicationEventPublisher.'],
        ['DomainEventToIntegrationEventMapper.java',
         'infrastructure/messaging',
         'Chuyển đổi DomainEvent sang IntegrationEvent + xác định topic và key.'],
        ['KafkaIntegrationEventPublisher.java',
         'infrastructure/messaging',
         'Listener AFTER_COMMIT - gọi mapper rồi gửi event qua Kafka.'],
        ['RolePermissionCacheInvalidator.java',
         'infrastructure/messaging',
         'Listener AFTER_COMMIT - refresh Redis cache khi có thay đổi permission.'],
    ],
    col_widths=[5.0, 4.0, 7.0])

# =====================================================
# CHUONG 9: NGUYEN TAC THIET KE
# =====================================================
doc.add_page_break()
add_heading(doc, '9. Nguyên tắc thiết kế áp dụng', level=1)

add_heading(doc, '9.1. Domain-Driven Design (DDD)', level=2)
add_bullet(doc,
    'Aggregate raise event khi state thay đổi - không trực tiếp gọi service khác. '
    'Decoupling triệt để.')
add_bullet(doc,
    'Event là "past tense" - mô tả "đã xảy ra", không phải "sẽ làm gì". Đặt tên đúng: '
    'UserCreatedEvent (không phải CreateUserCommand).')
add_bullet(doc,
    'Aggregate boundary rõ ràng - mỗi event thuộc về một aggregate duy nhất, xác định '
    'bởi aggregateId.')

add_heading(doc, '9.2. Hexagonal Architecture', level=2)
add_bullet(doc,
    'Domain layer không phụ thuộc framework (Spring, Kafka). Chỉ định nghĩa interface '
    '(port).')
add_bullet(doc,
    'Infrastructure layer cung cấp adapter implement các port. Đổi framework chỉ cần '
    'đổi adapter.')

add_heading(doc, '9.3. CQRS và Event-Driven Architecture', level=2)
add_bullet(doc,
    'Command (write) raise event. Event lan toả ra các consumer (read models) để cập '
    'nhật theo cách riêng.')
add_bullet(doc,
    'Eventually consistent - chấp nhận độ trễ nhỏ giữa các service, đổi lại scalability.')

add_heading(doc, '9.4. Outbox Pattern (chưa apply, có ghi chú)', level=2)
add_paragraph(doc,
    'Hiện tại code dùng @TransactionalEventListener(AFTER_COMMIT) làm "soft outbox". '
    'Đảm bảo event không bị publish khi transaction rollback. Tuy nhiên vẫn có rủi ro nhỏ: '
    'nếu service crash NGAY SAU commit nhưng TRƯỚC khi Kafka send xong, event sẽ mất. '
    'Comment trong code đã ghi rõ: cần nâng cấp lên Transactional Outbox Pattern nếu '
    'business yêu cầu at-least-once delivery cứng.')

add_heading(doc, '9.5. Single Responsibility Principle', level=2)
add_bullet(doc,
    'Mapper chỉ làm chuyển đổi - không gửi Kafka.')
add_bullet(doc,
    'Publisher chỉ làm I/O - không xử lý logic chuyển đổi.')
add_bullet(doc,
    'CacheInvalidator chỉ refresh cache - không can thiệp Kafka.')
add_bullet(doc,
    'Mỗi class có một lý do duy nhất để thay đổi.')

# =====================================================
# KET THUC
# =====================================================
doc.add_page_break()
add_heading(doc, '10. Kết luận', level=1)

add_paragraph(doc,
    'Hệ thống event trong auth-service được thiết kế theo các best practice của '
    'Domain-Driven Design và Hexagonal Architecture. Việc tách bạch Domain Event và '
    'Integration Event mang lại nhiều lợi ích:')

add_bullet(doc,
    'Domain layer "sạch" - không phụ thuộc framework, dễ test, dễ tiến hoá.')
add_bullet(doc,
    'Schema gửi ra ngoài ổn định - consumer service không bị break khi auth-service '
    'refactor nội bộ.')
add_bullet(doc,
    'Đảm bảo consistency - chỉ publish event khi DB commit thực sự, tránh "phantom event".')
add_bullet(doc,
    'Mở rộng dễ - thêm event mới chỉ cần thêm case trong mapper, không động đến code đã có.')

add_paragraph(doc,
    'Toàn bộ thiết kế này là implementation của các pattern kinh điển từ sách '
    '"Implementing Domain-Driven Design" của Vaughn Vernon và '
    '"Patterns of Enterprise Application Architecture" của Martin Fowler.')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('--- Hết tài liệu ---')
run.italic = True
run.font.size = Pt(10)

# Save
output_path = '/Users/admin/HieuTo/projectjn/ecommerce/auth-service/docs/auth-service-events-explanation.docx'
doc.save(output_path)
print(f"Da tao thanh cong: {output_path}")
