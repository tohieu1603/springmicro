#!/usr/bin/env python3
"""
Generates HIEU-Backend-Documentation-v2.docx in the style of
DacTa_MicroservicesPlatform.docx (formal Vietnamese spec format).

- 13 chapters (CHƯƠNG) + appendix (PHỤ LỤC)
- ~50 tables (version history, glossary, services, NFR, use cases, DB schema,
  API endpoints, communication matrix, security layers, metrics, env vars,
  service ports, error codes, roadmap, etc.)
- Reuses 10 original PNG figures extracted from the previous doc
- Plus many ASCII flow diagrams (the part the author especially liked)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC_IMG_DIR = "/tmp/be_images"
OUT_PATH = "/Users/admin/HieuTo/ecommerce/HIEU-Backend-Documentation-v2.docx"

# ============================================================
# Helpers
# ============================================================

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, italic=False, size=12, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def add_code(doc, text):
    """Monospace block (ASCII diagrams, code snippets)."""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(9)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
    return p

def add_image(doc, filename, width_cm=15):
    path = os.path.join(SRC_IMG_DIR, filename)
    if not os.path.exists(path):
        add_para(doc, f"[MISSING IMAGE: {filename}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_table(doc, headers, rows, col_widths_cm=None, header_color="2E5C8A", auto=True):
    n_cols = len(headers)
    t = doc.add_table(rows=1, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    # Header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            if i >= n_cols:
                break
            row[i].text = ''
            p = row[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Column widths
    if col_widths_cm:
        for row in t.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
    r.bold = True

def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(6)  # WD_BREAK.PAGE = 7 but enum may not be imported

# ============================================================
# Build document
# ============================================================
doc = Document()

# Set default styles
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

# Margins
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(2):
    doc.add_paragraph()
add_para(doc, "TÀI LIỆU ĐẶC TẢ KỸ THUẬT BACKEND", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ HIEU", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "KIẾN TRÚC MICROSERVICES", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "HIEU E-Commerce Backend Platform", italic=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()
add_para(doc, "Phiên bản: 2.0", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Công nghệ: Spring Boot 4.0.5 | Java 25 | gRPC | Kafka | PostgreSQL | Redis | Elasticsearch",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Ngày cập nhật: 14/05/2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(8):
    doc.add_paragraph()
add_para(doc, "© 2026 HIEU E-Commerce Team", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading(doc, "MỤC LỤC", level=1)
toc_items = [
    "LỊCH SỬ PHIÊN BẢN",
    "BẢNG CHÚ THÍCH THUẬT NGỮ VÀ VIẾT TẮT",
    "CHƯƠNG 1: TỔNG QUAN DỰ ÁN",
    "    1.1. Giới thiệu",
    "    1.2. Mục tiêu dự án",
    "    1.3. Phạm vi dự án",
    "    1.4. Đối tượng sử dụng",
    "    1.5. Cấu trúc tài liệu",
    "CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG",
    "    2.1. Kiến trúc tổng thể & Sơ đồ ngữ cảnh C4",
    "    2.2. Bốn tầng kiến trúc DDD",
    "    2.3. Cấu trúc thư mục chuẩn của một service",
    "    2.4. Topology hệ thống (15 services)",
    "    2.5. Stack công nghệ",
    "CHƯƠNG 3: ĐẶC TẢ YÊU CẦU",
    "    3.1. Yêu cầu chức năng theo service",
    "    3.2. Yêu cầu phi chức năng",
    "CHƯƠNG 4: USE CASE",
    "    4.1. Danh sách Actor",
    "    4.2. Tổng hợp Use Case",
    "    4.3. Chi tiết Use Case tiêu biểu",
    "CHƯƠNG 5: THIẾT KẾ CƠ SỞ DỮ LIỆU",
    "    5.1. Tổng quan database-per-service",
    "    5.2. Auth Service Database",
    "    5.3. Catalog Service Database",
    "    5.4. Order & Payment Database",
    "    5.5. Inventory & Reservation Database",
    "    5.6. Các bảng còn lại",
    "    5.7. Chỉ mục (Indexes)",
    "CHƯƠNG 6: THIẾT KẾ API",
    "    6.1. Quy ước chung",
    "    6.2. Auth Service API",
    "    6.3. User Profile API",
    "    6.4. Catalog API",
    "    6.5. Inventory API",
    "    6.6. Cart API",
    "    6.7. Order API",
    "    6.8. Payment API",
    "    6.9. Shipping API",
    "    6.10. Voucher API",
    "    6.11. Flash Sale API",
    "    6.12. Notification API",
    "CHƯƠNG 7: LUỒNG XỬ LÝ",
    "    7.1. Luồng đăng nhập và phát hành JWT",
    "    7.2. Luồng xử lý request qua API Gateway",
    "    7.3. Luồng addItem trong Cart Service",
    "    7.4. Saga đặt hàng (choreography)",
    "    7.5. Luồng thanh toán SePay",
    "    7.6. Luồng đổi/trả hàng",
    "    7.7. Luồng Flash Sale",
    "    7.8. State machine của Order",
    "CHƯƠNG 8: GIAO TIẾP GIỮA CÁC SERVICE",
    "    8.1. Tổng quan ba mô hình giao tiếp",
    "    8.2. Ma trận giao tiếp",
    "    8.3. Danh mục Kafka Topics",
    "    8.4. gRPC Service Definitions",
    "CHƯƠNG 9: BẢO MẬT",
    "    9.1. Tổng quan bảo mật nhiều lớp",
    "    9.2. JWT (JSON Web Token)",
    "    9.3. RBAC – Phân quyền theo Role + Permission",
    "    9.4. Refresh Token Rotation & Blacklist",
    "    9.5. Idempotency & Anti-replay",
    "    9.6. Rate Limiting & CORS",
    "    9.7. Bảo mật vận hành",
    "CHƯƠNG 10: KIỂM THỬ",
    "    10.1. Chiến lược kiểm thử (Test Pyramid)",
    "    10.2. Unit Test",
    "    10.3. Integration Test",
    "    10.4. Coverage và chất lượng",
    "CHƯƠNG 11: GIÁM SÁT VÀ VẬN HÀNH",
    "    11.1. Boot stack local",
    "    11.2. Migrations & Seeding",
    "    11.3. Structured Logging & Correlation ID",
    "    11.4. Metrics & Actuator",
    "    11.5. Khôi phục sau sự cố",
    "CHƯƠNG 12: TRIỂN KHAI",
    "    12.1. Docker Compose",
    "    12.2. Biến môi trường",
    "    12.3. Các lệnh khởi động",
    "    12.4. Dự kiến: Kubernetes",
    "    12.5. Dự kiến: CI/CD Pipeline",
    "CHƯƠNG 13: LỘ TRÌNH PHÁT TRIỂN",
    "    13.1. Hiện trạng triển khai",
    "    13.2. Công nghệ dự kiến bổ sung",
    "    13.3. Lộ trình theo giai đoạn",
    "PHỤ LỤC",
    "    A. Bảng cổng dịch vụ",
    "    B. Tham chiếu mã lỗi",
    "    C. Tài khoản và dữ liệu seed",
    "    D. Cấu hình mẫu",
]
for item in toc_items:
    add_para(doc, item, size=11)

doc.add_page_break()

# ============================================================
# VERSION HISTORY
# ============================================================
add_heading(doc, "LỊCH SỬ PHIÊN BẢN", level=1)
add_table_caption(doc, "Bảng 0.1: Lịch sử phiên bản tài liệu")
add_table(doc,
    ["Phiên bản", "Ngày", "Tác giả", "Mô tả thay đổi"],
    [
        ["0.1", "12/03/2026", "Backend Team", "Khởi tạo tài liệu với 5 service cơ bản (auth, catalog, cart, order, payment)"],
        ["0.5", "08/04/2026", "Backend Team", "Bổ sung inventory, shipping, voucher, notification; thêm sequence diagrams"],
        ["1.0", "13/05/2026", "Backend Team", "Hoàn thiện 15 service; thêm use case, saga đặt hàng, state machine, vận hành"],
        ["2.0", "14/05/2026", "Backend Team", "Tái cấu trúc theo chuẩn đặc tả – bổ sung 13 chương, bảng glossary, schema DB, ma trận giao tiếp, đặc tả use case chi tiết"],
    ],
    col_widths_cm=[2.0, 2.5, 3.0, 9.0])

# ============================================================
# GLOSSARY
# ============================================================
add_heading(doc, "BẢNG CHÚ THÍCH THUẬT NGỮ VÀ VIẾT TẮT", level=1)
add_para(doc, "Tài liệu sử dụng các thuật ngữ và từ viết tắt sau đây. Người đọc cần nắm các khái niệm này trước khi đi vào các chương tiếp theo.")

add_para(doc, "1. Thuật ngữ kỹ thuật", bold=True)
add_table_caption(doc, "Bảng 0.2: Thuật ngữ kỹ thuật")
add_table(doc,
    ["STT", "Thuật ngữ", "Tên đầy đủ", "Giải thích"],
    [
        ["1", "API", "Application Programming Interface", "Giao diện lập trình ứng dụng, cho phép các hệ thống giao tiếp với nhau"],
        ["2", "REST", "Representational State Transfer", "Kiến trúc API trên HTTP, dùng JSON cho payload"],
        ["3", "gRPC", "Google Remote Procedure Call", "Giao thức RPC dựa trên HTTP/2 + Protobuf, latency thấp"],
        ["4", "JWT", "JSON Web Token", "Chuẩn token tự đóng gói, ký bằng HMAC hoặc RSA"],
        ["5", "BCrypt", "-", "Thuật toán hash mật khẩu có salt, chống brute-force"],
        ["6", "HMAC", "Hash-based Message Authentication Code", "Cơ chế xác thực thông điệp bằng khóa bí mật + hash"],
        ["7", "TTL", "Time To Live", "Thời gian sống của một bản ghi (cache, token, reservation)"],
        ["8", "SSE", "Server-Sent Events", "Giao thức push một chiều từ server sang client trên HTTP"],
        ["9", "DLQ", "Dead Letter Queue", "Hàng đợi chứa message xử lý thất bại"],
        ["10", "RBAC", "Role-Based Access Control", "Mô hình phân quyền theo vai trò"],
    ],
    col_widths_cm=[1.0, 3.0, 4.5, 8.0])

add_para(doc, "2. Thuật ngữ kiến trúc", bold=True)
add_table_caption(doc, "Bảng 0.3: Thuật ngữ kiến trúc")
add_table(doc,
    ["STT", "Thuật ngữ", "Tên đầy đủ", "Giải thích"],
    [
        ["1", "Microservices", "-", "Kiến trúc chia hệ thống thành nhiều service nhỏ, độc lập, deploy riêng"],
        ["2", "DDD", "Domain-Driven Design", "Thiết kế phần mềm tập trung vào mô hình nghiệp vụ"],
        ["3", "Aggregate", "-", "Cụm entity có ranh giới giao dịch chung trong DDD"],
        ["4", "Saga", "-", "Chuỗi giao dịch local có compensation, thay cho distributed transaction"],
        ["5", "CQRS", "Command Query Responsibility Segregation", "Tách luồng ghi (Command) khỏi luồng đọc (Query)"],
        ["6", "Outbox Pattern", "-", "Lưu event vào bảng outbox cùng transaction nghiệp vụ, sau đó publisher đọc và phát đi – đảm bảo at-least-once"],
        ["7", "Idempotency", "-", "Tính chất một thao tác cho cùng kết quả khi gọi nhiều lần với cùng input"],
        ["8", "Choreography Saga", "-", "Saga không có orchestrator – mỗi service tự subscribe events và phản ứng"],
        ["9", "Service Discovery", "-", "Cơ chế tự đăng ký và tìm kiếm service đang chạy (Eureka)"],
        ["10", "API Gateway", "-", "Điểm vào duy nhất, định tuyến request đến service đích"],
        ["11", "Circuit Breaker", "-", "Cơ chế tạm dừng gọi service bị lỗi liên tiếp để bảo vệ hệ thống"],
    ],
    col_widths_cm=[1.0, 3.0, 4.5, 8.0])

add_para(doc, "3. Thuật ngữ sản phẩm/nghiệp vụ", bold=True)
add_table_caption(doc, "Bảng 0.4: Thuật ngữ nghiệp vụ thương mại điện tử")
add_table(doc,
    ["STT", "Thuật ngữ", "Tên đầy đủ", "Giải thích"],
    [
        ["1", "SPU", "Standard Product Unit", "Đơn vị sản phẩm chuẩn – đại diện cho một sản phẩm (ví dụ 'Áo thun HIEU')"],
        ["2", "SKU", "Stock Keeping Unit", "Đơn vị quản lý kho – biến thể cụ thể (Đen, Size L)"],
        ["3", "Cart", "Giỏ hàng", "Tập hợp các item user chọn nhưng chưa đặt"],
        ["4", "Order", "Đơn hàng", "Cam kết mua – có trạng thái, tổng tiền, người nhận"],
        ["5", "Reservation", "Đặt giữ kho", "Hành động tạm giữ tồn kho khi đặt hàng, có TTL"],
        ["6", "Voucher", "Phiếu giảm giá", "Mã giảm giá theo % hoặc số tiền cố định"],
        ["7", "Flash Sale", "-", "Chương trình giảm giá thời gian giới hạn với quota cứng"],
        ["8", "Refund", "Hoàn tiền", "Trả lại tiền cho khách khi return/cancel"],
        ["9", "Shipment", "Lô hàng vận chuyển", "Bản ghi giao vận với tracking number"],
        ["10", "VietQR", "-", "Chuẩn QR thanh toán liên ngân hàng tại Việt Nam"],
    ],
    col_widths_cm=[1.0, 3.0, 4.0, 8.5])

add_para(doc, "4. Thuật ngữ hạ tầng", bold=True)
add_table_caption(doc, "Bảng 0.5: Thuật ngữ hạ tầng")
add_table(doc,
    ["STT", "Thuật ngữ", "Tên đầy đủ", "Giải thích"],
    [
        ["1", "Docker", "-", "Nền tảng đóng gói ứng dụng thành container"],
        ["2", "K8s", "Kubernetes", "Hệ thống điều phối container, hỗ trợ auto-scale, rolling update"],
        ["3", "Kafka", "Apache Kafka", "Distributed log/event streaming – sử dụng làm event bus"],
        ["4", "KRaft", "Kafka Raft", "Chế độ Kafka không cần Zookeeper, dùng Raft consensus"],
        ["5", "PostgreSQL", "-", "Hệ quản trị CSDL quan hệ – dùng cho mọi service nghiệp vụ"],
        ["6", "Redis", "-", "In-memory data store – cache, session, blacklist, rate limit"],
        ["7", "Elasticsearch", "-", "Search engine + log store"],
        ["8", "Kibana", "-", "UI cho Elasticsearch – view dashboard log"],
        ["9", "Filebeat", "-", "Shipper log từ file lên Elasticsearch"],
        ["10", "MailHog", "-", "SMTP server giả lập cho dev"],
        ["11", "Eureka", "Netflix Eureka", "Service registry – các service đăng ký và discover qua đây"],
        ["12", "Flyway", "-", "Database migration tool – chạy V1__init.sql, V2__... khi service start"],
    ],
    col_widths_cm=[1.0, 3.0, 3.5, 9.0])

doc.add_page_break()

# ============================================================
# CHƯƠNG 1: TỔNG QUAN DỰ ÁN
# ============================================================
add_heading(doc, "CHƯƠNG 1: TỔNG QUAN DỰ ÁN", level=1)

add_heading(doc, "1.1. Giới thiệu", level=2)
add_para(doc, "Tài liệu này đặc tả chi tiết backend của hệ thống thương mại điện tử HIEU (HIEU E-Commerce Platform) – một nền tảng bán hàng thời trang phân khúc luxury được xây dựng từ đầu theo kiến trúc microservices. Hệ thống gồm 15 service triển khai độc lập, mỗi service phụ trách một nghiệp vụ nhỏ nhưng phối hợp với nhau qua REST, gRPC và Kafka để cung cấp đầy đủ luồng mua sắm: từ duyệt sản phẩm, thêm giỏ, đặt hàng, thanh toán, vận chuyển tới đổi/trả và đánh giá sau bán.")
add_para(doc, "Hệ thống được thiết kế theo các nguyên tắc sau:")
for item in [
    "Kiến trúc Microservices: 15 service độc lập, mỗi service có database riêng (database-per-service).",
    "Domain-Driven Design (DDD): 4 tầng rõ ràng – interfaces, application, domain, infrastructure.",
    "Choreography Saga: đặt hàng là saga phân tán không có orchestrator – mỗi service tự subscribe event và phản ứng.",
    "Idempotency: mọi command quan trọng (đặt hàng, thanh toán, hoàn tiền) đều có khóa idempotency để xử lý retry an toàn.",
    "Outbox Pattern: order-service ghi event vào bảng outbox cùng transaction nghiệp vụ, scheduler publish lên Kafka – đảm bảo at-least-once delivery.",
    "Service Discovery: tự đăng ký và tìm kiếm qua Netflix Eureka.",
    "Event-driven Communication: Kafka làm event bus xuyên 23 topic.",
    "Cloud-Native: đóng gói Docker; sẵn sàng triển khai Kubernetes; expose Actuator + Prometheus endpoints.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.2. Mục tiêu dự án", level=2)
add_para(doc, "Dự án hướng tới các mục tiêu chính sau:")
for item in [
    "Cung cấp một nền tảng e-commerce hoàn chỉnh cho phân khúc thời trang luxury với đầy đủ nghiệp vụ cốt lõi.",
    "Áp dụng kiến trúc microservices nhằm đảm bảo tính độc lập triển khai, mở rộng theo trục nghiệp vụ và cô lập sự cố.",
    "Tích hợp xác thực tập trung bằng JWT (HS256) – các service tự xác thực không cần round-trip về auth-service.",
    "Hỗ trợ ba mô hình giao tiếp: REST (FE ↔ Gateway), gRPC (inter-service sync), Kafka (inter-service async).",
    "Đảm bảo tính toàn vẹn dữ liệu qua saga + outbox + idempotency – kể cả khi Kafka delay hoặc service downstream bị lỗi.",
    "Tích hợp cổng thanh toán SePay (VietQR + bank transfer) và tracking vận chuyển GHN/GHTK.",
    "Thiết lập hệ thống giám sát qua Actuator, Filebeat, Elasticsearch, Kibana – tiền đề cho Prometheus + Grafana.",
    "Áp dụng test pyramid: unit test cho domain, integration test cho controller-service-repo, e2e qua TestContainers.",
    "Sẵn sàng cho roadmap mở rộng: Recommendation AI, Inventory replenishment, Multi-warehouse.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.3. Phạm vi dự án", level=2)
add_para(doc, "Hệ thống hiện tại gồm 15 service triển khai theo kiến trúc microservices. Tài liệu này không bao gồm phần frontend (đã có tài liệu riêng) ngoài khi cần làm rõ cách FE tương tác với BE.")
add_table_caption(doc, "Bảng 1.1: Danh sách các service trong dự án")
add_table(doc,
    ["STT", "Service", "Cổng", "Vai trò chính"],
    [
        ["1", "eureka-server", "8761", "Service registry – đăng ký và discover service"],
        ["2", "api-gateway", "8080", "Edge gateway – routing, CORS, rate limit, log trace"],
        ["3", "auth-service", "9091", "Xác thực, JWT, role/permission, refresh rotation"],
        ["4", "user-profile-service", "9099", "Hồ sơ khách hàng, sổ địa chỉ"],
        ["5", "catalog-service", "9093", "Sản phẩm, biến thể, danh mục, thuộc tính (gRPC + REST)"],
        ["6", "inventory-service", "9098", "Tồn kho, reservation, audit StockMovement"],
        ["7", "cart-service", "9094", "Giỏ hàng, pre-flight validation, idempotency"],
        ["8", "order-service", "9095", "Đặt hàng, saga coordinator, return request, outbox"],
        ["9", "payment-service", "8086", "Thanh toán SePay, refund, webhook"],
        ["10", "shipping-service", "8087", "Tạo và theo dõi shipment, public tracking"],
        ["11", "voucher-service", "8094", "Mã giảm giá, validate/apply/release"],
        ["12", "flash-sale-service", "8089", "Flash sale với atomic quota (Redis Lua)"],
        ["13", "notification-service", "8090", "Email + in-app notification, SSE realtime"],
        ["14", "search-service", "8092", "Full-text search trên Elasticsearch"],
        ["15", "analytics-service", "8095", "Dashboard KPI, log explorer, revenue chart"],
    ],
    col_widths_cm=[1.0, 4.5, 1.5, 10.0])

add_heading(doc, "1.4. Đối tượng sử dụng", level=2)
add_table_caption(doc, "Bảng 1.2: Các đối tượng sử dụng tài liệu")
add_table(doc,
    ["STT", "Đối tượng", "Mục đích sử dụng"],
    [
        ["1", "Kỹ sư backend mới", "Hiểu cấu trúc thư mục, domain model, quy ước, REST/Kafka contract"],
        ["2", "Kỹ sư DevOps", "Biết cổng dịch vụ, dependency, cách boot, cấu hình env, cách monitor"],
        ["3", "Kỹ sư QA", "Xây test plan dựa vào sequence diagram, use case và state machine"],
        ["4", "Tech Lead/Kiến trúc sư", "Review kiến trúc, đánh giá khả năng mở rộng, đề xuất cải tiến"],
        ["5", "Sản phẩm/Doanh nghiệp", "Hiểu phạm vi nghiệp vụ, các use case đã hỗ trợ"],
        ["6", "Đối tác tích hợp", "Tham chiếu API spec, mã lỗi, webhook contract (SePay, GHN)"],
    ],
    col_widths_cm=[1.0, 4.5, 11.0])

add_heading(doc, "1.5. Cấu trúc tài liệu", level=2)
add_table_caption(doc, "Bảng 1.3: Cấu trúc nội dung tài liệu")
add_table(doc,
    ["Chương", "Nội dung"],
    [
        ["Chương 1", "Tổng quan dự án – mục tiêu, phạm vi, đối tượng đọc"],
        ["Chương 2", "Kiến trúc hệ thống – C4, DDD, topology, stack"],
        ["Chương 3", "Đặc tả yêu cầu chức năng và phi chức năng"],
        ["Chương 4", "Use case theo từng actor (Customer, Admin, CSKH)"],
        ["Chương 5", "Thiết kế CSDL chi tiết từng service"],
        ["Chương 6", "Thiết kế API – endpoint, request/response, ApiResponse"],
        ["Chương 7", "Luồng xử lý – sequence diagram và sơ đồ trạng thái"],
        ["Chương 8", "Giao tiếp inter-service – REST/gRPC/Kafka, topic catalog"],
        ["Chương 9", "Bảo mật nhiều lớp – JWT, RBAC, idempotency, rate limit"],
        ["Chương 10", "Chiến lược kiểm thử – pyramid, coverage, framework"],
        ["Chương 11", "Giám sát và vận hành – logging, metrics, recovery"],
        ["Chương 12", "Triển khai – Docker Compose, env, K8s, CI/CD"],
        ["Chương 13", "Lộ trình – hiện trạng và roadmap"],
        ["Phụ lục", "Bảng cổng, mã lỗi, dữ liệu seed, cấu hình mẫu"],
    ],
    col_widths_cm=[2.5, 13.5])

doc.add_page_break()

# ============================================================
# CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG
# ============================================================
add_heading(doc, "CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG", level=1)

add_heading(doc, "2.1. Kiến trúc tổng thể & Sơ đồ ngữ cảnh C4", level=2)
add_para(doc, "Hệ thống áp dụng kiến trúc Microservices với API Gateway làm điểm vào duy nhất. Mỗi service hoạt động độc lập, có cơ sở dữ liệu riêng và giao tiếp với nhau qua REST, gRPC hoặc Kafka tuỳ theo ngữ cảnh.")
add_para(doc, "Các thành phần cốt lõi:")
for item in [
    "API Gateway (port 8080): Spring Cloud Gateway xử lý định tuyến theo path prefix, CORS allowlist, log trace ID, rate limit qua Redis counter.",
    "Eureka Server (port 8761): service registry – các service đăng ký lên đây mỗi lần boot, heartbeat mỗi 30 giây.",
    "Auth Service (port 9091): trung tâm xác thực – phát hành JWT, refresh rotation, RBAC, blacklist token.",
    "Catalog Service (port 9093): trung tâm sản phẩm – cung cấp REST cho FE và gRPC cho service nội bộ (cart, order).",
    "11 business services: cart, order, inventory, payment, shipping, voucher, flash-sale, notification, search, analytics, user-profile.",
    "Hạ tầng dùng chung: 11 database PostgreSQL (mỗi service một DB), Redis, Kafka KRaft, Elasticsearch + Kibana, MailHog (dev).",
]:
    add_bullet(doc, item)

add_para(doc, "Sơ đồ C4 Level 1 (ngữ cảnh hệ thống):")
add_image(doc, "image1.png", width_cm=15)
add_caption(doc, "Hình 2.1: Sơ đồ ngữ cảnh C4 Level 1 – các actor và hệ thống bên ngoài")

add_para(doc, "Bản đồ 15 service và hạ tầng phụ thuộc:")
add_image(doc, "image2.png", width_cm=15)
add_caption(doc, "Hình 2.2: Bản đồ 15 service và hạ tầng phụ thuộc")

add_para(doc, "Mô hình tổng quan dạng ASCII:")
add_code(doc, """                    +------------------+
                    |  Client (Next.js)|
                    +--------+---------+
                             |
                             v
                    +------------------+
                    |   API Gateway    |  :8080
                    | (Spring Cloud GW)|  - CORS / Path routing
                    |                  |  - JWT pass-through
                    |                  |  - Rate limit (Redis)
                    +--+--+--+--+--+--++
                       |  |  |  |  |  |
       +---------------+  |  |  |  |  +---------------+
       v                  v  v  v  v                  v
  +---------+      +----------+   ...   +----------+  +---------+
  |  Auth   |      | Catalog  |         | Analytics|  | Eureka  |
  |  :9091  |      |  :9093   |         |  :8095   |  |  :8761  |
  |  gRPC   |      |  gRPC    |         |          |  |         |
  +----+----+      +----+-----+         +----+-----+  +---------+
       |                |                    |
       v                v                    v
  +---------+      +----------+         +----------+
  |auth_db  |      |catalog_db|         |  ES logs |
  |Postgres |      | Postgres |         |Elastic   |
  +---------+      +----------+         +----------+

  Shared Infrastructure:
  +----------+  +-----------+  +-------------+  +------------+  +----------+
  |  Redis   |  |  Kafka    |  |Elasticsearch|  |  Kibana    |  | MailHog  |
  |  :6379   |  | (KRaft)   |  |   :9200     |  |   :5601    |  |  :8025   |
  | cache+   |  |  :9092    |  |             |  |            |  |          |
  | blacklist|  | 23 topics |  |             |  |            |  |          |
  +----------+  +-----------+  +-------------+  +------------+  +----------+""")
add_caption(doc, "Hình 2.3: Topology hệ thống – ASCII")

add_heading(doc, "2.2. Bốn tầng kiến trúc DDD", level=2)
add_para(doc, "Mỗi service được thiết kế theo Domain-Driven Design với 4 tầng. Phụ thuộc đi một chiều: tầng ngoài có thể phụ thuộc tầng trong, không bao giờ ngược lại. Domain hoàn toàn không phụ thuộc framework – về lý thuyết có thể đem domain sang chạy trên runtime khác.")
add_image(doc, "image3.png", width_cm=14)
add_caption(doc, "Hình 2.4: Bốn tầng DDD áp dụng cho mọi service")

add_table_caption(doc, "Bảng 2.1: Trách nhiệm của bốn tầng")
add_table(doc,
    ["Tầng", "Vai trò", "Ví dụ thành phần"],
    [
        ["interfaces", "Adapter inbound – nơi nhận request từ ngoài", "REST Controller, gRPC Service, Kafka Consumer"],
        ["application", "Use case + orchestration + DTO", "Command Handler, Query Handler, Service Facade"],
        ["domain", "Logic nghiệp vụ thuần – không framework", "Aggregate, Entity, Value Object, Domain Event, Repository interface"],
        ["infrastructure", "Adapter outbound – cài đặt cụ thể", "JPA Repository, Kafka Producer, Redis Adapter, JWT provider"],
    ],
    col_widths_cm=[3.0, 5.0, 8.0])

add_heading(doc, "2.3. Cấu trúc thư mục chuẩn của một service", level=2)
add_para(doc, "Quy ước cấu trúc thư mục áp dụng đồng nhất cho 13 service nghiệp vụ. Sự khác nhau duy nhất là phạm vi domain (auth có 6 aggregate, catalog có 4, cart có 1, …).")
add_code(doc, """auth-service/                              # ví dụ cho auth-service
├── pom.xml
├── src/main/java/com/hieu/auth_service/
│   ├── AuthServiceApplication.java        # @SpringBootApplication
│   ├── config/                            # Bean config, Security, CORS, Kafka
│   ├── interfaces/                        # ↘ Tầng giao tiếp (Adapters)
│   │   ├── rest/                          #   REST Controllers
│   │   ├── grpc/                          #   gRPC server (nếu có)
│   │   └── kafka/                         #   Consumers + Producers
│   ├── application/                       # ↘ Use case + DTO
│   │   ├── command/                       #   Command + Handler
│   │   ├── query/                         #   Query + Handler
│   │   └── dto/                           #   Request/Response DTO
│   ├── domain/                            # ↘ Pure domain
│   │   ├── models/                        #   Aggregate, Entity, VO
│   │   ├── events/                        #   Domain events
│   │   ├── repositories/                  #   Repository interfaces
│   │   └── services/                      #   Domain services
│   ├── infrastructure/                    # ↘ Cài đặt cụ thể
│   │   ├── persistence/jpa/               #   JPA entities + repos
│   │   ├── security/                      #   JWT provider, BCrypt
│   │   ├── messaging/                     #   Kafka producer impl
│   │   └── cache/                         #   Redis adapter
│   └── exceptions/                        # Global exception handler
└── src/main/resources/
    ├── application.yaml
    ├── db/migration/                      # Flyway V1__init.sql, …
    └── proto/                             # gRPC .proto files""")
add_caption(doc, "Hình 2.5: Cấu trúc thư mục chuẩn theo DDD")

add_heading(doc, "2.4. Topology hệ thống (15 services)", level=2)
add_table_caption(doc, "Bảng 2.2: Danh sách service và cổng")
add_table(doc,
    ["STT", "Service", "Port REST", "Port gRPC", "Database", "Kafka role"],
    [
        ["1", "eureka-server", "8761", "—", "—", "—"],
        ["2", "api-gateway", "8080", "—", "—", "—"],
        ["3", "auth-service", "9091", "9191", "postgres-auth", "Producer"],
        ["4", "user-profile-service", "9099", "—", "postgres-user", "Both"],
        ["5", "catalog-service", "9093", "9193", "postgres-catalog", "Producer"],
        ["6", "inventory-service", "9098", "—", "postgres-inventory", "Consumer"],
        ["7", "cart-service", "9094", "9194", "postgres-cart", "Consumer"],
        ["8", "order-service", "9095", "—", "postgres-order", "Both (Outbox)"],
        ["9", "payment-service", "8086", "—", "postgres-payment", "Producer"],
        ["10", "shipping-service", "8087", "—", "postgres-shipping", "Both"],
        ["11", "voucher-service", "8094", "—", "postgres-voucher", "—"],
        ["12", "flash-sale-service", "8089", "—", "postgres-flashsale", "—"],
        ["13", "notification-service", "8090", "—", "postgres-noti", "Consumer"],
        ["14", "search-service", "8092", "—", "—", "Consumer"],
        ["15", "analytics-service", "8095", "—", "—", "Consumer"],
    ],
    col_widths_cm=[1.0, 4.0, 2.0, 2.0, 3.5, 3.0])

add_heading(doc, "2.5. Stack công nghệ", level=2)
add_table_caption(doc, "Bảng 2.3: Stack công nghệ cốt lõi")
add_table(doc,
    ["Hạng mục", "Công nghệ", "Phiên bản", "Vai trò"],
    [
        ["Ngôn ngữ", "Java", "25 (LTS)", "Runtime cho toàn bộ service"],
        ["Framework", "Spring Boot", "4.0.5", "Web, JPA, Security, Cache, Actuator"],
        ["Cloud", "Spring Cloud", "2025.0.0", "Gateway, Eureka client, LoadBalancer"],
        ["Build", "Maven", "3.9+", "Build, test, package, dependency mgmt"],
        ["DB", "PostgreSQL", "16", "11 database (một per service)"],
        ["Migration", "Flyway", "10+", "Quản lý schema migration"],
        ["Cache/Session", "Redis", "7", "Idempotency key, blacklist token, rate limit, flash-sale quota"],
        ["Messaging", "Apache Kafka", "3.7 KRaft", "Event bus – 23 topics"],
        ["Search", "Elasticsearch", "8", "Full-text product search + log store"],
        ["Log shipper", "Filebeat", "8", "Ship log từ container lên Elasticsearch"],
        ["Log viewer", "Kibana", "8", "Dashboard log + alert"],
        ["Mail (dev)", "MailHog", "1.0.1", "SMTP giả lập cho dev"],
        ["RPC", "gRPC + Protobuf", "1.65+", "Inter-service sync calls"],
        ["Service Discovery", "Netflix Eureka", "Spring Cloud 2025.0", "Service registry"],
        ["Container", "Docker", "26", "Đóng gói và deploy"],
        ["Orchestrator", "Docker Compose / K8s (dự kiến)", "—", "Triển khai local / production"],
        ["Test", "JUnit 5 + Mockito + AssertJ + TestContainers", "—", "Unit + Integration + E2E"],
        ["Payment Gateway", "SePay (VietQR)", "API v1", "Tích hợp thanh toán"],
        ["Shipping", "GHN / GHTK (mock)", "—", "Tracking shipment"],
    ],
    col_widths_cm=[3.0, 4.5, 3.0, 5.5])

doc.add_page_break()

# ============================================================
# CHƯƠNG 3: ĐẶC TẢ YÊU CẦU
# ============================================================
add_heading(doc, "CHƯƠNG 3: ĐẶC TẢ YÊU CẦU", level=1)

add_heading(doc, "3.1. Yêu cầu chức năng theo service", level=2)
add_para(doc, "Phần này tổng hợp các yêu cầu chức năng (FR) của 13 service nghiệp vụ. Mỗi yêu cầu được gán mã định danh (FR-XXX) để tiện tham chiếu trong test plan và code review.")

# Auth FR
add_table_caption(doc, "Bảng 3.1: Yêu cầu chức năng – Auth Service")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-AUTH-01", "Đăng ký tài khoản", "User cung cấp username, email, password – hệ thống tạo user trạng thái ACTIVE, gán ROLE_CUSTOMER, phát event auth.user-registered"],
        ["FR-AUTH-02", "Đăng nhập", "User cung cấp username hoặc email + password; hệ thống phát cặp ACCESS + REFRESH token qua cookie HttpOnly"],
        ["FR-AUTH-03", "Refresh token", "Khi access hết hạn, gọi /refresh với REFRESH cookie – hệ thống phát cặp mới và blacklist refresh cũ (rotation)"],
        ["FR-AUTH-04", "Đăng xuất", "Revoke refresh token, xóa cookie"],
        ["FR-AUTH-05", "Đổi mật khẩu", "Yêu cầu mật khẩu cũ + mới; revoke toàn bộ refresh hiện có"],
        ["FR-AUTH-06", "Quản lý Role/Permission", "Admin CRUD role, permission; gán/thu hồi role cho user"],
        ["FR-AUTH-07", "Kiểm tra trạng thái đăng nhập", "Endpoint /api/v1/auth/me trả AuthenticatedUser từ JWT"],
        ["FR-AUTH-08", "Soft-delete user", "Đánh dấu status=DELETED thay cho xóa cứng để giữ audit trail"],
    ],
    col_widths_cm=[2.5, 4.0, 9.5])

add_table_caption(doc, "Bảng 3.2: Yêu cầu chức năng – Catalog Service")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-CAT-01", "CRUD Product", "Admin thêm/sửa/xóa sản phẩm với thumbnail, images, status"],
        ["FR-CAT-02", "Variant matrix", "Mỗi product có nhiều variant (SKU, giá, sale price, quantity)"],
        ["FR-CAT-03", "Attribute system", "Quản lý attr (size, color…) + attrVal; gán cho variant"],
        ["FR-CAT-04", "Category tree", "Parent-child, slug duy nhất, sortOrder"],
        ["FR-CAT-05", "Public listing", "FE gọi /api/v1/products với cursor pagination + filter"],
        ["FR-CAT-06", "gRPC GetProduct/CheckStock", "Inter-service cần gọi nhanh để validate cart/order"],
        ["FR-CAT-07", "Phát Kafka event", "product-created/updated/deleted/status-changed cho inventory + search"],
    ],
    col_widths_cm=[2.5, 4.5, 9.0])

add_table_caption(doc, "Bảng 3.3: Yêu cầu chức năng – Cart, Order, Inventory")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-CART-01", "Add/Update/Remove item", "Với pre-flight validation qua gRPC tới catalog"],
        ["FR-CART-02", "Idempotency", "Key qua Redis chống duplicate khi double-tap"],
        ["FR-CART-03", "Optimistic locking", "@Version + @Retryable x3 cho concurrent addItem cùng variant"],
        ["FR-CART-04", "Auto-clear", "Sau order.placed, consumer xóa cart của user"],
        ["FR-ORD-01", "Place from cart", "Order-service lấy cart, validate voucher, reserve inventory, tạo payment record"],
        ["FR-ORD-02", "Saga choreography", "Phát order.placed, downstream services tự subscribe"],
        ["FR-ORD-03", "Outbox", "Ghi event vào outbox cùng transaction, scheduler publish"],
        ["FR-ORD-04", "Return request", "Customer submit → admin approve → complete"],
        ["FR-ORD-05", "State machine", "Enforce transition hợp lệ qua method domain"],
        ["FR-INV-01", "Reserve/Confirm/Release", "Hai bước; reservation TTL 30 phút tự release"],
        ["FR-INV-02", "Audit StockMovement", "Append-only, mọi adjust ghi delta + before + after"],
        ["FR-INV-03", "Pessimistic locking", "Khi reserve để chống oversell"],
    ],
    col_widths_cm=[2.5, 4.5, 9.0])

add_table_caption(doc, "Bảng 3.4: Yêu cầu chức năng – Payment, Shipping, User Profile")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-PAY-01", "Tạo Payment", "Order-service POST với amount + method (COD/SEPAY_QR)"],
        ["FR-PAY-02", "Webhook SePay", "Verify HMAC, match orderNumber, mark COMPLETED, phát payment.completed"],
        ["FR-PAY-03", "Refund flow", "Tách 2 bước (request → process) cho audit"],
        ["FR-SHP-01", "Tạo Shipment", "Sau payment.completed, sinh trackingNumber"],
        ["FR-SHP-02", "Public tracking", "GET /tracking/{number} không cần auth"],
        ["FR-SHP-03", "Update status", "Manual (admin) hoặc carrier webhook"],
        ["FR-USR-01", "Auto-create profile", "Consume auth.user-registered → tạo profile rỗng"],
        ["FR-USR-02", "Manage addresses", "Nhiều địa chỉ giao hàng + đặt mặc định"],
    ],
    col_widths_cm=[2.5, 4.5, 9.0])

add_table_caption(doc, "Bảng 3.5: Yêu cầu chức năng – Voucher, Flash Sale, Notification")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-VCH-01", "CRUD voucher", "Admin tạo với type, value, min order, max discount, usage limit"],
        ["FR-VCH-02", "Validate", "Order-service gọi /validate khi checkout – ghi VoucherUsageRecord"],
        ["FR-VCH-03", "Release", "Khi order.cancelled – decrement usedCount"],
        ["FR-VCH-04", "Auto deactivate", "Scheduled job tắt voucher hết hạn"],
        ["FR-FS-01", "Flash sale", "Tạo với quota cứng cho mỗi variant"],
        ["FR-FS-02", "Atomic participate", "Redis Lua script kiểm tra + decrement quota single round-trip"],
        ["FR-FS-03", "Slot release", "Sau 15 phút không complete checkout – slot trả về pool"],
        ["FR-NOTI-01", "Consume events", "Sinh notification cho order/payment/shipping/auth events"],
        ["FR-NOTI-02", "Channel", "In-app + email (SMTP) + SMS (dự kiến)"],
        ["FR-NOTI-03", "SSE realtime", "GET /notifications/stream push event qua connection mở"],
        ["FR-NOTI-04", "Mark read", "Individual hoặc all"],
    ],
    col_widths_cm=[2.5, 4.5, 9.0])

add_table_caption(doc, "Bảng 3.6: Yêu cầu chức năng – Search, Analytics")
add_table(doc,
    ["Mã FR", "Tên yêu cầu", "Mô tả"],
    [
        ["FR-SRC-01", "Index product", "Consume catalog events → ES (analyzer tiếng Việt)"],
        ["FR-SRC-02", "Search + facet", "POST /search full-text + filter + facet aggregation"],
        ["FR-SRC-03", "Autocomplete", "Prefix suggest qua n-gram analyzer"],
        ["FR-SRC-04", "Reindex", "Admin force rebuild khi đổi mapping"],
        ["FR-ANL-01", "KPI summary", "totalRevenue, ordersToday, newCustomers, avgOrderValue"],
        ["FR-ANL-02", "Log explorer", "Query ES theo q, level, service, time range"],
        ["FR-ANL-03", "Revenue chart", "Aggregate theo ngày để vẽ chart"],
    ],
    col_widths_cm=[2.5, 4.5, 9.0])

add_heading(doc, "3.2. Yêu cầu phi chức năng", level=2)
add_table_caption(doc, "Bảng 3.7: Yêu cầu phi chức năng (NFR)")
add_table(doc,
    ["Mã NFR", "Loại", "Yêu cầu", "Tiêu chí đo lường"],
    [
        ["NFR-01", "Hiệu năng", "Latency p95 cho REST GET không quá 200ms (catalog list, product detail)", "Đo qua Actuator metric http.server.requests"],
        ["NFR-02", "Hiệu năng", "gRPC inter-service p95 không quá 50ms", "Đo qua Micrometer"],
        ["NFR-03", "Khả năng mở rộng", "Hỗ trợ scale theo trục nghiệp vụ – mỗi service triển khai nhiều instance", "Eureka discovery + Spring Cloud LB"],
        ["NFR-04", "Khả dụng", "Uptime mục tiêu 99.5% cho production", "Health check + auto-restart"],
        ["NFR-05", "Bảo mật", "Mật khẩu hash BCrypt cost ≥ 12; JWT_SECRET ≥ 32 ký tự", "Enforce ở JwtProperties + UserService"],
        ["NFR-06", "Bảo mật", "Mọi endpoint admin có @PreAuthorize role check", "Code review checklist"],
        ["NFR-07", "Bảo mật", "Webhook payment verify HMAC", "Reject nếu signature không match"],
        ["NFR-08", "Toàn vẹn dữ liệu", "Idempotency cho mọi command đặt hàng/thanh toán/hoàn tiền", "Redis key TTL 24h + UNIQUE constraint"],
        ["NFR-09", "Toàn vẹn", "Outbox đảm bảo at-least-once cho event order", "Bảng outbox_events + scheduler"],
        ["NFR-10", "Khả năng quan sát", "Mỗi request có Correlation ID propagate qua mọi service", "Header X-Request-Id + MDC"],
        ["NFR-11", "Khả năng quan sát", "Mọi service expose /actuator/health + /prometheus", "Kiểm tra qua boot script"],
        ["NFR-12", "Khả năng phục hồi", "Reservation TTL 30 phút tự release nếu không confirm", "Scheduled job mỗi 1 phút"],
        ["NFR-13", "Khả năng phục hồi", "Saga có compensation event cho từng bước", "order.cancelled → release inventory, voucher"],
        ["NFR-14", "Khả năng phục hồi", "Backup Postgres pg_dump hàng đêm, retention 7d + 4w + 6m", "Cron + S3"],
        ["NFR-15", "Tương thích", "REST API trả ApiResponse<T> chuẩn", "Wrapper với success, message, data"],
        ["NFR-16", "Kiểm thử", "Coverage tối thiểu: domain ≥ 80%, application ≥ 70%, overall ≥ 60%", "JaCoCo report"],
        ["NFR-17", "Vận hành", "Boot toàn stack < 5 phút trên máy dev", "Boot script + healthcheck"],
        ["NFR-18", "Tài liệu", "Mỗi service có README + Swagger UI ở /swagger-ui.html", "Springdoc OpenAPI"],
    ],
    col_widths_cm=[1.7, 2.5, 7.0, 5.0])

doc.add_page_break()

# ============================================================
# CHƯƠNG 4: USE CASE
# ============================================================
add_heading(doc, "CHƯƠNG 4: USE CASE", level=1)

add_heading(doc, "4.1. Danh sách Actor", level=2)
add_table_caption(doc, "Bảng 4.1: Danh sách Actor trong hệ thống")
add_table(doc,
    ["Mã", "Actor", "Vai trò"],
    [
        ["A-01", "Khách (Guest)", "Truy cập không đăng nhập – chỉ xem sản phẩm, search"],
        ["A-02", "Khách hàng (Customer)", "Đã đăng nhập – mua hàng, quản lý profile, review"],
        ["A-03", "Quản trị viên (Admin)", "Toàn quyền – CRUD sản phẩm, quản lý đơn, voucher, flash sale, user"],
        ["A-04", "Nhân viên CSKH (Staff)", "Tra cứu đơn, xử lý đổi/trả, hỗ trợ khách"],
        ["A-05", "Cổng thanh toán SePay", "External – gửi webhook khi bank ghi nhận tiền"],
        ["A-06", "Đối tác vận chuyển (GHN/GHTK)", "External – nhận tracking + push status"],
        ["A-07", "Scheduled Job", "Internal – release reservation, deactivate voucher, retry outbox"],
    ],
    col_widths_cm=[1.5, 5.0, 9.5])

add_heading(doc, "4.2. Tổng hợp Use Case", level=2)
add_image(doc, "image9.png", width_cm=15)
add_caption(doc, "Hình 4.1: Sơ đồ use case Khách hàng – storefront")
add_image(doc, "image10.png", width_cm=15)
add_caption(doc, "Hình 4.2: Sơ đồ use case Quản trị viên – admin panel")

add_table_caption(doc, "Bảng 4.2: Tổng hợp Use Case toàn hệ thống")
add_table(doc,
    ["Mã UC", "Tên Use Case", "Actor", "Tiền điều kiện", "Hậu điều kiện"],
    [
        ["UC-01", "Đăng ký tài khoản", "Guest", "Chưa có tài khoản với email/username này", "User được tạo, ROLE_CUSTOMER gán, phát event"],
        ["UC-02", "Đăng nhập", "Customer/Admin", "Đã có tài khoản ACTIVE", "Cặp JWT cookie set, audit LOGIN_SUCCESS"],
        ["UC-03", "Refresh token", "Đã login", "REFRESH cookie hợp lệ", "Cặp mới, REFRESH cũ blacklist"],
        ["UC-04", "Đăng xuất", "Đã login", "—", "Cookie xóa, refresh blacklist"],
        ["UC-05", "Đổi mật khẩu", "Đã login", "Mật khẩu cũ đúng", "Hash mới lưu, mọi refresh revoked"],
        ["UC-06", "Xem trang chủ", "Guest", "—", "Render HTML SSR + danh mục + featured product"],
        ["UC-07", "Duyệt danh mục", "Guest/Customer", "—", "List sản phẩm cursor pagination"],
        ["UC-08", "Tìm kiếm", "Guest/Customer", "Elasticsearch lên", "Ranked list từ ES + facet"],
        ["UC-09", "Xem chi tiết sản phẩm", "Guest/Customer", "Product status=ACTIVE", "Trả product + variant + attr"],
        ["UC-10", "Thêm vào giỏ", "Customer", "Đã login, product ACTIVE, variant đủ stock", "Cart item tạo/cộng dồn, idempotent"],
        ["UC-11", "Áp voucher", "Customer", "Voucher còn hạn, đạt min order", "Discount tính, VoucherUsageRecord tạo"],
        ["UC-12", "Đặt hàng từ giỏ", "Customer", "Cart không rỗng, địa chỉ giao có sẵn", "Order PENDING → INVENTORY_RESERVED → PAYMENT_PENDING; phát order.placed"],
        ["UC-13", "Thanh toán SePay", "Customer", "Order có Payment PENDING", "Bank ghi nhận → webhook → Payment COMPLETED → order CONFIRMED"],
        ["UC-14", "Xem đơn của tôi", "Customer", "Đã login", "List đơn theo cursor"],
        ["UC-15", "Yêu cầu đổi/trả", "Customer", "Đơn DELIVERED, ≤ 30 ngày", "ReturnRequest PENDING; email admin + customer"],
        ["UC-16", "Hủy đơn", "Customer", "Đơn ở trạng thái có thể hủy (PENDING/PAYMENT_PENDING)", "Order CANCELLED; reservation release; voucher release"],
        ["UC-17", "Đánh giá sản phẩm", "Customer", "Đã mua + DELIVERED", "Review record tạo (dự kiến – chưa triển khai)"],
        ["UC-18", "CRUD sản phẩm", "Admin", "Đã login admin", "Product create/update/delete; phát Kafka event"],
        ["UC-19", "CRUD voucher", "Admin", "—", "Voucher create/update/delete; auto deactivate khi hết hạn"],
        ["UC-20", "Tạo flash sale", "Admin", "—", "FlashSale DRAFT → ACTIVE theo schedule"],
        ["UC-21", "Approve đổi/trả", "Admin/Staff", "ReturnRequest PENDING", "Status APPROVED/REJECTED; trigger refund flow"],
        ["UC-22", "Process refund", "Admin", "Refund đã approved", "Gọi SePay API hoàn tiền; mark REFUNDED"],
        ["UC-23", "Tra cứu đơn", "Staff", "—", "Search theo orderNumber, phone, email"],
        ["UC-24", "Xem dashboard KPI", "Admin", "Analytics-service lên", "Hiển thị tổng doanh thu, đơn hôm nay, KH mới"],
        ["UC-25", "Public tracking shipment", "Guest", "Có trackingNumber + phone verify", "Trả status + events của shipment"],
    ],
    col_widths_cm=[1.5, 3.5, 2.5, 4.0, 4.5])

add_heading(doc, "4.3. Chi tiết Use Case tiêu biểu", level=2)

add_para(doc, "UC-02: Đăng nhập", bold=True)
add_table_caption(doc, "Bảng 4.3: Đặc tả Use Case UC-02 – Đăng nhập")
add_table(doc,
    ["Mục", "Mô tả"],
    [
        ["Mã UC", "UC-02"],
        ["Tên", "Đăng nhập"],
        ["Actor", "Customer / Admin / Staff"],
        ["Mô tả", "User cung cấp credential, nhận cặp JWT cookie để dùng cho các request tiếp theo"],
        ["Tiền điều kiện", "Tài khoản tồn tại, status=ACTIVE; chưa bị rate limit (5 lần/phút/IP)"],
        ["Luồng chính", "1. Client POST /api/v1/auth/login {usernameOrEmail, password}\n2. Auth-service tìm user bằng username hoặc email\n3. Kiểm tra status=ACTIVE; nếu LOCKED/DELETED → 403\n4. BCrypt.matches(password, user.passwordHash); fail → 401 + tăng counter Redis\n5. Phát hành ACCESS (TTL 15p) + REFRESH (TTL 7d) ký HS256\n6. Lưu refreshTokenHash vào DB; trả Set-Cookie HttpOnly\n7. Phát Kafka auth.session.events.v1 = LOGIN_SUCCESS"],
        ["Luồng phụ", "Sai mật khẩu → 401; quá 5 lần/phút → 429; status LOCKED → 403"],
        ["Hậu điều kiện", "Browser có ACCESS + REFRESH cookie; audit trail ghi log"],
        ["Exceptions", "Auth-service down → 503; Redis down → fallback đếm bằng DB"],
        ["NFR liên quan", "NFR-01 (latency p95 ≤ 200ms), NFR-05 (BCrypt cost 12)"],
    ],
    col_widths_cm=[3.5, 12.5])

add_para(doc, "UC-12: Đặt hàng từ giỏ", bold=True)
add_table_caption(doc, "Bảng 4.4: Đặc tả Use Case UC-12 – Đặt hàng từ giỏ")
add_table(doc,
    ["Mục", "Mô tả"],
    [
        ["Mã UC", "UC-12"],
        ["Tên", "Đặt hàng từ giỏ"],
        ["Actor", "Customer"],
        ["Mô tả", "Customer hoàn tất checkout từ giỏ – saga đa bước phối hợp 4 service"],
        ["Tiền điều kiện", "Đã đăng nhập; cart không rỗng; địa chỉ giao có sẵn"],
        ["Luồng chính", "1. POST /api/v1/orders/from-cart {addressId, paymentMethod, voucherCode?, idempotencyKey}\n2. Order-service kiểm tra idempotencyKey – nếu đã xử lý, trả response cũ\n3. Lấy cart qua gRPC tới cart-service\n4. Validate voucher qua REST tới voucher-service (ghi VoucherUsageRecord)\n5. Reserve inventory qua REST tới inventory-service – nhận reservationId\n6. Tạo Payment PENDING qua REST tới payment-service – nhận payUrl + qrCodeUrl\n7. INSERT order PENDING + outbox event order.placed (cùng transaction)\n8. Scheduler poll outbox → publish Kafka\n9. Cart-service consume order.placed → xóa cart"],
        ["Luồng phụ", "Voucher invalid → 400; inventory không đủ → 409 + release voucher; payment-service down → rollback reservation + voucher"],
        ["Hậu điều kiện", "Order PENDING, reservation HOLD, voucher used, payment record PENDING"],
        ["Exceptions", "Idempotency hit → trả response cũ; database conflict → 409"],
        ["NFR liên quan", "NFR-08 (idempotency), NFR-09 (outbox at-least-once), NFR-13 (compensation)"],
    ],
    col_widths_cm=[3.5, 12.5])

add_para(doc, "UC-15: Yêu cầu đổi/trả hàng", bold=True)
add_table_caption(doc, "Bảng 4.5: Đặc tả Use Case UC-15 – Yêu cầu đổi/trả")
add_table(doc,
    ["Mục", "Mô tả"],
    [
        ["Mã UC", "UC-15"],
        ["Tên", "Yêu cầu đổi/trả hàng"],
        ["Actor", "Customer (yêu cầu); Admin/Staff (xử lý)"],
        ["Mô tả", "Customer submit yêu cầu sau khi nhận hàng – admin/staff review và phê duyệt"],
        ["Tiền điều kiện", "Đơn DELIVERED, chưa quá 30 ngày kể từ ngày giao"],
        ["Luồng chính", "1. POST /api/v1/orders/return-requests/{orderId} {reason, description, contactEmail}\n2. Order-service kiểm tra điều kiện + tạo ReturnRequest PENDING\n3. Phát Kafka order.return-requested\n4. Notification-service consume → email cho admin + customer\n5. Admin POST /approve hoặc /reject (kèm note)\n6. Nếu APPROVED – customer ship hàng về kho\n7. Admin POST /complete sau khi nhận hàng → trigger refund + restock"],
        ["Luồng phụ", "Quá 30 ngày → 400; đơn không DELIVERED → 400; reject → status REJECTED"],
        ["Hậu điều kiện", "ReturnRequest có status cuối (APPROVED/REJECTED/COMPLETED); refund phát động nếu approved"],
        ["NFR liên quan", "NFR-08 (idempotency), NFR-13 (saga compensation)"],
    ],
    col_widths_cm=[3.5, 12.5])

doc.add_page_break()

# ============================================================
# CHƯƠNG 5: THIẾT KẾ CƠ SỞ DỮ LIỆU
# ============================================================
add_heading(doc, "CHƯƠNG 5: THIẾT KẾ CƠ SỞ DỮ LIỆU", level=1)

add_heading(doc, "5.1. Tổng quan database-per-service", level=2)
add_para(doc, "Hệ thống áp dụng nguyên tắc Database per Service – mỗi service nghiệp vụ có cơ sở dữ liệu PostgreSQL riêng, đảm bảo loose coupling và không có foreign key cross-service. Tham chiếu giữa các service là logical (chỉ lưu ID dạng UUID hoặc Long), không enforce ở DB layer.")
add_table_caption(doc, "Bảng 5.1: Danh sách cơ sở dữ liệu")
add_table(doc,
    ["DB Name", "Service sở hữu", "Cổng host", "Số bảng chính"],
    [
        ["auth_db", "auth-service", "5433", "6 (users, roles, permissions, user_roles, role_permissions, refresh_tokens)"],
        ["user_db", "user-profile-service", "5434", "2 (user_profiles, addresses)"],
        ["catalog_db", "catalog-service", "5435", "6 (categories, products, variants, attrs, attr_vals, variant_attrs)"],
        ["inventory_db", "inventory-service", "5436", "3 (inventories, stock_reservations, stock_movements)"],
        ["cart_db", "cart-service", "5437", "1 (cart_items)"],
        ["order_db", "order-service", "5438", "5 (orders, order_items, return_requests, order_idempotency, outbox_events)"],
        ["payment_db", "payment-service", "5439", "2 (payments, refund_requests)"],
        ["shipping_db", "shipping-service", "5440", "2 (shipments, shipment_events)"],
        ["voucher_db", "voucher-service", "5441", "2 (vouchers, voucher_usage)"],
        ["flashsale_db", "flash-sale-service", "5442", "3 (flash_sales, flash_sale_items, flash_sale_participations)"],
        ["notification_db", "notification-service", "5443", "1 (notifications)"],
    ],
    col_widths_cm=[3.0, 4.5, 2.0, 6.5])
add_para(doc, "Ghi chú: Trong môi trường dev, có thể dùng H2 in-memory cho khởi động nhanh. Production luôn dùng PostgreSQL 16 qua Docker.", italic=True)

add_heading(doc, "5.2. Auth Service Database", level=2)
add_para(doc, "Cơ sở dữ liệu auth_db gồm 6 bảng chính. Mô hình RBAC chuẩn: user N-N role, role N-N permission.")

add_para(doc, "Bảng users:", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", "Định danh user"],
        ["username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Username đăng nhập"],
        ["email", "VARCHAR(120)", "UNIQUE, NOT NULL", "Email"],
        ["password_hash", "VARCHAR(80)", "NOT NULL", "BCrypt hash"],
        ["status", "VARCHAR(20)", "NOT NULL DEFAULT 'ACTIVE'", "ACTIVE | LOCKED | DELETED"],
        ["created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", "Thời điểm tạo"],
        ["updated_at", "TIMESTAMP", "NOT NULL", "Cập nhật cuối"],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_para(doc, "Bảng refresh_tokens (rotation):", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", ""],
        ["user_id", "UUID", "FK users(id), NOT NULL", "Chủ sở hữu"],
        ["token_hash", "VARCHAR(255)", "UNIQUE, NOT NULL", "SHA-256 hash của refresh token"],
        ["expires_at", "TIMESTAMP", "NOT NULL", "Thời điểm hết hạn"],
        ["revoked", "BOOLEAN", "NOT NULL DEFAULT FALSE", "Đã thu hồi chưa"],
        ["created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", "Thời điểm tạo"],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])
add_para(doc, "Các bảng còn lại (roles, permissions, user_roles, role_permissions) giữ chuẩn RBAC – user_roles và role_permissions là bảng join N-N với composite PK.")

add_heading(doc, "5.3. Catalog Service Database", level=2)
add_para(doc, "Catalog gộp danh mục, sản phẩm, biến thể, thuộc tính. 6 bảng: categories, products, variants, attrs, attr_vals, variant_attrs.")

add_para(doc, "Bảng products (SPU):", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "BIGSERIAL", "PK", ""],
        ["name", "VARCHAR(255)", "NOT NULL", "Tên sản phẩm"],
        ["slug", "VARCHAR(255)", "UNIQUE, NOT NULL", "Slug SEO"],
        ["description", "TEXT", "", "Mô tả dài"],
        ["brand", "VARCHAR(100)", "", "Nhãn hàng"],
        ["category_id", "BIGINT", "FK categories(id)", "Danh mục"],
        ["thumbnail", "TEXT", "", "URL ảnh đại diện"],
        ["images", "JSONB", "", "Mảng URL ảnh khác"],
        ["status", "VARCHAR(20)", "NOT NULL DEFAULT 'DRAFT'", "DRAFT | ACTIVE | ARCHIVED | DELETED"],
        ["created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", ""],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_para(doc, "Bảng variants (SKU):", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "BIGSERIAL", "PK", ""],
        ["product_id", "BIGINT", "FK products(id), NOT NULL", "Sản phẩm cha"],
        ["sku", "VARCHAR(100)", "UNIQUE, NOT NULL", "Mã SKU"],
        ["price", "NUMERIC(15,2)", "NOT NULL", "Giá niêm yết"],
        ["sale_price", "NUMERIC(15,2)", "", "Giá khuyến mãi (NULL = không sale)"],
        ["quantity", "INTEGER", "NOT NULL DEFAULT 0", "Tồn kho (cache, source of truth ở inventory_db)"],
        ["image", "TEXT", "", "Ảnh variant"],
        ["status", "VARCHAR(20)", "NOT NULL", "ACTIVE | INACTIVE"],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_heading(doc, "5.4. Order & Payment Database", level=2)
add_para(doc, "Order là aggregate phức tạp nhất – có 5 bảng. Đáng chú ý: outbox_events đảm bảo at-least-once publish, order_idempotency chống duplicate, return_requests tách riêng cho audit.")

add_para(doc, "Bảng orders:", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", ""],
        ["order_number", "VARCHAR(20)", "UNIQUE, NOT NULL", "ORD-YYYYMMDD-NNNNNN"],
        ["user_id", "UUID", "NOT NULL", "User đặt"],
        ["status", "VARCHAR(30)", "NOT NULL", "PENDING|INVENTORY_RESERVED|PAYMENT_PENDING|PAID|CONFIRMED|SHIPPED|DELIVERED|CANCELLED|FAILED|RETURN_REQUESTED|RETURN_APPROVED|RETURNED"],
        ["subtotal_amount", "NUMERIC(15,2)", "NOT NULL", "Tổng giá trị item"],
        ["discount_amount", "NUMERIC(15,2)", "NOT NULL DEFAULT 0", "Giảm từ voucher"],
        ["shipping_fee", "NUMERIC(15,2)", "NOT NULL DEFAULT 0", "Phí ship"],
        ["total_amount", "NUMERIC(15,2)", "NOT NULL", "Tổng thanh toán"],
        ["voucher_code", "VARCHAR(50)", "", ""],
        ["recipient_name", "VARCHAR(255)", "", ""],
        ["recipient_phone", "VARCHAR(20)", "", ""],
        ["shipping_address", "TEXT", "", ""],
        ["payment_id", "UUID", "", "Liên kết payment"],
        ["reservation_id", "UUID", "", "Liên kết inventory"],
        ["shipment_id", "UUID", "", "Liên kết shipping"],
        ["idempotency_key", "VARCHAR(64)", "UNIQUE", ""],
        ["created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", ""],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_para(doc, "Bảng outbox_events (at-least-once publish):", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "BIGSERIAL", "PK", ""],
        ["aggregate_type", "VARCHAR(50)", "NOT NULL", "ORDER, RETURN, …"],
        ["aggregate_id", "UUID", "NOT NULL", ""],
        ["event_type", "VARCHAR(60)", "NOT NULL", "order.placed, order.confirmed…"],
        ["payload", "JSONB", "NOT NULL", "Nội dung event"],
        ["created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", ""],
        ["published_at", "TIMESTAMP", "", "NULL nếu chưa publish"],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_para(doc, "Bảng payments:", bold=True)
add_table(doc,
    ["Cột", "Kiểu", "Ràng buộc", "Mô tả"],
    [
        ["id", "UUID", "PK", ""],
        ["order_id", "UUID", "NOT NULL", ""],
        ["amount", "NUMERIC(15,2)", "NOT NULL", ""],
        ["currency", "VARCHAR(3)", "NOT NULL DEFAULT 'VND'", ""],
        ["method", "VARCHAR(20)", "NOT NULL", "COD | SEPAY_QR | BANK_TRANSFER"],
        ["status", "VARCHAR(20)", "NOT NULL", "PENDING|COMPLETED|FAILED|REFUNDED|CANCELLED"],
        ["pay_url", "TEXT", "", "Link thanh toán"],
        ["qr_code_url", "TEXT", "", ""],
        ["external_ref", "VARCHAR(120)", "UNIQUE", "Webhook idempotency"],
        ["completed_at", "TIMESTAMP", "", ""],
    ],
    col_widths_cm=[3.0, 2.5, 4.5, 6.0])

add_heading(doc, "5.5. Inventory & Reservation Database", level=2)
add_para(doc, "Inventory tách 3 bảng: cân bằng hiện tại (inventories), reservation tạm (stock_reservations), audit append-only (stock_movements).")
add_table_caption(doc, "Bảng 5.2: Tóm tắt 3 bảng của inventory_db")
add_table(doc,
    ["Bảng", "Vai trò", "Chú ý"],
    [
        ["inventories", "Cân bằng tồn kho hiện tại theo SKU", "Có column available = quantity - reserved; @Version cho optimistic lock"],
        ["stock_reservations", "Đặt giữ kho tạm 30 phút", "TTL → scheduled job release"],
        ["stock_movements", "Audit append-only mọi adjust", "Source of truth – quantity có thể rebuild từ history"],
    ],
    col_widths_cm=[4.0, 6.0, 6.0])

add_heading(doc, "5.6. Các bảng còn lại", level=2)
add_table_caption(doc, "Bảng 5.3: Tóm tắt schema các service nhỏ hơn")
add_table(doc,
    ["Service", "Bảng chính", "Mô tả"],
    [
        ["user-profile", "user_profiles, addresses", "Hồ sơ + nhiều địa chỉ; default_address_id"],
        ["cart", "cart_items", "Một row mỗi (user, variant); @Version cho optimistic lock"],
        ["voucher", "vouchers, voucher_usage", "voucher_usage UNIQUE(voucher_id, order_id) – idempotent"],
        ["flash_sale", "flash_sales, flash_sale_items, flash_sale_participations", "soldQuota cập nhật atomic qua Redis"],
        ["shipping", "shipments, shipment_events", "events lưu lịch sử status change"],
        ["notification", "notifications", "Có channel + status + isRead; index theo (userId, createdAt)"],
    ],
    col_widths_cm=[3.5, 5.5, 7.0])

add_heading(doc, "5.7. Chỉ mục (Indexes)", level=2)
add_table_caption(doc, "Bảng 5.4: Các chỉ mục quan trọng")
add_table(doc,
    ["Bảng", "Chỉ mục", "Mục đích"],
    [
        ["users", "UNIQUE(username), UNIQUE(email)", "Lookup khi đăng nhập"],
        ["refresh_tokens", "UNIQUE(token_hash)", "Validate refresh nhanh"],
        ["refresh_tokens", "INDEX(user_id)", "List token theo user"],
        ["products", "UNIQUE(slug)", "Lookup theo slug FE"],
        ["products", "INDEX(category_id, status)", "Listing theo category"],
        ["variants", "UNIQUE(sku)", "Lookup nhanh"],
        ["orders", "UNIQUE(order_number)", "Tra cứu"],
        ["orders", "UNIQUE(idempotency_key)", "Anti-duplicate"],
        ["orders", "INDEX(user_id, created_at DESC, id)", "Cursor pagination /orders/my"],
        ["outbox_events", "INDEX(published_at) WHERE published_at IS NULL", "Scheduler quét pending events nhanh"],
        ["voucher_usage", "UNIQUE(voucher_id, order_id)", "Idempotent apply"],
        ["payments", "UNIQUE(external_ref)", "Webhook idempotency"],
        ["inventories", "UNIQUE(sku)", "Lookup tồn kho"],
        ["stock_reservations", "INDEX(expires_at) WHERE status='PENDING'", "Scheduler release expired"],
        ["notifications", "INDEX(user_id, created_at DESC)", "Feed cursor pagination"],
    ],
    col_widths_cm=[3.5, 6.0, 6.5])

doc.add_page_break()

# ============================================================
# CHƯƠNG 6: THIẾT KẾ API
# ============================================================
add_heading(doc, "CHƯƠNG 6: THIẾT KẾ API", level=1)

add_heading(doc, "6.1. Quy ước chung", level=2)
for item in [
    "Base URL: http://localhost:8080 (qua API Gateway).",
    "Content-Type: application/json.",
    "Authentication: cookie HttpOnly (ACCESS, REFRESH) – không dùng Authorization header trên browser; tuỳ chọn Bearer cho gRPC client hoặc Postman.",
    "Response format: ApiResponse<T> = { success: boolean, message: string, data: T }.",
    "Error format: { code: 'AUTH-401', status: 401, message: '...', timestamp: '2026-05-14T08:23:00Z' }.",
    "Pagination: cursor-based qua query param ?cursor=&size=20 (mặc định) – trả {items, nextCursor, pageSize, totalElements}.",
    "Idempotency: header X-Idempotency-Key (UUID v4) bắt buộc cho POST /orders, /payments, /refunds.",
    "Versioning: hiện tại chưa versioning; dự kiến /api/v1/... trong roadmap.",
    "Correlation ID: header X-Request-Id từ Gateway propagate qua mọi service – nếu client không gửi, Gateway tự generate.",
]:
    add_bullet(doc, item)

add_heading(doc, "6.2. Auth Service API", level=2)
add_para(doc, "API công khai (không yêu cầu JWT):", bold=True)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/auth/register", "Đăng ký user mới"],
        ["POST", "/api/v1/auth/login", "Đăng nhập, trả Set-Cookie ACCESS + REFRESH"],
        ["POST", "/api/v1/auth/refresh", "Cấp ACCESS mới từ REFRESH cookie"],
    ],
    col_widths_cm=[2.0, 6.0, 8.0])
add_para(doc, "API bảo mật (yêu cầu cookie ACCESS):", bold=True)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["GET", "/api/v1/auth/me", "Trả AuthenticatedUser của session hiện tại"],
        ["POST", "/api/v1/auth/logout", "Revoke refresh, xóa cookie"],
        ["POST", "/api/v1/auth/password/change", "Đổi mật khẩu (yêu cầu mật khẩu cũ)"],
        ["GET", "/api/v1/auth/users", "Admin – list user phân trang"],
        ["PATCH", "/api/v1/auth/users/{id}/lock", "Admin – khóa user"],
        ["POST", "/api/v1/auth/roles", "Admin – tạo role"],
        ["POST", "/api/v1/auth/permissions", "Admin – tạo permission"],
        ["POST", "/api/v1/auth/users/{id}/roles", "Admin – gán role cho user"],
        ["POST", "/api/v1/auth/roles/{id}/permissions", "Admin – gán permission cho role"],
    ],
    col_widths_cm=[2.0, 6.5, 7.5])

add_heading(doc, "6.3. User Profile API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["GET", "/api/v1/user-profiles/me", "Profile của user đang đăng nhập"],
        ["PATCH", "/api/v1/user-profiles/me", "Update profile (tên, sinh nhật, avatar)"],
        ["GET", "/api/v1/user-profiles/me/addresses", "Danh sách địa chỉ"],
        ["POST", "/api/v1/user-profiles/me/addresses", "Thêm địa chỉ"],
        ["PATCH", "/api/v1/user-profiles/me/addresses/{id}", "Sửa địa chỉ"],
        ["DELETE", "/api/v1/user-profiles/me/addresses/{id}", "Xóa địa chỉ"],
        ["POST", "/api/v1/user-profiles/me/addresses/{id}/default", "Đặt làm địa chỉ mặc định"],
        ["GET", "/api/v1/user-profiles/by-user/{userId}", "(Internal) Lookup profile theo userId"],
        ["GET", "/api/v1/user-profiles/by-email/{email}", "(Internal) Lookup theo email"],
    ],
    col_widths_cm=[2.0, 7.0, 7.0])

add_heading(doc, "6.4. Catalog API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["GET", "/api/v1/products", "List với cursor + filter (category, brand, price range)"],
        ["GET", "/api/v1/products/{id}", "Chi tiết product + variants"],
        ["GET", "/api/v1/products/by-slug/{slug}", "Chi tiết theo slug (SEO)"],
        ["POST", "/api/v1/products", "Admin – tạo product mới"],
        ["PATCH", "/api/v1/products/{id}", "Admin – sửa"],
        ["DELETE", "/api/v1/products/{id}", "Admin – soft-delete"],
        ["GET", "/api/v1/categories", "Tree danh mục"],
        ["POST", "/api/v1/categories", "Admin – tạo danh mục"],
        ["GET", "/api/v1/attrs", "Danh sách attr"],
        ["POST", "/api/v1/attrs", "Admin – tạo attr"],
        ["GET", "/api/v1/variants/by-sku/{sku}", "Lookup variant"],
    ],
    col_widths_cm=[2.0, 6.5, 7.5])

add_heading(doc, "6.5. Inventory API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/inventory/", "Admin – tạo bản ghi inventory"],
        ["GET", "/api/v1/inventory/{productId}", "Tồn kho theo productId"],
        ["GET", "/api/v1/inventory/sku/{sku}", "Tồn kho theo SKU"],
        ["PATCH", "/api/v1/inventory/sku/{sku}/set-stock", "Admin – set tồn kho mới (yêu cầu reason)"],
        ["POST", "/api/v1/inventory/reservations", "(Internal) Reserve cho saga"],
        ["POST", "/api/v1/inventory/reservations/{id}/confirm", "(Internal) Confirm reservation"],
        ["POST", "/api/v1/inventory/reservations/{id}/release", "(Internal) Release reservation"],
        ["GET", "/api/v1/inventory/movements", "Audit – list movements"],
    ],
    col_widths_cm=[2.0, 7.0, 7.0])

add_heading(doc, "6.6. Cart API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["GET", "/api/v1/cart", "Cart hiện tại + warnings"],
        ["POST", "/api/v1/cart/items", "Add or upsert item (X-Idempotency-Key required)"],
        ["PUT", "/api/v1/cart/items/{variantId}", "Update qty (qty=0 = delete)"],
        ["DELETE", "/api/v1/cart/items/{variantId}", "Xóa item"],
        ["DELETE", "/api/v1/cart", "Clear toàn bộ"],
    ],
    col_widths_cm=[2.0, 6.5, 7.5])

add_heading(doc, "6.7. Order API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/orders", "Admin – tạo đơn manual với items inline"],
        ["POST", "/api/v1/orders/from-cart", "Customer – đặt từ giỏ (X-Idempotency-Key required)"],
        ["GET", "/api/v1/orders/{id}", "Chi tiết (owner hoặc admin)"],
        ["GET", "/api/v1/orders/my", "Customer – đơn của tôi (cursor)"],
        ["GET", "/api/v1/orders", "Admin – tất cả đơn (cursor + filter status)"],
        ["POST", "/api/v1/orders/{id}/cancel", "Customer – hủy đơn ở trạng thái cho phép"],
        ["POST", "/api/v1/orders/return-requests/{orderId}", "Customer – submit yêu cầu đổi/trả"],
        ["GET", "/api/v1/orders/return-requests", "Admin – list return requests"],
        ["POST", "/api/v1/orders/return-requests/{id}/approve", "Admin – approve"],
        ["POST", "/api/v1/orders/return-requests/{id}/reject", "Admin – reject"],
        ["POST", "/api/v1/orders/return-requests/{id}/complete", "Admin – mark completed"],
    ],
    col_widths_cm=[2.0, 7.5, 6.5])

add_heading(doc, "6.8. Payment API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/payments", "(Internal) Order-service tạo Payment PENDING"],
        ["POST", "/api/v1/payments/sepay/webhook", "Webhook từ SePay – verify HMAC"],
        ["GET", "/api/v1/payments/{id}", "Chi tiết"],
        ["GET", "/api/v1/payments/order/{orderId}", "Payment của order"],
        ["POST", "/api/v1/payments/{id}/confirm", "Admin – manual confirm"],
        ["POST", "/api/v1/payments/{id}/cancel", "Cancel khi user hủy"],
        ["POST", "/api/v1/payments/{id}/refund", "Tạo refund request (PENDING)"],
        ["POST", "/api/v1/payments/{id}/process-refund", "Admin – process refund qua SePay API"],
    ],
    col_widths_cm=[2.0, 7.5, 6.5])

add_heading(doc, "6.9. Shipping API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/shipments", "Admin – tạo shipment"],
        ["POST", "/api/v1/shipments/internal", "(Internal) Order-service tạo sau payment.completed"],
        ["GET", "/api/v1/shipments/{id}", "Chi tiết"],
        ["GET", "/api/v1/shipments/order/{orderId}", "Shipment của order"],
        ["PATCH", "/api/v1/shipments/{id}/status", "Admin – update status"],
        ["GET", "/tracking/{number}", "Public tracking (yêu cầu phone verify)"],
    ],
    col_widths_cm=[2.0, 7.0, 7.0])

add_heading(doc, "6.10. Voucher API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/vouchers", "Admin – tạo"],
        ["GET", "/api/v1/vouchers", "List all (admin)"],
        ["GET", "/api/v1/vouchers/active", "List active (FE hiển thị)"],
        ["GET", "/api/v1/vouchers/code/{code}", "FE preview voucher (không apply)"],
        ["POST", "/api/v1/vouchers/validate", "(Internal) Order-service gọi khi checkout"],
        ["POST", "/api/v1/vouchers/release", "(Internal) Khi order cancelled"],
        ["PATCH", "/api/v1/vouchers/{id}", "Admin – sửa"],
        ["DELETE", "/api/v1/vouchers/{id}", "Admin – soft-delete"],
    ],
    col_widths_cm=[2.0, 6.5, 7.5])

add_heading(doc, "6.11. Flash Sale API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/flash-sales", "Admin – tạo flash sale"],
        ["GET", "/api/v1/flash-sales", "List all"],
        ["GET", "/api/v1/flash-sales/active", "Đang chạy"],
        ["GET", "/api/v1/flash-sales/{id}", "Chi tiết"],
        ["GET", "/api/v1/flash-sales/{id}/availability", "Số lượng còn lại realtime"],
        ["POST", "/api/v1/flash-sales/{id}/participate", "User – claim slot (atomic Lua)"],
        ["POST", "/api/v1/flash-sales/{id}/activate", "Admin – activate"],
        ["POST", "/api/v1/flash-sales/{id}/end", "Admin – kết thúc sớm"],
    ],
    col_widths_cm=[2.0, 7.0, 7.0])

add_heading(doc, "6.12. Notification API", level=2)
add_table(doc,
    ["Method", "Path", "Mục đích"],
    [
        ["POST", "/api/v1/notifications/send", "Admin – gửi thông báo tới user/group"],
        ["GET", "/api/v1/notifications/my", "Offset paginated"],
        ["GET", "/api/v1/notifications/my/feed", "Cursor paginated cho infinite scroll"],
        ["GET", "/api/v1/notifications/stream", "SSE realtime push (keep-alive 30s)"],
        ["POST", "/api/v1/notifications/{id}/read", "Mark read individual"],
        ["POST", "/api/v1/notifications/read-all", "Mark read all"],
        ["GET", "/api/v1/notifications/unread-count", "Số notification chưa đọc"],
    ],
    col_widths_cm=[2.0, 7.0, 7.0])

doc.add_page_break()

# ============================================================
# CHƯƠNG 7: LUỒNG XỬ LÝ
# ============================================================
add_heading(doc, "CHƯƠNG 7: LUỒNG XỬ LÝ", level=1)
add_para(doc, "Chương này tổng hợp các luồng xử lý nghiệp vụ quan trọng dưới hai dạng: hình ảnh sequence diagram (đã có) và sơ đồ ASCII bổ sung để tiện diff trong git.")

add_heading(doc, "7.1. Luồng đăng nhập và phát hành JWT", level=2)
add_para(doc, "Auth-service phát hành hai token: ACCESS_TOKEN (TTL 15 phút) và REFRESH_TOKEN (TTL 7 ngày). Cả hai ký HS256 với secret chia sẻ giữa các service – biến môi trường JWT_SECRET. Mỗi service tự xác thực JWT trong filter chain, không cần gọi ngược về auth-service.")
add_image(doc, "image4.png", width_cm=15)
add_caption(doc, "Hình 7.1: Sequence diagram đăng nhập + phát hành JWT")

add_code(doc, """Client          API Gateway          Auth Service          Database          Kafka
  |                  |                     |                    |                |
  |-- POST /login -->|                     |                    |                |
  |                  |-- Forward --------->|                    |                |
  |                  |                     |-- Find user ----->|                |
  |                  |                     |<-- User row ------|                |
  |                  |                     |-- Validate status  |                |
  |                  |                     |-- BCrypt.matches   |                |
  |                  |                     |-- Generate ACCESS  |                |
  |                  |                     |   (TTL 15p, jti)   |                |
  |                  |                     |-- Generate REFRESH |                |
  |                  |                     |   (TTL 7d)         |                |
  |                  |                     |-- Save refresh -->|                |
  |                  |                     |-- Publish LOGIN_SUCCESS ----------->|
  |                  |<-- 200 + cookies ---|                    |                |
  |<-- Set-Cookie ---|                     |                    |                |
""")
add_caption(doc, "Hình 7.2: Luồng đăng nhập – ASCII")

add_heading(doc, "7.2. Luồng xử lý request qua API Gateway", level=2)
add_para(doc, "API Gateway là điểm vào duy nhất, áp dụng các filter (CORS, path routing, trace ID, rate limit). Sau Gateway, mỗi service tự xác thực JWT qua local filter, không cần round-trip ngược về auth-service.")
add_code(doc, """Client         API Gateway       JWT Filter        Downstream Service
  |                 |                |                     |
  |--GET /api/v1/xxx ->|                |                     |
  | (Cookie ACCESS) |-- Route match->|                     |
  |                 |-- Add trace ID-|                     |
  |                 |                |-- Extract cookie    |
  |                 |                |-- Verify signature  |
  |                 |                |-- Check expiry      |
  |                 |                |-- Extract claims    |
  |                 |<-- Add headers-|                     |
  |                 |  X-User-Id     |                     |
  |                 |  X-Username    |                     |
  |                 |  X-Request-Id  |                     |
  |                 |-- Forward ------------------------- >|
  |                 |                |                     |-- local JWT verify
  |                 |                |                     |-- @PreAuthorize check
  |                 |                |                     |-- Process logic
  |                 |<-- Response ------------------------ |
  |<-- 200 + data --|                |                     |
""")
add_caption(doc, "Hình 7.3: Luồng xử lý request qua Gateway")

add_heading(doc, "7.3. Luồng addItem trong Cart Service", level=2)
add_para(doc, "Cart-service là điểm tích hợp đầu tiên của saga checkout. Validation gắt: mỗi addItem/updateItem đều round-trip tới catalog qua gRPC để xác thực status + stock thực tế. Hai lớp validation (pre-flight gRPC + BE re-check) đảm bảo stale tab không push được state sai.")
add_image(doc, "image6.png", width_cm=15)
add_caption(doc, "Hình 7.4: Pipeline xác thực addItem trong cart-service")

add_code(doc, """Client          Cart Service          Redis           Catalog Service (gRPC)        Cart DB
  |                    |                  |                       |                       |
  |-- POST addItem --->|                  |                       |                       |
  |  X-Idem-Key: k     |                  |                       |                       |
  |                    |-- GET cart:idem:k|                       |                       |
  |                    |<-- Cached value -|                       |                       |
  |     (if hit) <-----|                  |                       |                       |
  |                    |-- gRPC GetProduct(productId) ----------->|                       |
  |                    |<-- Product (status, variants) -----------|                       |
  |                    |-- Validate status = ACTIVE               |                       |
  |                    |-- Find variant in variants list          |                       |
  |                    |-- Validate variant ACTIVE + stock >= qty |                       |
  |                    |-- @Retryable upsert cart_item ---------------------------------->|
  |                    |   (@Version optimistic lock)             |                       |
  |                    |<-- Updated row --------------------------------------------------|
  |                    |-- SET cart:idem:k = result (TTL 24h) --->|                       |
  |<-- 200 CartDTO ----|                                          |                       |
""")
add_caption(doc, "Hình 7.5: Sequence diagram addItem – ASCII")

add_heading(doc, "7.4. Saga đặt hàng (choreography)", level=2)
add_para(doc, "Đặt hàng là saga distributed: cần thay đổi state ở 4 service (order, voucher, inventory, payment). Lựa chọn kiến trúc là choreography (không có orchestrator trung tâm): order-service thực thi local transaction, phát event order.placed, các service downstream (notification, search, analytics) tự subscribe.")
add_image(doc, "image5.png", width_cm=15)
add_caption(doc, "Hình 7.6: Saga đặt hàng – sequence diagram đầy đủ")
add_para(doc, "Mỗi bước được thiết kế để có thể rollback bằng compensation event:")
for item in [
    "Reserve inventory: nếu sau đó payment fail thì release qua order.cancelled.",
    "Apply voucher: ghi VoucherUsageRecord ngay; nếu order cancelled thì release qua API.",
    "Tạo Payment PENDING: chưa charge thực; nếu user không pay trong N phút thì auto cancel.",
    "Tạo Shipment: chỉ tạo sau payment.completed; không có compensation cần thiết.",
]:
    add_bullet(doc, item)

add_code(doc, """Client     Order Svc     Voucher     Inventory     Payment      Cart     Kafka     Shipping  Notification
  |          |             |            |             |            |        |           |          |
  |-POST/-->|              |            |             |            |        |           |          |
  | from-   |--gRPC GetCart----------------------------------------|        |           |          |
  | cart    |<--CartDTO----------------------------------------------        |           |          |
  |         |              |            |             |            |        |           |          |
  |         |--validate--->|            |             |            |        |           |          |
  |         |<--discount---|            |             |            |        |           |          |
  |         |              |            |             |            |        |           |          |
  |         |--reserve---------------->|              |            |        |           |          |
  |         |<--reservationId-----------|             |            |        |           |          |
  |         |              |            |             |            |        |           |          |
  |         |--create Payment------------------------>|            |        |           |          |
  |         |<--payUrl + qrCode-----------------------|            |        |           |          |
  |         |              |            |             |            |        |           |          |
  |         |-- TX: INSERT order + outbox.order.placed ----|       |        |           |          |
  |         |              |            |             |            |        |           |          |
  |<-200----|              |            |             |            |        |           |          |
  |         |-- scheduler poll outbox + publish to Kafka -------------------|           |          |
  |         |              |            |             |            |<--clear cart-------|          |
  |         |              |            |             |            |        |--order.placed----->-|--->email
  |         |              |            |             |            |        |           |          |
  | (later) |              |            |             |--webhook---|        |           |          |
  |         |              |            |             |--mark COMPLETED+publish payment.completed-->|
  |         |<--consume payment.completed-----------------------------------|           |          |
  |         |--confirm------------>|    |             |            |        |           |          |
  |         |--create Shipment-------------------------------------|------->|           |          |
  |         |--publish order.confirmed------------------------------------->|           |          |
""")
add_caption(doc, "Hình 7.7: Saga đặt hàng – ASCII đầy đủ (đặt → pay → confirm)")

add_heading(doc, "7.5. Luồng thanh toán SePay", level=2)
add_para(doc, "Khi user chọn SEPAY_QR, payment-service tạo bản ghi PENDING kèm payUrl + qrCodeUrl. User quét VietQR và chuyển khoản. SePay nhận biết tiền vào tài khoản tổng, gửi webhook về payment-service. Service verify HMAC, match orderNumber trong nội dung chuyển khoản, mark COMPLETED và phát Kafka event.")
add_code(doc, """User       FE          Payment Svc          SePay Gateway        Bank        Order Svc
 |          |                |                       |                |              |
 |--checkout->|              |                       |                |              |
 |          |--call /from-cart-->                    |                |              |
 |          |             (saga creates Payment)     |                |              |
 |          |<--payUrl + qr-|                        |                |              |
 |--scan QR via banking app-------------------------------------------|              |
 |                                                  |--bank notif--->|              |
 |                                                  |<--credited ----|              |
 |                          |<--webhook /sepay/webhook--|             |              |
 |                          |-- verify HMAC          |                |              |
 |                          |-- match orderNumber    |                |              |
 |                          |-- mark COMPLETED       |                |              |
 |                          |-- publish payment.completed ----------------->         |
 |                          |                                                        |--confirm reservation
 |                          |                                                        |--create shipment
 |                          |                                                        |--publish order.confirmed
""")
add_caption(doc, "Hình 7.8: Luồng thanh toán SePay – ASCII")

add_heading(doc, "7.6. Luồng đổi/trả hàng", level=2)
add_code(doc, """Customer    Order Svc        Kafka        Notification        Admin         Payment Svc
  |             |                 |               |                |                |
  |--POST /return-requests------->|               |                |                |
  |             |--validate cond.--                |                |                |
  |             |--INSERT ReturnRequest PENDING    |                |                |
  |             |--publish order.return-requested---->|              |                |
  |             |                 |                 |--email admin-->|              |
  |             |                 |                 |--email customer (acked)        |
  |<--200-------|                 |                 |                |                |
  |             |                 |                 |                |                |
  |             |<--POST /approve--------------------------|         |                |
  |             |--status=APPROVED                          |        |                |
  |             |--publish order.return-approved ---------->|        |                |
  |                                            (customer ships back)                  |
  |             |<--POST /complete ----------------------------------|                |
  |             |--status=COMPLETED                                                   |
  |             |--publish order.returned -------------------------------------------->
  |             |                                                                      |--create refund PENDING
  |                                                                                    |
  |             |<--POST /process-refund (admin)------------------------|              |
  |                                                                                    |--call SePay refund API
  |                                                                                    |--mark REFUNDED
""")
add_caption(doc, "Hình 7.9: Luồng đổi/trả – ASCII")

add_heading(doc, "7.7. Luồng Flash Sale", level=2)
add_code(doc, """Admin       FE           Flash-Sale Svc       Redis           Order Svc      Kafka
  |          |                  |                   |                |              |
  |--create-->                  |                   |                |              |
  |   FS DRAFT                  |                   |                |              |
  |<--id-----                   |                   |                |              |
  |              (at startTime) |                   |                |              |
  |                             |--scheduler activate                |              |
  |                             |--status=ACTIVE                     |              |
  |                             |--SET fs:{id}:qty:{vId}=quota ----->|              |
  |          |                  |                   |                |              |
  |          |--GET /availability                   |                |              |
  |          |<--remaining-------|                  |                |              |
  |          |--POST /participate->                 |                |              |
  |                             |--EVAL Lua atomic:                  |              |
  |                             |   if qty>0 then DECR + INSERT       |              |
  |                             |<--ok or quota_exhausted             |              |
  |          |<--flashPrice + ttl                                     |              |
  |          |--checkout (within 15min) ----------------------------> |              |
  |                                                                  |--use flashPrice
  |                                                                  |--place order
""")
add_caption(doc, "Hình 7.10: Luồng Flash Sale với atomic quota – ASCII")

add_heading(doc, "7.8. State machine của Order", level=2)
add_para(doc, "Mọi chuyển trạng thái đều được enforce ở method domain (state guard throws DomainException nếu transition không hợp lệ). Không thể nhảy từ PENDING thẳng tới SHIPPED – phải đi qua INVENTORY_RESERVED → PAYMENT_PENDING → PAID → CONFIRMED → SHIPPED tuần tự.")
add_image(doc, "image8.png", width_cm=15)
add_caption(doc, "Hình 7.11: State machine của Order")

add_code(doc, """                       +---------+
                       | PENDING |
                       +----+----+
                            | reserve inventory ok
                            v
                +---------------------+
                | INVENTORY_RESERVED  |
                +----------+----------+
                           | create payment
                           v
                  +-------------------+   payment.failed
                  |  PAYMENT_PENDING  |--------+
                  +-------+-----------+        |
                          | payment.completed   |
                          v                    v
                       +------+          +-----------+
                       | PAID |          | CANCELLED |
                       +--+---+          +-----------+
                          | order.confirmed     ^
                          v                     |
                     +-----------+              |
                     | CONFIRMED |              | user/admin cancel
                     +-----+-----+              | (only from PENDING / PAYMENT_PENDING)
                           | create shipment    |
                           v                    |
                      +---------+               |
                      | SHIPPED |               |
                      +----+----+               |
                           | shipping.delivered |
                           v                    |
                     +-----------+              |
                     | DELIVERED |--------+     |
                     +-----------+        |     |
                           |  return req  v     |
                           v        +---------------+
                  +------------------+ RETURN_REQUESTED
                  |    APPROVED      |
                  +--------+---------+
                           v
                       +----------+
                       | RETURNED |
                       +----------+
""")
add_caption(doc, "Hình 7.12: State machine của Order – ASCII")

doc.add_page_break()

# ============================================================
# CHƯƠNG 8: GIAO TIẾP GIỮA CÁC SERVICE
# ============================================================
add_heading(doc, "CHƯƠNG 8: GIAO TIẾP GIỮA CÁC SERVICE", level=1)

add_heading(doc, "8.1. Tổng quan ba mô hình giao tiếp", level=2)
add_para(doc, "Hệ thống sử dụng ba phương thức giao tiếp, mỗi cái cho một mục đích cụ thể.")
add_table_caption(doc, "Bảng 8.1: So sánh REST – gRPC – Kafka")
add_table(doc,
    ["Tiêu chí", "REST", "gRPC", "Kafka"],
    [
        ["Sync/Async", "Sync", "Sync", "Async"],
        ["Tốc độ", "Trung bình", "Rất nhanh", "Tùy consumer"],
        ["Format", "JSON", "Protobuf", "JSON / Avro"],
        ["Coupling", "Trung bình", "Trung bình", "Rất thấp"],
        ["Use case", "Client ↔ Gateway ↔ service", "Inter-service sync (latency-sensitive)", "Inter-service async (event)"],
        ["Schema enforcement", "Manual (OpenAPI)", "Strict (.proto)", "Manual (event class share)"],
        ["Retry", "Client side", "Stub/interceptor", "Consumer auto + DLQ"],
    ],
    col_widths_cm=[3.5, 3.5, 3.5, 5.5])

add_heading(doc, "8.2. Ma trận giao tiếp", level=2)
add_table_caption(doc, "Bảng 8.2: Ma trận giao tiếp giữa các service")
add_table(doc,
    ["From → To", "Mô hình", "Mục đích"],
    [
        ["client → api-gateway", "REST", "Mọi request từ web/mobile"],
        ["api-gateway → service", "REST (LB qua Eureka)", "Forward sau khi route + add trace ID"],
        ["cart-service → catalog-service", "gRPC", "GetProduct, CheckStock pre-flight"],
        ["order-service → cart-service", "gRPC", "GetCart khi place from cart"],
        ["order-service → catalog-service", "gRPC", "GetProduct cho admin manual order"],
        ["order-service → voucher-service", "REST", "Validate + apply + release"],
        ["order-service → inventory-service", "REST", "Reserve + Confirm + Release"],
        ["order-service → payment-service", "REST", "Tạo Payment PENDING"],
        ["payment-service → SePay", "REST (HTTPS)", "Process refund qua API ngoài"],
        ["SePay → payment-service", "REST webhook", "Notify khi bank ghi nhận tiền"],
        ["auth-service → Kafka", "Producer", "auth.user-registered, auth.session.events.v1"],
        ["catalog-service → Kafka", "Producer", "catalog.product-* events"],
        ["order-service → Kafka", "Producer (Outbox)", "order.placed, order.confirmed,…"],
        ["payment-service → Kafka", "Producer", "payment.completed, payment.failed"],
        ["shipping-service → Kafka", "Producer", "shipping.status-changed, shipping.delivered"],
        ["user-profile-service → Kafka", "Both", "Consume auth.user-registered; produce user.profile-upserted"],
        ["inventory-service → Kafka", "Consumer", "catalog.product-created/deleted"],
        ["cart-service → Kafka", "Consumer", "order.placed (clear cart), catalog.product-deleted"],
        ["search-service → Kafka", "Consumer", "catalog.product-* → re-index"],
        ["notification-service → Kafka", "Consumer", "order/payment/shipping/auth events"],
        ["analytics-service → Kafka/ES", "Consumer + ES read", "auth.session.events.v1; aggregate"],
    ],
    col_widths_cm=[5.0, 3.5, 7.5])

add_heading(doc, "8.3. Danh mục Kafka Topics", level=2)
add_image(doc, "image7.png", width_cm=15)
add_caption(doc, "Hình 8.1: Producer → Topic → Consumer toàn hệ thống")
add_table_caption(doc, "Bảng 8.3: Danh mục Kafka topics")
add_table(doc,
    ["Topic", "Producer", "Consumer", "Schema (rút gọn)"],
    [
        ["auth.user-registered", "auth", "user-profile, notification", "{userId, username, email, createdAt}"],
        ["auth.user.events.v1", "auth", "audit", "{type, userId, payload, occurredAt}"],
        ["auth.session.events.v1", "auth", "analytics", "{type, userId, ip, ua, occurredAt}"],
        ["catalog.product-created", "catalog", "inventory, search", "{productId, name, sku[], status}"],
        ["catalog.product-updated", "catalog", "search", "{productId, changes}"],
        ["catalog.product-deleted", "catalog", "inventory, cart, search", "{productId}"],
        ["catalog.product-status-changed", "catalog", "cart, search", "{productId, oldStatus, newStatus}"],
        ["order.placed", "order (outbox)", "cart, notification, analytics", "{orderId, userId, items[], total, voucherCode}"],
        ["order.confirmed", "order", "notification, search", "{orderId, paidAt}"],
        ["order.shipped", "order", "notification", "{orderId, shipmentId, trackingNumber}"],
        ["order.delivered", "order", "notification", "{orderId, deliveredAt}"],
        ["order.cancelled", "order", "voucher, inventory, notification", "{orderId, reason}"],
        ["order.failed", "order", "notification", "{orderId, reason}"],
        ["order.return-requested", "order", "notification", "{orderId, reason}"],
        ["order.return-approved", "order", "notification, payment", "{orderId}"],
        ["order.return-rejected", "order", "notification", "{orderId}"],
        ["order.returned", "order", "inventory, payment, notification", "{orderId}"],
        ["payment.completed", "payment", "order, notification", "{paymentId, orderId, amount}"],
        ["payment.failed", "payment", "order, notification", "{paymentId, orderId, reason}"],
        ["payment.refunded", "payment (dự kiến)", "order, notification", "{paymentId, orderId, amount}"],
        ["shipping.status-changed", "shipping", "order, notification", "{shipmentId, orderId, status}"],
        ["shipping.delivered", "shipping", "order, notification", "{shipmentId, orderId, deliveredAt}"],
        ["user.profile-upserted", "user-profile", "search", "{userId, profile}"],
    ],
    col_widths_cm=[4.5, 2.5, 4.0, 5.0])

add_heading(doc, "8.4. gRPC Service Definitions", level=2)
add_para(doc, "Auth-service expose gRPC server cho user lookup, validate token, tăng tokenVersion. Catalog cung cấp product/variant lookup. Cart cung cấp GetCart cho order-service. Các proto file lưu trong module grpc-lib (share-jar).")
add_para(doc, "Catalog Service gRPC (port 9193):", bold=True)
add_code(doc, """service CatalogService {
  rpc GetProduct (GetProductRequest) returns (Product);
  rpc GetVariantBySku (SkuRequest) returns (Variant);
  rpc CheckStock (StockRequest) returns (StockResponse);
  rpc ListVariantsByProduct (GetProductRequest) returns (VariantList);
}
message Product { int64 id; string name; string slug; string status; repeated Variant variants; }
message Variant { int64 id; string sku; double price; double salePrice; int32 quantity; string status; }
message StockResponse { bool available; int32 quantity; string status; }""")

add_para(doc, "Auth Service gRPC (port 9191):", bold=True)
add_code(doc, """service AuthService {
  rpc ValidateToken (TokenRequest) returns (TokenValidation);
  rpc GetUserById (UserIdRequest) returns (UserSummary);
  rpc IncrementTokenVersion (UserIdRequest) returns (Empty);
}
message TokenValidation { bool valid; string userId; repeated string roles; }""")

add_para(doc, "Cart Service gRPC (port 9194):", bold=True)
add_code(doc, """service CartService {
  rpc GetCart (UserIdRequest) returns (Cart);
}
message Cart { repeated CartItem items; double total; }
message CartItem { int64 variantId; string sku; int32 quantity; double unitPrice; }""")

doc.add_page_break()

# ============================================================
# CHƯƠNG 9: BẢO MẬT
# ============================================================
add_heading(doc, "CHƯƠNG 9: BẢO MẬT", level=1)

add_heading(doc, "9.1. Tổng quan bảo mật nhiều lớp", level=2)
add_para(doc, "Hệ thống áp dụng defense-in-depth với nhiều lớp bảo vệ độc lập. Vô hiệu hóa một lớp không làm sập toàn bộ.")
add_table_caption(doc, "Bảng 9.1: Các lớp bảo mật")
add_table(doc,
    ["Lớp", "Vị trí", "Biện pháp"],
    [
        ["1. Network", "API Gateway", "HTTPS bắt buộc (prod), CORS allowlist, IP allowlist webhook"],
        ["2. Authentication", "Gateway + service", "JWT HS256, cookie HttpOnly + SameSite=Strict"],
        ["3. Authorization", "Service filter + @PreAuthorize", "Role check + permission check granular"],
        ["4. Rate Limiting", "Gateway", "100 req/phút/IP cho /api/v1/auth/*; 1000 req/phút cho /api/v1/products/*"],
        ["5. Input Validation", "Controller DTO", "@Valid, @NotBlank, @Size, @Email, @Pattern"],
        ["6. Business Rule", "Domain layer", "State guard, invariant check trong aggregate"],
        ["7. Idempotency", "Service layer", "Redis key + UNIQUE constraint chống replay"],
        ["8. Persistence", "Database", "BCrypt hash, UNIQUE constraint, FK, prepared statement"],
        ["9. Audit & Detection", "Kafka + ES", "Log mọi LOGIN_SUCCESS/FAILED, refresh, revoke; alert qua Kibana Watcher"],
    ],
    col_widths_cm=[4.0, 4.5, 7.5])

add_heading(doc, "9.2. JWT (JSON Web Token)", level=2)
add_table_caption(doc, "Bảng 9.2: Cấu hình JWT")
add_table(doc,
    ["Tham số", "Giá trị", "Ghi chú"],
    [
        ["Thuật toán ký", "HS256", "Shared secret giữa các service"],
        ["Secret length", "≥ 32 ký tự", "Enforce ở JwtProperties.afterPropertiesSet"],
        ["Access token TTL", "15 phút", "Đủ ngắn để giảm rủi ro token leak"],
        ["Refresh token TTL", "7 ngày", "Đủ dài cho UX, có rotation chống replay"],
        ["Phân phối", "Cookie HttpOnly + SameSite=Strict", "Tránh XSS đánh cắp"],
        ["Cookie names", "ACCESS, REFRESH", "Cookie ACCESS đọc được bởi mọi service qua Gateway"],
    ],
    col_widths_cm=[4.0, 5.0, 7.0])

add_para(doc, "Cấu trúc JWT payload:", bold=True)
add_code(doc, """{
  "sub": "john_doe",            // username
  "userId": "550e8400-e29b...", // UUID
  "email": "john@hieu.vn",
  "roles": ["ROLE_CUSTOMER"],
  "jti": "uuid-v4",             // JWT ID – dùng cho blacklist
  "tokenVersion": 0,            // Revoke-all version
  "iat": 1747200000,            // Issued at
  "exp": 1747200900             // Expiration (iat + 900s)
}""")
add_para(doc, "Lưu ý: JWT KHÔNG chứa permissions để giữ kích thước nhỏ. Permissions được resolve runtime qua RBAC lookup cache (Redis hoặc gRPC tới auth-service).", italic=True)

add_heading(doc, "9.3. RBAC – Phân quyền theo Role + Permission", level=2)
add_para(doc, "Mô hình phân quyền chuẩn RBAC: User N-N Role, Role N-N Permission.")
add_code(doc, """User (1) ---< (N) UserRole (N) >--- (1) Role
Role (1) ---< (N) RolePermission (N) >--- (1) Permission

Ví dụ:
  admin    -> ROLE_ADMIN    -> [USER_*, PRODUCT_*, ORDER_*, VOUCHER_*, …]
  staff    -> ROLE_STAFF    -> [ORDER_READ_ALL, RETURN_APPROVE, …]
  customer -> ROLE_CUSTOMER -> [PRODUCT_READ, CART_*, ORDER_OWN_*]""")

add_table_caption(doc, "Bảng 9.3: Các role/permission tiêu biểu")
add_table(doc,
    ["Role", "Permission tiêu biểu", "Mô tả"],
    [
        ["ROLE_ADMIN", "*", "Toàn quyền – chỉ admin/CTO có"],
        ["ROLE_STAFF", "ORDER_READ_ALL, RETURN_APPROVE, NOTIFICATION_SEND", "Hỗ trợ KH, xử lý đổi/trả"],
        ["ROLE_CUSTOMER", "PRODUCT_READ, CART_*, ORDER_OWN_*", "Khách hàng mua sắm"],
    ],
    col_widths_cm=[3.5, 6.5, 6.0])

add_heading(doc, "9.4. Refresh Token Rotation & Blacklist", level=2)
add_table_caption(doc, "Bảng 9.4: Bốn cơ chế kiểm tra token")
add_table(doc,
    ["Lớp", "Kiểm tra", "Vị trí", "Khi nào fail"],
    [
        ["1. Signature", "Verify HS256 với JWT_SECRET", "JwtAuthenticationFilter (mọi service)", "Token giả mạo hoặc secret đổi"],
        ["2. Expiry", "exp > now", "Filter", "Token hết hạn"],
        ["3. Blacklist (jti)", "EXISTS in Redis SET blacklist:{jti}", "Filter", "Token đã logout/revoke"],
        ["4. Token Version", "JWT.tokenVersion == users.token_version", "gRPC hoặc cache lookup", "Admin revoke-all sau khi token phát hành"],
    ],
    col_widths_cm=[2.5, 5.5, 4.5, 3.5])

add_para(doc, "Cơ chế revoke:")
for item in [
    "Logout: thêm jti vào Redis blacklist với TTL = remaining token lifetime + ghi DB.",
    "Revoke all (đổi mật khẩu, security incident): tăng users.token_version – mọi JWT cũ bị reject ở lớp 4.",
    "Refresh token rotation: mỗi lần refresh, token cũ blacklist + token mới phát hành.",
    "Reuse detection: nếu refresh đã blacklist mà vẫn dùng – revoke cả family (tăng tokenVersion).",
    "Admin revoke: endpoint /revoke-all/{userId} cho admin – chỉ admin quyền cao có.",
]:
    add_bullet(doc, item)

add_heading(doc, "9.5. Idempotency & Anti-replay", level=2)
add_para(doc, "Mọi command quan trọng đều tuân theo nguyên tắc idempotency. Client gửi kèm X-Idempotency-Key (UUIDv4); service lưu (key → result) vào Redis với TTL 24h. Nếu key đã tồn tại, trả lại kết quả cũ thay vì thực thi lần hai.")
add_table_caption(doc, "Bảng 9.5: Idempotency theo service")
add_table(doc,
    ["Service", "Cơ chế", "TTL"],
    [
        ["cart-service", "Redis HSET cart:idem:{key} → CartDTO", "24h"],
        ["order-service", "Bảng order_idempotency UNIQUE(idempotency_key)", "Permanent (DB)"],
        ["payment-service", "UNIQUE(external_ref) chống xử lý webhook nhiều lần", "Permanent (DB)"],
        ["voucher-service", "UNIQUE(voucher_id, order_id) chống apply nhiều lần", "Permanent (DB)"],
    ],
    col_widths_cm=[4.0, 8.0, 4.0])

add_heading(doc, "9.6. Rate Limiting & CORS", level=2)
add_table_caption(doc, "Bảng 9.6: Cấu hình rate limit qua Gateway")
add_table(doc,
    ["Pattern", "Giới hạn", "Mục đích"],
    [
        ["/api/v1/auth/login", "5 req/phút/IP", "Chống brute force mật khẩu"],
        ["/api/v1/auth/register", "10 req/giờ/IP", "Chống spam đăng ký"],
        ["/api/v1/auth/*", "100 req/phút/IP", "Chống abuse auth flow"],
        ["/api/v1/products/*", "1000 req/phút/IP", "Public, cần response nhanh"],
        ["/api/v1/orders/from-cart", "30 req/phút/user", "Chống multi-click"],
        ["/api/v1/payments/sepay/webhook", "Không limit nhưng IP allowlist", "Tránh giả mạo"],
        ["Default", "300 req/phút/IP", "Bảo vệ tổng thể"],
    ],
    col_widths_cm=[6.0, 4.5, 5.5])

add_para(doc, "CORS Configuration:", bold=True)
add_code(doc, """spring.cloud.gateway.globalcors:
  cors-configurations:
    '[/**]':
      allowedOriginPatterns:
        - https://hieu.vn
        - https://admin.hieu.vn
        - http://localhost:3000  # dev only
      allowedMethods: [GET, POST, PATCH, PUT, DELETE, OPTIONS]
      allowedHeaders: ['*']
      allowCredentials: true
      maxAge: 3600""")

add_heading(doc, "9.7. Bảo mật vận hành", level=2)
for item in [
    "JWT_SECRET dùng chung phải ≥ 32 ký tự, rotate định kỳ (yêu cầu re-login toàn hệ thống).",
    "BCrypt cost 12 – đủ chậm tránh brute force trên dump database.",
    "PostgreSQL: mỗi service có user riêng, password đặt qua env, không bao giờ commit; production bật pg_hba md5/scram-sha-256.",
    "Kafka: dev chưa bật ACL; production cần SASL/SCRAM + ACL theo topic (producer/consumer riêng).",
    "Redis: dev không password; production bật requirepass + TLS.",
    "Webhook SePay: verify HMAC header X-Signature; IP allowlist từ SePay dashboard.",
    "Secret management: production dùng Kubernetes Secret hoặc HashiCorp Vault.",
    "Audit log: mọi action admin (lock user, grant role, create voucher) đều ghi vào auth.user.events.v1.",
    "Vulnerability scan: OWASP Dependency Check + Snyk trong CI/CD.",
    "Pen-test định kỳ: mỗi 6 tháng, do bên thứ 3 thực hiện.",
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# CHƯƠNG 10: KIỂM THỬ
# ============================================================
add_heading(doc, "CHƯƠNG 10: KIỂM THỬ", level=1)

add_heading(doc, "10.1. Chiến lược kiểm thử (Test Pyramid)", level=2)
add_para(doc, "Dự án áp dụng test pyramid – ưu tiên unit test (nhiều, nhanh, rẻ) hơn integration test, hơn e2e (ít, chậm, đắt). Coverage mục tiêu phân bổ:")
add_code(doc, """                          /\\
                         /  \\
                        / E2E\\        5%  (TestContainers, slow)
                       /------\\
                      /        \\
                     /  Integration\\   25% (H2 + @SpringBootTest)
                    /----------------\\
                   /                  \\
                  /       Unit         \\ 70% (Mockito, no Spring)
                 /----------------------\\
""")
add_caption(doc, "Hình 10.1: Test pyramid của dự án")

add_table_caption(doc, "Bảng 10.1: Framework kiểm thử")
add_table(doc,
    ["Framework", "Mục đích"],
    [
        ["JUnit 5", "Test runner chính"],
        ["Mockito", "Mock dependency cho unit test"],
        ["AssertJ", "Fluent assertion"],
        ["Spring Boot Test", "@WebMvcTest, @DataJpaTest, @SpringBootTest"],
        ["spring-security-test", "@WithMockUser cho authorization test"],
        ["TestContainers", "PostgreSQL + Kafka container thật cho integration test"],
        ["grpc-test", "Test gRPC server inproc"],
        ["WireMock", "Stub external HTTP (SePay, GHN)"],
        ["JaCoCo", "Coverage report"],
    ],
    col_widths_cm=[5.0, 11.0])

add_heading(doc, "10.2. Unit Test", level=2)
add_para(doc, "Đối tượng unit test chính:")
for item in [
    "Domain entity & value object: invariant, factory, equality.",
    "Domain service: business rule không cần DB.",
    "Command handler: gọi domain method, không Spring context.",
    "Query handler: chuyển entity sang DTO.",
    "Mapper: chuyển đổi DTO ↔ domain.",
    "Util: JWT util, BCrypt wrapper, idempotency hash.",
]:
    add_bullet(doc, item)

add_table_caption(doc, "Bảng 10.2: Phân bổ unit test theo service")
add_table(doc,
    ["Service", "Số class test mục tiêu", "Coverage mục tiêu"],
    [
        ["auth-service", "30+", "≥ 80%"],
        ["catalog-service", "25+", "≥ 75%"],
        ["order-service", "30+", "≥ 80%"],
        ["payment-service", "15+", "≥ 70%"],
        ["inventory-service", "20+", "≥ 80%"],
        ["cart-service", "15+", "≥ 80%"],
        ["voucher-service", "15+", "≥ 75%"],
        ["flash-sale-service", "12+", "≥ 75%"],
        ["shipping-service", "10+", "≥ 70%"],
        ["notification-service", "10+", "≥ 60%"],
        ["user-profile-service", "8+", "≥ 70%"],
        ["search-service", "8+", "≥ 60%"],
        ["analytics-service", "5+", "≥ 50%"],
    ],
    col_widths_cm=[5.0, 5.0, 6.0])

add_heading(doc, "10.3. Integration Test", level=2)
add_para(doc, "Integration test kiểm tra:")
for item in [
    "Controller → Service → Repository flow với H2 in-memory hoặc PostgreSQL TestContainer.",
    "Spring Security filter chain với @WebMvcTest + @WithMockUser.",
    "JPA entity mapping, query performance, lazy loading.",
    "gRPC service endpoint với grpc-test inproc.",
    "Kafka producer/consumer với EmbeddedKafka hoặc Kafka TestContainer.",
    "Outbox publisher: scheduler đọc DB → publish thật vào Kafka.",
    "Webhook contract: WireMock stub SePay → payment-service xử lý.",
]:
    add_bullet(doc, item)

add_heading(doc, "10.4. Coverage và chất lượng", level=2)
add_para(doc, "Mục tiêu coverage tổng thể:")
add_table_caption(doc, "Bảng 10.3: Coverage mục tiêu theo tầng")
add_table(doc,
    ["Tầng", "Mức tối thiểu", "Lý do"],
    [
        ["domain", "≥ 80%", "Business logic quan trọng nhất, dễ test, ít dependency"],
        ["application", "≥ 70%", "Command/Query handler – orchestration"],
        ["infrastructure", "≥ 50%", "Phần lớn là wrapper – test integration đủ"],
        ["interfaces", "≥ 60%", "Controller test qua @WebMvcTest"],
        ["Overall", "≥ 60%", "Cân bằng giữa nỗ lực và lợi ích"],
    ],
    col_widths_cm=[3.5, 3.5, 9.0])
add_para(doc, "Công cụ: JaCoCo tích hợp Maven, report xuất ra target/site/jacoco/index.html. Chạy test: mvn test -pl {module} hoặc mvn test cho toàn dự án.")

doc.add_page_break()

# ============================================================
# CHƯƠNG 11: GIÁM SÁT VÀ VẬN HÀNH
# ============================================================
add_heading(doc, "CHƯƠNG 11: GIÁM SÁT VÀ VẬN HÀNH", level=1)

add_heading(doc, "11.1. Boot stack local", level=2)
add_code(doc, """# Bước 1: hạ tầng (Postgres x 11, Redis, Kafka KRaft, ES, Kibana, MailHog)
cd /Users/admin/HieuTo/ecommerce
docker compose up -d

# Bước 2: services Spring Boot (cần Java 25 + Maven)
./boot-bg.sh        # boot 15 service nền, ~5 phút sẵn sàng

# Bước 3: seed dữ liệu (idempotent)
./seed.sh           # users + categories + attrs + vouchers
python3 seed-hieu-luxe.py --wipe   # 25 sản phẩm HIEU

# Bước 4: kiểm tra
curl http://localhost:8080/api/v1/products | jq '.totalElements'
curl http://localhost:9091/actuator/health""")

add_para(doc, "Thứ tự boot quan trọng:", bold=True)
for item in [
    "eureka-server boot đầu tiên (chờ 25 giây để registry sẵn sàng).",
    "auth-service tiếp theo – các service khác cần JWT_SECRET trùng + có thể đăng ký Eureka.",
    "Catalog, inventory, cart, order, payment, shipping, voucher, flash-sale, notification: song song được.",
    "api-gateway boot cuối cùng để có Eureka đã đầy đủ instance.",
]:
    add_bullet(doc, item)

add_heading(doc, "11.2. Migrations & Seeding", level=2)
for item in [
    "Flyway tự chạy migration khi service start (spring.flyway.locations=classpath:db/migration).",
    "Convention: V{n}__init_{service}_schema.sql, V{n+1}__{desc}.sql.",
    "Không bao giờ sửa migration đã merge – luôn tạo migration mới.",
    "seed.sh dùng curl gọi qua gateway, idempotent (409 trên duplicate được swallow).",
    "Reset toàn bộ: docker compose down -v + rebuild + chạy lại seed.",
]:
    add_bullet(doc, item)

add_heading(doc, "11.3. Structured Logging & Correlation ID", level=2)
add_para(doc, "Mỗi service log dạng JSON qua Logback + Logstash encoder; Filebeat tail file log; Elasticsearch index analytics-events-*; Kibana view dashboard.")
add_code(doc, """%d{yyyy-MM-dd HH:mm:ss.SSS} [%thread] [%X{traceId},%X{spanId}] %-5level %logger{36} - %msg%n

Correlation ID: header X-Request-Id từ Gateway → propagate qua mọi service
qua MDC (Mapped Diagnostic Context). Mọi log line có traceId để correlate
request xuyên service.""")
add_table_caption(doc, "Bảng 11.1: Log levels và quy ước")
add_table(doc,
    ["Level", "Khi nào dùng", "Ví dụ"],
    [
        ["ERROR", "Lỗi nghiêm trọng cần xử lý ngay", "DB connection lost; webhook signature invalid"],
        ["WARN", "Bất thường nhưng chưa fail", "Redis down → fallback DB; retry attempt"],
        ["INFO", "Sự kiện quan trọng", "Login success; Order placed; Payment completed"],
        ["DEBUG", "Chi tiết luồng (dev only)", "Cache hit/miss; gRPC stub call"],
        ["TRACE", "Chi tiết nhất (hiếm)", "Toàn bộ payload Kafka"],
    ],
    col_widths_cm=[2.0, 5.5, 8.5])

add_heading(doc, "11.4. Metrics & Actuator", level=2)
add_para(doc, "Mỗi service expose /actuator/health, /info, /metrics, /prometheus.")
add_table_caption(doc, "Bảng 11.2: Metrics quan trọng")
add_table(doc,
    ["Metric", "Mô tả", "Alert ngưỡng"],
    [
        ["http.server.requests p95", "Latency REST 95th percentile", "> 500ms 5 phút"],
        ["http.server.requests count{status=5xx}", "Số 5xx error", "> 10/phút"],
        ["jvm.memory.used", "Memory sử dụng", "> 85% heap"],
        ["jvm.gc.pause", "GC pause time", "> 200ms"],
        ["hikaricp.connections.active", "Active DB connection", "> 80% pool"],
        ["kafka.consumer.lag", "Consumer lag", "> 1000 message"],
        ["resilience4j.circuitbreaker.calls", "Circuit breaker open count", "Open > 1 phút"],
        ["custom.order.placed", "Số order placed", "Drop > 50% so với giờ trước"],
        ["custom.payment.failed.rate", "Tỉ lệ payment fail", "> 10%"],
    ],
    col_widths_cm=[5.5, 5.5, 5.0])

add_table_caption(doc, "Bảng 11.3: Actuator endpoints")
add_table(doc,
    ["Endpoint", "Mục đích"],
    [
        ["/actuator/health", "Health check tổng – up/down/degraded"],
        ["/actuator/health/liveness", "K8s liveness probe – chỉ kiểm tra JVM còn sống"],
        ["/actuator/health/readiness", "K8s readiness probe – DB + Kafka + Eureka up"],
        ["/actuator/info", "Build info (commit, version, build time)"],
        ["/actuator/prometheus", "Scrape endpoint cho Prometheus"],
        ["/actuator/loggers", "Đổi log level runtime (require admin)"],
        ["/actuator/env", "Xem env vars (require admin, mask secret)"],
    ],
    col_widths_cm=[5.5, 10.5])

add_heading(doc, "11.5. Khôi phục sau sự cố", level=2)
for item in [
    "Backup Postgres: pg_dump hàng đêm; retention 7 ngày + 4 tuần + 6 tháng (3-2-1 rule).",
    "Kafka topic retention: 7 ngày – đủ để replay event nếu consumer mất state.",
    "Idempotency record giữ 24h – đủ cho retry an toàn.",
    "Order-service outbox: nếu Kafka publish fail, scheduled job retry mỗi 10s với exponential backoff.",
    "Saga compensation: order-service chủ động consume payment.failed → release inventory + cancel order + release voucher.",
    "Database-per-service: hỏng một DB chỉ làm gãy một service; các service khác hoạt động bình thường (theo dependency chain).",
    "Disaster Recovery (RTO ≤ 4h, RPO ≤ 1h): backup off-site S3; failover Postgres replica.",
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# CHƯƠNG 12: TRIỂN KHAI
# ============================================================
add_heading(doc, "CHƯƠNG 12: TRIỂN KHAI", level=1)

add_heading(doc, "12.1. Docker Compose", level=2)
add_para(doc, "Hệ thống chia thành 3 file Docker Compose tương ứng 3 nhóm dịch vụ:")
add_table_caption(doc, "Bảng 12.1: Các file Docker Compose")
add_table(doc,
    ["File", "Nội dung", "Lệnh khởi động"],
    [
        ["docker-compose.yml", "Toàn bộ hạ tầng + 15 service", "docker compose up -d"],
        ["docker-compose-infra.yml", "Chỉ hạ tầng (DB, Redis, Kafka, ES, Kibana, MailHog)", "docker compose -f docker-compose-infra.yml up -d"],
        ["docker-compose-monitoring.yml", "Filebeat, Prometheus (dự kiến), Grafana (dự kiến)", "docker compose -f docker-compose-monitoring.yml up -d"],
    ],
    col_widths_cm=[5.0, 7.5, 3.5])

add_heading(doc, "12.2. Biến môi trường", level=2)
add_table_caption(doc, "Bảng 12.2: Biến môi trường chính")
add_table(doc,
    ["Biến", "Mô tả", "Mặc định dev"],
    [
        ["JWT_SECRET", "Secret HS256 cho JWT (≥32 ký tự)", "hieuJwtSecretKey123456789ABCDEFGHIJKLMNOP"],
        ["JWT_EXPIRATION", "TTL access token (ms)", "900000 (15 phút)"],
        ["JWT_REFRESH_EXPIRATION", "TTL refresh token (ms)", "604800000 (7 ngày)"],
        ["SPRING_DATASOURCE_URL", "JDBC URL DB", "jdbc:postgresql://postgres-auth:5432/auth_db"],
        ["SPRING_DATASOURCE_USERNAME", "User DB", "postgres"],
        ["SPRING_DATASOURCE_PASSWORD", "Password DB", "postgres (dev)"],
        ["EUREKA_SERVER_URL", "URL Eureka", "http://eureka-server:8761/eureka/"],
        ["SPRING_REDIS_HOST", "Host Redis", "redis"],
        ["SPRING_REDIS_PORT", "Port Redis", "6379"],
        ["KAFKA_BOOTSTRAP_SERVERS", "Kafka broker list", "kafka:9092"],
        ["ELASTICSEARCH_URI", "URL ES", "http://elasticsearch:9200"],
        ["SEPAY_WEBHOOK_SECRET", "Secret HMAC cho webhook SePay", "(production only)"],
        ["SMTP_HOST", "SMTP server", "mailhog (dev) / smtp.gmail.com (prod)"],
        ["SMTP_PORT", "SMTP port", "1025 (dev) / 587 (prod)"],
        ["SPRING_PROFILES_ACTIVE", "Profile config", "dev / prod"],
    ],
    col_widths_cm=[5.0, 6.0, 5.0])

add_para(doc, "Quản lý secret:", bold=True)
for item in [
    "Production: Kubernetes Secret hoặc HashiCorp Vault – không bao giờ commit secret vào git.",
    "Dev/CI: .env.example commit; .env (thật) gitignore.",
    "Application.yaml dùng cú pháp ${ENV_VAR:default} để fallback giá trị mặc định.",
]:
    add_bullet(doc, item)

add_heading(doc, "12.3. Các lệnh khởi động", level=2)
add_table_caption(doc, "Bảng 12.3: Các lệnh thường dùng")
add_table(doc,
    ["Lệnh", "Mục đích"],
    [
        ["docker compose up -d", "Khởi động toàn bộ stack ở chế độ nền"],
        ["docker compose down -v", "Dừng + xóa volume (reset dữ liệu)"],
        ["docker compose logs -f auth-service", "Xem log realtime của một service"],
        ["./boot-bg.sh", "Build và boot 15 service Spring Boot (dev local, không qua Docker)"],
        ["./dev.sh", "Mode dev nhanh – chỉ boot service đang code"],
        ["./stop.sh", "Dừng tất cả service local"],
        ["./test-api.sh", "Chạy smoke test API qua curl"],
        ["mvn -pl auth-service test", "Chạy unit test cho auth-service"],
        ["mvn clean install -DskipTests", "Build toàn dự án không test"],
    ],
    col_widths_cm=[7.0, 9.0])

add_heading(doc, "12.4. Dự kiến: Kubernetes", level=2)
add_para(doc, "Trong giai đoạn tiếp theo, hệ thống sẽ triển khai trên Kubernetes:")
for item in [
    "Deployment cho mỗi service với replicaCount cấu hình được (mặc định 2).",
    "Service (ClusterIP) cho giao tiếp nội bộ.",
    "Ingress Controller (NGINX) thay thế api-gateway cho routing bên ngoài.",
    "ConfigMap cho cấu hình không nhạy cảm (URLs, feature flag).",
    "Secret cho JWT_SECRET, database password, SePay webhook secret.",
    "Horizontal Pod Autoscaler (HPA) tự động scale theo CPU/memory/custom metric (orders/min).",
    "Persistent Volume Claims (PVC) cho database storage.",
    "Helm Charts để quản lý deployment template per environment.",
    "Istio service mesh cho mTLS giữa các service (tùy chọn).",
    "Cert-Manager + Let's Encrypt cho TLS cert tự động renew.",
]:
    add_bullet(doc, item)

add_heading(doc, "12.5. Dự kiến: CI/CD Pipeline", level=2)
add_para(doc, "Dự kiến CI/CD qua GitHub Actions hoặc Jenkins:")
for i, item in enumerate([
    "Code Push: developer push lên branch feature/* hoặc main.",
    "Build: Maven package – Lombok annotation processing, tạo JAR và Docker image.",
    "Unit Test: mvn test; fail pipeline nếu test fail.",
    "Code Quality: SonarQube scan code smell, bug, security hotspot.",
    "Security Scan: OWASP Dependency Check + Snyk SAST/SCA.",
    "Docker Build & Push: build image, tag :git-sha + :latest, push registry.",
    "Deploy Staging: tự động deploy lên K8s staging cluster qua kubectl/helm.",
    "Integration Test: chạy e2e test trên staging với dữ liệu seed.",
    "Manual Approval: reviewer approve để deploy lên production (chỉ branch main).",
    "Deploy Production: rolling update với canary deployment (5% → 25% → 100%).",
    "Monitor: theo dõi error rate; auto-rollback nếu vượt 5% trong 5 phút.",
]):
    add_number(doc, item)

doc.add_page_break()

# ============================================================
# CHƯƠNG 13: LỘ TRÌNH PHÁT TRIỂN
# ============================================================
add_heading(doc, "CHƯƠNG 13: LỘ TRÌNH PHÁT TRIỂN", level=1)

add_heading(doc, "13.1. Hiện trạng triển khai", level=2)
add_para(doc, "Tính đến phiên bản hiện tại (2.0), hệ thống đã hoàn thành phần code cơ bản với kiến trúc DDD + Choreography Saga nhất quán. Các tính năng đã implement:")
for item in [
    "Xác thực và phân quyền với JWT cookie + RBAC + refresh rotation + blacklist.",
    "Quản lý sản phẩm, danh mục, biến thể, thuộc tính trong catalog-service (gRPC + REST).",
    "Nghiệp vụ e-commerce cốt lõi: cart → order → payment → shipping → inventory với saga + outbox + idempotency.",
    "Tính năng tăng trưởng: voucher, flash sale (atomic Lua quota), notification (in-app + email + SSE).",
    "Tính năng thông minh: search (Elasticsearch), analytics dashboard, log explorer.",
    "Đổi/trả: ReturnRequest flow với approve/reject/complete.",
    "Tích hợp SePay (VietQR webhook).",
    "Unit test cho domain và security critical path.",
    "Docker Compose cho infrastructure + service.",
    "Filebeat → Elasticsearch → Kibana stack cho log.",
]:
    add_bullet(doc, item)

add_heading(doc, "13.2. Công nghệ dự kiến bổ sung", level=2)
add_table_caption(doc, "Bảng 13.1: Công nghệ dự kiến trong roadmap")
add_table(doc,
    ["Công nghệ", "Mục đích", "Mức độ ưu tiên"],
    [
        ["Prometheus + Grafana", "Thay Filebeat-only stack cho metrics + dashboard", "Cao"],
        ["Loki", "Log aggregation + alert", "Cao"],
        ["Jaeger / Zipkin", "Distributed tracing thay cho correlation ID thủ công", "Cao"],
        ["Resilience4j circuit breaker", "Bảo vệ inter-service call", "Cao"],
        ["Kubernetes", "Production deployment thay Docker Compose", "Cao"],
        ["Helm Charts", "Quản lý K8s deployment", "Cao"],
        ["GitHub Actions / Jenkins", "CI/CD pipeline", "Cao"],
        ["SonarQube", "Code quality scan", "Trung bình"],
        ["Recommendation AI (ML)", "Đề xuất sản phẩm cá nhân hoá", "Trung bình"],
        ["Multi-warehouse inventory", "Hỗ trợ nhiều kho địa lý", "Trung bình"],
        ["GraphQL gateway", "Cho mobile/partner app", "Thấp"],
        ["Vault", "Secret management production", "Cao"],
        ["Istio service mesh", "mTLS, traffic split, canary", "Thấp"],
    ],
    col_widths_cm=[5.0, 7.5, 3.5])

add_heading(doc, "13.3. Lộ trình theo giai đoạn", level=2)
add_table_caption(doc, "Bảng 13.2: Lộ trình phát triển dự án")
add_table(doc,
    ["Giai đoạn", "Thời gian", "Mục tiêu chính"],
    [
        ["Phase 1 (đã xong)", "Q1-Q2 2026", "15 service core + saga đặt hàng + SePay + Elasticsearch search"],
        ["Phase 2 (đang làm)", "Q3 2026", "Resilience4j, Prometheus/Grafana, Loki, Distributed tracing"],
        ["Phase 3", "Q4 2026", "Kubernetes deploy, Helm, CI/CD pipeline, SonarQube"],
        ["Phase 4", "Q1 2027", "Recommendation AI, multi-warehouse, performance tuning"],
        ["Phase 5", "Q2 2027", "Mobile app + GraphQL gateway"],
        ["Phase 6", "Q3 2027 trở đi", "Istio, multi-region active-active, advanced AI"],
    ],
    col_widths_cm=[3.5, 3.5, 9.0])

doc.add_page_break()

# ============================================================
# PHỤ LỤC
# ============================================================
add_heading(doc, "PHỤ LỤC", level=1)

add_heading(doc, "A. Bảng cổng dịch vụ", level=2)
add_table_caption(doc, "Bảng A.1: Tổng hợp cổng dịch vụ và database")
add_table(doc,
    ["Service", "Cổng host", "Cổng container", "DB", "DB port host"],
    [
        ["eureka-server", "8761", "—", "—", "—"],
        ["api-gateway", "8080", "—", "—", "—"],
        ["auth-service (REST)", "9091", "8080", "postgres-auth", "5433"],
        ["auth-service (gRPC)", "9191", "9090", "—", "—"],
        ["user-profile-service", "9099", "8080", "postgres-user", "5434"],
        ["catalog-service (REST)", "9093", "8080", "postgres-catalog", "5435"],
        ["catalog-service (gRPC)", "9193", "9090", "—", "—"],
        ["inventory-service", "9098", "8080", "postgres-inventory", "5436"],
        ["cart-service (REST)", "9094", "8080", "postgres-cart", "5437"],
        ["cart-service (gRPC)", "9194", "9090", "—", "—"],
        ["order-service", "9095", "8080", "postgres-order", "5438"],
        ["payment-service", "8086", "8080", "postgres-payment", "5439"],
        ["shipping-service", "8087", "8080", "postgres-shipping", "5440"],
        ["voucher-service", "8094", "8080", "postgres-voucher", "5441"],
        ["flash-sale-service", "8089", "8080", "postgres-flashsale", "5442"],
        ["notification-service", "8090", "8080", "postgres-noti", "5443"],
        ["search-service", "8092", "8080", "—", "—"],
        ["analytics-service", "8095", "8080", "—", "—"],
        ["Redis", "6379", "6379", "—", "—"],
        ["Kafka", "9092", "9092", "—", "—"],
        ["Elasticsearch", "9200", "9200", "—", "—"],
        ["Kibana", "5601", "5601", "—", "—"],
        ["MailHog UI", "8025", "8025", "—", "—"],
        ["MailHog SMTP", "1025", "1025", "—", "—"],
    ],
    col_widths_cm=[4.5, 2.5, 2.5, 4.0, 2.5])

add_heading(doc, "B. Tham chiếu mã lỗi", level=2)
add_table_caption(doc, "Bảng B.1: Mã lỗi chuẩn")
add_table(doc,
    ["HTTP", "Code", "Ý nghĩa"],
    [
        ["400", "APP-400", "Malformed JSON / validation error"],
        ["400", "CART-5400", "Cart validation (qty âm, quá 999, sku không tồn tại)"],
        ["400", "ORD-5400", "Order validation (cart rỗng, address không hợp lệ)"],
        ["400", "VCH-5400", "Voucher invalid (hết hạn, không đủ min order)"],
        ["401", "AUTH-401", "Chưa đăng nhập / cookie không hợp lệ"],
        ["403", "AUTH-403", "Không đủ quyền (thiếu role/permission)"],
        ["403", "AUTH-LOCKED", "Tài khoản bị lock"],
        ["404", "APP-404", "Resource không tồn tại"],
        ["404", "PRD-404", "Product không tồn tại / đã xóa"],
        ["409", "AUTH-409", "Username/email đã tồn tại"],
        ["409", "INV-409", "Stock không đủ"],
        ["409", "ORD-409", "Order conflict (idempotency hit với payload khác)"],
        ["429", "APP-429", "Rate limit exceeded"],
        ["500", "APP-500", "Internal server error"],
        ["502", "GW-502", "Gateway: backend service không phản hồi"],
        ["503", "APP-503", "Service tạm dừng (maintenance hoặc dependency down)"],
        ["504", "GW-504", "Gateway timeout (backend chậm > 30s)"],
    ],
    col_widths_cm=[2.0, 3.0, 11.0])

add_heading(doc, "C. Tài khoản và dữ liệu seed", level=2)
add_table_caption(doc, "Bảng C.1: Tài khoản mặc định sau seed")
add_table(doc,
    ["Username", "Password", "Role", "Mục đích"],
    [
        ["admin", "Admin@2026", "ROLE_ADMIN", "Full quyền – quản trị, test backoffice"],
        ["staff", "Staff@2026", "ROLE_STAFF", "CSKH – test luồng đổi/trả, tra cứu"],
        ["customer", "Customer@2026", "ROLE_CUSTOMER", "Khách hàng – test luồng mua sắm"],
        ["test1..test5", "Test@2026", "ROLE_CUSTOMER", "Tài khoản test load/concurrent"],
    ],
    col_widths_cm=[3.0, 3.0, 3.5, 6.5])

add_table_caption(doc, "Bảng C.2: Dữ liệu seed khác")
add_table(doc,
    ["Loại dữ liệu", "Nguồn", "Số lượng", "Ghi chú"],
    [
        ["Categories", "seed.sh", "5 (QuanAo, DienThoai, Laptop, GiayDep, PhuKien)", "ID cố định 10-14"],
        ["Attrs", "seed.sh", "4 (Size, Color, Storage, RAM)", "Size=1, Color=2, Storage=3, RAM=4"],
        ["Attr Values", "seed.sh", "Size: S=1..XXL=5; Color: Đen=6..Xám=10", "ID cố định"],
        ["Vouchers", "seed.sh", "WELCOME10 (10%), SUMMER50K (50,000đ), VIP20 (20%, cap 500K)", "Active"],
        ["Products HIEU Luxe", "seed-hieu-luxe.py --wipe", "25 sản phẩm (5 túi, 5 giày, 2 đồng hồ, 3 da nhỏ, 5 may mặc, 5 phụ kiện)", "ID 26-50"],
        ["Stock", "seed-hieu-luxe.py", "Mỗi variant: random 5-100", "Auto qua catalog event → inventory"],
    ],
    col_widths_cm=[3.5, 3.5, 5.0, 4.0])

add_heading(doc, "D. Cấu hình mẫu", level=2)
add_para(doc, "D.1. application.yaml mẫu (auth-service):", bold=True)
add_code(doc, """spring:
  application:
    name: auth-service
  datasource:
    url: ${SPRING_DATASOURCE_URL:jdbc:postgresql://localhost:5433/auth_db}
    username: ${SPRING_DATASOURCE_USERNAME:postgres}
    password: ${SPRING_DATASOURCE_PASSWORD:postgres}
  jpa:
    hibernate:
      ddl-auto: validate
  flyway:
    locations: classpath:db/migration
  data:
    redis:
      host: ${SPRING_REDIS_HOST:localhost}
      port: ${SPRING_REDIS_PORT:6379}
  kafka:
    bootstrap-servers: ${KAFKA_BOOTSTRAP_SERVERS:localhost:9092}

eureka:
  client:
    service-url:
      defaultZone: ${EUREKA_SERVER_URL:http://localhost:8761/eureka/}

server:
  port: 9091

jwt:
  secret: ${JWT_SECRET}
  expiration: ${JWT_EXPIRATION:900000}
  refresh-expiration: ${JWT_REFRESH_EXPIRATION:604800000}

management:
  endpoints:
    web:
      exposure:
        include: health,info,metrics,prometheus,loggers
  endpoint:
    health:
      show-details: when-authorized""")

add_para(doc, "D.2. API Gateway route config (snippet):", bold=True)
add_code(doc, """spring:
  cloud:
    gateway:
      routes:
        - id: auth-public
          uri: lb://AUTH-SERVICE
          predicates:
            - Path=/api/v1/auth/register,/api/v1/auth/login,/api/v1/auth/refresh

        - id: auth-protected
          uri: lb://AUTH-SERVICE
          predicates:
            - Path=/api/v1/auth/**

        - id: catalog
          uri: lb://CATALOG-SERVICE
          predicates:
            - Path=/api/v1/products/**,/api/v1/categories/**,/api/v1/attrs/**

        - id: cart
          uri: lb://CART-SERVICE
          predicates:
            - Path=/api/v1/cart/**

        - id: order
          uri: lb://ORDER-SERVICE
          predicates:
            - Path=/api/v1/orders/**

        - id: payment
          uri: lb://PAYMENT-SERVICE
          predicates:
            - Path=/api/v1/payments/**""")

add_para(doc, "D.3. Docker Compose hạ tầng (snippet):", bold=True)
add_code(doc, """version: "3.8"
services:
  postgres-auth:
    image: postgres:16-alpine
    environment:
      POSTGRES_DB: auth_db
      POSTGRES_USER: postgres
      POSTGRES_PASSWORD: postgres
    ports: ["5433:5432"]
    volumes: [auth-db-data:/var/lib/postgresql/data]

  redis:
    image: redis:7-alpine
    ports: ["6379:6379"]

  kafka:
    image: bitnami/kafka:3.7
    environment:
      KAFKA_CFG_NODE_ID: 0
      KAFKA_CFG_PROCESS_ROLES: controller,broker
      KAFKA_CFG_LISTENERS: PLAINTEXT://:9092,CONTROLLER://:9093
      KAFKA_CFG_ADVERTISED_LISTENERS: PLAINTEXT://kafka:9092
      KAFKA_CFG_CONTROLLER_LISTENER_NAMES: CONTROLLER
      KAFKA_CFG_LISTENER_SECURITY_PROTOCOL_MAP: CONTROLLER:PLAINTEXT,PLAINTEXT:PLAINTEXT
      KAFKA_CFG_CONTROLLER_QUORUM_VOTERS: 0@kafka:9093
    ports: ["9092:9092"]

  elasticsearch:
    image: docker.elastic.co/elasticsearch/elasticsearch:8.13.0
    environment:
      discovery.type: single-node
      xpack.security.enabled: false
    ports: ["9200:9200"]

  kibana:
    image: docker.elastic.co/kibana/kibana:8.13.0
    ports: ["5601:5601"]
    depends_on: [elasticsearch]

  mailhog:
    image: mailhog/mailhog
    ports:
      - "1025:1025"
      - "8025:8025"
""")

# ============================================================
# Save
# ============================================================
doc.save(OUT_PATH)
print(f"✅ Generated: {OUT_PATH}")

# Stats
out = Document(OUT_PATH)
print(f"   Paragraphs: {len(out.paragraphs)}")
print(f"   Tables: {len(out.tables)}")
img = sum(1 for r in out.part.rels.values() if 'image' in r.target_ref)
print(f"   Images: {img}")
