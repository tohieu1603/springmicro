"""
Sinh file docx tài liệu chi tiết về auth-service.
Viết bằng tiếng Việt, format đẹp với heading, bảng, code block.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ─────────────────────────────────────────────────────────────

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code(doc, code, lang=''):
    """Thêm code block với font monospace, nền xám nhạt."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Nền xám nhạt
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        set_cell_background(cell, '2E5C8A')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def add_note(doc, text):
    """Box ghi chú."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('💡 ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8DC')
    pPr.append(shd)


# ── Document ────────────────────────────────────────────────────────────

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============================================================
# TRANG BÌA
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\nAUTH-SERVICE\n')
run.font.name = 'Calibri'
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Tài liệu thiết kế chi tiết & giải thích luồng nghiệp vụ\n\n')
run.font.size = Pt(16)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Hexagonal Architecture · CQRS · Event-Driven · Spring Boot\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ============================================================
# MỤC LỤC (manual TOC)
# ============================================================
add_heading(doc, 'Mục lục', level=1)

toc_items = [
    '1. Tổng quan auth-service',
    '2. Kiến trúc tổng thể - Hexagonal Architecture',
    '3. Tầng Domain - "Trái tim" nghiệp vụ',
    '4. Tầng Application - Điều phối use case (CQRS)',
    '5. Tầng Infrastructure - Adapter ra ngoài thế giới',
    '6. Tầng Interfaces - Cổng giao tiếp HTTP/gRPC',
    '7. Bảo mật - JWT, BCrypt, Token Rotation, Blacklist',
    '8. Hệ thống Event - Domain Event vs Integration Event',
    '9. Luồng nghiệp vụ chi tiết',
    '10. Giải thích các khái niệm Java/Spring trong code',
    '11. Quyết định thiết kế - Tại sao chọn cách này?',
    '12. Phụ lục - Cheatsheet & FAQ',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)
    for r in p.runs:
        r.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. TỔNG QUAN
# ============================================================
add_heading(doc, '1. Tổng quan auth-service', level=1)

add_para(doc,
    'auth-service là microservice đảm nhiệm toàn bộ chức năng xác thực (authentication) và phân quyền '
    '(authorization) trong hệ thống ecommerce. Đây là "cánh cổng" mà mọi service khác tin tưởng để '
    'kiểm chứng danh tính người dùng và quyền truy cập tài nguyên.')

add_heading(doc, '1.1 Trách nhiệm chính', level=2)
for item in [
    'Đăng ký tài khoản (Register) — tạo user mới, hash mật khẩu, gán role mặc định.',
    'Đăng nhập (Login) — xác thực username/email + password, phát hành JWT access + refresh token.',
    'Làm mới token (Refresh) — xoay vòng refresh token theo pattern Rotation + Family Revocation.',
    'Đăng xuất (Logout) — đưa access token vào blacklist Redis, thu hồi refresh token.',
    'Đổi mật khẩu (Change Password) — bump tokenVersion → vô hiệu hóa mọi JWT cũ tức thì.',
    'Quản lý user — phân quyền (assign/unassign role), khóa/mở tài khoản, đổi email.',
    'Kiểm tra phân quyền — cung cấp API check role / check permission cho các service khác qua gRPC.',
    'Phát event qua Kafka — thông báo "user vừa được tạo", "mật khẩu vừa đổi"... cho hệ thống bên ngoài.',
]:
    add_bullet(doc, item)

add_heading(doc, '1.2 Công nghệ sử dụng', level=2)

add_table(doc,
    headers=['Hạng mục', 'Công nghệ'],
    rows=[
        ['Ngôn ngữ', 'Java 21'],
        ['Framework', 'Spring Boot 3 (Web, Data JPA, Security, Validation)'],
        ['Database', 'PostgreSQL (chính), H2 (test/dev)'],
        ['Cache', 'Redis (blacklist token, cache role/permission)'],
        ['Message Broker', 'Apache Kafka (integration events ra ngoài service)'],
        ['JWT Library', 'jjwt (io.jsonwebtoken) — HS256'],
        ['Password Hashing', 'BCrypt (cost 12, theo khuyến nghị OWASP)'],
        ['ORM', 'Hibernate JPA'],
        ['Migration', 'Flyway'],
        ['Boilerplate', 'Lombok (@Slf4j, @RequiredArgsConstructor, @Getter...)'],
        ['API Docs', 'SpringDoc OpenAPI / Swagger UI'],
        ['Build Tool', 'Maven / Gradle'],
        ['Communication', 'REST (public), gRPC (internal service-to-service)'],
    ],
    col_widths=[5, 11])

doc.add_page_break()

# ============================================================
# 2. KIẾN TRÚC TỔNG THỂ
# ============================================================
add_heading(doc, '2. Kiến trúc tổng thể — Hexagonal Architecture', level=1)

add_para(doc,
    'auth-service áp dụng Hexagonal Architecture (còn gọi là Ports & Adapters). Đây là kiểu thiết kế '
    'tách phần "nghiệp vụ thuần" ra khỏi phần "công nghệ" (database, framework, broker...). Mục tiêu '
    'duy nhất: nếu mai mốt thay PostgreSQL bằng MongoDB, hoặc thay Spring bằng Quarkus, thì phần '
    'NGHIỆP VỤ KHÔNG ĐỔI 1 DÒNG.')

add_heading(doc, '2.1 Bốn tầng trong codebase', level=2)

add_code(doc, """auth-service/
├── domain/           ← Trái tim — POJO Java thuần, KHÔNG import framework
│   ├── models/       ← Aggregate Root: User, RefreshToken, Role, Permission
│   ├── events/       ← DomainEvent, DomainEventPublisher (interface = port)
│   ├── repositories/ ← UserRepository, RoleRepository... (interface = port)
│   ├── services/     ← PasswordEncoderPort, TokenProviderPort (interface = port)
│   └── shared/       ← AggregateRoot, DomainException
│
├── application/      ← Điều phối use case — orchestrate domain
│   ├── command/      ← LoginCommand, RegisterUserCommand... (input)
│   ├── query/        ← GetUserByIdQuery... (input cho read)
│   ├── handler/      ← LoginHandler, RegisterUserHandler... (xử lý command)
│   ├── dto/          ← UserDTO, AuthResponseDTO (output cho UI)
│   ├── port/         ← TokenBlacklistPort, RolePermissionCachePort
│   └── events/       ← IntegrationEvent (hợp đồng public với service khác)
│
├── infrastructure/   ← Adapter — implement các port của domain/application
│   ├── persistence/  ← UserRepositoryImpl + JPA entities + mapper
│   ├── security/     ← JwtTokenProvider, BCryptPasswordEncoderAdapter
│   ├── cache/        ← RedisRolePermissionCacheAdapter
│   ├── events/       ← SpringDomainEventPublisher
│   └── messaging/    ← KafkaIntegrationEventPublisher
│
└── interfaces/       ← Cổng giao tiếp ngoài — HTTP, gRPC
    ├── rest/         ← AuthController, UserController + filters
    └── grpc/         ← AuthGrpcService""")

add_heading(doc, '2.2 Hướng phụ thuộc — Quy tắc vàng', level=2)

add_para(doc,
    'Phụ thuộc CHỈ ĐƯỢC ĐI 1 CHIỀU từ ngoài vào trong:', bold=True)

add_code(doc, """interfaces → application → domain
              ↑
   infrastructure (implement port của domain)""")

add_para(doc, 'Cụ thể:', bold=True)
for item in [
    'domain KHÔNG được import bất cứ thứ gì từ application, infrastructure, interfaces.',
    'application có thể import domain, nhưng KHÔNG được import infrastructure trực tiếp.',
    'infrastructure import cả 3 — đây là "nơi nhận việc bẩn".',
    'interfaces import application + domain để adapt HTTP → command/query.',
]:
    add_bullet(doc, item)

add_note(doc,
    'Quy tắc này đảo chiều phụ thuộc (Dependency Inversion). Domain không phụ thuộc framework, mà '
    'framework "đăng ký" làm việc cho domain qua interface (port).')

add_heading(doc, '2.3 Vì sao tách phức tạp thế?', level=2)

add_table(doc,
    headers=['Tình huống', 'Kiến trúc thường', 'Hexagonal'],
    rows=[
        ['Đổi DB Postgres → Mongo', 'Phải sửa cả service + entity + repo', 'Chỉ viết lại UserRepositoryImpl'],
        ['Thay JWT bằng Paseto', 'Đụng vào mọi handler có token', 'Chỉ thay JwtTokenProvider'],
        ['Test domain logic', 'Phải start Spring + DB', 'JUnit thuần, ms-level'],
        ['Thay Kafka bằng RabbitMQ', 'Toàn bộ event flow refactor', 'Chỉ viết RabbitIntegrationEventPublisher'],
        ['Đổi BCrypt → Argon2', 'Replace mọi nơi gọi BCrypt', 'Thay 1 class adapter'],
    ],
    col_widths=[5, 5.5, 5.5])

doc.add_page_break()

# ============================================================
# 3. TẦNG DOMAIN
# ============================================================
add_heading(doc, '3. Tầng Domain — "Trái tim" nghiệp vụ', level=1)

add_para(doc,
    'Đây là phần CHỈ chứa quy tắc nghiệp vụ thuần. Không Spring, không Hibernate, không HTTP. '
    'Mục tiêu: viết được toàn bộ logic User/Token mà không cần biết Spring tồn tại.')

add_heading(doc, '3.1 Aggregate Root — Đơn vị nhất quán', level=2)

add_para(doc,
    'Aggregate Root là entity "cấp cao nhất" của 1 nhóm object, kiểm soát mọi thao tác lên nhóm đó. '
    'Trong auth-service có 4 aggregate root:')
for item in [
    'User — quản lý identity, credentials, roles, account status, tokenVersion.',
    'RefreshToken — quản lý 1 token cụ thể với family + generation.',
    'Role — nhóm quyền (ROLE_USER, ROLE_ADMIN...).',
    'Permission — quyền nguyên tử (USER_READ, ORDER_CREATE...).',
]:
    add_bullet(doc, item)

add_heading(doc, '3.2 Lớp AggregateRoot — Quản lý DomainEvent', level=2)

add_para(doc,
    'Mọi aggregate kế thừa lớp này. Nó cung cấp cơ chế "ghi nhớ event trong quá trình mutate" và '
    'cho phép infrastructure "rút ruột" event sau khi save thành công.')

add_code(doc, """public abstract class AggregateRoot {
    private final transient List<DomainEvent> domainEvents = new ArrayList<>();

    protected final void registerEvent(DomainEvent event) {
        if (event == null) return;
        this.domainEvents.add(event);
    }

    public final List<DomainEvent> pullDomainEvents() {
        if (domainEvents.isEmpty()) return List.of();
        List<DomainEvent> snapshot = List.copyOf(domainEvents);
        domainEvents.clear();
        return snapshot;
    }
}""")

add_para(doc, 'Tại sao thiết kế thế?', bold=True)
for item in [
    'Aggregate ghi event vào buffer khi state thay đổi (registerEvent).',
    'Sau khi repository save() thành công vào DB → gọi pullDomainEvents() để lấy event ra publish.',
    'Một lần "pull" duy nhất → tránh bug "quên clearEvents" gây duplicate event.',
    'transient để Hibernate không persist list này vào DB (nó chỉ tồn tại trong RAM).',
]:
    add_bullet(doc, item)

add_heading(doc, '3.3 User aggregate — Ví dụ cụ thể', level=2)

add_code(doc, """public final class User extends AggregateRoot {
    private UserId id;
    private Username username;
    private Email email;
    private Password password;
    private Set<RoleId> roles;
    private int tokenVersion;        // ← khóa quan trọng để invalidate JWT
    ...

    public static User register(Username u, Email e, Password rawPwd,
                                PersonName name, PasswordEncoderPort encoder) {
        User user = new User();
        user.id = UserId.generate();
        user.password = Password.createEncoded(encoder.encode(rawPwd.value()));
        ...
        user.registerEvent(new UserCreatedEvent(user.id.value(), ...));
        return user;
    }

    public void changePassword(Password old, Password newPwd, PasswordEncoderPort encoder) {
        if (!encoder.matches(old.value(), password.value())) {
            throw new InvalidCredentialsException();
        }
        password = Password.createEncoded(encoder.encode(newPwd.value()));
        incrementTokenVersion();   // ← KEY: bump version → mọi JWT cũ tự invalid
        registerEvent(new PasswordChangedEvent(...));
    }
}""")

add_para(doc, 'Điểm đáng chú ý:', bold=True)
for item in [
    'Constructor private — không cho new từ ngoài. Tạo User qua factory register() / reconstitute().',
    'Mọi mutation đi qua phương thức nghiệp vụ — đảm bảo invariants không bị phá.',
    'Plaintext password KHÔNG BAO GIỜ chạm tới persistence — hash ngay tại register().',
    'incrementTokenVersion() là "núm vặn" để invalidate mọi JWT cũ tức khắc khi đổi mật khẩu.',
    'registerEvent() raise event nhưng KHÔNG publish — chờ repository drain sau khi save xong.',
]:
    add_bullet(doc, item)

add_heading(doc, '3.4 Value Object (VO) — Bất biến, ép buộc invariant', level=2)

add_para(doc,
    'Value Object là object nhỏ chỉ chứa giá trị (không có ID). Mục đích: ép buộc quy tắc kiểm tra '
    'NGAY KHI tạo, không để giá trị invalid lan ra hệ thống.')

add_code(doc, """// VO Email - kiểm tra format ngay khi tạo
public record Email(String value) {
    public Email {
        if (value == null || !value.contains("@"))
            throw new IllegalArgumentException("Invalid email");
    }
    public static Email of(String raw) { return new Email(raw.toLowerCase().trim()); }
}

// VO Password - phân biệt raw vs encoded
public record Password(String value, boolean needsEncoding) {
    public static Password createRaw(String raw) { return new Password(raw, true); }
    public static Password createEncoded(String hash) { return new Password(hash, false); }
}""")

add_para(doc, 'Lợi ích:', bold=True)
for item in [
    'Compile-time safety: hàm nhận Email không thể bị truyền nhầm String/Username.',
    'Không có Email invalid trong RAM — kiểm tra một lần tại VO, dùng yên tâm.',
    'Tự documenting — đọc signature "register(Username, Email, Password)" là hiểu ngay.',
    'Password phân biệt raw/encoded tránh bug đáng sợ: lưu plaintext vào DB.',
]:
    add_bullet(doc, item)

add_heading(doc, '3.5 Domain Event — Sự kiện trong nội bộ JVM', level=2)

add_para(doc,
    'Khi entity thay đổi state quan trọng, nó "raise" event để các phần khác biết. Event là POJO '
    'thuần, không qua mạng, không serialize — chỉ chạy trong JVM hiện tại.')

add_code(doc, """public abstract class DomainEvent {
    private final UUID eventId = UUID.randomUUID();
    private final Instant occurredOn = Instant.now();
    public abstract String aggregateId();
    public String eventType() { return getClass().getSimpleName(); }
}

// Event cụ thể:
public class PasswordChangedEvent extends DomainEvent {
    private final String userId;
    private final String username;
    @Override public String aggregateId() { return userId; }
}""")

add_para(doc,
    'DomainEventPublisher là interface (port) — domain không biết publish thế nào, infra lo:')

add_code(doc, """// domain/events/DomainEventPublisher.java
public interface DomainEventPublisher {
    void publish(DomainEvent event);
}

// infrastructure/events/SpringDomainEventPublisher.java
@Component
public class SpringDomainEventPublisher implements DomainEventPublisher {
    private final ApplicationEventPublisher delegate;   // ← của Spring
    public void publish(DomainEvent event) {
        delegate.publishEvent(event);
    }
}""")

add_heading(doc, '3.6 Port — Interface "yêu cầu công việc" từ domain', level=2)

add_para(doc,
    'Domain cần các "công cụ" để hoàn thành nghiệp vụ (mã hóa password, sinh JWT, lưu DB...). Nhưng '
    'domain không được biết các công cụ đó được làm bằng gì. Giải pháp: domain định nghĩa INTERFACE '
    '(gọi là "port"), infra implement.')

add_table(doc,
    headers=['Port (interface)', 'Vị trí', 'Implementation', 'Vị trí'],
    rows=[
        ['UserRepository', 'domain/repositories/', 'UserRepositoryImpl', 'infrastructure/persistence/impl/'],
        ['PasswordEncoderPort', 'domain/services/', 'BCryptPasswordEncoderAdapter', 'infrastructure/security/'],
        ['TokenProviderPort', 'domain/services/', 'JwtTokenProvider', 'infrastructure/security/'],
        ['DomainEventPublisher', 'domain/events/', 'SpringDomainEventPublisher', 'infrastructure/events/'],
        ['TokenBlacklistPort', 'application/port/', 'TokenBlacklistService', 'infrastructure/security/'],
        ['RolePermissionCachePort', 'application/port/', 'RedisRolePermissionCacheAdapter', 'infrastructure/cache/'],
    ],
    col_widths=[4, 3.5, 4.5, 4])

doc.add_page_break()

# ============================================================
# 4. TẦNG APPLICATION
# ============================================================
add_heading(doc, '4. Tầng Application — Điều phối use case (CQRS)', level=1)

add_para(doc,
    'Application là tầng "đạo diễn" — không chứa logic nghiệp vụ chi tiết (đó là việc của domain), '
    'mà điều phối các bước để hoàn thành 1 use case: gọi repo lấy dữ liệu → gọi domain method → '
    'lưu lại → trả response.')

add_heading(doc, '4.1 CQRS — Tách Command và Query', level=2)

add_para(doc,
    'CQRS (Command Query Responsibility Segregation) phân biệt rõ:')

add_table(doc,
    headers=['Loại', 'Mục đích', 'Ví dụ'],
    rows=[
        ['Command', 'Thay đổi state (write)', 'LoginCommand, RegisterUserCommand, ChangePasswordCommand'],
        ['Query', 'Đọc dữ liệu (read)', 'GetUserByIdQuery, ListUsersQuery, CheckPermissionQuery'],
    ],
    col_widths=[3, 5, 8])

add_heading(doc, '4.2 Pattern Command + Handler', level=2)

add_para(doc,
    '1 Command (input) ↔ 1 Handler (logic). Mỗi handler xử lý đúng 1 use case → đơn nhiệm, dễ test.')

add_code(doc, """// Command - chỉ là input data, không có logic
public record LoginCommand(
        String usernameOrEmail,
        String rawPassword
) implements Command<AuthResponseDTO> { }

// Handler - chứa logic gọi domain + repo
@Service
@RequiredArgsConstructor
public class LoginHandler implements CommandHandler<LoginCommand, AuthResponseDTO> {

    private final UserRepository userRepository;
    private final RoleRepository roleRepository;
    private final PasswordEncoderPort passwordEncoder;
    private final TokenProviderPort tokenProvider;
    ...

    @Override
    @Transactional
    public AuthResponseDTO handle(LoginCommand command) {
        User user = lookup(command.usernameOrEmail()).orElseThrow(InvalidCredentialsException::new);
        if (!user.authenticate(Password.createRaw(command.rawPassword()), passwordEncoder)) {
            throw new InvalidCredentialsException();
        }
        User saved = userRepository.save(user);

        // Issue tokens
        var issued = tokenProvider.issueAccessToken(saved, roleNames);
        RefreshToken refresh = tokenDomainService.issueForUser(saved, refreshExpiryDays);
        refreshTokenRepository.save(refresh);

        return AuthResponseDTO.bearer(issued.token(), refresh.getValue().value(), ...);
    }
}""")

add_para(doc, 'Lợi ích:', bold=True)
for item in [
    'Mỗi handler chỉ làm 1 việc — đọc 50 dòng là hiểu toàn bộ use case.',
    'Test đơn giản: new LoginHandler(mockUserRepo, mockEncoder, ...).handle(loginCmd).',
    'Controller trở nên "siêu mỏng" — chỉ adapt HTTP → Command, không có logic.',
    'Dễ thêm cross-cutting (logging, metrics, retry) qua AOP/decorator.',
]:
    add_bullet(doc, item)

add_heading(doc, '4.3 Integration Event — Hợp đồng public với service khác', level=2)

add_para(doc,
    'Khác với DomainEvent (in-process), IntegrationEvent là event GỬI RA NGOÀI service qua Kafka. '
    'Nó cần stable, có schema version, format phẳng để service khác (có thể viết bằng Python/Node) đọc được.')

add_code(doc, """public interface IntegrationEvent {
    UUID eventId();         // unique id - dùng để dedupe ở consumer
    String eventType();     // "auth.user.created.v1" - stable string
    Instant occurredOn();
    String aggregateId();
    int schemaVersion();    // bump khi breaking change
}

public record UserLifecycleIntegrationEvent(
    UUID eventId, String eventType, Instant occurredOn,
    String aggregateId, int schemaVersion,
    Map<String, Object> payload    // flexible, theo eventType
) implements IntegrationEvent { }""")

add_heading(doc, '4.4 KafkaTopics — Single source of truth cho tên topic', level=2)

add_code(doc, """public final class KafkaTopics {
    private KafkaTopics() {}   // ← cấm new

    public static final String AUTH_USER_EVENTS    = "auth.user.events.v1";
    public static final String AUTH_SESSION_EVENTS = "auth.session.events.v1";
}""")

add_para(doc, 'Vì sao tách thành utility class?', bold=True)
for item in [
    'Topic name dùng ở nhiều nơi (publisher, consumer, config) — đặt 1 chỗ, đổi 1 lần.',
    'final class + private constructor = utility class, không tạo instance được.',
    'Hậu tố .v1 trong tên topic → schema versioning, đổi breaking change → tạo .v2 chạy song song.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# 5. TẦNG INFRASTRUCTURE
# ============================================================
add_heading(doc, '5. Tầng Infrastructure — Adapter ra ngoài thế giới', level=1)

add_para(doc,
    'Đây là tầng "thực thi" — implement mọi port của domain/application bằng công nghệ cụ thể.')

add_heading(doc, '5.1 Persistence Adapter — JPA/Hibernate', level=2)

add_para(doc, 'Cấu trúc:', bold=True)
add_code(doc, """infrastructure/persistence/
├── jpa/
│   ├── entities/     ← UserJpaEntity, RefreshTokenJpaEntity (Hibernate @Entity)
│   └── repositories/ ← UserJpaRepository extends JpaRepository
├── mapper/           ← UserJpaMapper (convert domain ↔ entity)
└── impl/             ← UserRepositoryImpl implements UserRepository""")

add_para(doc,
    'Tại sao tách JpaEntity và Domain model?', bold=True)
for item in [
    'JpaEntity bị "trói" với Hibernate (@Entity, @Id, @Column, getter/setter cho lazy loading...).',
    'Domain model là POJO sạch, có invariants, không bị Hibernate "đụng" vào.',
    'Mapper là cầu nối — viết tay hoặc dùng MapStruct.',
    'Nếu mai mốt đổi sang MongoDB: chỉ thay JpaEntity + UserRepositoryImpl, domain không đổi.',
]:
    add_bullet(doc, item)

add_para(doc, 'Ví dụ UserRepositoryImpl:', bold=True)
add_code(doc, """@Repository
@RequiredArgsConstructor
public class UserRepositoryImpl implements UserRepository {

    private final UserJpaRepository jpaRepository;
    private final UserJpaMapper mapper;
    private final DomainEventPublisher eventPublisher;

    @Override
    @Transactional
    public User save(User user) {
        UserJpaEntity entity = mapper.toJpaEntity(user, ...);
        UserJpaEntity saved = jpaRepository.save(entity);

        // KEY: Drain events SAU KHI persist thành công
        user.pullDomainEvents().forEach(eventPublisher::publish);

        return mapper.toDomain(saved);
    }
}""")

add_note(doc,
    'Drain events SAU khi save thành công đảm bảo: nếu DB rollback → event không bị publish nhầm. '
    'Kết hợp với @TransactionalEventListener(AFTER_COMMIT) ở Kafka publisher → at-least-once delivery '
    'thật sự (event chỉ ra Kafka khi transaction đã commit).')

add_heading(doc, '5.2 Security Adapter — JWT + BCrypt', level=2)

add_code(doc, """// JwtTokenProvider implements TokenProviderPort
@Component
public class JwtTokenProvider implements TokenProviderPort {
    private static final MacAlgorithm ALGORITHM = Jwts.SIG.HS256;
    private final SecretKey key;

    @Override
    public IssuedAccessToken issueAccessToken(User user, Set<String> roles) {
        Instant now = Instant.now();
        String tokenId = UUID.randomUUID().toString();

        Map<String, Object> claims = new HashMap<>();
        claims.put("jti", tokenId);
        claims.put("userId", user.getId().value());
        claims.put("tokenVersion", user.getTokenVersion());  // ← KEY
        claims.put("roles", roles);

        String token = Jwts.builder()
                .id(tokenId)
                .claims(claims)
                .issuer(props.issuer())
                .signWith(key, ALGORITHM)
                .compact();

        return new IssuedAccessToken(token, tokenId, exp, ttlSeconds);
    }
}""")

add_para(doc, 'Các custom claims trong JWT:', bold=True)
add_table(doc,
    headers=['Claim', 'Ý nghĩa', 'Mục đích bảo mật'],
    rows=[
        ['jti', 'JWT ID duy nhất', 'Key để blacklist khi logout/revoke'],
        ['userId', 'UUID của user', 'Lookup user nhanh không cần parse subject'],
        ['tokenVersion', 'Phiên bản token của user', 'Đổi mật khẩu → bump version → JWT cũ tự invalid'],
        ['roles', 'Danh sách role tại thời điểm phát hành', 'Authorize không cần query DB mỗi request'],
        ['issuer', 'auth-service', 'Chặn token được mint bởi service khác cùng secret'],
        ['exp', 'Thời điểm hết hạn', 'Hết hạn tự động'],
    ],
    col_widths=[3, 5, 8])

add_heading(doc, '5.3 TokenBlacklistService — Redis + DB hybrid', level=2)

add_para(doc,
    'Khi user logout, access token chưa hết hạn — phải đưa vào blacklist để filter chặn. Pattern dùng '
    'Redis (nhanh) + DB (đảm bảo durable):')

add_code(doc, """@Service
public class TokenBlacklistService implements TokenBlacklistPort {

    private static final String BLACKLIST_PREFIX = "blacklist:";
    private final StringRedisTemplate redisTemplate;
    private final TokenRevocationJpaRepository repository;

    @Override
    @Transactional
    public void revoke(String tokenId, String userId, Instant expiresAt, String reason) {
        long ttl = Duration.between(Instant.now(), expiresAt).getSeconds();
        if (ttl > 0) {
            try {
                redisTemplate.opsForValue().set(
                    BLACKLIST_PREFIX + tokenId, reason, Duration.ofSeconds(ttl));
            } catch (Exception e) {
                log.warn("Redis down, fallback to DB only");
            }
        }
        repository.save(TokenRevocationJpaEntity.forRevocation(...));
    }

    @Override
    public boolean isRevoked(String tokenId) {
        try {
            return Boolean.TRUE.equals(redisTemplate.hasKey(BLACKLIST_PREFIX + tokenId));
        } catch (Exception e) {
            return repository.existsById(tokenId);  // fallback
        }
    }

    @EventListener(ApplicationReadyEvent.class)
    public void warmUpRedisFromDb() {
        // Khởi động → load mọi revocation chưa hết hạn từ DB lên Redis
    }
}""")

add_para(doc, 'Tại sao hybrid?', bold=True)
for item in [
    'Redis: tốc độ check ~1ms, mọi request đều phải check → cần nhanh.',
    'DB: source of truth — Redis chết hoặc bị xóa data → vẫn còn DB.',
    'TTL trong Redis = thời gian còn lại của JWT → tự dọn dẹp, không cần cleanup thủ công.',
    'warmUpRedisFromDb() ở startup → Redis instance mới không bỏ sót revocation cũ.',
    'Scheduled cleanup mỗi giờ → xóa các row DB đã hết hạn tự nhiên.',
]:
    add_bullet(doc, item)

add_heading(doc, '5.4 KafkaIntegrationEventPublisher — Bridge domain → Kafka', level=2)

add_code(doc, """@Component
@RequiredArgsConstructor
public class KafkaIntegrationEventPublisher {

    private final KafkaTemplate<String, Object> kafkaTemplate;
    private final DomainEventToIntegrationEventMapper mapper;

    @TransactionalEventListener(phase = TransactionPhase.AFTER_COMMIT)
    public void onDomainEvent(DomainEvent event) {
        var routed = mapper.map(event);
        if (routed == null) return;

        try {
            kafkaTemplate.send(routed.topic(), routed.key(), routed.event());
        } catch (Exception e) {
            log.warn("Failed to publish: {}", e.getMessage());
        }
    }
}""")

add_para(doc, 'Điểm hay:', bold=True)
for item in [
    '@TransactionalEventListener(AFTER_COMMIT) → chỉ chạy KHI transaction DB đã commit thành công.',
    'Tránh bug: DB rollback nhưng Kafka đã publish → service khác xử lý event "ma".',
    'Mapper riêng (DomainEventToIntegrationEventMapper) → switch pattern chuyển domain event → integration event.',
    'Lỗi gửi Kafka → log warn, không retry → trade-off đơn giản. Nếu cần at-least-once → upgrade lên Transactional Outbox pattern.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# 6. TẦNG INTERFACES
# ============================================================
add_heading(doc, '6. Tầng Interfaces — Cổng giao tiếp HTTP/gRPC', level=1)

add_heading(doc, '6.1 REST Controller — Siêu mỏng', level=2)

add_para(doc,
    'Controller chỉ làm 3 việc: nhận HTTP request, biến thành Command, gọi handler, map response. '
    'KHÔNG có business logic — đó là việc của handler.')

add_code(doc, """@RestController
@RequestMapping("/api/v1/auth")
@RequiredArgsConstructor
public class AuthController {

    private final CommandHandler<LoginCommand, AuthResponseDTO> loginHandler;
    private final AuthCookieWriter cookieWriter;

    @PostMapping("/login")
    public ResponseEntity<AuthResponseDTO> login(@Valid @RequestBody LoginRequest req) {
        AuthResponseDTO tokens = loginHandler.handle(
            new LoginCommand(req.usernameOrEmail(), req.password()));
        return cookieWriter.writeTokens(tokens);   // ← Set HttpOnly cookies
    }
    ...
}""")

add_para(doc, 'Endpoints chính:', bold=True)
add_table(doc,
    headers=['HTTP Method', 'Path', 'Mục đích', 'Public?'],
    rows=[
        ['POST', '/api/v1/auth/register', 'Đăng ký user mới', 'Yes'],
        ['POST', '/api/v1/auth/login', 'Đăng nhập username/email + password', 'Yes'],
        ['POST', '/api/v1/auth/refresh', 'Xoay vòng refresh token', 'Yes'],
        ['POST', '/api/v1/auth/logout', 'Đăng xuất, blacklist access token', 'Auth required'],
        ['POST', '/api/v1/auth/change-password', 'Đổi mật khẩu, bump tokenVersion', 'Auth required'],
        ['GET', '/api/v1/auth/me', 'Thông tin user hiện tại', 'Auth required'],
        ['GET', '/api/v1/users/...', 'Quản lý user (list, view, update)', 'Auth + role'],
    ],
    col_widths=[2.5, 4.5, 6, 3])

add_heading(doc, '6.2 Cookie-based Auth + Bearer fallback', level=2)

add_para(doc,
    'Khác với JWT thông thường trả về body, service này set HttpOnly cookies (ACCESS_TOKEN, '
    'REFRESH_TOKEN). Tại sao?')
for item in [
    'HttpOnly cookie KHÔNG đọc được từ JavaScript → chống XSS đánh cắp token.',
    'Browser tự động đính kèm cookie mỗi request → frontend không phải tự xử Authorization header.',
    'Fallback: JwtAuthenticationFilter vẫn chấp nhận Authorization: Bearer ... → mobile/CLI/gRPC vẫn dùng được.',
]:
    add_bullet(doc, item)

add_heading(doc, '6.3 JwtAuthenticationFilter — Trạm gác mỗi request', level=2)

add_code(doc, """@Component
public class JwtAuthenticationFilter extends OncePerRequestFilter {

    @Override
    protected void doFilterInternal(HttpServletRequest req, HttpServletResponse res,
                                    FilterChain chain) {
        String token = extractBearer(req);   // cookie trước, header sau

        if (token != null && SecurityContextHolder.getContext().getAuthentication() == null) {
            try {
                // 1. Parse + verify signature
                AccessClaims claims = tokenProvider.parseAccessToken(token);

                // 2. Check blacklist (Redis)
                if (tokenBlacklist.isRevoked(claims.tokenId())) {
                    res.sendError(401, "Token revoked");
                    return;
                }

                // 3. Load user details (status, authorities)
                UserDetails userDetails = userDetailsService.loadUserById(claims.userId());

                // 4. Check tokenVersion - reject nếu user đã đổi mật khẩu sau khi mint
                if (userDetails instanceof AuthUserDetails aud
                        && aud.tokenVersion() != claims.tokenVersion()) {
                    res.sendError(401, "Token superseded");
                    return;
                }

                // 5. Set SecurityContext
                SecurityContextHolder.getContext().setAuthentication(...);
            } catch (Exception ignored) { /* log debug */ }
        }
        chain.doFilter(req, res);
    }
}""")

add_para(doc, 'Filter chạy qua 5 bước kiểm tra:', bold=True)
for item in [
    'Parse + verify chữ ký HMAC (bằng secret key).',
    'Check blacklist Redis — đã revoke chưa?',
    'Load AuthUserDetails từ DB qua CustomUserDetailsService.',
    'So tokenVersion trong JWT với version hiện tại của user — đổi mật khẩu sẽ bump version, JWT cũ tự fail.',
    'Set Authentication vào SecurityContext → controller dùng @AuthenticationPrincipal lấy user.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# 7. BẢO MẬT
# ============================================================
add_heading(doc, '7. Bảo mật — Cơ chế phòng thủ nhiều lớp', level=1)

add_heading(doc, '7.1 BCrypt cost 12 — Password hashing', level=2)
add_para(doc,
    'BCrypt là thuật toán hash chậm có chủ đích — chống brute force. Cost 12 là khuyến nghị OWASP, '
    'mỗi lần hash mất ~250ms trên server hiện đại. Attacker brute force trên GPU cũng chỉ vài '
    'nghìn hash/giây → không khả thi.')

add_heading(doc, '7.2 JWT tokenVersion — Vô hiệu hóa token tức thì', level=2)
add_para(doc,
    'Vấn đề kinh điển của JWT: phát rồi không thu hồi được tới khi hết hạn. Giải pháp đặt số version '
    'trong JWT + lưu version hiện tại trong DB của user.')
add_code(doc, """// Khi đổi mật khẩu:
user.changePassword(oldPwd, newPwd, encoder);  // → incrementTokenVersion() trong User

// Mọi request sau đó:
if (jwt.tokenVersion != user.currentTokenVersion) → 401 Unauthorized

// Hệ quả: Tất cả JWT phát trước đó (dù chưa hết hạn) đều tự động invalid.""")

add_heading(doc, '7.3 Refresh Token Rotation + Family Revocation', level=2)
add_para(doc,
    'Mỗi lần refresh: revoke token cũ, phát token mới cùng "family". Nếu attacker cướp được token cũ '
    'và đem dùng → server thấy "token này đã bị revoke" → revoke CẢ FAMILY → cả user thật và attacker '
    'đều bị đẩy ra → user phải login lại.')

add_code(doc, """public RefreshToken rotate(RefreshToken old, int expiryDays, RefreshTokenRepository repo) {
    if (old.isReuseAttempt()) {
        revokeFamily(old.getFamily(), repo);   // ← bắt được kẻ tấn công
        throw new TokenReuseDetectedException(...);
    }

    old.verifyValidity();
    old.revoke(RevokedReason.NORMAL);
    repo.save(old);

    return RefreshToken.rotate(old, expiryDays);  // gen+1, cùng family
}""")

add_heading(doc, '7.4 Pessimistic Lock chống race condition', level=2)
add_para(doc,
    'Khi 2 request refresh cùng 1 token chạy đồng thời (network retry, double-click) → có thể cả 2 '
    'cùng pass → 2 access token mới được phát. Giải pháp: SELECT FOR UPDATE.')
add_code(doc, """RefreshToken presented = refreshTokenRepository
    .findByTokenValueForUpdate(TokenValue.of(command.refreshToken()))  // ← SELECT FOR UPDATE
    .orElseThrow(...);""")

add_heading(doc, '7.5 Rate Limiting', level=2)
add_para(doc,
    'RateLimitFilter chạy TRƯỚC JwtAuthenticationFilter — chặn brute-force ngay tầng đầu, không tốn '
    'CPU parse JWT cho request bị throttle.')

add_heading(doc, '7.6 Issuer Validation', level=2)
add_para(doc,
    'JWT có claim issuer = "auth-service". Khi parse, requireIssuer("auth-service") chặn token mint '
    'bởi service khác dùng chung secret (ví dụ payment-service dùng cùng HS256 key cho internal token).')

doc.add_page_break()

# ============================================================
# 8. HỆ THỐNG EVENT
# ============================================================
add_heading(doc, '8. Hệ thống Event — Domain Event vs Integration Event', level=1)

add_para(doc,
    'Hai loại event hoàn toàn khác mục đích. Đây là sự phân biệt quan trọng nhất khi đọc code này.')

add_heading(doc, '8.1 Bảng so sánh', level=2)
add_table(doc,
    headers=['Tiêu chí', 'Domain Event', 'Integration Event'],
    rows=[
        ['Phạm vi', 'In-process (cùng JVM)', 'Cross-service (qua Kafka)'],
        ['Vị trí định nghĩa', 'domain/events/', 'application/events/'],
        ['Kênh truyền', 'Spring ApplicationEventPublisher', 'Kafka (KafkaTemplate)'],
        ['Format', 'POJO Java', 'JSON (record)'],
        ['eventType', 'Class name (UserCreatedEvent)', 'String stable ("auth.user.created.v1")'],
        ['Ổn định', 'Tự do refactor', 'Schema versioned, không phá'],
        ['Consumer', 'Cùng codebase', 'Service khác (có thể khác ngôn ngữ)'],
        ['Vai trò', 'Decouple TRONG service', 'Decouple GIỮA service'],
    ],
    col_widths=[3, 6, 6])

add_heading(doc, '8.2 Luồng từ domain event → Kafka', level=2)

add_code(doc, """┌─────────────────────────────────────────────────────────────────────┐
│  1. User.changePassword() → registerEvent(PasswordChangedEvent)     │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
┌─────────────────────────────────────────────────────────────────────┐
│  2. UserRepositoryImpl.save(user)                                   │
│     → jpaRepository.save(entity)                                    │
│     → user.pullDomainEvents().forEach(eventPublisher::publish)      │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
┌─────────────────────────────────────────────────────────────────────┐
│  3. SpringDomainEventPublisher.publish(event)                       │
│     → ApplicationEventPublisher.publishEvent(event) [Spring]        │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
                  (Transaction commits)
                            ↓
┌─────────────────────────────────────────────────────────────────────┐
│  4. @TransactionalEventListener(AFTER_COMMIT) lắng nghe             │
│     → KafkaIntegrationEventPublisher.onDomainEvent(event)           │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
┌─────────────────────────────────────────────────────────────────────┐
│  5. DomainEventToIntegrationEventMapper.map(event)                  │
│     → Routed(topic, key, integrationEvent)                          │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
┌─────────────────────────────────────────────────────────────────────┐
│  6. kafkaTemplate.send(topic, key, integrationEvent)                │
│     → Kafka topic "auth.user.events.v1"                             │
└─────────────────────────────────────────────────────────────────────┘
                            ↓
        Service khác (notification, billing, audit) consume""")

add_heading(doc, '8.3 Tại sao tách 2 record Integration Event?', level=2)

add_para(doc,
    'Có 2 record với field giống hệt nhau: UserLifecycleIntegrationEvent và SessionIntegrationEvent. '
    'Vì sao không gộp 1?')

add_table(doc,
    headers=['Event', 'Kafka Topic', 'Ai consume?'],
    rows=[
        ['UserLifecycleIntegrationEvent', 'auth.user.events.v1', 'Notification, Billing, Profile services'],
        ['SessionIntegrationEvent', 'auth.session.events.v1', 'Audit, Security monitoring, Analytics'],
    ],
    col_widths=[5, 4.5, 7])

add_para(doc, 'Lý do tách:', bold=True)
for item in [
    'Routing tới topic khác nhau → consumer pattern khác, scale khác.',
    'Method overloading: publisher dispatch dựa trên kiểu → type-safe, không nhầm topic.',
    'Service consume audit không cần subscribe user lifecycle (tiết kiệm bandwidth).',
    'Schema riêng biệt → mai sau evolve độc lập.',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ============================================================
# 9. LUỒNG NGHIỆP VỤ
# ============================================================
add_heading(doc, '9. Luồng nghiệp vụ chi tiết', level=1)

add_heading(doc, '9.1 Luồng Register (đăng ký)', level=2)
add_code(doc, """┌─ HTTP POST /api/v1/auth/register
│  Body: {username, email, password, firstName, lastName}
│
├→ AuthController.register()
│  └→ new RegisterUserCommand(...)
│
├→ RegisterUserHandler.handle(cmd)
│  ├→ Username.of() / Email.of() — VO ép buộc format
│  ├→ userRepository.existsByUsername() → nếu có → throw UserAlreadyExistsException
│  ├→ userRepository.existsByEmail() → idem
│  ├→ User.register(username, email, Password.createRaw(rawPwd), name, encoder)
│  │   ├→ encoder.encode(rawPwd)  ← BCrypt hash, plaintext không xuống DB
│  │   └→ registerEvent(new UserCreatedEvent(...))
│  ├→ roleRepository.findByName("ROLE_USER")
│  │  └→ user.assignRole(role.getId())
│  │       └→ registerEvent(new RoleAssignedEvent(...))
│  ├→ userRepository.save(user)
│  │   ├→ jpaRepository.save(userJpaEntity)
│  │   └→ pullDomainEvents() → publish [UserCreatedEvent, RoleAssignedEvent]
│  ├→ tokenProvider.issueAccessToken(user, roles)  ← JWT HS256
│  ├→ tokenDomainService.issueForUser(user, 7)     ← Refresh token, gen=0, family mới
│  └→ refreshTokenRepository.save(refresh)
│      └→ registerEvent(new TokenCreatedEvent(...))
│
├→ AuthCookieWriter.writeTokens()
│  └→ Set-Cookie: ACCESS_TOKEN=...; HttpOnly; Secure; SameSite=Lax
│  └→ Set-Cookie: REFRESH_TOKEN=...; HttpOnly; Secure; SameSite=Lax
│
├→ (Transaction commits)
├→ @TransactionalEventListener(AFTER_COMMIT) → KafkaIntegrationEventPublisher
│  ├→ map(UserCreatedEvent) → "auth.user.created.v1" → Kafka "auth.user.events.v1"
│  ├→ map(RoleAssignedEvent) → "auth.user.role_assigned.v1" → ...
│  └→ map(TokenCreatedEvent) → "auth.token.created.v1" → Kafka "auth.session.events.v1"
│
└→ 200 OK + body {accessToken, refreshToken, user: {...}}""")

add_heading(doc, '9.2 Luồng Login', level=2)
add_code(doc, """┌─ HTTP POST /api/v1/auth/login {usernameOrEmail, password}
├→ LoginHandler.handle(cmd)
│  ├→ lookup(usernameOrEmail) → user (hoặc throw InvalidCredentials)
│  ├→ user.authenticate(rawPwd, encoder)
│  │   ├→ ensureAuthenticatable() — check disabled/locked/expired
│  │   ├→ encoder.matches(rawPwd, password.hash) — BCrypt verify (~250ms)
│  │   └→ if ok: recordLogin() → registerEvent(UserLoggedInEvent)
│  ├→ userRepository.save(user)
│  ├→ tokenProvider.issueAccessToken(user, roles) — JWT mới
│  ├→ tokenDomainService.issueForUser(user, 7) — refresh token mới
│  └→ refreshTokenRepository.save(refresh)
│
├→ Set HttpOnly cookies
└→ Sau commit: Kafka → "auth.user.logged_in.v1" + "auth.token.created.v1"

Timing safety: Lookup failure VÀ password mismatch đều throw cùng InvalidCredentialsException
→ attacker không phân biệt được "username không tồn tại" vs "password sai" qua timing.""")

add_heading(doc, '9.3 Luồng Refresh Token (Rotation + Family Revocation)', level=2)
add_code(doc, """┌─ HTTP POST /api/v1/auth/refresh (refresh token đọc từ cookie)
├→ RefreshTokenHandler.handle(cmd)
│  ├→ refreshTokenRepository.findByTokenValueForUpdate(value)  ← SELECT FOR UPDATE
│  │   → presented token (hoặc throw)
│  │
│  ├→ tokenDomainService.rotate(presented, 7, repo)
│  │   │
│  │   ├→ if presented.isReuseAttempt():   ← token đã revoked NHƯNG vẫn được dùng
│  │   │   ├→ revokeFamily(presented.family, repo)  ← revoke CẢ FAMILY
│  │   │   └→ throw TokenReuseDetectedException
│  │   │       → response 401, user phải login lại
│  │   │
│  │   ├→ presented.verifyValidity()  ← check expired
│  │   ├→ presented.revoke(NORMAL) → registerEvent(TokenRevokedEvent)
│  │   ├→ repo.save(presented)
│  │   └→ return RefreshToken.rotate(presented, 7)
│  │        → gen+1, cùng family, value mới
│  │        → registerEvent(TokenRotatedEvent)
│  │
│  ├→ refreshTokenRepository.save(rotated)
│  ├→ user lookup + roles
│  ├→ tokenProvider.issueAccessToken(user, roles)
│  └→ return AuthResponseDTO
│
└→ Set HttpOnly cookies với token mới
   Sau commit: Kafka → "auth.token.revoked.v1" + "auth.token.rotated.v1"

KEY INSIGHT: pessimistic lock đảm bảo nếu 2 request refresh đồng thời:
- Request 1: lock, revoke old, issue new → commit
- Request 2: chờ lock → thấy old đã revoked → trigger reuse detection → revoke family""")

add_heading(doc, '9.4 Luồng Logout', level=2)
add_code(doc, """┌─ HTTP POST /api/v1/auth/logout (auth required)
├→ Đọc accessToken từ cookie/header
├→ Đọc refreshToken từ cookie
│
├→ LogoutHandler.handle(LogoutCommand)
│  ├→ Parse access token → lấy jti, exp
│  ├→ tokenBlacklist.revoke(jti, userId, exp, "LOGOUT")
│  │   ├→ Redis SET blacklist:{jti} = "LOGOUT" EX <ttl>
│  │   └→ DB INSERT token_revocations
│  ├→ Tìm refresh token → revoke(NORMAL) → save
│  │   └→ registerEvent(TokenRevokedEvent)
│
├→ cookieWriter.expire() → Set-Cookie cleared
└→ 204 No Content""")

add_heading(doc, '9.5 Luồng Change Password', level=2)
add_code(doc, """┌─ HTTP POST /api/v1/auth/change-password (auth required)
│  Body: {oldPassword, newPassword}
│
├→ AuthController: parse current access token để lấy jti+exp
│
├→ ChangePasswordHandler.handle(cmd)
│  ├→ userRepository.findById(userId)
│  ├→ user.changePassword(old, new, encoder)
│  │   ├→ encoder.matches(old, password.hash) — verify old
│  │   ├→ password = encoder.encode(new)
│  │   ├→ incrementTokenVersion()   ← KEY: bump version
│  │   └→ registerEvent(PasswordChangedEvent)
│  ├→ userRepository.save(user) → events published
│  ├→ tokenBlacklist.revoke(currentJti, ...)  ← blacklist token hiện tại ngay
│  └→ refreshTokenRepository.revokeAllByUser(userId)  ← thu hồi mọi refresh token
│
└→ HỆ QUẢ: tất cả JWT cũ (mọi device user đang dùng) đều invalid:
   - tokenVersion mismatch → JwtAuthenticationFilter reject
   - Refresh token đã revoke → không refresh được nữa
   → User phải login lại trên mọi device""")

doc.add_page_break()

# ============================================================
# 10. GIẢI THÍCH KHÁI NIỆM JAVA/SPRING
# ============================================================
add_heading(doc, '10. Giải thích các khái niệm Java/Spring trong code', level=1)

add_heading(doc, '10.1 @Component vs @Service vs @Repository vs @Controller', level=2)
add_para(doc,
    'Về kỹ thuật, 4 annotation này GIỐNG NHAU — đều là @Component, Spring scan và tạo bean. Khác '
    'nhau ở Ý NGHĨA NGỮ NGHĨA giúp người đọc code hiểu vai trò class.')

add_table(doc,
    headers=['Annotation', 'Dùng cho', 'Ví dụ'],
    rows=[
        ['@Component', 'Class generic, adapter, helper', 'JwtTokenProvider, SpringDomainEventPublisher'],
        ['@Service', 'Class chứa business logic (application/domain service)', 'LoginHandler, TokenBlacklistService'],
        ['@Repository', 'Class truy cập DB. Bonus: tự dịch JDBC exception → Spring DataAccessException', 'UserRepositoryImpl'],
        ['@RestController', 'Class nhận HTTP request', 'AuthController, UserController'],
        ['@Configuration', 'Class chứa @Bean methods', 'SecurityConfig, KafkaConfig'],
    ],
    col_widths=[3, 6, 7])

add_heading(doc, '10.2 static, final, và các combo', level=2)

add_para(doc, 'Hiểu từng modifier:', bold=True)
for item in [
    'static = thuộc về CLASS, không thuộc instance. Cả class share 1 ô nhớ.',
    'final = không gán lại được sau khi khởi tạo.',
    'private = chỉ class này thấy được.',
]:
    add_bullet(doc, item)

add_para(doc, 'Quyết định nhanh — Hỏi 2 câu:', bold=True)
add_table(doc,
    headers=['Dùng chung cả class?', 'Đổi giá trị được?', 'Combo', 'Ví dụ'],
    rows=[
        ['Có', 'Không', 'static final', 'BLACKLIST_PREFIX = "blacklist:"'],
        ['Có', 'Có', 'static (no final)', 'Hiếm — counter toàn cục'],
        ['Không', 'Không', 'final (no static)', 'private final UserRepository repo (DI)'],
        ['Không', 'Có', '(không gì cả)', 'private int retryCount'],
    ],
    col_widths=[3.5, 3, 3, 6.5])

add_para(doc, 'Trong context Spring service:', bold=True)
add_code(doc, """@Service
@RequiredArgsConstructor
public class TokenBlacklistService {

    // Hằng số: dùng chung, không đổi → private static final
    private static final String BLACKLIST_PREFIX = "blacklist:";

    // Dependency inject: mỗi instance riêng, không đổi sau khi inject → private final
    private final StringRedisTemplate redisTemplate;
    private final TokenRevocationJpaRepository repository;

    // State thay đổi (counter, flag) → private (không final)
    private int failureCount;
}""")

add_para(doc, 'Vì sao "private final" tốt cho DI?', bold=True)
for item in [
    'Thread-safe miễn phí — Java Memory Model đảm bảo final field được safely published.',
    'Fail-fast — Spring không inject được → app crash lúc start, không phải runtime.',
    'Tương thích Lombok @RequiredArgsConstructor — gen constructor cho mọi field final.',
    'Dễ test — new MyService(mockA, mockB) trong unit test.',
]:
    add_bullet(doc, item)

add_heading(doc, '10.3 Lombok @Slf4j vs Logger thủ công', level=2)
add_para(doc,
    'Hai cách hoàn toàn TƯƠNG ĐƯƠNG. @Slf4j gen ra dòng:')
add_code(doc, """private static final Logger log = LoggerFactory.getLogger(MyClass.class);""")
add_para(doc,
    'Trong codebase đang trộn 2 style. Khuyến nghị: dùng @Slf4j cho tất cả để gọn (3 dòng → 1 dòng annotation).')

add_heading(doc, '10.4 record vs class', level=2)
add_para(doc,
    'record là cú pháp Java 14+ cho class chứa data bất biến. 1 dòng record tương đương 30 dòng '
    'class viết tay (constructor, getter, equals, hashCode, toString).')
add_code(doc, """// 6 dòng - tự động sinh constructor, getter, equals, hashCode, toString
public record UserLifecycleIntegrationEvent(
    UUID eventId, String eventType, Instant occurredOn,
    String aggregateId, int schemaVersion, Map<String, Object> payload
) implements IntegrationEvent { }""")

add_para(doc, 'Dùng record khi nào?', bold=True)
for item in [
    'DTO (data transfer object) — passing data giữa các tầng.',
    'Value Object đơn giản (Email, Username).',
    'Command/Query — input cho handler.',
    'Event — integration event, domain event đơn giản.',
    'KHÔNG dùng khi: object cần mutate, cần kế thừa class khác (record là implicit final).',
]:
    add_bullet(doc, item)

add_heading(doc, '10.5 Utility Class Pattern (KafkaTopics)', level=2)
add_code(doc, """public final class KafkaTopics {       // 1. final → không kế thừa
    private KafkaTopics() {}           // 2. private constructor → không new
    public static final String AUTH_USER_EVENTS = "auth.user.events.v1";  // 3. chỉ static
}""")
add_para(doc,
    'Bộ 3 đặc điểm này định nghĩa "utility class". JDK dùng pattern này cho Math, Collections, Arrays...')

add_heading(doc, '10.6 opsForValue / opsForHash / opsForSet (Redis)', level=2)
add_para(doc, 'Mỗi loại data type của Redis có 1 "ops" tương ứng trong Spring Data Redis:', bold=True)
add_table(doc,
    headers=['Method', 'Redis type', 'Lệnh tương đương'],
    rows=[
        ['opsForValue()', 'String', 'SET, GET, INCR, SETEX'],
        ['opsForHash()', 'Hash', 'HSET, HGET, HGETALL'],
        ['opsForList()', 'List', 'LPUSH, RPOP, LRANGE'],
        ['opsForSet()', 'Set', 'SADD, SMEMBERS, SISMEMBER'],
        ['opsForZSet()', 'Sorted Set', 'ZADD, ZRANGE, ZRANGEBYSCORE'],
    ],
    col_widths=[4, 3.5, 8.5])

add_heading(doc, '10.7 @Transactional và @TransactionalEventListener', level=2)
add_para(doc, '@Transactional bao bọc method trong transaction DB:', bold=True)
for item in [
    'Method gọi method @Transactional khác → cùng transaction (PROPAGATION_REQUIRED).',
    'Throw RuntimeException → rollback.',
    '@Transactional(readOnly = true) → hint cho Hibernate skip dirty checking → nhanh hơn.',
]:
    add_bullet(doc, item)

add_para(doc, '@TransactionalEventListener thì khác:', bold=True)
add_code(doc, """@TransactionalEventListener(phase = TransactionPhase.AFTER_COMMIT)
public void onDomainEvent(DomainEvent event) { ... }""")
add_para(doc,
    'Listener chỉ chạy SAU KHI transaction commit thành công. Nếu rollback → event không bao giờ fire. '
    'Đây là cơ chế đảm bảo "side effect chỉ xảy ra khi DB chắc chắn lưu rồi".')

doc.add_page_break()

# ============================================================
# 11. QUYẾT ĐỊNH THIẾT KẾ
# ============================================================
add_heading(doc, '11. Quyết định thiết kế — Tại sao chọn cách này?', level=1)

decisions = [
    ('Vì sao tách Domain và JpaEntity?',
     'Domain là POJO sạch chứa quy tắc nghiệp vụ; JpaEntity bị "trói" với Hibernate. Tách ra để '
     'domain test được không cần Spring và đổi DB không phá nghiệp vụ.'),

    ('Vì sao có Value Object (Email, Username, Password)?',
     'Ép buộc invariant tại điểm tạo, type-safe ở compile time, tránh "stringly typed" code dễ bug. '
     'Password có createRaw/createEncoded ngăn lưu plaintext vào DB.'),

    ('Vì sao tách DomainEvent và IntegrationEvent?',
     'DomainEvent in-process (cùng JVM, có thể chứa POJO/class name) — free refactor. '
     'IntegrationEvent là hợp đồng public với service khác qua Kafka — schema versioned, không phá.'),

    ('Vì sao publish event AFTER_COMMIT?',
     'Tránh bug kinh điển: DB rollback nhưng event đã gửi → service khác xử lý event "ma". '
     'AFTER_COMMIT đảm bảo event ra Kafka chỉ khi DB chắc chắn lưu rồi.'),

    ('Vì sao có tokenVersion trong JWT?',
     'JWT stateless không thu hồi được. tokenVersion là cách "tắt" mọi JWT cũ tức khắc khi đổi mật '
     'khẩu hoặc admin revoke — chỉ cần bump số trong DB, mọi JWT cũ tự fail.'),

    ('Vì sao Refresh Token Rotation + Family Revocation?',
     'Phát hiện và chống token theft. Token cũ bị dùng lại → revoke cả family → kẻ tấn công lẫn '
     'user thật đều bị đẩy ra → user phải login lại trên kênh sạch.'),

    ('Vì sao Pessimistic Lock khi refresh?',
     'Tránh race condition khi 2 request refresh đồng thời cùng 1 token → cả 2 cùng pass → leak '
     'security. SELECT FOR UPDATE serialize 2 request → request thứ 2 thấy token đã revoke → trigger '
     'reuse detection.'),

    ('Vì sao Cookie HttpOnly thay vì localStorage?',
     'localStorage đọc được từ JS → XSS đánh cắp token. HttpOnly cookie JS không đọc được. '
     'Vẫn fallback Bearer header cho mobile/CLI/gRPC.'),

    ('Vì sao BCrypt cost 12 thay vì default 10?',
     'OWASP khuyến nghị tối thiểu 12 cho server hiện đại. Cost 12 ~250ms/hash → user không thấy '
     'chậm, attacker brute force trên GPU không khả thi.'),

    ('Vì sao Hybrid Redis + DB cho Token Blacklist?',
     'Redis: ~1ms check, mọi request đều phải check → cần nhanh. DB: source of truth — Redis chết '
     'không mất revocation. WarmUp on startup → Redis instance mới không bỏ sót.'),

    ('Vì sao CQRS (Command + Handler) thay vì 1 Service class to?',
     'Mỗi handler 1 use case → đơn nhiệm, dễ test, controller mỏng. Tránh "GodService" hàng nghìn '
     'dòng. Dễ thêm cross-cutting (logging, metrics) qua decorator.'),

    ('Vì sao Map<String, Object> trong IntegrationEvent payload?',
     'Flexibility — thêm event type mới không cần tạo record mới. Trade-off: mất type safety. '
     'Hợp lý vì integration event đổi nhanh và consumer thường viết bằng ngôn ngữ khác.'),

    ('Vì sao Issuer Validation trong JWT?',
     'Nhiều service share cùng HS256 secret → có thể giả mạo token nếu không check issuer. '
     'requireIssuer("auth-service") chặn token mint bởi service khác cùng key.'),

    ('Vì sao Rate Limit chạy TRƯỚC JWT filter?',
     'Brute force attack không nên tốn CPU parse JWT. Throttle ngay tầng đầu → request bị chặn '
     'không bao giờ đụng tới logic auth → tiết kiệm resource, tăng độ chịu tải.'),
]

for q, a in decisions:
    add_para(doc, q, bold=True, size=12)
    add_para(doc, a)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 12. PHỤ LỤC
# ============================================================
add_heading(doc, '12. Phụ lục — Cheatsheet & FAQ', level=1)

add_heading(doc, '12.1 Cheatsheet: Khi nào dùng modifier nào?', level=2)
add_table(doc,
    headers=['Tình huống', 'Modifier'],
    rows=[
        ['Hằng số dùng chung cả class (prefix, timeout, magic number)', 'private static final'],
        ['Hằng số public chia sẻ ra ngoài', 'public static final hoặc enum'],
        ['Dependency được Spring inject (constructor injection)', 'private final'],
        ['Logger', 'private static final (hoặc dùng @Slf4j)'],
        ['Pre-compiled Pattern, ObjectMapper, etc.', 'private static final'],
        ['Counter / state thay đổi theo thời gian', 'private (no final)'],
        ['Counter cho multi-thread', 'private volatile / AtomicXxx'],
        ['Cache shared, multi-thread', 'private final ConcurrentXxx'],
    ],
    col_widths=[10, 6])

add_heading(doc, '12.2 Cheatsheet: Khi nào dùng Annotation Spring nào?', level=2)
add_table(doc,
    headers=['Class này là gì?', 'Annotation'],
    rows=[
        ['Business logic (use case handler, domain service Spring wrapper)', '@Service'],
        ['DB access (repository implementation)', '@Repository'],
        ['HTTP endpoint', '@RestController'],
        ['Bean configuration (@Bean methods)', '@Configuration'],
        ['Adapter generic (Redis, Kafka, event publisher...)', '@Component'],
        ['Filter HTTP', '@Component (gắn vào OncePerRequestFilter)'],
    ],
    col_widths=[10, 6])

add_heading(doc, '12.3 Cheatsheet: Tầng nào làm gì?', level=2)
add_table(doc,
    headers=['Tầng', 'Code mẫu thuộc tầng này'],
    rows=[
        ['domain/models/', 'User, RefreshToken (Aggregate Root, business methods)'],
        ['domain/services/', 'TokenDomainService (cross-aggregate rules)'],
        ['domain/repositories/', 'UserRepository (interface)'],
        ['domain/events/', 'DomainEvent, DomainEventPublisher (interface)'],
        ['application/command/', 'LoginCommand (record - input data)'],
        ['application/handler/', 'LoginHandler (orchestrate use case)'],
        ['application/dto/', 'UserDTO, AuthResponseDTO (output)'],
        ['application/events/', 'IntegrationEvent (hợp đồng với service khác)'],
        ['infrastructure/persistence/', 'UserRepositoryImpl (JPA), mapper, JpaEntity'],
        ['infrastructure/security/', 'JwtTokenProvider, BCryptAdapter, TokenBlacklistService'],
        ['infrastructure/messaging/', 'KafkaIntegrationEventPublisher'],
        ['infrastructure/events/', 'SpringDomainEventPublisher'],
        ['interfaces/rest/', 'AuthController, JwtAuthenticationFilter'],
        ['interfaces/grpc/', 'AuthGrpcService'],
        ['config/', 'SecurityConfig, KafkaConfig, DataSeeder'],
    ],
    col_widths=[6, 10])

add_heading(doc, '12.4 FAQ', level=2)

faqs = [
    ('Q: Sao có 2 file DomainEvent.java và DomainEventPublisher.java?',
     'DomainEvent là "mẫu giấy thông báo" (abstract class — mọi event kế thừa). '
     'DomainEventPublisher là "người đưa thư" (interface — port mà infra implement). '
     'Hai vai trò khác nhau, đặt cùng folder domain/events/.'),

    ('Q: DomainEventPublisher đặt ở domain nhưng dùng cho infrastructure?',
     'Interface ĐẶT ở domain (vì là yêu cầu của domain), nhưng impl ở infrastructure. '
     'Đây là Dependency Inversion: domain ra yêu cầu, infra nhận việc.'),

    ('Q: Vì sao có cả Spring event publisher VÀ Kafka publisher?',
     'Spring publisher (in-process) → KafkaIntegrationEventPublisher lắng nghe '
     'bằng @TransactionalEventListener(AFTER_COMMIT) → map sang Integration Event → gửi Kafka. '
     'Hai bước riêng: trong nhà trước, ra ngoài sau.'),

    ('Q: tokenVersion để làm gì?',
     'Chèn vào JWT khi mint. Đổi mật khẩu → bump trong DB. JwtAuthenticationFilter so JWT.tokenVersion '
     'với user.currentTokenVersion → khác = reject. Là cách tắt mọi JWT cũ tức thì.'),

    ('Q: Vì sao Refresh Token có "family" và "generation"?',
     'Family = ID cố định cho cả 1 phiên login (token gốc + mọi token rotation từ nó). '
     'Generation = thứ tự rotation. Token bị cướp & dùng lại → server thấy 1 token đã revoke '
     '(generation cũ) → revoke cả family → user thật + attacker đều bị đẩy ra.'),

    ('Q: Vì sao logout cần blacklist mà không chỉ xóa cookie?',
     'JWT stateless: server không lưu state. Xóa cookie chỉ làm cho browser quên token, NHƯNG nếu '
     'kẻ xấu đã copy token (XSS) thì vẫn dùng được. Blacklist Redis + DB chặn token đó ở phía server.'),

    ('Q: Vì sao có cả gRPC và REST?',
     'REST: cho client browser, mobile (public). gRPC: cho service-to-service nội bộ (nhanh hơn, '
     'type-safe qua proto). order-service kiểm permission user → gRPC, không qua HTTP.'),

    ('Q: Vì sao seal class với "public final"?',
     'Ngăn ai đó extends User và phá invariants bằng cách override. Aggregate Root nên là final để '
     'đảm bảo quy tắc trong class là duy nhất.'),

    ('Q: AggregateRoot vs Entity khác gì?',
     'Aggregate Root là entity "chính" của 1 nhóm, kiểm soát mọi thao tác lên cả nhóm. Ví dụ User '
     'aggregate gồm User + roles (collection của RoleId). Mọi thay đổi đi qua User. Entity con không '
     'được truy cập trực tiếp từ ngoài aggregate.'),

    ('Q: @Transactional có propagation gì?',
     'Mặc định REQUIRED — join transaction hiện có, hoặc tạo mới. Đa số trường hợp dùng default. '
     'Khi cần "luôn tạo transaction mới" (ví dụ retry độc lập) dùng REQUIRES_NEW.'),
]

for q, a in faqs:
    add_para(doc, q, bold=True)
    add_para(doc, 'A: ' + a)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# KẾT THÚC
# ============================================================
add_heading(doc, 'Lời kết', level=1)

add_para(doc,
    'auth-service là codebase được thiết kế tương đối chỉn chu theo các pattern enterprise Java hiện đại:')
for item in [
    'Hexagonal Architecture — domain tách rời framework.',
    'CQRS với Command + Handler — đơn nhiệm, test dễ.',
    'Domain Event + Integration Event — decouple in-process và cross-service.',
    'Refresh Token Rotation + Family Revocation — chống token theft.',
    'tokenVersion + Redis Blacklist — JWT vẫn revoke được tức thì.',
    'Pessimistic Lock + AFTER_COMMIT Listener — race condition + dual-write safe.',
]:
    add_bullet(doc, item)

add_para(doc,
    '\nNếu nắm vững các pattern trong service này, bạn đã có nền tảng vững để đọc/viết bất cứ Spring '
    'microservice enterprise nào. Mỗi quyết định thiết kế đều có lý do cụ thể — không phải "trang trí" '
    'mà giải quyết bug hoặc trade-off cụ thể.')

add_para(doc, '\n— Hết —', italic=True)


# ── Save ────────────────────────────────────────────────────────────────
output_path = '/Users/admin/HieuTo/ecommerce/auth-service-documentation.docx'
doc.save(output_path)
print(f"✅ Saved: {output_path}")
