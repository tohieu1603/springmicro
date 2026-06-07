"""
Mega-detailed docx generator for auth-service.
Documents EVERY file, EVERY class, EVERY method - what it does, where it's located, why.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Calibri'
    return h


def P(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def CODE(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)


def BULLET(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def TABLE(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        set_cell_bg(cell, '2E5C8A')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def NOTE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run('GHI CHÚ: ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8DC')
    pPr.append(shd)


def FILE_HEADER(doc, filename, path):
    """Tiêu đề file kèm đường dẫn."""
    H(doc, filename, level=3)
    p = doc.add_paragraph()
    run = p.add_run('Vị trí: ' + path)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True


def METHOD(doc, signature, purpose, why=None, location=None):
    """Mô tả 1 method/field."""
    p = doc.add_paragraph()
    run = p.add_run('• ' + signature)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6F)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('  → ' + purpose)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)

    if why:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run('  Lý do: ' + why)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── Document ────────────────────────────────────────────────────────────

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


# ============================================================
# TRANG BÌA
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\nAUTH-SERVICE\n')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('TÀI LIỆU KỸ THUẬT CHI TIẾT\n')
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Mọi class, mọi method - làm gì, ở đâu, vì sao\n\n')
run.font.size = Pt(14)
run.italic = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Hexagonal Architecture · CQRS · Event-Driven · Spring Boot\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()


# ============================================================
# MỤC LỤC
# ============================================================
H(doc, 'Mục lục', level=1)

toc_items = [
    'PHẦN I - TỔNG QUAN',
    '  1. Giới thiệu auth-service',
    '  2. Kiến trúc Hexagonal - 4 tầng',
    '  3. Cây thư mục đầy đủ',
    '',
    'PHẦN II - TẦNG DOMAIN (Trái tim nghiệp vụ)',
    '  4. domain/shared - Base class chung',
    '  5. domain/events - Hệ thống event in-process',
    '  6. domain/models/user - User aggregate + VOs + Events + Exceptions',
    '  7. domain/models/refreshtoken - RefreshToken aggregate',
    '  8. domain/models/role - Role aggregate',
    '  9. domain/models/permission - Permission entity',
    '  10. domain/services - Domain services + Ports',
    '  11. domain/repositories - Repository interfaces',
    '',
    'PHẦN III - TẦNG APPLICATION (CQRS)',
    '  12. application/common - Interface chung',
    '  13. application/command - 9 Command records',
    '  14. application/query - 5 Query records',
    '  15. application/handler - 14 Handler classes',
    '  16. application/dto - Read models',
    '  17. application/mapper - Domain → DTO',
    '  18. application/port - Outbound ports',
    '  19. application/events - Integration events ra Kafka',
    '',
    'PHẦN IV - TẦNG INFRASTRUCTURE (Adapter)',
    '  20. infrastructure/security - JWT, BCrypt, UserDetails',
    '  21. infrastructure/persistence - JPA adapter',
    '  22. infrastructure/cache - Redis cache',
    '  23. infrastructure/events - Spring event publisher',
    '  24. infrastructure/messaging - Kafka bridge',
    '',
    'PHẦN V - TẦNG INTERFACES (Cổng giao tiếp)',
    '  25. interfaces/rest - Controllers',
    '  26. interfaces/rest/filter - Filters',
    '  27. interfaces/rest/support - Cookie writer',
    '  28. interfaces/rest/dto - Request DTOs',
    '  29. interfaces/grpc - gRPC service',
    '',
    'PHẦN VI - CONFIG & EXCEPTION',
    '  30. config/* - Spring configurations',
    '  31. exceptions/* - Global exception handler',
    '',
    'PHẦN VII - LUỒNG NGHIỆP VỤ ĐẦY ĐỦ',
    '  32. Register flow',
    '  33. Login flow',
    '  34. Refresh token flow',
    '  35. Logout flow',
    '  36. Change password flow',
    '  37. Role assignment flow',
    '  38. gRPC verify token flow',
    '',
    'PHẦN VIII - QUYẾT ĐỊNH THIẾT KẾ & FAQ',
    '  39. 25 quyết định thiết kế quan trọng',
    '  40. Cheat sheet & FAQ',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.3)
    for r in p.runs:
        r.font.size = Pt(10.5)

doc.add_page_break()

# ============================================================
# PHẦN I - TỔNG QUAN
# ============================================================
H(doc, 'PHẦN I - TỔNG QUAN', level=1)

H(doc, '1. Giới thiệu auth-service', level=2)
P(doc,
  'auth-service là microservice phụ trách xác thực (authentication) và phân quyền (authorization) '
  'của hệ thống ecommerce. Service này là "cánh cổng" mà mọi microservice khác tin tưởng để '
  'kiểm chứng danh tính và quyền của người dùng. 154 file Java, ~10000 dòng code, áp dụng '
  'Hexagonal Architecture + CQRS + Event-Driven design.')

P(doc, 'Trách nhiệm:', bold=True)
for x in [
    'Đăng ký, đăng nhập, đăng xuất, đổi mật khẩu',
    'Phát hành JWT access token + Refresh token',
    'Xoay vòng refresh token (Rotation + Family Revocation)',
    'Blacklist token bằng Redis + DB hybrid',
    'Quản lý user, role, permission',
    'Kiểm tra phân quyền cho service khác qua gRPC',
    'Phát integration event qua Kafka khi user/role thay đổi',
]:
    BULLET(doc, x)

H(doc, '2. Kiến trúc Hexagonal - 4 tầng', level=2)
P(doc,
  'Code chia thành 4 tầng, phụ thuộc CHỈ đi 1 chiều từ ngoài vào trong: '
  'interfaces → application → domain. infrastructure cung cấp adapter implement port của domain.')

TABLE(doc,
      headers=['Tầng', 'Vai trò', 'Phụ thuộc framework?'],
      rows=[
          ['domain/', 'Quy tắc nghiệp vụ thuần (POJO Java)', 'KHÔNG'],
          ['application/', 'Điều phối use case (CQRS)', 'Spring (annotation only)'],
          ['infrastructure/', 'Implement port: JPA, Redis, Kafka, JWT', 'Toàn bộ stack'],
          ['interfaces/', 'HTTP REST, gRPC, filter', 'Spring Web, gRPC'],
      ],
      col_widths=[3.5, 8, 5])

H(doc, '3. Cây thư mục đầy đủ', level=2)
CODE(doc, '''auth-service/src/main/java/com/hieu/auth_service/
├── AuthServiceApplication.java                  ← Main class
│
├── domain/                                       ← TẦNG NGHIỆP VỤ
│   ├── shared/
│   │   ├── AggregateRoot.java                   ← Base cho mọi aggregate
│   │   └── DomainException.java                 ← Exception gốc
│   ├── events/
│   │   ├── DomainEvent.java                     ← Base event
│   │   └── DomainEventPublisher.java            ← Port publish event
│   ├── repositories/                            ← 4 interface repository
│   │   ├── UserRepository.java
│   │   ├── RoleRepository.java
│   │   ├── PermissionRepository.java
│   │   └── RefreshTokenRepository.java
│   ├── services/
│   │   ├── PasswordEncoderPort.java             ← Port hash password
│   │   ├── TokenProviderPort.java               ← Port JWT
│   │   ├── AuthenticationDomainService.java    ← Logic authorize
│   │   └── TokenDomainService.java              ← Logic rotation
│   └── models/
│       ├── user/                                ← User aggregate
│       │   ├── User.java
│       │   ├── vo/    (6 file: UserId, Username, Email, Password, AccountStatus, PersonName)
│       │   ├── events/ (7 event: UserCreated, PasswordChanged...)
│       │   └── exceptions/ (4 exception)
│       ├── refreshtoken/                        ← RefreshToken aggregate
│       │   ├── RefreshToken.java
│       │   ├── vo/    (6 VO: TokenId, TokenValue, TokenFamily...)
│       │   ├── events/ (3 event)
│       │   └── exceptions/ (4 exception)
│       ├── role/                                ← Role aggregate
│       │   ├── Role.java
│       │   ├── vo/    (RoleId, RoleName)
│       │   └── events/ (PermissionGranted, PermissionRevoked)
│       └── permission/                          ← Permission entity
│           ├── Permission.java
│           └── vo/    (PermissionId, PermissionName)
│
├── application/                                  ← TẦNG ĐIỀU PHỐI
│   ├── common/   (Command, Query, CommandHandler, QueryHandler, CursorCodec...)
│   ├── command/  (9 record command)
│   ├── query/    (5 record query)
│   ├── handler/  (14 class handler)
│   ├── dto/      (UserDTO, AuthResponseDTO, PageDTO, RoleDTO, PermissionDTO)
│   ├── mapper/   (UserDtoMapper)
│   ├── port/     (TokenBlacklistPort, RolePermissionCachePort)
│   └── events/   (IntegrationEvent, UserLifecycle/SessionIntegrationEvent, KafkaTopics)
│
├── infrastructure/                               ← TẦNG ADAPTER
│   ├── security/   (JwtTokenProvider, BCryptAdapter, AuthUserDetails, CustomUserDetailsService, TokenBlacklistService, JwtProperties)
│   ├── persistence/
│   │   ├── jpa/entities/      (UserJpaEntity, RoleJpaEntity, RefreshTokenJpaEntity...)
│   │   ├── jpa/repositories/  (UserJpaRepository extends JpaRepository...)
│   │   ├── mapper/            (UserJpaMapper, RefreshTokenJpaMapper...)
│   │   └── impl/              (UserRepositoryImpl, RefreshTokenRepositoryImpl...)
│   ├── cache/      (RedisRolePermissionCacheAdapter)
│   ├── events/     (SpringDomainEventPublisher)
│   └── messaging/  (KafkaIntegrationEventPublisher, DomainEventToIntegrationEventMapper, RolePermissionCacheInvalidator)
│
├── interfaces/                                   ← TẦNG GIAO TIẾP
│   ├── rest/
│   │   ├── AuthController.java       ← /api/v1/auth/*
│   │   ├── UserController.java       ← /api/v1/users/*
│   │   ├── dto/      (LoginRequest, RegisterRequest, ChangePasswordRequest, AuthMeResponse...)
│   │   ├── filter/   (JwtAuthenticationFilter, RateLimitFilter, JwtAuthenticationEntryPoint)
│   │   ├── support/  (AuthCookieWriter)
│   │   └── advice/
│   └── grpc/
│       ├── AuthGrpcService.java
│       └── client/   (GrpcClientConfig, AuthServiceGrpcClient)
│
├── config/                                       ← Spring configuration
│   ├── SecurityConfig.java          ← Spring Security
│   ├── KafkaConfig.java             ← Kafka topics
│   ├── DomainServiceConfig.java     ← @Bean cho domain services
│   ├── AsyncConfig.java             ← Virtual thread executor
│   ├── CorsConfig.java              ← CORS
│   ├── OpenApiConfig.java           ← Swagger
│   └── DataSeeder.java              ← Seed dữ liệu khởi tạo
│
└── exceptions/
    ├── GlobalExceptionHandler.java   ← @RestControllerAdvice
    └── ErrorResponse.java''')

doc.add_page_break()

# ============================================================
# PHẦN II - TẦNG DOMAIN
# ============================================================
H(doc, 'PHẦN II - TẦNG DOMAIN', level=1)

P(doc,
  'Domain là tầng trong cùng, chứa quy tắc nghiệp vụ THUẦN. Quy tắc: KHÔNG import Spring, '
  'KHÔNG import Hibernate, KHÔNG import HTTP. Mục tiêu: domain test được bằng JUnit thuần, '
  'không cần startup Spring; đổi DB hoặc framework không phá nghiệp vụ.')

# ── 4. domain/shared ──
H(doc, '4. domain/shared', level=2)

FILE_HEADER(doc, 'AggregateRoot.java', 'domain/shared/AggregateRoot.java')
P(doc, 'Base class trừu tượng cho mọi aggregate root. Quản lý buffer DomainEvent.', bold=True)

P(doc, 'Field:', bold=True)
METHOD(doc, 'private final transient List<DomainEvent> domainEvents',
       'Danh sách event đang chờ publish.',
       'transient để Hibernate không persist list này vào DB - chỉ tồn tại trong RAM.')

P(doc, 'Method:', bold=True)
METHOD(doc, 'protected final void registerEvent(DomainEvent event)',
       'Aggregate gọi method này khi state thay đổi quan trọng để ghi event vào buffer.',
       'protected để chỉ subclass (User, RefreshToken, Role) gọi được. final để subclass không override.')
METHOD(doc, 'public final List<DomainEvent> pullDomainEvents()',
       'Rút (atomic) toàn bộ event ra khỏi buffer và clear buffer. Trả List.copyOf - bất biến.',
       'Pattern "drain once": gọi 1 lần duy nhất bên trong repository.save() sau khi persist thành công. Tránh bug "quên clearEvents" gây duplicate event.')
METHOD(doc, 'public final List<DomainEvent> peekDomainEvents()',
       'Xem event không destruktif - dùng cho test/log.',
       'Trả Collections.unmodifiableList để consumer không thay đổi được.')
METHOD(doc, 'protected final void clearDomainEvents()',
       'Xóa buffer mà không publish - dùng cho rollback scenarios.',
       'Hiếm khi dùng - chỉ khi xử lý compensation manually.')

FILE_HEADER(doc, 'DomainException.java', 'domain/shared/DomainException.java')
P(doc,
  'Exception gốc cho mọi domain exception. Có method code() trả về error code stable '
  '(ví dụ "AUTH-1001") để client/test branch theo code thay vì message.')

# ── 5. domain/events ──
H(doc, '5. domain/events', level=2)

FILE_HEADER(doc, 'DomainEvent.java', 'domain/events/DomainEvent.java')
P(doc, 'Abstract class - mọi event in-process kế thừa.', bold=True)
P(doc, 'Field:', bold=True)
METHOD(doc, 'private final UUID eventId = UUID.randomUUID()',
       'ID duy nhất của event, sinh tự động khi raise.',
       'final + khởi tạo inline đảm bảo eventId bất biến và không null.')
METHOD(doc, 'private final Instant occurredOn = Instant.now()',
       'Thời điểm event xảy ra.',
       'Snapshot lúc raise, không phải lúc consume.')
P(doc, 'Method:', bold=True)
METHOD(doc, 'public final UUID eventId()',
       'Getter eventId. final - subclass không override.', None)
METHOD(doc, 'public final Instant occurredOn()',
       'Getter occurredOn.', None)
METHOD(doc, 'public String eventType()',
       'Trả về tên class đơn giản (ví dụ "PasswordChangedEvent").',
       'Mặc định lấy getClass().getSimpleName() - rẻ, ổn định cross-redeploy của cùng class. Subclass có thể override nếu cần tên khác.')
METHOD(doc, 'public abstract String aggregateId()',
       'Subclass bắt buộc trả ID của aggregate raise event này (userId, tokenId...).',
       'Dùng cho routing/partitioning xuống Kafka - cùng aggregate vào cùng partition để giữ thứ tự.')
METHOD(doc, '@Override public String toString()',
       'Format chuẩn cho log: ClassName{eventId=..., aggregateId=..., occurredOn=...}.', None)

FILE_HEADER(doc, 'DomainEventPublisher.java', 'domain/events/DomainEventPublisher.java')
P(doc, 'Interface (PORT) - khai báo cách publish event. Đặt ở domain để domain không phụ thuộc infrastructure.', bold=True)
P(doc, 'Method:', bold=True)
METHOD(doc, 'void publish(DomainEvent event)',
       'Publish 1 event.',
       'Domain chỉ biết "có cách nào đó publish event". Infrastructure (SpringDomainEventPublisher) lo phần thực thi.')
METHOD(doc, 'default void publishAll(Iterable<? extends DomainEvent> events)',
       'Default method publish nhiều event bằng cách lặp.',
       'Default để implementation có thể override với batch publishing nếu cần.')

# ── 6. domain/models/user ──
H(doc, '6. domain/models/user', level=2)

H(doc, '6.1 User.java - Aggregate Root', level=3)
P(doc, 'Vị trí: domain/models/user/User.java')
P(doc, 'Aggregate root quản lý: identity, credentials, account status, roles, tokenVersion.', bold=True)

P(doc, 'Field:', bold=True)
TABLE(doc,
      headers=['Field', 'Kiểu', 'Mục đích'],
      rows=[
          ['id', 'UserId (UUID)', 'Định danh duy nhất, immutable'],
          ['username', 'Username', 'Tên đăng nhập, unique'],
          ['email', 'Email', 'Email, unique'],
          ['password', 'Password', 'Mật khẩu đã hash'],
          ['personName', 'PersonName', 'Họ + tên'],
          ['accountStatus', 'AccountStatus', '4 flag trạng thái + lastLogin'],
          ['roles', 'Set<RoleId>', 'Các role được gán'],
          ['tokenVersion', 'int', 'KEY: invalidate mọi JWT cũ khi bump'],
          ['createdAt, updatedAt', 'Instant', 'Audit timestamp'],
      ],
      col_widths=[3.5, 3.5, 9])

P(doc, 'Factory method:', bold=True)
METHOD(doc, 'public static User register(Username, Email, Password rawPwd, PersonName, PasswordEncoderPort encoder)',
       'Tạo user mới. Hash mật khẩu ngay tại đây, raise UserCreatedEvent.',
       'Plaintext password KHÔNG bao giờ chạm tới persistence layer - hash ngay khi tạo. Encoder inject qua tham số để domain không phụ thuộc Spring.')
METHOD(doc, 'public static User reconstitute(UserId, Username, ..., int tokenVersion, Instant createdAt, Instant updatedAt)',
       'Khôi phục User từ DB. KHÔNG raise event - đây là rehydration, không phải tạo mới.',
       'Tách biệt giữa "tạo mới" (register) và "load từ DB" (reconstitute) là pattern chuẩn DDD.')

P(doc, 'Method authentication:', bold=True)
METHOD(doc, 'public boolean authenticate(Password rawPassword, PasswordEncoderPort encoder)',
       'Verify mật khẩu. Gọi ensureAuthenticatable() check status trước, sau đó so hash. Nếu OK gọi recordLogin().',
       'Logic xác thực nằm trong aggregate (không phải service ngoài) - tuân thủ "Tell, Don\'t Ask".')
METHOD(doc, 'public void ensureAuthenticatable()',
       'Throw AccountNotUsableException nếu account disabled/locked/expired/credentials expired.',
       'Public để TokenDomainService gọi standalone khi rotate token (không cần re-verify password mà vẫn check status).')
METHOD(doc, 'private void recordLogin()',
       'Cập nhật lastLogin trong AccountStatus, raise UserLoggedInEvent.',
       'private vì chỉ authenticate() được gọi - đảm bảo không bypass verify password.')

P(doc, 'Method credentials:', bold=True)
METHOD(doc, 'public void changePassword(Password old, Password newPwd, PasswordEncoderPort encoder)',
       'Verify mật khẩu cũ, hash mật khẩu mới, BUMP tokenVersion, raise PasswordChangedEvent.',
       'Bump tokenVersion là KEY: mọi JWT cũ tự động invalid ngay khi handler save xong (filter sẽ reject).')
METHOD(doc, 'public void updateEmail(Email newEmail)',
       'Đổi email. Idempotent: không-op nếu trùng email cũ. Raise EmailChangedEvent.',
       'Idempotent giúp retry an toàn.')
METHOD(doc, 'public void updatePersonName(PersonName newName)',
       'Đổi họ tên. Không raise event vì không phải security-critical.',
       'Quy tắc: chỉ raise event cho thay đổi đáng quan tâm bên ngoài service.')

P(doc, 'Method status:', bold=True)
METHOD(doc, 'public void lock() / unlock() / disable() / enable()',
       'Chuyển trạng thái account. Mỗi method raise AccountStatusChangedEvent với Transition enum tương ứng.',
       'Method nhỏ + 1 method private transitionStatus() gộp logic chung - tránh duplicate code.')
METHOD(doc, 'public boolean isActive()',
       'Trả true nếu cả 4 flag đều positive.',
       'Convenience method - tránh client phải kiểm tra 4 flag riêng.')

P(doc, 'Method role management:', bold=True)
METHOD(doc, 'public void assignRole(RoleId roleId)',
       'Thêm role. Set.add() trả false nếu đã có - skip event.',
       'Idempotent. Raise RoleAssignedEvent chỉ khi thực sự thay đổi.')
METHOD(doc, 'public void unassignRole(RoleId roleId)',
       'Bỏ role. Tương tự, idempotent.', None)
METHOD(doc, 'public boolean hasRole(RoleId roleId)',
       'Check user có role không.', None)
METHOD(doc, 'public Set<RoleId> getRoles()',
       'Trả Collections.unmodifiableSet - ngăn caller mutate.',
       'Defensive: nếu trả trực tiếp Set internal, caller có thể add/remove phá invariants.')

P(doc, 'Method token version:', bold=True)
METHOD(doc, 'public void incrementTokenVersion()',
       'Tăng tokenVersion lên 1. Cập nhật updatedAt.',
       'Public để admin có thể force revoke (admin endpoint gọi method này → bump → mọi JWT cũ invalid).')

# ── 6.2 User VOs ──
H(doc, '6.2 Value Objects của User', level=3)

FILE_HEADER(doc, 'UserId.java', 'domain/models/user/vo/UserId.java')
P(doc, 'record UserId(String value) - định danh user bằng UUID dạng String.', bold=True)
METHOD(doc, 'Compact constructor public UserId { ... }',
       'Validate: không null/empty, parse được UUID.fromString.',
       'Chặn corrupted string từ DB/JWT/external lọt vào aggregate. UUID format validation là defense-in-depth.')
METHOD(doc, 'public static UserId of(String value)',
       'Factory method - thay vì new UserId(...) trực tiếp.',
       'Convention DDD: VO dùng factory of() đẹp hơn constructor.')
METHOD(doc, 'public static UserId generate()',
       'Sinh UUID mới ngẫu nhiên.',
       'Dùng trong User.register().')

FILE_HEADER(doc, 'Username.java', 'domain/models/user/vo/Username.java')
P(doc, 'record Username(String value) - tên đăng nhập.', bold=True)
P(doc, 'Constraint trong compact constructor:', bold=True)
for x in ['Không null/empty (sau trim)', 'Độ dài 3-50 ký tự',
          'Chỉ chứa [a-zA-Z0-9_-] (regex)', 'Trim whitespace tự động']:
    BULLET(doc, x)
METHOD(doc, 'public static Username of(String value)', 'Factory method.', None)

FILE_HEADER(doc, 'Email.java', 'domain/models/user/vo/Email.java')
P(doc, 'record Email(String value) - email address với validation.', bold=True)
P(doc, 'Constraint:', bold=True)
for x in ['Regex match RFC-style email format', 'Tự động trim + toLowerCase',
          'Throw IllegalArgumentException nếu sai format']:
    BULLET(doc, x)
METHOD(doc, 'public String domain()',
       'Trả phần sau @ (ví dụ "gmail.com").',
       'Method ngắn hơn getter style - DDD prefer noun, không "getXxx".')
METHOD(doc, 'public String localPart()',
       'Trả phần trước @.', None)

FILE_HEADER(doc, 'Password.java', 'domain/models/user/vo/Password.java')
P(doc, 'record Password(String value, boolean encoded) - phân biệt raw vs hashed.', bold=True)
P(doc, 'Field thứ 2 "encoded" là KEY:', bold=True)
for x in [
    'encoded=false → đây là plaintext, sẽ validate độ dài + complexity',
    'encoded=true → đây là BCrypt hash, skip validation (vì hash không match raw rule)',
    'Phân biệt giúp tránh bug nguy hiểm: lưu plaintext vào DB',
]:
    BULLET(doc, x)
P(doc, 'Method:', bold=True)
METHOD(doc, 'public static Password createRaw(String rawPassword)',
       'Tạo Password instance từ plaintext - sẽ validate complexity (≥8 ký tự, có chữ + số).', None)
METHOD(doc, 'public static Password createEncoded(String hash)',
       'Tạo Password từ hash đã có sẵn (load từ DB, hoặc sau khi encode).',
       'Phân biệt rõ 2 nguồn để typesystem giúp kiểm soát.')
METHOD(doc, 'public boolean needsEncoding()',
       'Trả !encoded - convenience cho User.changePassword().',
       'Tên method readable hơn là gọi !encoded() trực tiếp.')
METHOD(doc, 'private static void validateRawPassword(String)',
       'Static helper validate complexity.',
       'Static vì không phụ thuộc instance state.')
METHOD(doc, '@Override public String toString()',
       'Trả "Password{encoded=..., length=...}" - KHÔNG IN value.',
       'Security: ngăn vô tình log plaintext password.')

FILE_HEADER(doc, 'AccountStatus.java', 'domain/models/user/vo/AccountStatus.java')
P(doc, 'record với 5 field: 4 boolean status + lastLogin.', bold=True)
TABLE(doc,
      headers=['Field', 'Ý nghĩa khi true'],
      rows=[
          ['enabled', 'Account đang hoạt động (false = disabled)'],
          ['accountNonExpired', 'Account chưa hết hạn'],
          ['accountNonLocked', 'Account chưa bị lock'],
          ['credentialsNonExpired', 'Credentials còn hiệu lực'],
          ['lastLogin', 'Thời điểm login cuối (nullable)'],
      ],
      col_widths=[6, 10])
NOTE(doc,
     'Tên field theo convention "non" của Spring Security UserDetails. true là trạng thái "OK". '
     'Khi viết logic nên hỏi "isLocked?" → !accountNonLocked, không phải lock=true.')

P(doc, 'Factory:', bold=True)
METHOD(doc, 'public static AccountStatus createActive()', 'Status mặc định cho user mới: tất cả 4 flag true.', None)
METHOD(doc, 'public static AccountStatus createDisabled()', 'enabled=false, các flag khác true.', None)
METHOD(doc, 'public static AccountStatus createLocked()', 'accountNonLocked=false, các flag khác true.', None)

P(doc, 'Transition (immutable):', bold=True)
METHOD(doc, 'public AccountStatus lock()',
       'Trả AccountStatus mới với accountNonLocked=false. KHÔNG mutate this.',
       'Immutable VO: mỗi transition tạo instance mới. Đảm bảo thread-safety và không có side effect bất ngờ.')
METHOD(doc, 'public AccountStatus unlock() / disable() / enable() / withLastLogin(Instant)',
       'Cùng pattern: trả copy mới với 1 field thay đổi.', None)
METHOD(doc, 'public boolean isActive()',
       'Trả true nếu cả 4 flag positive.', None)

FILE_HEADER(doc, 'PersonName.java', 'domain/models/user/vo/PersonName.java')
P(doc, 'record PersonName(String firstName, String lastName).', bold=True)
METHOD(doc, 'public String fullName()',
       'Trả firstName + " " + lastName.',
       'Cleaner hơn caller phải tự concat.')

# ── 6.3 User Events ──
H(doc, '6.3 Domain Events của User', level=3)
P(doc, 'Tất cả kế thừa DomainEvent, override aggregateId() trả userId.value().')

TABLE(doc,
      headers=['Event', 'Khi nào raise', 'Field đặc thù'],
      rows=[
          ['UserCreatedEvent', 'register()', 'userId, username, email'],
          ['UserLoggedInEvent', 'authenticate() OK', 'userId, username'],
          ['PasswordChangedEvent', 'changePassword()', 'userId, username'],
          ['EmailChangedEvent', 'updateEmail() (đổi)', 'userId, oldEmail, newEmail'],
          ['AccountStatusChangedEvent', 'lock/unlock/enable/disable', 'userId, username, transition (enum LOCKED|UNLOCKED|...)'],
          ['RoleAssignedEvent', 'assignRole() add thật', 'userId, roleId'],
          ['RoleRemovedEvent', 'unassignRole() remove thật', 'userId, roleId'],
      ],
      col_widths=[5, 5.5, 6.5])

# ── 6.4 User Exceptions ──
H(doc, '6.4 Exceptions của User', level=3)
TABLE(doc,
      headers=['Exception', 'HTTP', 'Khi nào throw'],
      rows=[
          ['UserNotFoundException', '404', 'Lookup userId không tồn tại'],
          ['UserAlreadyExistsException', '409', 'register() với username/email đã có'],
          ['InvalidCredentialsException', '401', 'Lookup fail HOẶC password sai (cùng exception để timing-safe)'],
          ['AccountNotUsableException', '403', 'Account disabled/locked/expired - có enum Reason'],
      ],
      col_widths=[6, 1.5, 9])

NOTE(doc,
     'Tất cả đều extends DomainException và có code() trả error code stable. '
     'GlobalExceptionHandler dùng pattern matching mapping → HTTP status.')

# ── 7. RefreshToken aggregate ──
H(doc, '7. domain/models/refreshtoken - RefreshToken aggregate', level=2)

FILE_HEADER(doc, 'RefreshToken.java', 'domain/models/refreshtoken/RefreshToken.java')
P(doc,
  'Implement pattern Refresh-Token Rotation + Family Revocation. Mỗi login tạo root token '
  '(generation 0, family mới). Mỗi refresh: revoke token cũ, phát token mới cùng family, '
  'generation+1. Nếu token đã revoke bị dùng lại → revoke cả family (phát hiện token theft).',
  bold=True)

P(doc, 'Field:', bold=True)
TABLE(doc,
      headers=['Field', 'Kiểu', 'Mục đích'],
      rows=[
          ['id', 'TokenId (UUID)', 'PK của token'],
          ['value', 'TokenValue', 'Chuỗi opaque gửi cho client'],
          ['userId', 'UserId', 'Chủ sở hữu'],
          ['family', 'TokenFamily (UUID)', 'Group cho cả chuỗi rotation'],
          ['generation', 'GenerationNumber (int)', 'Thứ tự rotation: 0 → 1 → 2...'],
          ['expiry', 'TokenExpiry', 'Thời điểm hết hạn'],
          ['revoked', 'boolean', 'Đã thu hồi?'],
          ['reason', 'RevokedReason', 'Lý do revoke (NORMAL, REUSE_DETECTED...)'],
          ['createdAt, revokedAt', 'Instant', 'Audit'],
      ],
      col_widths=[3.5, 4.5, 8])

P(doc, 'Factory method:', bold=True)
METHOD(doc, 'public static RefreshToken create(UserId userId, int expiryDays)',
       'Tạo root token mới (generation 0, family ngẫu nhiên).',
       'Gọi khi login → tạo phiên mới hoàn toàn, family riêng biệt.')
METHOD(doc, 'public static RefreshToken rotate(RefreshToken old, int expiryDays)',
       'Tạo next-generation token: cùng userId, cùng family, generation+1.',
       'Pattern "rotation": token mới kế thừa lineage của token cũ. Family giữ nguyên để link audit trail.')
METHOD(doc, 'public static RefreshToken reconstitute(...)',
       'Khôi phục từ DB. KHÔNG raise event.',
       'Tách biệt creation vs hydration.')

P(doc, 'Method validation:', bold=True)
METHOD(doc, 'public void verifyValidity()',
       'Throw TokenRevokedException nếu revoked, TokenExpiredException nếu expired.',
       'Throw exception khác nhau cho 2 lý do khác nhau để handler/client phân biệt.')
METHOD(doc, 'public boolean isValid()',
       'Trả true nếu !revoked && !expired.', None)
METHOD(doc, 'public boolean isReuseAttempt()',
       'KEY METHOD: trả true nếu (revoked && reason != NORMAL). Báo hiệu token đã revoke nhưng bị dùng lại.',
       'NORMAL = revoke do rotation thường. Nếu thấy reason khác (REUSE_DETECTED, FAMILY_REVOKED) → kẻ tấn công đang replay token.')
METHOD(doc, 'public boolean belongsTo(UserId u)',
       'Check token có thuộc user này không.',
       'Ngăn user A dùng token của user B.')
METHOD(doc, 'public boolean willExpireSoon(long seconds)',
       'Kiểm tra token sắp expired trong N giây.', None)
METHOD(doc, 'public long getRemainingSeconds()',
       'Thời gian còn lại (0 nếu invalid).', None)
METHOD(doc, 'public boolean wasRevokedForSecurity()',
       'reason là security-related (REUSE_DETECTED, FAMILY_REVOKED).', None)

P(doc, 'Method state transition:', bold=True)
METHOD(doc, 'public void revoke(RevokedReason reason)',
       'Set revoked=true, revokedAt=now, raise TokenRevokedEvent. Idempotent: skip nếu đã revoked.',
       'Idempotent giúp tránh duplicate event khi gọi nhiều lần.')
METHOD(doc, 'public void revoke()',
       'Shortcut: revoke(NORMAL).',
       'Convenience cho rotation thường.')

H(doc, '7.1 RefreshToken Value Objects', level=3)
FILE_HEADER(doc, 'TokenId.java', 'vo/TokenId.java')
P(doc, 'record TokenId(String value) - UUID identifier. Validate UUID format.')

FILE_HEADER(doc, 'TokenValue.java', 'vo/TokenValue.java')
P(doc, 'record TokenValue(String value) - chuỗi gửi client.')
METHOD(doc, 'public static TokenValue generate()',
       'Sinh UUID ngẫu nhiên làm token value.',
       'Đơn giản; trong production có thể dùng cryptographically secure random nếu cần entropy cao hơn.')
METHOD(doc, '@Override public String toString()',
       'Trả "TokenValue{<8 ký tự đầu>...}" - KHÔNG in full value.',
       'Security: tránh log lộ token.')

FILE_HEADER(doc, 'TokenFamily.java', 'vo/TokenFamily.java')
P(doc, 'record TokenFamily(String value) - UUID identifier cho family.')
METHOD(doc, 'public boolean isSameFamily(TokenFamily other)',
       'So sánh 2 family.', None)

FILE_HEADER(doc, 'GenerationNumber.java', 'vo/GenerationNumber.java')
P(doc, 'record GenerationNumber(int value) - thứ tự rotation, ≥ 0.')
METHOD(doc, 'public static GenerationNumber root()',
       'Trả GenerationNumber(0) - cho root token.',
       'Đặt tên "root" hơn là "zero" để semantic rõ ràng.')
METHOD(doc, 'public boolean isRoot()',
       'Check generation == 0.', None)
METHOD(doc, 'public GenerationNumber next()',
       'Trả instance mới với value+1.',
       'Immutable: KHÔNG mutate this, tạo instance mới.')
METHOD(doc, 'public boolean isAfter/isBefore/isSameAs',
       'Comparison helpers.', None)
METHOD(doc, 'public int difference(GenerationNumber other)',
       'Signed distance.',
       'Hữu ích trong logic phát hiện attack: nếu generation chênh nhau quá nhiều → có thể là replay.')

FILE_HEADER(doc, 'TokenExpiry.java', 'vo/TokenExpiry.java')
P(doc, 'record TokenExpiry(Instant expiryDate).')
METHOD(doc, 'public static TokenExpiry fromDaysFromNow(int days)',
       'Tạo expiry = Instant.now() + N ngày.',
       'Factory method có ý nghĩa hơn là caller tự tính.')
METHOD(doc, 'public static TokenExpiry fromHoursFromNow(int hours)',
       'Tương tự cho hours.', None)
METHOD(doc, 'public boolean isExpired()',
       'Instant.now().isAfter(expiryDate).', None)
METHOD(doc, 'public long getRemainingSeconds()',
       'Giây còn lại, 0 nếu expired.', None)
METHOD(doc, 'public boolean willExpireWithin(long seconds)',
       'Sẽ hết hạn trong N giây không?',
       'Dùng để trigger pre-emptive refresh client-side.')

FILE_HEADER(doc, 'RevokedReason.java', 'vo/RevokedReason.java')
P(doc, 'record với 5 constant predefined + factory.', bold=True)
P(doc, 'Constants:', bold=True)
TABLE(doc,
      headers=['Reason', 'Khi nào'],
      rows=[
          ['NORMAL', 'Rotation thường (token cũ revoke sau khi rotate)'],
          ['REUSE_DETECTED', 'Token đã revoke bị dùng lại - phát hiện theft'],
          ['FAMILY_REVOKED', 'Cascade revoke do reuse detection'],
          ['USER_INITIATED', 'User chủ động logout'],
          ['EXPIRED', 'Hết hạn tự nhiên'],
      ],
      col_widths=[5, 11])
METHOD(doc, 'public static RevokedReason of(String value)',
       'Factory: switch trả constant nếu match, hoặc tạo mới.',
       'Cho phép custom reason nếu cần, nhưng prefer constants.')
METHOD(doc, 'public boolean isSecurityRelated()',
       'true nếu là REUSE_DETECTED hoặc FAMILY_REVOKED.', None)
METHOD(doc, 'public boolean isNormal()',
       'true nếu là NORMAL.', None)
METHOD(doc, 'public boolean shouldRevokeFamily()',
       'true nếu là FAMILY_REVOKED.',
       'Policy method - logic chuyển vào VO thay vì rải rác.')

H(doc, '7.2 RefreshToken Events', level=3)
TABLE(doc,
      headers=['Event', 'Khi nào', 'Field'],
      rows=[
          ['TokenCreatedEvent', 'create() - root token', 'tokenId, userId, family, generation, expiry'],
          ['TokenRotatedEvent', 'rotate() - next generation', 'userId, family, oldTokenId, newTokenId, newGeneration'],
          ['TokenRevokedEvent', 'revoke()', 'tokenId, userId, family, reason'],
      ],
      col_widths=[5, 5, 6])

H(doc, '7.3 RefreshToken Exceptions', level=3)
TABLE(doc,
      headers=['Exception', 'HTTP', 'Khi nào'],
      rows=[
          ['TokenExpiredException', '401', 'Token quá hạn'],
          ['TokenRevokedException', '401', 'Token đã revoke'],
          ['TokenReuseDetectedException', '401', 'Phát hiện token theft → revoke cả family'],
          ['TokenOwnershipException', '403', 'Token không thuộc user'],
      ],
      col_widths=[6, 1.5, 8.5])

# ── 8. Role aggregate ──
H(doc, '8. domain/models/role - Role aggregate', level=2)

FILE_HEADER(doc, 'Role.java', 'domain/models/role/Role.java')
P(doc, 'Role aggregate root - 1 nhóm có tên chứa các Permission. Ví dụ: ROLE_USER, ROLE_ADMIN.', bold=True)
P(doc, 'Field:', bold=True)
TABLE(doc,
      headers=['Field', 'Kiểu', 'Mục đích'],
      rows=[
          ['id', 'RoleId', 'PK UUID'],
          ['name', 'RoleName', 'Tên unique (ROLE_USER, ROLE_ADMIN)'],
          ['description', 'String', 'Mô tả'],
          ['permissions', 'Set<PermissionId>', 'Các permission được grant'],
          ['createdAt, updatedAt', 'Instant', 'Audit'],
      ],
      col_widths=[3.5, 4, 8.5])

P(doc, 'Factory:', bold=True)
METHOD(doc, 'public static Role create(RoleName name, String description)',
       'Tạo Role mới với UUID ngẫu nhiên.', None)
METHOD(doc, 'public static Role reconstitute(...)',
       'Khôi phục từ DB.', None)

P(doc, 'Method:', bold=True)
METHOD(doc, 'public void updateDescription(String newDescription)',
       'Đổi description. Idempotent - không-op nếu trùng.', None)
METHOD(doc, 'public void grantPermission(PermissionId permission)',
       'Add permission. Raise PermissionGrantedEvent nếu thực sự add.',
       'Idempotent: Set.add() trả false nếu đã có.')
METHOD(doc, 'public void revokePermission(PermissionId permission)',
       'Remove permission. Raise PermissionRevokedEvent.', None)
METHOD(doc, 'public boolean hasPermission(PermissionId permission)',
       'Check role có permission này không.', None)
METHOD(doc, 'public Set<PermissionId> getPermissions()',
       'Trả Collections.unmodifiableSet.',
       'Defensive copy ngăn caller mutate state internal.')

H(doc, '8.1 Role Value Objects', level=3)
FILE_HEADER(doc, 'RoleId.java', 'vo/RoleId.java')
P(doc, 'record RoleId(String value) - UUID.')

FILE_HEADER(doc, 'RoleName.java', 'vo/RoleName.java')
P(doc, 'record RoleName(String value).')
P(doc, 'Convention: bắt đầu với "ROLE_" để Spring Security nhận dạng (hasRole("ADMIN") sẽ check "ROLE_ADMIN").')

H(doc, '8.2 Role Events', level=3)
TABLE(doc,
      headers=['Event', 'Khi nào'],
      rows=[
          ['PermissionGrantedEvent', 'grantPermission() add thật'],
          ['PermissionRevokedEvent', 'revokePermission() remove thật'],
      ],
      col_widths=[6, 10])

NOTE(doc,
     '2 event này còn được RolePermissionCacheInvalidator lắng nghe để invalidate Redis cache role→permissions. '
     'Mapper cũng map ra Kafka để service khác cập nhật cache local của họ.')

# ── 9. Permission ──
H(doc, '9. domain/models/permission - Permission entity', level=2)

FILE_HEADER(doc, 'Permission.java', 'domain/models/permission/Permission.java')
P(doc, 'Entity (KHÔNG aggregate root). Permission được sở hữu/quản lý qua Role aggregate.', bold=True)
P(doc, 'Field: id, name (resource + action), description, audit.')

P(doc, 'Method:', bold=True)
METHOD(doc, 'public static Permission create(String resource, String action, String description)',
       'Tạo permission mới. Ví dụ: ("user", "read", "View user info").', None)
METHOD(doc, 'public static Permission reconstitute(...)', 'Khôi phục từ DB.', None)
METHOD(doc, 'public void updateDescription(String description)', 'Update.', None)
METHOD(doc, 'public boolean grants(String resource, String action)',
       'Check permission cho phép action lên resource cụ thể.',
       'Mô hình ABAC đơn giản dựa trên resource+action.')

H(doc, '9.1 Permission Value Objects', level=3)
FILE_HEADER(doc, 'PermissionId.java', 'vo/PermissionId.java')
P(doc, 'record PermissionId(String value) - UUID.')

FILE_HEADER(doc, 'PermissionName.java', 'vo/PermissionName.java')
P(doc, 'record PermissionName(String resource, String action, String value).', bold=True)
P(doc, 'Format value: "resource:action" (ví dụ "user:read").')
METHOD(doc, 'public static PermissionName of(String resource, String action)',
       'Build với value = resource + ":" + action.', None)
METHOD(doc, 'public boolean isForResource(String resource)',
       'Match resource.', None)
METHOD(doc, 'public boolean allowsAction(String action)',
       'Match action.', None)

# ── 10. Domain services ──
H(doc, '10. domain/services - Domain services & Ports', level=2)

FILE_HEADER(doc, 'PasswordEncoderPort.java', 'domain/services/PasswordEncoderPort.java')
P(doc, 'Interface (PORT) - hash password. Implement bởi BCryptPasswordEncoderAdapter.', bold=True)
METHOD(doc, 'String encode(String rawPassword)',
       'Hash password.', None)
METHOD(doc, 'boolean matches(String raw, String encoded)',
       'So sánh raw với hash.', None)
NOTE(doc,
     'Domain dùng interface này để hash password mà KHÔNG biết về BCrypt/Argon2/etc. '
     'Đổi thuật toán mai sau chỉ cần thay adapter, domain không đụng.')

FILE_HEADER(doc, 'TokenProviderPort.java', 'domain/services/TokenProviderPort.java')
P(doc, 'Interface (PORT) - phát hành + parse access token. Implement bởi JwtTokenProvider.', bold=True)
P(doc, 'Inner record:', bold=True)
METHOD(doc, 'record IssuedAccessToken(String token, String tokenId, Instant expiresAt, long expiresInSeconds)',
       'Output của issue.', None)
METHOD(doc, 'record AccessClaims(String tokenId, String userId, String username, int tokenVersion, Set<String> roles, Instant expiresAt)',
       'Parse result.',
       'tokenVersion để check JwtAuthenticationFilter so với DB.')
P(doc, 'Method:', bold=True)
METHOD(doc, 'IssuedAccessToken issueAccessToken(User user, Set<String> roles)',
       'Mint token mới cho user.', None)
METHOD(doc, 'AccessClaims parseAccessToken(String token)',
       'Parse + verify signature + expiry. Throw nếu invalid.', None)

FILE_HEADER(doc, 'AuthenticationDomainService.java', 'domain/services/AuthenticationDomainService.java')
P(doc, 'Stateless service - logic phân quyền cross-aggregate (User + Role + Permission).', bold=True)
NOTE(doc,
     'Tại sao là domain service, không là method trong User? Vì logic này cần Collection<Role> và '
     'Collection<Permission> - những thứ không thuộc về 1 aggregate đơn lẻ. Khi logic vượt phạm vi '
     '1 aggregate → tạo domain service.')

P(doc, 'Method:', bold=True)
METHOD(doc, 'public boolean hasPermission(User, List<Role> userRoles, List<Permission> allSystemPermissions, String permissionName)',
       'User có permission này (qua role) không?',
       'Tách parameter rõ ràng giúp service không cần inject repository.')
METHOD(doc, 'public boolean hasRole(User, List<Role> userRoles, String roleName)',
       'User có role không?', None)
METHOD(doc, 'public Set<String> getPermissionNames(User, List<Role>, List<Permission>)',
       'Lấy mọi permission name user có.', None)
METHOD(doc, 'public boolean canAccessResource(User, List<Role>, List<Permission>, String resource, String action)',
       'ABAC check: user có quyền action lên resource không?', None)
METHOD(doc, 'private Set<String> extractPermissionIds(Collection<Role>)',
       'Helper: flatMap permissions từ tất cả roles.', None)

FILE_HEADER(doc, 'TokenDomainService.java', 'domain/services/TokenDomainService.java')
P(doc,
  'Stateless service - logic refresh token cross-aggregate (User + RefreshToken). '
  'Implement Rotation + Family Revocation pattern.', bold=True)

P(doc, 'Method:', bold=True)
METHOD(doc, 'public RefreshToken issueForUser(User user, int expiryDays)',
       'Phát hành root token. Gọi user.ensureAuthenticatable() trước.',
       'Đảm bảo không issue token cho account locked/disabled - kể cả khi caller skip kiểm tra.')
METHOD(doc, 'public RefreshToken rotate(RefreshToken old, int expiryDays, RefreshTokenRepository repository)',
       'KEY METHOD: Rotation + reuse detection.',
       'Bước 1: nếu old.isReuseAttempt() → revokeFamily() + throw TokenReuseDetectedException. '
       'Bước 2: old.verifyValidity(). '
       'Bước 3: old.revoke(NORMAL), save. '
       'Bước 4: tạo RefreshToken.rotate(old) với gen+1.')
METHOD(doc, 'public void revokeFamily(TokenFamily family, RefreshTokenRepository repository)',
       'Cascade revoke mọi token live trong family với reason=FAMILY_REVOKED.',
       'Khi phát hiện token theft, không thể biết token nào là của attacker → revoke hết, user phải login lại.')
METHOD(doc, 'public void validateOwnership(RefreshToken token, UserId userId)',
       'Throw TokenOwnershipException nếu token không thuộc user.',
       'Defense-in-depth: ngăn user A dùng token leak của user B.')

# ── 11. Repositories interface ──
H(doc, '11. domain/repositories - Repository interfaces', level=2)
P(doc, '4 interface, đặt ở domain để aggregate có thể "yêu cầu" persistence mà không biết JPA.')

FILE_HEADER(doc, 'UserRepository.java', 'domain/repositories/UserRepository.java')
TABLE(doc,
      headers=['Method', 'Mục đích'],
      rows=[
          ['save(User)', 'Insert hoặc update'],
          ['findById(UserId)', 'Lookup theo ID'],
          ['findByUsername(Username)', 'Lookup theo username'],
          ['findByEmail(Email)', 'Lookup theo email'],
          ['existsByUsername(Username)', 'Check tồn tại'],
          ['existsByEmail(Email)', 'Check tồn tại'],
          ['delete(User)', 'Xóa'],
          ['findByIdWithRoles(UserId)', 'Lookup + fetch roles (LEFT JOIN FETCH)'],
          ['findByUsernameWithRoles(Username)', 'Login flow - load roles cùng lúc'],
          ['findAfterCursor(Instant cursorCreatedAt, String cursorId, int limit)', 'Keyset pagination'],
      ],
      col_widths=[7, 9])

FILE_HEADER(doc, 'RefreshTokenRepository.java', 'domain/repositories/RefreshTokenRepository.java')
TABLE(doc,
      headers=['Method', 'Mục đích'],
      rows=[
          ['save(RefreshToken)', 'Persist'],
          ['findById(TokenId)', 'Lookup'],
          ['findByTokenValue(TokenValue)', 'Lookup theo value'],
          ['findByTokenValueForUpdate(TokenValue)', 'PESSIMISTIC LOCK - cho rotation flow'],
          ['findByUserId(UserId)', 'Liệt kê token của user'],
          ['findValidTokensByUserId(UserId)', 'Chỉ token còn valid'],
          ['findByFamily(TokenFamily)', 'Family revocation'],
          ['delete(RefreshToken)', 'Xóa'],
          ['deleteByUserId(UserId)', 'Xóa khi delete user'],
          ['deleteExpiredTokens()', 'Cleanup job'],
          ['revokeAllTokensForUser(UserId)', 'Change password - revoke hàng loạt'],
      ],
      col_widths=[6.5, 9.5])

NOTE(doc,
     'findByTokenValueForUpdate là RẤT QUAN TRỌNG. SELECT FOR UPDATE serialize 2 request refresh đồng '
     'thời cùng token → ngăn race condition gây 2 access token cùng lúc. Không có nó: cả 2 request '
     'cùng pass, mỗi cái issue 1 token mới = double session leak.')

FILE_HEADER(doc, 'RoleRepository.java', 'domain/repositories/RoleRepository.java')
TABLE(doc,
      headers=['Method', 'Mục đích'],
      rows=[
          ['save(Role)', 'Persist'],
          ['findById/ByName', 'Lookup'],
          ['findByIdIn(Set<RoleId>)', 'Bulk lookup'],
          ['findByIdWithPermissions / findByNameWithPermissions', 'Fetch permissions cùng lúc'],
          ['existsByName(RoleName)', 'Check duplicate'],
          ['findAll()', 'List toàn bộ - dùng cho cache warm-up'],
      ],
      col_widths=[7, 9])

FILE_HEADER(doc, 'PermissionRepository.java', 'domain/repositories/PermissionRepository.java')
P(doc, 'Tương tự với save, findById, findByIdIn(Set<PermissionId>), findAll...')

doc.add_page_break()

# ============================================================
# PHẦN III - TẦNG APPLICATION
# ============================================================
H(doc, 'PHẦN III - TẦNG APPLICATION (CQRS)', level=1)

P(doc,
  'Application là tầng "đạo diễn" - không chứa quy tắc nghiệp vụ, mà điều phối domain để hoàn '
  'thành use case. Áp dụng CQRS: tách Command (write) và Query (read).')

# ── 12. application/common ──
H(doc, '12. application/common - Interface chung', level=2)

FILE_HEADER(doc, 'Command.java', 'application/common/Command.java')
P(doc, 'Marker interface Command<R>. Mỗi command record implements Command<ResultType>.', bold=True)
P(doc, 'Chỉ là marker - không có method. Generic R xác định kiểu trả về của handler.')

FILE_HEADER(doc, 'CommandHandler.java', 'application/common/CommandHandler.java')
P(doc, 'Functional interface CommandHandler<C extends Command<R>, R>.', bold=True)
METHOD(doc, 'R handle(C command)',
       'Method duy nhất - xử lý command, trả result.',
       '@FunctionalInterface để có thể implement bằng lambda nếu cần (chủ yếu trong test).')

FILE_HEADER(doc, 'Query.java', 'application/common/Query.java')
P(doc, 'Marker interface Query<R>. Tương tự Command nhưng cho read-only.')

FILE_HEADER(doc, 'QueryHandler.java', 'application/common/QueryHandler.java')
P(doc, 'Functional interface QueryHandler<Q extends Query<R>, R>.', bold=True)
NOTE(doc, 'Handler bắt buộc chạy với @Transactional(readOnly = true) để Hibernate skip dirty checking → nhanh hơn.')

FILE_HEADER(doc, 'CursorCodec.java', 'application/common/CursorCodec.java')
P(doc, 'Utility class encode/decode cursor cho keyset pagination.', bold=True)
P(doc, 'Format: base64("<epochMicros>|<id>"). Client thấy là opaque, không forge được.')

P(doc, 'Inner record:', bold=True)
METHOD(doc, 'public record Cursor(Instant createdAt, String id)',
       'Decoded pair. Compact constructor validate non-null.', None)

P(doc, 'Method:', bold=True)
METHOD(doc, 'public static String encode(Instant createdAt, String id)',
       'Encode thành base64 URL-safe. Trả null nếu input null.',
       'Dùng URL-safe variant để cursor có thể đặt trong query string không cần URL-encode.')
METHOD(doc, 'public static Cursor decode(String opaque)',
       'Decode + parse + validate. Throw IllegalArgumentException nếu malformed.',
       'Validation ngăn client forge cursor để skip ahead.')
NOTE(doc, 'Dùng epoch MICROS thay vì millis vì Postgres lưu timestamp đến microsecond - tránh tie-breaking issues.')

FILE_HEADER(doc, 'ApplicationException.java', 'application/common/ApplicationException.java')
P(doc, 'Exception gốc cho lỗi application layer (validation, mapping...).')

FILE_HEADER(doc, 'ValidationException.java', 'application/common/ValidationException.java')
P(doc, 'Throw khi input command/query không hợp lệ vượt mức Bean Validation.')

# ── 13. application/command ──
H(doc, '13. application/command - 9 Command records', level=2)

TABLE(doc,
      headers=['Command', 'Field', 'Handler', 'Output'],
      rows=[
          ['RegisterUserCommand', 'username, email, rawPassword, firstName, lastName', 'RegisterUserHandler', 'AuthResponseDTO'],
          ['LoginCommand', 'usernameOrEmail, rawPassword', 'LoginHandler', 'AuthResponseDTO'],
          ['RefreshTokenCommand', 'refreshToken', 'RefreshTokenHandler', 'AuthResponseDTO'],
          ['LogoutCommand', 'accessToken, refreshToken', 'LogoutHandler', 'Void'],
          ['ChangePasswordCommand', 'userId, oldRawPassword, newRawPassword, currentAccessTokenJti, currentAccessTokenExp', 'ChangePasswordHandler', 'Void'],
          ['UpdateEmailCommand', 'userId, newEmail', 'UpdateEmailHandler', 'UserDTO'],
          ['AssignRoleCommand', 'userId, roleName', 'AssignRoleHandler', 'Void'],
          ['UnassignRoleCommand', 'userId, roleName', 'UnassignRoleHandler', 'Void'],
          ['ChangeAccountStatusCommand', 'userId, Transition (LOCK/UNLOCK/ENABLE/DISABLE)', 'ChangeAccountStatusHandler', 'Void'],
      ],
      col_widths=[4, 5, 4, 3])

NOTE(doc,
     'Tất cả command là Java record - immutable, equals/hashCode/toString tự động. '
     'Constructor compact có thể validate (chuẩn hóa input, không null...).')

# ── 14. application/query ──
H(doc, '14. application/query - 5 Query records', level=2)
TABLE(doc,
      headers=['Query', 'Field', 'Handler', 'Output'],
      rows=[
          ['GetUserByIdQuery', 'userId', 'GetUserByIdHandler', 'UserDTO'],
          ['GetUserByUsernameQuery', 'username', 'GetUserByUsernameHandler', 'UserDTO'],
          ['ListUsersQuery', 'cursor, limit', 'ListUsersHandler', 'PageDTO<UserDTO>'],
          ['CheckRoleQuery', 'userId, roleName', 'CheckRoleHandler', 'Boolean'],
          ['CheckPermissionQuery', 'userId, permissionName', 'CheckPermissionHandler', 'Boolean'],
      ],
      col_widths=[4.5, 4.5, 4, 3])

# ── 15. application/handler ──
H(doc, '15. application/handler - 14 Handler classes', level=2)

P(doc, 'Mỗi handler là @Service Spring bean, implement CommandHandler hoặc QueryHandler. '
       '@Transactional bao bọc method handle.')

# RegisterUserHandler
H(doc, '15.1 RegisterUserHandler', level=3)
P(doc, 'Vị trí: application/handler/RegisterUserHandler.java')
P(doc, 'Field:', bold=True)
for x in ['private static final String DEFAULT_ROLE = "ROLE_USER" (hằng số class)',
          'private final UserRepository userRepository (DI)',
          'private final RoleRepository roleRepository (DI)',
          'private final RefreshTokenRepository refreshTokenRepository (DI)',
          'private final PasswordEncoderPort passwordEncoder (DI)',
          'private final TokenProviderPort tokenProvider (DI)',
          'private final TokenDomainService tokenDomainService (DI)',
          'private final UserDtoMapper userDtoMapper (DI)',
          '@Value("${jwt.refresh-expiration-days:7}") private int refreshExpiryDays']:
    BULLET(doc, x)

P(doc, 'Method handle(RegisterUserCommand):', bold=True)
CODE(doc, '''1. Tạo Username.of(), Email.of() - VO ép buộc format validation
2. Check existsByUsername / existsByEmail → throw UserAlreadyExistsException
3. User.register(...) - hash password, raise UserCreatedEvent
4. roleRepository.findByName("ROLE_USER") → assignRole → raise RoleAssignedEvent
5. userRepository.save(user) - drain events sau khi persist
6. roleRepository.findByIdIn(user.getRoles()) - lấy Role objects
7. tokenProvider.issueAccessToken(user, roleNames) - mint JWT
8. tokenDomainService.issueForUser(user, refreshExpiryDays) - tạo RefreshToken
9. refreshTokenRepository.save(refresh) - persist + raise TokenCreatedEvent
10. Return AuthResponseDTO.bearer(...)''')

# LoginHandler
H(doc, '15.2 LoginHandler', level=3)
P(doc, 'Vị trí: application/handler/LoginHandler.java')
P(doc, 'Method handle(LoginCommand):', bold=True)
CODE(doc, '''1. lookup(usernameOrEmail) - heuristic "@" → findByEmail, else findByUsername
   → Optional<User>; nếu empty throw InvalidCredentialsException (SAME exception
     cho cả "không tồn tại" và "password sai" → timing-safe)
2. user.authenticate(rawPassword, encoder)
   → ensureAuthenticatable() throw AccountNotUsableException nếu lock/disable/expired
   → encoder.matches() BCrypt verify ~250ms
   → if OK: recordLogin() - update lastLogin, raise UserLoggedInEvent
3. userRepository.save(user) - lastLogin được persist, event publish
4. roleRepository.findByIdIn - lấy roles
5. tokenProvider.issueAccessToken - JWT mới
6. tokenDomainService.issueForUser - RefreshToken mới (family mới hoàn toàn)
7. refreshTokenRepository.save
8. Return AuthResponseDTO.bearer(...)''')

NOTE(doc,
     'Method lookup() có try-catch IllegalArgumentException - nếu input malformed (format Email invalid) '
     'cũng trả Optional.empty → throw InvalidCredentials. Tránh leak thông tin qua các response code khác nhau.')

# RefreshTokenHandler
H(doc, '15.3 RefreshTokenHandler', level=3)
P(doc, 'Vị trí: application/handler/RefreshTokenHandler.java')
P(doc, 'Method handle(RefreshTokenCommand):', bold=True)
CODE(doc, '''1. refreshTokenRepository.findByTokenValueForUpdate(value)  ← SELECT FOR UPDATE
   → presented token; nếu không tìm thấy throw TokenRevokedException(null,null)
2. tokenDomainService.rotate(presented, refreshExpiryDays, repository)
   → Nếu presented.isReuseAttempt():
     • revokeFamily() - revoke mọi token live trong family với reason=FAMILY_REVOKED
     • throw TokenReuseDetectedException
   → presented.verifyValidity() - check expired
   → presented.revoke(NORMAL) - raise TokenRevokedEvent
   → repo.save(presented)
   → Return RefreshToken.rotate(presented, ...) - gen+1, raise TokenRotatedEvent
3. refreshTokenRepository.save(rotated)
4. userRepository.findById(userId) - lấy user để mint JWT mới
5. roleRepository.findByIdIn - roles
6. tokenProvider.issueAccessToken - JWT mới
7. Return AuthResponseDTO''')

# LogoutHandler
H(doc, '15.4 LogoutHandler', level=3)
P(doc, 'Vị trí: application/handler/LogoutHandler.java')
P(doc, 'Method handle(LogoutCommand):', bold=True)
CODE(doc, '''1. refreshTokenRepository.findByTokenValue(refreshToken).ifPresent(this::revokeRefresh)
   → revokeRefresh: token.revoke(USER_INITIATED), save
   (Idempotent - missing token = user đã logout rồi)

2. if (accessToken != null && !blank):
   try {
     claims = tokenProvider.parseAccessToken(accessToken)
     tokenBlacklist.revoke(claims.tokenId(), userId, expiresAt, "LOGOUT")
   } catch (Exception) {
     log warn - access token expired/malformed thì bỏ qua
     (KHÔNG được rollback refresh revoke đã commit)
   }''')

NOTE(doc,
     'H5 comment: try-catch RIÊNG cho access token blacklist sau khi refresh revoke đã commit. '
     'Nếu gộp vào 1 try-catch lớn: expired access token sẽ rollback refresh revoke → user vẫn login được.')

# ChangePasswordHandler
H(doc, '15.5 ChangePasswordHandler', level=3)
P(doc, 'Vị trí: application/handler/ChangePasswordHandler.java')
P(doc, 'Method handle(ChangePasswordCommand):', bold=True)
CODE(doc, '''1. userRepository.findById(userId) → throw UserNotFoundException nếu không có
2. user.changePassword(oldRaw, newRaw, encoder)
   → encoder.matches(old) - verify mật khẩu cũ
   → password = encoder.encode(new)
   → incrementTokenVersion()  ← KEY
   → raise PasswordChangedEvent
3. userRepository.save(user) - persist version mới
4. refreshTokenRepository.revokeAllTokensForUser(userId) - bulk revoke
5. tokenBlacklist.revoke(jti, userId, exp, "PASSWORD_CHANGED")
   → Đóng race window: trong vài ms giữa bump version và Redis update, JWT cũ
     vẫn có thể được dùng. Blacklist ngay để filter reject sớm hơn.''')

# UpdateEmailHandler
H(doc, '15.6 UpdateEmailHandler', level=3)
P(doc, 'Vị trí: application/handler/UpdateEmailHandler.java')
CODE(doc, '''1. findById → user
2. Email.of(newEmail) - VO validation
3. Nếu newEmail != currentEmail && existsByEmail → throw UserAlreadyExistsException
4. user.updateEmail(newEmail) - raise EmailChangedEvent
5. userRepository.save
6. Return UserDTO mapped''')

# AssignRole / UnassignRole
H(doc, '15.7 AssignRoleHandler / UnassignRoleHandler', level=3)
CODE(doc, '''AssignRoleHandler:
  findById(userId) → user
  findByName(roleName) → role (throw IllegalArgumentException nếu không tồn tại)
  user.assignRole(role.getId()) → raise RoleAssignedEvent nếu add thật
  save

UnassignRoleHandler: tương tự, gọi user.unassignRole''')

# ChangeAccountStatusHandler
H(doc, '15.8 ChangeAccountStatusHandler', level=3)
CODE(doc, '''switch (command.transition()) {
    case LOCK    -> user.lock();
    case UNLOCK  -> user.unlock();
    case ENABLE  -> user.enable();
    case DISABLE -> user.disable();
}
userRepository.save(user)

Mỗi method raise AccountStatusChangedEvent với Transition enum tương ứng.''')

# Query handlers
H(doc, '15.9 GetUserByIdHandler (Query)', level=3)
CODE(doc, '''@Transactional(readOnly = true)

1. userRepository.findByIdWithRoles(userId) → user (throw nếu không có)
2. roleRepository.findByIdIn(user.getRoles()) → List<Role>
3. flatMap permissions từ roles → Set<PermissionId>
4. permissionRepository.findByIdIn(permissionIds) → List<Permission>
5. userDtoMapper.toDto(user, roles, permissions) → UserDTO đầy đủ''')

H(doc, '15.10 GetUserByUsernameHandler', level=3)
P(doc,
  'Lightweight - SKIP permission resolution. Dùng cho "admin lookup who is this", không phải authz.')

H(doc, '15.11 ListUsersHandler (Keyset pagination)', level=3)
CODE(doc, '''private static final int MAX_LIMIT = 100;

1. pageSize = Math.clamp(query.limit(), 1, MAX_LIMIT)
2. CursorCodec.decode(query.cursor()) → Cursor (null cho trang đầu)
3. userRepository.findAfterCursor(cursorCreatedAt, cursorId, pageSize + 1)
   → Over-fetch +1 để detect hasNext mà không cần count query riêng
4. hasNext = rows.size() > pageSize
   pageRows = hasNext ? rows.subList(0, pageSize) : rows
5. Map từng user → UserDTO (gọi findByIdIn từng row)
6. Nếu hasNext: encode lastRow.createdAt + id → nextCursor
7. Return PageDTO(items, nextCursor, pageSize, -1)
   (-1 = totalElements unknown, không count để tránh full scan)''')

H(doc, '15.12 CheckRoleHandler', level=3)
CODE(doc, '''@Transactional(readOnly = true)

1. userRepository.findByIdWithRoles(userId)
2. user.getRoles() → Set<RoleId>
3. roleRepository.findByIdIn(roleIds) → roles
4. Trả true nếu có role nào có name == query.roleName

Optional chain - nếu user không có thì trả false (không throw).''')

H(doc, '15.13 CheckPermissionHandler', level=3)
CODE(doc, '''@Transactional(readOnly = true)

1. userRepository.findByIdWithRoles(userId)
2. user.getRoles() → Set<RoleId>
3. roleRepository.findByIdIn → flatMap permissions → Set<PermissionId>
4. permissionRepository.findByIdIn → check anyMatch name == query.permissionName

Tối đa 3 DB round trips. Có thể optimize nếu cache hot.''')

# ── 16. application/dto ──
H(doc, '16. application/dto - Read models', level=2)

FILE_HEADER(doc, 'UserDTO.java', 'application/dto/UserDTO.java')
P(doc, 'record với 14 field. @JsonInclude(NON_NULL) - skip null khi serialize.')
TABLE(doc,
      headers=['Field', 'Kiểu'],
      rows=[
          ['id', 'String (UUID)'],
          ['username, email', 'String'],
          ['firstName, lastName', 'String'],
          ['enabled, accountNonExpired, accountNonLocked, credentialsNonExpired', 'boolean'],
          ['roles', 'Set<String>'],
          ['permissions', 'Set<String>'],
          ['createdAt, updatedAt, lastLogin', 'Instant'],
      ],
      col_widths=[10, 6])
NOTE(doc, 'id là String thay vì Long - tránh bug truncate UUID khi serialize JSON.')

FILE_HEADER(doc, 'AuthResponseDTO.java', 'application/dto/AuthResponseDTO.java')
P(doc, 'record(accessToken, refreshToken, tokenType, expiresInSeconds, UserDTO user).', bold=True)
METHOD(doc, 'public static AuthResponseDTO bearer(String access, String refresh, long ttl, UserDTO user)',
       'Factory tự set tokenType = "Bearer".',
       'Tránh caller hard-code "Bearer" string ở khắp nơi.')
NOTE(doc, 'AuthCookieWriter.writeTokens() nhân DTO này, KHÔNG trả accessToken/refreshToken trong body, chỉ set cookies.')

FILE_HEADER(doc, 'PageDTO.java', 'application/dto/PageDTO.java')
P(doc, 'record(List<T> items, String nextCursor, int pageSize, long totalElements).', bold=True)
P(doc, 'Generic envelope - tránh leak Spring Data Page<T> (chứa thông tin nội bộ không cần) ra API.')
METHOD(doc, 'public static <T> PageDTO<T> of(List<T> items, String nextCursor, int pageSize, long total)',
       'Factory. Dùng List.copyOf để immutable.',
       'totalElements = -1 nghĩa là "không count" - tránh expensive COUNT(*) trên bảng lớn.')

FILE_HEADER(doc, 'RoleDTO.java', 'application/dto/RoleDTO.java')
P(doc, 'record(id, name, description, Set<String> permissions, createdAt, updatedAt).')

FILE_HEADER(doc, 'PermissionDTO.java', 'application/dto/PermissionDTO.java')
P(doc, 'record(id, name, resource, action, description, createdAt, updatedAt).')

# ── 17. application/mapper ──
H(doc, '17. application/mapper - Domain → DTO', level=2)

FILE_HEADER(doc, 'UserDtoMapper.java', 'application/mapper/UserDtoMapper.java')
P(doc, '@Component (singleton Spring bean) - convert Domain → DTO.', bold=True)
METHOD(doc, 'public UserDTO toDto(User, Collection<Role>, Collection<Permission>)',
       'Full projection - user + roles + permissions.',
       'Caller chịu trách nhiệm pre-fetch roles + permissions để tránh N+1 trong mapper.')
METHOD(doc, 'public UserDTO toDto(User, Collection<Role>)',
       'Overload không có permissions - cheap projection.', None)
METHOD(doc, 'public RoleDTO toDto(Role, Set<String> grantedPermissions)',
       'Map Role + permission names.', None)
METHOD(doc, 'public PermissionDTO toDto(Permission)',
       'Map Permission.', None)

# ── 18. application/port ──
H(doc, '18. application/port - Outbound ports', level=2)

FILE_HEADER(doc, 'TokenBlacklistPort.java', 'application/port/TokenBlacklistPort.java')
METHOD(doc, 'void revoke(String tokenId, String userId, Instant expiresAt, String reason)',
       'Đưa token vào blacklist. expiresAt để TTL Redis tự dọn dẹp.', None)
METHOD(doc, 'boolean isRevoked(String tokenId)',
       'Check token bị revoke chưa. Filter gọi mỗi request.',
       'Cần nhanh - implementation dùng Redis fast-path, DB fallback.')

FILE_HEADER(doc, 'RolePermissionCachePort.java', 'application/port/RolePermissionCachePort.java')
METHOD(doc, 'Set<String> get(String roleName)', 'Cache lookup. Trả null nếu miss.', None)
METHOD(doc, 'void put(String roleName, Set<String> permissionNames)', 'Cache write.', None)
METHOD(doc, 'void evict(String roleName)', 'Xóa 1 entry.', None)
METHOD(doc, 'void evictAll()', 'Wipe toàn bộ - dùng cho bulk permission changes.', None)

NOTE(doc, '2 port này đặt ở application/ (không phải domain/) vì cache không phải nghiệp vụ thuần, '
          'mà là optimization của application use case.')

# ── 19. application/events ──
H(doc, '19. application/events - Integration events ra Kafka', level=2)

FILE_HEADER(doc, 'IntegrationEvent.java', 'application/events/IntegrationEvent.java')
P(doc, 'Interface contract cho mọi event ra Kafka.', bold=True)
METHOD(doc, 'UUID eventId()', 'Unique id - consumer dedupe.', None)
METHOD(doc, 'String eventType()',
       'Stable string ("auth.user.created.v1").',
       'KHÔNG dùng class name - service consume có thể là Python/Node.')
METHOD(doc, 'Instant occurredOn()', 'Thời điểm raise ở source.', None)
METHOD(doc, 'String aggregateId()', 'Cho Kafka partitioning.', None)
METHOD(doc, 'int schemaVersion()', 'Bump khi breaking change payload.', None)

FILE_HEADER(doc, 'KafkaTopics.java', 'application/events/KafkaTopics.java')
P(doc, 'final class utility - chứa hằng số tên topic.', bold=True)
CODE(doc, '''public final class KafkaTopics {
    private KafkaTopics() {}                                     // CẤM new
    public static final String AUTH_USER_EVENTS    = "auth.user.events.v1";
    public static final String AUTH_SESSION_EVENTS = "auth.session.events.v1";
}''')

NOTE(doc, 'Bộ 3 đặc trưng utility class: (1) final class - không kế thừa, (2) private constructor - không new, '
          '(3) chỉ static members. Pattern này JDK dùng cho Math, Collections, Arrays.')

FILE_HEADER(doc, 'UserLifecycleIntegrationEvent.java', 'application/events/UserLifecycleIntegrationEvent.java')
P(doc, 'record implements IntegrationEvent. 6 field: eventId, eventType, occurredOn, aggregateId, schemaVersion, payload (Map<String,Object>).')
P(doc, 'Gửi vào topic "auth.user.events.v1". Cho: UserCreated, EmailChanged, PasswordChanged, RoleAssigned, RoleRemoved, AccountStatusChanged.')

FILE_HEADER(doc, 'SessionIntegrationEvent.java', 'application/events/SessionIntegrationEvent.java')
P(doc, 'Cùng cấu trúc field. Gửi vào topic "auth.session.events.v1". Cho: UserLoggedIn, TokenCreated, TokenRotated, TokenRevoked.')

NOTE(doc, 'TẠI SAO TÁCH 2 RECORD CÙNG FIELD? Routing tới topic khác → consumer khác → tách để type-safe method overloading.')

doc.add_page_break()

# ============================================================
# PHẦN IV - TẦNG INFRASTRUCTURE
# ============================================================
H(doc, 'PHẦN IV - TẦNG INFRASTRUCTURE', level=1)

# ── 20. infrastructure/security ──
H(doc, '20. infrastructure/security - JWT, BCrypt, UserDetails', level=2)

FILE_HEADER(doc, 'JwtProperties.java', 'infrastructure/security/JwtProperties.java')
P(doc, '@ConfigurationProperties(prefix = "jwt") - bind YAML.', bold=True)
P(doc, 'record(String secret, long accessExpirationSeconds, int refreshExpirationDays, String issuer).')
P(doc, 'Compact constructor validation (fail-fast at startup):', bold=True)
for x in ['secret phải ≥ 32 chars (yêu cầu HS256)',
          'Nếu prod profile + secret bắt đầu "change-me" → throw (chặn placeholder lọt prod)',
          'accessExpirationSeconds ≤ 0 → default 900 (15 phút)',
          'refreshExpirationDays ≤ 0 → default 7 ngày',
          'issuer null/blank → default "hieu.com"']:
    BULLET(doc, x)

FILE_HEADER(doc, 'JwtTokenProvider.java', 'infrastructure/security/JwtTokenProvider.java')
P(doc, '@Component implements TokenProviderPort. JWT HS256 implementation dùng jjwt library.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['private static final MacAlgorithm ALGORITHM = Jwts.SIG.HS256 (ép HS256)',
          'private final JwtProperties props',
          'private final SecretKey key (HMAC key)']:
    BULLET(doc, x)
P(doc, 'Constructor: Keys.hmacShaKeyFor(props.secret().getBytes()) build SecretKey 1 lần.')
P(doc, 'Method issueAccessToken(User user, Set<String> roles):', bold=True)
CODE(doc, '''Instant now = Instant.now();
Instant exp = now.plusSeconds(props.accessExpirationSeconds());
String tokenId = UUID.randomUUID().toString();

Map<String, Object> claims = new HashMap<>();
claims.put("jti", tokenId);
claims.put("userId", user.getId().value());
claims.put("email", user.getEmail().value());
claims.put("tokenVersion", user.getTokenVersion());
claims.put("roles", roles);

String token = Jwts.builder()
    .id(tokenId).claims(claims)
    .subject(user.getUsername().value())
    .issuer(props.issuer())
    .issuedAt(Date.from(now))
    .expiration(Date.from(exp))
    .signWith(key, ALGORITHM)
    .compact();

return new IssuedAccessToken(token, tokenId, exp, ttlSeconds);''')

P(doc, 'Method parseAccessToken(String token):', bold=True)
CODE(doc, '''Claims c = parseClaims(token);  // verify signature + issuer
Set<String> roles = Set.copyOf(c.getOrDefault("roles", Set.of()));
int tv = ((Number) c.getOrDefault("tokenVersion", 0)).intValue();
String userId = c.get("userId", String.class);
return new AccessClaims(c.getId(), userId, c.getSubject(), tv, roles, c.getExpiration().toInstant());''')

P(doc, 'Method isSignatureValid(String token):', bold=True)
P(doc, 'try/catch chỉ verify signature, không throw - dùng trong filter cho cheap check.')

P(doc, 'Private method parseClaims:', bold=True)
CODE(doc, '''Jwts.parser()
    .verifyWith(key)
    .requireIssuer(props.issuer())   ← KEY: chặn token mint bởi service khác cùng secret
    .build()
    .parseSignedClaims(token)
    .getPayload();''')

FILE_HEADER(doc, 'BCryptPasswordEncoderAdapter.java', 'infrastructure/security/BCryptPasswordEncoderAdapter.java')
P(doc, '@Component implements PasswordEncoderPort. Delegate sang Spring Security PasswordEncoder bean.')
METHOD(doc, 'encode(rawPassword)', 'delegate.encode() - BCrypt hash.', None)
METHOD(doc, 'matches(raw, hash)', 'delegate.matches() - constant-time compare.', None)
NOTE(doc, 'Đây là adapter thin - thực sự BCryptPasswordEncoder bean định nghĩa trong SecurityConfig với cost=12.')

FILE_HEADER(doc, 'AuthUserDetails.java', 'infrastructure/security/AuthUserDetails.java')
P(doc, 'final class implements UserDetails. Principal đưa vào SecurityContext.', bold=True)
P(doc, 'Field (đều final):', bold=True)
TABLE(doc,
      headers=['Field', 'Kiểu', 'Lý do thêm'],
      rows=[
          ['userId', 'String', 'Stable UUID - dễ lookup hơn username'],
          ['username', 'String', 'UserDetails contract'],
          ['email', 'String', 'Cho controller @AuthenticationPrincipal'],
          ['passwordHash', 'String', 'UserDetails contract (dù không dùng)'],
          ['tokenVersion', 'int', 'KEY: filter so với JWT.tokenVersion'],
          ['authorities', 'Set<GrantedAuthority>', 'Defensive copy - immutable'],
          ['4 boolean status flags', 'boolean', 'UserDetails contract'],
      ],
      col_widths=[3.5, 4, 8.5])

METHOD(doc, 'public static AuthUserDetails from(User user, Set<String> roleNames, Set<String> permissionNames)',
       'Build từ domain User + resolved role/permission names. Tạo SimpleGrantedAuthority cho từng role/permission.',
       'roles và permissions đều thành GrantedAuthority - controller dùng hasRole/hasAuthority đều được.')

P(doc, 'UserDetails contract method: getAuthorities, getPassword, getUsername, isEnabled, isAccountNonExpired, isAccountNonLocked, isCredentialsNonExpired. equals/hashCode dựa trên userId.')

FILE_HEADER(doc, 'CustomUserDetailsService.java', 'infrastructure/security/CustomUserDetailsService.java')
P(doc, '@Service implements UserDetailsService. Wired vào Spring Security.', bold=True)

METHOD(doc, 'public UserDetails loadUserByUsername(String username)',
       'Spring Security contract. findByUsernameWithRoles → toUserDetails.',
       '@Transactional(readOnly=true). Throw UsernameNotFoundException với generic message ("Authentication failed") để tránh username enumeration.')

METHOD(doc, 'public UserDetails loadUserById(String userIdRaw)',
       'Custom method - filter gọi trong JWT flow vì JWT đã có userId, không cần lookup theo username.',
       'Tránh round-trip thừa.')

METHOD(doc, 'private AuthUserDetails toUserDetails(User user)',
       'Assemble principal: roles từ DB, permissions qua resolvePermissions() (Redis cached).', None)

METHOD(doc, 'private Set<String> resolvePermissions(Role role)',
       'KEY: cache-first lookup.',
       'Cache hit → return; miss → DB query + populate cache. Treats null từ port (Redis down) như miss → vẫn serve từ DB.')

FILE_HEADER(doc, 'TokenBlacklistService.java', 'infrastructure/security/TokenBlacklistService.java')
P(doc, '@Service implements TokenBlacklistPort. Hybrid Redis + DB.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['private static final String BLACKLIST_PREFIX = "blacklist:" (key Redis)',
          'private final StringRedisTemplate redisTemplate',
          'private final TokenRevocationJpaRepository tokenRevocationRepository']:
    BULLET(doc, x)

P(doc, 'Method revoke(tokenId, userId, expiresAt, reason):', bold=True)
CODE(doc, '''@Transactional
ttl = Duration.between(now, expiresAt).getSeconds()
if (ttl > 0) {
    try {
        redisTemplate.opsForValue().set(
            BLACKLIST_PREFIX + tokenId, reason, Duration.ofSeconds(ttl));
    } catch (Exception e) {
        log warn - Redis down, fallback DB
    }
}
tokenRevocationRepository.save(TokenRevocationJpaEntity.forRevocation(...));
// DB là source of truth - luôn save kể cả Redis OK''')

P(doc, 'Method isRevoked(String tokenId):', bold=True)
CODE(doc, '''try {
    return redisTemplate.hasKey(BLACKLIST_PREFIX + tokenId);
} catch (Exception e) {
    return tokenRevocationRepository.existsById(tokenId);   // fallback DB
}''')

METHOD(doc, '@EventListener(ApplicationReadyEvent.class) public void warmUpRedisFromDb()',
       'Khởi động: load mọi revocation chưa hết hạn từ DB lên Redis.',
       'Đảm bảo Redis instance mới (sau restart, scale-out) không bỏ sót revocation cũ.')

METHOD(doc, '@Scheduled(fixedRate = 3_600_000L) public void cleanupExpiredRevocations()',
       'Mỗi giờ xóa row DB có expiresAt < now.',
       'TTL Redis tự xóa, nhưng DB không có TTL → cần job thủ công.')

# ── 21. infrastructure/persistence ──
H(doc, '21. infrastructure/persistence - JPA adapter', level=2)

P(doc, 'Cấu trúc 3 lớp:', bold=True)
CODE(doc, '''jpa/
├── entities/      ← UserJpaEntity, RoleJpaEntity... (@Entity, @Id, @Column)
└── repositories/  ← UserJpaRepository extends JpaRepository (interface Spring Data)

mapper/            ← UserJpaMapper, RefreshTokenJpaMapper (convert Domain ↔ Entity)

impl/              ← UserRepositoryImpl implements UserRepository (adapter)''')

NOTE(doc, 'TẠI SAO TÁCH JpaEntity VÀ Domain Model? '
          'JpaEntity bị "trói" với Hibernate (annotation, lazy loading, setter cho proxy). '
          'Domain model là POJO sạch, có invariants. Tách ra để Hibernate không "đụng" vào domain.')

FILE_HEADER(doc, 'UserRepositoryImpl.java', 'infrastructure/persistence/impl/UserRepositoryImpl.java')
P(doc, '@Repository implements UserRepository.', bold=True)
P(doc, 'Field DI: UserJpaRepository, RoleJpaRepository, UserJpaMapper, DomainEventPublisher.')

P(doc, 'Method save(User user):', bold=True)
CODE(doc, '''@Transactional
1. isNew = !jpaRepository.existsById(user.getId().value())
   → Manual ID aggregate: phải check để Hibernate INSERT vs UPDATE đúng
2. roleEntities = roleRepository.findByIdIn(roleIds) - resolve managed entities
3. UserJpaEntity saved = jpaRepository.save(mapper.toJpaEntity(user, roleEntities, isNew))
4. user.pullDomainEvents().forEach(eventPublisher::publish)  ← KEY
   → Drain events SAU KHI persist thành công
   → Spring listener @AFTER_COMMIT chờ commit DB rồi mới publish Kafka
5. Return mapper.toDomain(saved)''')

P(doc, 'Method find*: tất cả @Transactional(readOnly=true), delegate JpaRepository, map về domain.')

METHOD(doc, 'findAfterCursor(Instant, String, int)',
       'Two-phase keyset pagination.',
       'Phase 1: findFirstPageIds hoặc findIdsAfterCursor (chỉ lấy ID). Phase 2: findAllByIdInWithRoles. '
       'Tránh N+1 mà không LEFT JOIN FETCH trên unbounded set (gây Cartesian explosion).')

FILE_HEADER(doc, 'RefreshTokenRepositoryImpl.java', 'infrastructure/persistence/impl/RefreshTokenRepositoryImpl.java')
P(doc, '@Repository implements RefreshTokenRepository.', bold=True)
P(doc, 'Method save(RefreshToken token):', bold=True)
CODE(doc, '''@Transactional
1. UserJpaEntity userRef = userJpaRepository.findById(userId)
   → throw IllegalArgumentException nếu user không tồn tại
   → Cần managed entity cho FK user_id
2. isNew = !jpaRepository.existsById(token.getId().value())
3. saved = jpaRepository.save(mapper.toJpaEntity(token, userRef, isNew))
4. token.pullDomainEvents().forEach(eventPublisher::publish)
5. return mapper.toDomain(saved)''')

METHOD(doc, 'findByTokenValueForUpdate(TokenValue)',
       'Pessimistic lock - SELECT FOR UPDATE.',
       'Implementation gọi jpaRepository.findByTokenForUpdate() có @Lock(PESSIMISTIC_WRITE).')

METHOD(doc, 'revokeAllTokensForUser(UserId)',
       'Bulk UPDATE - revokedAt = now, revoked = true cho mọi token chưa revoke của user.',
       'Tránh load N row vào memory rồi save từng cái. Dùng JPQL @Modifying query.')

METHOD(doc, 'deleteExpiredTokens()',
       'DELETE WHERE expiresAt < now. Trả số row deleted.',
       'Dùng cho scheduled cleanup job.')

# JPA entities và mappers
H(doc, '21.1 JPA Entities', level=3)
P(doc, 'Mỗi domain aggregate có 1 entity tương ứng:')
TABLE(doc,
      headers=['Domain', 'JpaEntity', 'Bảng DB'],
      rows=[
          ['User', 'UserJpaEntity', 'users'],
          ['Role', 'RoleJpaEntity', 'roles'],
          ['Permission', 'PermissionJpaEntity', 'permissions'],
          ['RefreshToken', 'RefreshTokenJpaEntity', 'refresh_tokens'],
          ['(token blacklist)', 'TokenRevocationJpaEntity', 'token_revocations'],
      ],
      col_widths=[5, 6, 5])

P(doc, 'Đặc điểm:', bold=True)
for x in ['@Entity, @Table với tên bảng explicit',
          '@Id manually set (UUID String) - không dùng IDENTITY auto-gen',
          'JoinTable cho many-to-many (user_roles, role_permissions)',
          'Audit field createdAt/updatedAt @Column(updatable=false) cho created',
          '@Version cho optimistic locking nếu cần']:
    BULLET(doc, x)

H(doc, '21.2 JPA Mappers', level=3)
P(doc, 'Class @Component convert Domain ↔ Entity.')
P(doc, 'Method chung:')
for x in ['toJpaEntity(Domain, dependencies..., boolean isNew) - convert sang entity',
          'toDomain(JpaEntity) - rebuild aggregate qua reconstitute factory']:
    BULLET(doc, x)

# ── 22. infrastructure/cache ──
H(doc, '22. infrastructure/cache - Redis cache', level=2)

FILE_HEADER(doc, 'RedisRolePermissionCacheAdapter.java', 'infrastructure/cache/RedisRolePermissionCacheAdapter.java')
P(doc, '@Component implements RolePermissionCachePort. Lưu Redis SET tại key role_permissions:{ROLE_NAME}.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['static final String KEY_PREFIX = "role_permissions:"',
          'private static final Duration TTL = Duration.ofMinutes(5)',
          'private final StringRedisTemplate redis',
          'private final RoleRepository roleRepository',
          'private final PermissionRepository permissionRepository']:
    BULLET(doc, x)

P(doc, 'Method get(String roleName):', bold=True)
CODE(doc, '''try {
    Set<String> members = redis.opsForSet().members(KEY_PREFIX + roleName);
    return (members == null || members.isEmpty()) ? null : members;
} catch (Exception e) {
    log warn - Redis down → return null (caller fallback DB)
}''')

P(doc, 'Method put(String roleName, Set<String> permissionNames):', bold=True)
CODE(doc, '''if (permissionNames null/empty) { evict(roleName); return; }
try {
    String key = KEY_PREFIX + roleName;
    redis.delete(key);
    redis.opsForSet().add(key, permissionNames.toArray(String[]::new));
    redis.expire(key, TTL);  // 5 phút TTL
} catch (Exception) { log warn }''')

P(doc, 'Method evict(String roleName):', bold=True)
P(doc, 'redis.delete(KEY_PREFIX + roleName). Best-effort - log warn nếu Redis down.')

P(doc, 'Method evictAll():', bold=True)
CODE(doc, '''Dùng SCAN (không phải KEYS - tránh O(N) blocking trên large keyspace)
try (var cursor = redis.keyCommands().scan(
        ScanOptions.scanOptions().match(KEY_PREFIX + "*").count(100).build())) {
    cursor.forEachRemaining(b -> keys.add(new String(b)));
}
if (!keys.isEmpty()) redis.delete(keys);''')

METHOD(doc, '@EventListener(ApplicationReadyEvent.class) public void warmUp()',
       'Khởi động: load mọi role + permissions vào Redis.',
       'Request đầu tiên không phải cold-miss → DB query.')

# ── 23. infrastructure/events ──
H(doc, '23. infrastructure/events - Spring event publisher', level=2)

FILE_HEADER(doc, 'SpringDomainEventPublisher.java', 'infrastructure/events/SpringDomainEventPublisher.java')
P(doc, '@Component implements DomainEventPublisher. Delegate sang Spring ApplicationEventPublisher.', bold=True)
P(doc, 'Field: private final ApplicationEventPublisher delegate.')

METHOD(doc, 'public void publish(DomainEvent event)',
       'log.debug + delegate.publishEvent(event).',
       'KHÔNG publish vào Kafka trực tiếp - chỉ in-process. KafkaIntegrationEventPublisher lắng nghe sự kiện này và bridge sang Kafka.')

# ── 24. infrastructure/messaging ──
H(doc, '24. infrastructure/messaging - Kafka bridge', level=2)

FILE_HEADER(doc, 'KafkaIntegrationEventPublisher.java', 'infrastructure/messaging/KafkaIntegrationEventPublisher.java')
P(doc, '@Component - bridge in-process DomainEvent → Kafka.', bold=True)
P(doc, 'Field: KafkaTemplate<String, Object> kafkaTemplate, DomainEventToIntegrationEventMapper mapper.')

METHOD(doc, '@TransactionalEventListener(phase = AFTER_COMMIT) public void onDomainEvent(DomainEvent event)',
       'KEY: chỉ fire SAU khi transaction DB commit thành công.',
       'Tránh bug "DB rollback nhưng Kafka đã gửi" → service khác nhận event "ma".')

CODE(doc, '''var routed = mapper.map(event);
if (routed == null) return;  // event nội bộ, không cần bắn ra ngoài

try {
    kafkaTemplate.send(routed.topic(), routed.key(), routed.event());
} catch (Exception e) {
    log warn - best-effort, không retry
    // Nếu cần at-least-once: upgrade lên Transactional Outbox pattern
}''')

FILE_HEADER(doc, 'DomainEventToIntegrationEventMapper.java', 'infrastructure/messaging/DomainEventToIntegrationEventMapper.java')
P(doc, '@Component - convert DomainEvent → Routed(topic, key, IntegrationEvent).', bold=True)
P(doc, 'Inner record: Routed(String topic, String key, IntegrationEvent event).')

METHOD(doc, 'public Routed map(DomainEvent event)',
       'Switch pattern matching trên kiểu event. Trả null cho event nội bộ không cần bắn ra.', None)

P(doc, 'Routing rules:', bold=True)
TABLE(doc,
      headers=['DomainEvent', 'eventType (string)', 'Kafka Topic'],
      rows=[
          ['UserCreatedEvent', 'auth.user.created.v1', 'AUTH_USER_EVENTS'],
          ['EmailChangedEvent', 'auth.user.email_changed.v1', 'AUTH_USER_EVENTS'],
          ['PasswordChangedEvent', 'auth.user.password_changed.v1', 'AUTH_USER_EVENTS'],
          ['AccountStatusChangedEvent', 'auth.user.status_changed.v1', 'AUTH_USER_EVENTS'],
          ['RoleAssignedEvent', 'auth.user.role_assigned.v1', 'AUTH_USER_EVENTS'],
          ['RoleRemovedEvent', 'auth.user.role_removed.v1', 'AUTH_USER_EVENTS'],
          ['PermissionGrantedEvent', 'auth.role.permission_granted.v1', 'AUTH_USER_EVENTS'],
          ['PermissionRevokedEvent', 'auth.role.permission_revoked.v1', 'AUTH_USER_EVENTS'],
          ['UserLoggedInEvent', 'auth.user.logged_in.v1', 'AUTH_SESSION_EVENTS'],
          ['TokenCreatedEvent', 'auth.token.created.v1', 'AUTH_SESSION_EVENTS'],
          ['TokenRotatedEvent', 'auth.token.rotated.v1', 'AUTH_SESSION_EVENTS'],
          ['TokenRevokedEvent', 'auth.token.revoked.v1', 'AUTH_SESSION_EVENTS'],
      ],
      col_widths=[5, 6, 5])

P(doc, 'Private helpers:', bold=True)
METHOD(doc, 'private Routed lifecycle(DomainEvent src, String type, Map payload)',
       'Build UserLifecycleIntegrationEvent + route topic USER_EVENTS.', None)
METHOD(doc, 'private Routed session(DomainEvent src, String type, Map payload)',
       'Build SessionIntegrationEvent + route topic SESSION_EVENTS.', None)
METHOD(doc, 'private static Map<String,Object> fields(Object... keyValuePairs)',
       'Builder linh hoạt cho payload map, skip null values.',
       'Tránh boilerplate Map.of() khi có nhiều optional field.')

FILE_HEADER(doc, 'RolePermissionCacheInvalidator.java', 'infrastructure/messaging/RolePermissionCacheInvalidator.java')
P(doc, '@Component - lắng nghe in-process event PermissionGranted/Revoked → rebuild Redis cache.', bold=True)
METHOD(doc, '@TransactionalEventListener(AFTER_COMMIT) public void onPermissionGranted(event)',
       'refresh(event).', None)
METHOD(doc, '@TransactionalEventListener(AFTER_COMMIT) public void onPermissionRevoked(event)',
       'refresh(event).', None)
METHOD(doc, 'private void refresh(DomainEvent event)',
       'Reload role + permissions từ DB → cache.put().',
       'Single atomic replace cheaper hơn partial update. Try-catch + evictAll() làm fallback safe.')

NOTE(doc, 'Class này chỉ xử lý in-process invalidation. Sibling auth-service instance trên JVM khác consume Kafka event riêng (pending outbox pattern); copy Redis của họ self-heal trong 5-minute TTL.')

doc.add_page_break()

# ============================================================
# PHẦN V - TẦNG INTERFACES
# ============================================================
H(doc, 'PHẦN V - TẦNG INTERFACES', level=1)

# ── 25. AuthController ──
H(doc, '25. interfaces/rest - REST Controllers', level=2)

FILE_HEADER(doc, 'AuthController.java', 'interfaces/rest/AuthController.java')
P(doc, '@RestController @RequestMapping("/api/v1/auth"). Public auth endpoints.', bold=True)
P(doc, 'Field DI:', bold=True)
for x in ['CommandHandler<RegisterUserCommand, AuthResponseDTO> registerHandler',
          'CommandHandler<LoginCommand, AuthResponseDTO> loginHandler',
          'CommandHandler<RefreshTokenCommand, AuthResponseDTO> refreshHandler',
          'CommandHandler<LogoutCommand, Void> logoutHandler',
          'CommandHandler<ChangePasswordCommand, Void> changePasswordHandler',
          'AuthCookieWriter cookieWriter',
          'TokenProviderPort tokenProvider']:
    BULLET(doc, x)

P(doc, 'Endpoints:', bold=True)

TABLE(doc,
      headers=['Method', 'Path', 'Mô tả', 'Auth?'],
      rows=[
          ['GET', '/api/v1/auth/me', 'Thông tin principal hiện tại', 'YES'],
          ['POST', '/api/v1/auth/register', 'Đăng ký user mới', 'NO'],
          ['POST', '/api/v1/auth/login', 'Đăng nhập', 'NO'],
          ['POST', '/api/v1/auth/refresh', 'Xoay vòng refresh token', 'NO (token là auth)'],
          ['POST', '/api/v1/auth/logout', 'Đăng xuất, blacklist token', 'YES'],
          ['POST', '/api/v1/auth/change-password', 'Đổi mật khẩu', 'YES'],
      ],
      col_widths=[2, 4.5, 7, 2.5])

P(doc, 'Method me(@AuthenticationPrincipal AuthUserDetails principal):', bold=True)
P(doc, 'Đọc principal từ SecurityContext (do JwtAuthenticationFilter set). Return AuthMeResponse.from(principal).')

P(doc, 'Method register(@Valid @RequestBody RegisterRequest):', bold=True)
CODE(doc, '''AuthResponseDTO tokens = registerHandler.handle(new RegisterUserCommand(
        req.username(), req.email(), req.password(),
        req.firstName(), req.lastName()));
return cookieWriter.writeTokens(tokens);  // Set-Cookie headers''')

P(doc, 'Method login(@Valid @RequestBody LoginRequest):', bold=True)
P(doc, 'Build LoginCommand, gọi handler, writeTokens. Tương tự register.')

P(doc, 'Method refresh(HttpServletRequest, @RequestBody(required=false) RefreshBody):', bold=True)
CODE(doc, '''1. readRefreshToken(httpRequest, request)
   → Cookie-first: REFRESH_TOKEN cookie
   → Fallback: body.refreshToken() (cho mobile/CLI)
2. tokens = refreshHandler.handle(new RefreshTokenCommand(refreshToken))
3. cookieWriter.writeTokens(tokens) - Set-Cookie với token mới''')

P(doc, 'Inner record: public record RefreshBody(String refreshToken) - body cho non-cookie client.')

P(doc, 'Method logout(HttpServletRequest):', bold=True)
CODE(doc, '''1. accessToken = readAccessToken(req)   // cookie or Authorization header
2. refreshToken = readCookie(req, REFRESH_COOKIE)
3. logoutHandler.handle(new LogoutCommand(accessToken, refreshToken))
4. headers.add expired cookies (cookieWriter.expire)
5. Return 204 No Content''')

P(doc, 'Method changePassword(@AuthenticationPrincipal, @Valid @RequestBody, HttpServletRequest):', bold=True)
CODE(doc, '''1. Parse current access token để lấy jti + exp
   → try { claims = tokenProvider.parseAccessToken(rawToken); jti = claims.tokenId(); exp = claims.expiresAt(); }
   → catch: jti vẫn null - handler sẽ log warn và skip blacklist
2. changePasswordHandler.handle(new ChangePasswordCommand(
       principal.userId(), req.oldPassword(), req.newPassword(), jti, exp))
3. Set expired cookies, return 204''')

P(doc, 'Private helpers:', bold=True)
METHOD(doc, 'private static String readAccessToken(HttpServletRequest)',
       'Cookie ACCESS_TOKEN first, else Authorization Bearer header.', None)
METHOD(doc, 'private static String readRefreshToken(HttpServletRequest, RefreshBody)',
       'Cookie REFRESH_TOKEN first, else body field.', None)
METHOD(doc, 'private static String readCookie(HttpServletRequest, String name)',
       'Generic cookie reader.', None)

FILE_HEADER(doc, 'UserController.java', 'interfaces/rest/UserController.java')
P(doc, '@RestController @RequestMapping("/api/v1/users") @SecurityRequirement(name="bearerAuth"). User management.', bold=True)
TABLE(doc,
      headers=['Method', 'Path', 'Mô tả', 'PreAuthorize'],
      rows=[
          ['GET', '/api/v1/users/me', 'Profile của user hiện tại', '(authenticated)'],
          ['PATCH', '/api/v1/users/me/email', 'Đổi email user hiện tại', '(authenticated)'],
          ['GET', '/api/v1/users/{userId}', 'Lookup user khác', "hasRole('ADMIN')"],
          ['GET', '/api/v1/users', 'List users với cursor pagination', "hasRole('ADMIN')"],
          ['GET', '/api/v1/users/{userId}/has-role/{roleName}', 'Check role', '(authenticated)'],
          ['GET', '/api/v1/users/{userId}/has-permission/{permissionName}', 'Check permission', '(authenticated)'],
          ['POST', '/api/v1/users/{userId}/roles', 'Gán role', "hasRole('ADMIN')"],
          ['DELETE', '/api/v1/users/{userId}/roles/{roleName}', 'Bỏ role', "hasRole('ADMIN')"],
          ['POST', '/api/v1/users/{userId}/status/{transition}', 'Thay đổi status (LOCK/UNLOCK...)', "hasRole('ADMIN')"],
      ],
      col_widths=[2, 6.5, 5, 2.5])

# ── 26. Filters ──
H(doc, '26. interfaces/rest/filter - Filters', level=2)

FILE_HEADER(doc, 'JwtAuthenticationFilter.java', 'interfaces/rest/filter/JwtAuthenticationFilter.java')
P(doc, '@Component extends OncePerRequestFilter. Authenticate mỗi request bằng JWT.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['private static final Logger log',
          'private static final String HEADER = "Authorization"',
          'private static final String PREFIX = "Bearer "',
          'private final TokenProviderPort tokenProvider',
          'private final TokenBlacklistPort tokenBlacklist',
          'private final CustomUserDetailsService userDetailsService']:
    BULLET(doc, x)

P(doc, 'Method doFilterInternal(req, res, chain) - chạy 5 bước:', bold=True)
CODE(doc, '''1. token = extractBearer(request)
   → Cookie ACCESS_TOKEN first, else Authorization: Bearer header

2. Nếu token != null && SecurityContext chưa có authentication:
   try {
     2.1. claims = tokenProvider.parseAccessToken(token)  // verify signature + issuer + expiry
     2.2. if (tokenBlacklist.isRevoked(claims.tokenId())) {
            response.sendError(401, "Token revoked"); return;
          }
     2.3. userDetails = userDetailsService.loadUserById(claims.userId())
     2.4. if (userDetails.tokenVersion != claims.tokenVersion) {
            response.sendError(401, "Token superseded"); return;
          }  // KEY: reject JWT cũ sau khi user đổi mật khẩu
     2.5. Set SecurityContextHolder.getContext().setAuthentication(...)
   } catch (UsernameNotFoundException) {
     log debug, clear context
   } catch (Exception parseError) {
     log debug (expired/malformed token là normal trên public endpoint), clear context
   }

3. chain.doFilter(req, res)  // tiếp tục filter chain''')

NOTE(doc, 'Missing/invalid token sẽ silent fail - filter không reject. SecurityConfig dùng authorizeHttpRequests '
          'để decide endpoint nào cần authenticated; nếu protected mà chưa auth → JwtAuthenticationEntryPoint trả 401.')

FILE_HEADER(doc, 'JwtAuthenticationEntryPoint.java', 'interfaces/rest/filter/JwtAuthenticationEntryPoint.java')
P(doc, '@Component implements AuthenticationEntryPoint.', bold=True)
P(doc, 'Spring Security trigger method commence() khi request đến protected endpoint mà chưa auth.')
METHOD(doc, 'public void commence(req, res, AuthenticationException)',
       'Return JSON 401 với ErrorResponse format (match GlobalExceptionHandler schema).',
       'Mặc định Spring Security trả HTML 401 - không phù hợp JSON API. Override để clients chỉ cần parse 1 schema.')

FILE_HEADER(doc, 'RateLimitFilter.java', 'interfaces/rest/filter/RateLimitFilter.java')
P(doc, '@Component extends OncePerRequestFilter. Bucket4j-based per-IP rate limit.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['private static final String LOGIN_PATH = "/api/v1/auth/login"',
          'private static final String REGISTER_PATH = "/api/v1/auth/register"',
          'ConcurrentHashMap<String, Bucket> loginBuckets',
          'ConcurrentHashMap<String, Bucket> registerBuckets',
          'private final ObjectMapper objectMapper']:
    BULLET(doc, x)

P(doc, 'Method doFilterInternal:', bold=True)
CODE(doc, '''path = request.getRequestURI()
ip = request.getRemoteAddr()

if (path == LOGIN_PATH)    bucket = loginBuckets.computeIfAbsent(ip, k -> newBucket(5))
else if (path == REGISTER) bucket = registerBuckets.computeIfAbsent(ip, k -> newBucket(3))

if (bucket != null && !bucket.tryConsume(1)) {
    response.setStatus(429)
    write JSON {code: "AUTH-1014", message: "Too many requests"}
    return  // KHÔNG gọi chain.doFilter → request bị chặn
}
chain.doFilter(request, response)''')

METHOD(doc, 'private static Bucket newBucket(int reqPerMinute)',
       'Tạo bucket với bandwidth = N tokens/phút, refill greedy.',
       'Greedy refill = continuous refill thay vì batch end-of-period - smooth hơn.')

NOTE(doc, 'RateLimitFilter chạy TRƯỚC JwtAuthenticationFilter (SecurityConfig.addFilterBefore). '
          'Request bị throttle không bao giờ đụng tới logic auth → tiết kiệm CPU + tăng độ chịu tải.')

# ── 27. Support ──
H(doc, '27. interfaces/rest/support - Cookie writer', level=2)

FILE_HEADER(doc, 'AuthCookieWriter.java', 'interfaces/rest/support/AuthCookieWriter.java')
P(doc, '@Component - viết HttpOnly cookie cho access + refresh token.', bold=True)
P(doc, 'Field:', bold=True)
for x in ['public static final String ACCESS_COOKIE = "ACCESS_TOKEN"',
          'public static final String REFRESH_COOKIE = "REFRESH_TOKEN"',
          'private static final String REFRESH_PATH = "/api/v1/auth"',
          '@Value("${auth.cookie.secure:true}") private boolean secure',
          '@Value("${auth.cookie.same-site:Lax}") private String sameSite',
          '@Value("${jwt.refresh-expiration-days:7}") private int refreshExpirationDays']:
    BULLET(doc, x)

METHOD(doc, 'public ResponseEntity<AuthResponseDTO> writeTokens(AuthResponseDTO source)',
       'Build response 200 với Set-Cookie headers. Body bỏ token strings, chỉ giữ user profile.',
       'Token không trả body để giảm risk vô tình log/console.log capture token.')

P(doc, 'Đặc điểm cookie:', bold=True)
TABLE(doc,
      headers=['Cookie', 'Path', 'Max-Age', 'HttpOnly', 'Secure', 'SameSite'],
      rows=[
          ['ACCESS_TOKEN', '/', 'JWT TTL (~15min)', 'true', '${auth.cookie.secure}', '${auth.cookie.same-site}'],
          ['REFRESH_TOKEN', '/api/v1/auth', '7 days', 'true', '...', '...'],
      ],
      col_widths=[3.5, 2, 3, 2, 3, 2.5])

NOTE(doc, 'REFRESH path="/api/v1/auth" để cookie chỉ gửi với auth-specific request - giảm exposure trên mọi request thường.')

METHOD(doc, 'public void expire(HttpHeaders headers)',
       'Set Set-Cookie với maxAge=0 cho cả 2 cookie → browser xóa.',
       'Dùng cho logout endpoint.')

# ── 28. REST DTOs ──
H(doc, '28. interfaces/rest/dto - Request DTOs', level=2)
TABLE(doc,
      headers=['DTO', 'Field', 'Dùng cho'],
      rows=[
          ['RegisterRequest', 'username, email, password, firstName, lastName (validation annotations)', '/register'],
          ['LoginRequest', 'usernameOrEmail, password', '/login'],
          ['ChangePasswordRequest', 'oldPassword, newPassword', '/change-password'],
          ['UpdateEmailRequest', 'newEmail (@Email validation)', 'PATCH /users/me/email'],
          ['RoleAssignmentRequest', 'roleName', 'POST /users/{id}/roles'],
          ['AuthMeResponse', 'static from(AuthUserDetails) → projection', 'GET /auth/me'],
      ],
      col_widths=[5, 8, 3])

NOTE(doc, 'DTO Request dùng @Valid + Bean Validation annotations (@NotBlank, @Size, @Email...). '
          'GlobalExceptionHandler.handleValidation map MethodArgumentNotValidException → 400 với field errors.')

# ── 29. gRPC ──
H(doc, '29. interfaces/grpc - gRPC service', level=2)

FILE_HEADER(doc, 'AuthGrpcService.java', 'interfaces/grpc/AuthGrpcService.java')
P(doc, '@GrpcService extends AuthServiceGrpc.AuthServiceImplBase. Service-to-service auth API.', bold=True)
P(doc, 'Field: TokenProviderPort, TokenBlacklistPort, 3 QueryHandler.')

P(doc, 'Method:', bold=True)
METHOD(doc, 'public void verifyToken(VerifyTokenRequest req, StreamObserver<VerifyTokenResponse> observer)',
       'Parse + verify access token. Check blacklist.',
       'Invalid/expired/revoked → valid=false (KHÔNG gRPC error). Caller chỉ cần check 1 Boolean.')

METHOD(doc, 'public void checkRole(CheckRoleRequest, StreamObserver)',
       'Delegate sang CheckRoleHandler. Trả CheckRoleResponse.hasRole.', None)
METHOD(doc, 'public void checkPermission(CheckPermissionRequest, StreamObserver)',
       'Delegate sang CheckPermissionHandler.', None)
METHOD(doc, 'public void getUser(GetUserRequest, StreamObserver)',
       'Delegate sang GetUserByIdHandler. UserNotFoundException → found=false.',
       'Sidestep gRPC NOT_FOUND status bikeshedding cho lookup operations.')
METHOD(doc, 'private static String nullSafe(String s)',
       'Protobuf string không null được - normalize null → "".', None)

doc.add_page_break()

# ============================================================
# PHẦN VI - CONFIG & EXCEPTION
# ============================================================
H(doc, 'PHẦN VI - CONFIG & EXCEPTION', level=1)

H(doc, '30. config/* - Spring configurations', level=2)

FILE_HEADER(doc, 'SecurityConfig.java', 'config/SecurityConfig.java')
P(doc, '@Configuration @EnableWebSecurity @EnableMethodSecurity. Spring Security chain.', bold=True)
P(doc, 'Field DI: CustomUserDetailsService, JwtAuthenticationFilter, JwtAuthenticationEntryPoint, RateLimitFilter.')

P(doc, 'Bean:', bold=True)

METHOD(doc, '@Bean @Order(1) @Profile("!prod") SecurityFilterChain devEndpoints(HttpSecurity)',
       'Dev-only chain permit Swagger UI, OpenAPI, H2 console.',
       '@Order(1) ưu tiên match trước main chain. Frame-options SAMEORIGIN cho H2 console iframe.')

METHOD(doc, '@Bean @Order(2) SecurityFilterChain securityFilterChain(HttpSecurity)',
       'Main stateless JWT chain.', None)

CODE(doc, '''http
.csrf(disable)
.sessionManagement(STATELESS)  // không tạo session
.exceptionHandling(authenticationEntryPoint)
.authorizeHttpRequests(
    .dispatcherTypeMatchers(ERROR).permitAll()  // /error path
    .requestMatchers(POST, "/login","/register","/refresh").permitAll()
    .requestMatchers("/actuator/health","/actuator/info").permitAll()
    .requestMatchers("/actuator/**").hasAuthority("ROLE_ADMIN")
    .anyRequest().authenticated())
.authenticationProvider(daoProvider)
.addFilterBefore(rateLimitFilter,    UsernamePasswordAuthenticationFilter.class)
.addFilterBefore(jwtAuthenticationFilter, UsernamePasswordAuthenticationFilter.class)''')

METHOD(doc, '@Bean AuthenticationProvider authenticationProvider()',
       'DaoAuthenticationProvider wired to CustomUserDetailsService + BCrypt encoder.',
       'AuthenticationManager khi cần programmatic auth (test, future flows).')

METHOD(doc, '@Bean AuthenticationManager authenticationManager(AuthenticationConfiguration)',
       'Expose standard AuthenticationManager.', None)

METHOD(doc, '@Bean PasswordEncoder passwordEncoder()',
       'BCryptPasswordEncoder(12) - cost 12 theo OWASP.',
       'Single source of truth - BCryptPasswordEncoderAdapter delegate sang bean này.')

FILE_HEADER(doc, 'KafkaConfig.java', 'config/KafkaConfig.java')
P(doc, '@Configuration @EnableKafka. Declare topics.', bold=True)
METHOD(doc, '@Bean NewTopic userEventsTopic()',
       'AUTH_USER_EVENTS topic, 3 partitions, 1 replica.', None)
METHOD(doc, '@Bean NewTopic sessionEventsTopic()',
       'AUTH_SESSION_EVENTS topic, 3 partitions, 1 replica.', None)
NOTE(doc, 'KafkaAdmin tự pickup bean và ensure topic tồn tại (idempotent). @EnableKafka cho phép future in-process @KafkaListener.')

FILE_HEADER(doc, 'DomainServiceConfig.java', 'config/DomainServiceConfig.java')
P(doc, '@Configuration @EnableConfigurationProperties(JwtProperties.class). Wire domain services không có Spring annotation.', bold=True)
METHOD(doc, '@Bean AuthenticationDomainService authenticationDomainService()',
       'new AuthenticationDomainService() - đăng ký singleton.',
       'Domain class là POJO không có @Component để giữ tầng domain sạch framework. Config này là điểm DUY NHẤT Spring biết về domain services.')
METHOD(doc, '@Bean TokenDomainService tokenDomainService()',
       'new TokenDomainService().', None)

FILE_HEADER(doc, 'AsyncConfig.java', 'config/AsyncConfig.java')
P(doc, '@Configuration @EnableAsync. Cấu hình virtual thread cho @Async.', bold=True)
METHOD(doc, '@Bean TaskExecutor taskExecutor()',
       'Adapter wrap virtual-thread-per-task executor.',
       'Java 21 virtual threads: zero pool-sizing, low overhead. spring.threads.virtual.enabled=true handle Tomcat; bean này handle @Async dispatch.')

FILE_HEADER(doc, 'CorsConfig.java', 'config/CorsConfig.java')
P(doc, 'CORS configuration cho cross-origin request từ frontend.')

FILE_HEADER(doc, 'OpenApiConfig.java', 'config/OpenApiConfig.java')
P(doc, 'SpringDoc OpenAPI - generate Swagger UI tại /swagger-ui.html.')

FILE_HEADER(doc, 'DataSeeder.java', 'config/DataSeeder.java')
P(doc, '@Component CommandLineRunner. Seed dữ liệu khởi tạo (roles, permissions, admin user) vào DB lần đầu chạy.')

H(doc, '31. exceptions/* - Global exception handler', level=2)

FILE_HEADER(doc, 'GlobalExceptionHandler.java', 'exceptions/GlobalExceptionHandler.java')
P(doc, '@RestControllerAdvice. Central mapping exception → HTTP response.', bold=True)

P(doc, 'Method handleDomain(DomainException ex):', bold=True)
CODE(doc, '''Pattern matching switch trên kiểu domain exception:
case UserNotFoundException       → 404 NOT_FOUND
case UserAlreadyExistsException  → 409 CONFLICT
case InvalidCredentialsException → 401 UNAUTHORIZED
case TokenExpiredException       → 401
case TokenRevokedException       → 401
case TokenReuseDetectedException → 401
case TokenOwnershipException     → 403 FORBIDDEN
case AccountNotUsableException   → 403
default                          → 400 BAD_REQUEST

Body chứa ex.code() - stable error code cho client.''')

P(doc, 'Other handlers:', bold=True)
METHOD(doc, 'handleValidation(MethodArgumentNotValidException)',
       '@Valid fail → 400 với list field errors.', None)
METHOD(doc, 'handleBadCredentials(BadCredentialsException)',
       'Spring Security wrong credentials → 401.', None)
METHOD(doc, 'handleAuth(AuthenticationException)',
       'Generic auth fail → 401.', None)
METHOD(doc, 'handleAccessDenied(AccessDeniedException)',
       '@PreAuthorize fail → 403.', None)
METHOD(doc, 'handleJwtExpired(ExpiredJwtException)',
       'jjwt expired → 401 với code TOKEN_EXPIRED.', None)
METHOD(doc, 'handleJwt(SignatureException, MalformedJwtException)',
       'JWT invalid → 401 với TOKEN_INVALID.', None)
METHOD(doc, 'handleNotReadable(HttpMessageNotReadableException)',
       'JSON malformed → 400 thay vì 500.', None)
METHOD(doc, 'handleAll(Exception)',
       'Fallback → 500 với generic message.',
       'log.error đầy đủ stack trace nhưng không expose ra response.')

FILE_HEADER(doc, 'ErrorResponse.java', 'exceptions/ErrorResponse.java')
P(doc, 'POJO với @Builder. Field: timestamp, status, error (code), message, path, validationErrors (optional).')

doc.add_page_break()

# ============================================================
# PHẦN VII - LUỒNG NGHIỆP VỤ
# ============================================================
H(doc, 'PHẦN VII - LUỒNG NGHIỆP VỤ ĐẦY ĐỦ', level=1)

H(doc, '32. Register flow', level=2)
CODE(doc, '''┌─ HTTP POST /api/v1/auth/register
│  Body: {username, email, password, firstName, lastName}
│  Headers: Content-Type: application/json
│
├→ RateLimitFilter check IP (3 req/phút)
│   Quá → 429 + body {"code":"AUTH-1014","message":"Too many requests"}
│
├→ JwtAuthenticationFilter: no token → skip (endpoint public)
│
├→ AuthController.register(@Valid @RequestBody RegisterRequest)
│   @Valid → MethodArgumentNotValidException nếu fail
│
├→ RegisterUserHandler.handle(RegisterUserCommand)
│   @Transactional bắt đầu
│   │
│   ├→ Username.of(cmd.username())  // VO validate
│   ├→ Email.of(cmd.email())
│   ├→ userRepository.existsByUsername(username) → false
│   ├→ userRepository.existsByEmail(email) → false
│   │
│   ├→ User.register(username, email, Password.createRaw(rawPwd), name, encoder)
│   │   ├→ UserId.generate() - UUID
│   │   ├→ encoder.encode(rawPwd) - BCrypt cost 12 (~250ms)
│   │   ├→ password = Password.createEncoded(hash)
│   │   ├→ accountStatus = AccountStatus.createActive()
│   │   ├→ tokenVersion = 1
│   │   ├→ registerEvent(new UserCreatedEvent(userId, username, email))
│   │   └→ return User
│   │
│   ├→ roleRepository.findByName("ROLE_USER")
│   │   .ifPresent(role -> user.assignRole(role.getId()))
│   │       └→ registerEvent(new RoleAssignedEvent(userId, roleId))
│   │
│   ├→ userRepository.save(user)  // UserRepositoryImpl
│   │   ├→ isNew = !jpaRepo.existsById(...) → true
│   │   ├→ roleEntities = roleRepo.findByIdIn(roleIds)
│   │   ├→ saved = jpaRepo.save(mapper.toJpaEntity(...))
│   │   │   ↓ Hibernate INSERT users, INSERT user_roles
│   │   └→ user.pullDomainEvents().forEach(eventPublisher::publish)
│   │       ↓ Spring ApplicationEventPublisher (in-process, sync)
│   │       ↓ UserCreatedEvent + RoleAssignedEvent queued cho TransactionalEventListener
│   │
│   ├→ roleRepository.findByIdIn(savedRoles) → List<Role>
│   ├→ tokenProvider.issueAccessToken(user, roleNames)
│   │   ├→ tokenId = UUID.randomUUID()
│   │   ├→ Jwts.builder().claims({jti, userId, email, tokenVersion, roles})
│   │   │       .subject(username).issuer(...).expiration(now+15min).signWith(key, HS256)
│   │   └→ return IssuedAccessToken(token, jti, exp, ttl)
│   │
│   ├→ tokenDomainService.issueForUser(user, 7)
│   │   ├→ user.ensureAuthenticatable() - kiểm tra status
│   │   ├→ RefreshToken.create(userId, 7)
│   │   │   ├→ TokenId.generate(), TokenValue.generate(), TokenFamily.generate()
│   │   │   ├→ GenerationNumber.root() = 0
│   │   │   ├→ TokenExpiry.fromDaysFromNow(7)
│   │   │   └→ registerEvent(new TokenCreatedEvent(...))
│   │   └→ return RefreshToken
│   │
│   ├→ refreshTokenRepository.save(refresh)
│   │   ↓ Hibernate INSERT refresh_tokens
│   │   └→ refresh.pullDomainEvents() → publish TokenCreatedEvent
│   │
│   └→ return AuthResponseDTO.bearer(jwt, refreshValue, ttl, userDto)
│
├→ AuthCookieWriter.writeTokens(tokens)
│   ├→ ResponseCookie ACCESS_TOKEN: HttpOnly, Secure, SameSite=Lax, Path=/, Max-Age=15min
│   ├→ ResponseCookie REFRESH_TOKEN: HttpOnly, Secure, Path=/api/v1/auth, Max-Age=7days
│   ├→ Body: AuthResponseDTO(null, null, "Bearer", ttl, userDto)  ← token strings xóa
│   └→ ResponseEntity 200 OK
│
├→ Transaction commits → COMMIT users + user_roles + refresh_tokens
│
├→ @TransactionalEventListener(AFTER_COMMIT) fire:
│   ├→ KafkaIntegrationEventPublisher.onDomainEvent(UserCreatedEvent)
│   │   ├→ mapper.map(event) → Routed("auth.user.events.v1", userId, UserLifecycleIntegrationEvent)
│   │   └→ kafkaTemplate.send(...)  → Kafka topic
│   ├→ Tương tự cho RoleAssignedEvent → "auth.user.role_assigned.v1"
│   └→ TokenCreatedEvent → "auth.session.events.v1" / "auth.token.created.v1"
│
└→ Service khác (notification, billing) consume Kafka → side effects (gửi email welcome, init account...)''')

H(doc, '33. Login flow', level=2)
CODE(doc, '''┌─ POST /api/v1/auth/login {usernameOrEmail, password}
├→ RateLimitFilter (5 req/phút)
├→ LoginHandler.handle(LoginCommand) @Transactional
│   │
│   ├→ lookup(usernameOrEmail)
│   │   ├→ contains "@" → userRepository.findByEmail(Email.of(...))
│   │   ├→ else → userRepository.findByUsername(Username.of(...))
│   │   └→ Empty? throw InvalidCredentialsException
│   │       (SAME exception cho cả "không tồn tại" và "password sai" - timing-safe)
│   │
│   ├→ user.authenticate(Password.createRaw(rawPwd), encoder)
│   │   ├→ ensureAuthenticatable() - throw AccountNotUsableException nếu lock/disable/expired
│   │   ├→ encoder.matches(rawPwd, password.hash) - BCrypt verify ~250ms (constant-time)
│   │   ├→ if NOT matches → throw InvalidCredentialsException
│   │   └→ if matches → recordLogin()
│   │       ├→ accountStatus = accountStatus.withLastLogin(now)
│   │       └→ registerEvent(new UserLoggedInEvent(userId, username))
│   │
│   ├→ userRepository.save(user) - persist lastLogin
│   │   └→ publish UserLoggedInEvent
│   │
│   ├→ roleRepository.findByIdIn(roles)
│   ├→ tokenProvider.issueAccessToken(user, roleNames) - JWT mới
│   ├→ tokenDomainService.issueForUser(user, 7) - RefreshToken root mới (family mới)
│   ├→ refreshTokenRepository.save(refresh)
│   └→ return AuthResponseDTO.bearer(...)
│
├→ cookieWriter.writeTokens(tokens)
├→ Transaction commits
└→ Kafka: "auth.user.logged_in.v1" + "auth.token.created.v1"''')

H(doc, '34. Refresh token flow', level=2)
CODE(doc, '''┌─ POST /api/v1/auth/refresh
│  Cookie: REFRESH_TOKEN=<value>  (or body.refreshToken cho mobile)
│
├→ AuthController.refresh()
│   ├→ readRefreshToken(httpRequest, body) - cookie first, fallback body
│
├→ RefreshTokenHandler.handle(RefreshTokenCommand) @Transactional
│   │
│   ├→ refreshTokenRepository.findByTokenValueForUpdate(value)  ← SELECT ... FOR UPDATE
│   │   → Pessimistic lock trên row token
│   │   → 2 request đồng thời cùng token: request 2 chờ request 1 xong
│   │
│   ├→ presented = result.orElseThrow(TokenRevokedException::new)
│   │
│   ├→ tokenDomainService.rotate(presented, 7, repository)
│   │   ├→ if (presented.isReuseAttempt())  // revoked && reason != NORMAL
│   │   │   ├→ revokeFamily(presented.family, repo)
│   │   │   │   ├→ findByFamily → mọi token live trong family
│   │   │   │   └→ for each: revoke(FAMILY_REVOKED), save, raise TokenRevokedEvent
│   │   │   └→ throw TokenReuseDetectedException
│   │   │       → response 401, user phải login lại
│   │   │
│   │   ├→ presented.verifyValidity()
│   │   │   ├→ throw TokenRevokedException nếu revoked
│   │   │   └→ throw TokenExpiredException nếu expired
│   │   │
│   │   ├→ presented.revoke(RevokedReason.NORMAL)
│   │   │   └→ registerEvent(new TokenRevokedEvent(...))
│   │   ├→ repository.save(presented)
│   │   │
│   │   └→ return RefreshToken.rotate(presented, 7)
│   │       ├→ id mới, value mới, family GIỮ NGUYÊN, generation+1, expiry mới
│   │       └→ registerEvent(new TokenRotatedEvent(userId, family, oldId, newId, gen))
│   │
│   ├→ refreshTokenRepository.save(rotated) - publish TokenRotatedEvent
│   ├→ userRepository.findById(presented.userId) → user
│   ├→ tokenProvider.issueAccessToken(user, roles) - JWT mới
│   └→ return AuthResponseDTO
│
├→ cookieWriter.writeTokens(...) - cookie mới
├→ Transaction commits
└→ Kafka: "auth.token.revoked.v1" + "auth.token.rotated.v1"

KEY INSIGHT về Pessimistic Lock:
2 request refresh đồng thời cùng token:
- Request 1: lock row, revoke old, issue new → commit
- Request 2: chờ lock, lock đc, đọc lại → thấy old.revoked=true, reason=NORMAL
            → isReuseAttempt() = false → verifyValidity() throw TokenRevokedException
            → 401, user reload page''')

H(doc, '35. Logout flow', level=2)
CODE(doc, '''┌─ POST /api/v1/auth/logout (auth required)
├→ JwtAuthenticationFilter authenticate request
├→ AuthController.logout(HttpServletRequest)
│   ├→ accessToken = readAccessToken(req)
│   ├→ refreshToken = readCookie(req, REFRESH_COOKIE)
│
├→ LogoutHandler.handle(LogoutCommand) @Transactional
│   ├→ if refreshToken present:
│   │   ├→ repo.findByTokenValue(value).ifPresent(t -> {
│   │   │       t.revoke(USER_INITIATED) → raise TokenRevokedEvent
│   │   │       repo.save(t)
│   │   │   })
│   │   (Idempotent - missing = user đã logout)
│   │
│   ├→ if accessToken present: try {
│   │       claims = tokenProvider.parseAccessToken(accessToken)
│   │       tokenBlacklist.revoke(jti, userId, exp, "LOGOUT")
│   │           ├→ Redis SET blacklist:{jti} = "LOGOUT" EX <ttl>
│   │           └→ DB INSERT token_revocations
│   │   } catch (Exception) {
│   │       log warn - expired/malformed → bỏ qua
│   │       (KHÔNG rollback refresh revoke đã commit)
│   │   }
│   └→ return null
│
├→ cookieWriter.expire(headers) - Set-Cookie với maxAge=0
├→ ResponseEntity 204 No Content
├→ Transaction commits
└→ Kafka: "auth.token.revoked.v1"''')

H(doc, '36. Change password flow', level=2)
CODE(doc, '''┌─ POST /api/v1/auth/change-password (auth required)
│  Body: {oldPassword, newPassword}
│
├→ AuthController.changePassword(@AuthenticationPrincipal, request, httpReq)
│   ├→ rawToken = readAccessToken(httpReq)
│   ├→ try {
│   │     claims = tokenProvider.parseAccessToken(rawToken)
│   │     jti = claims.tokenId(); exp = claims.expiresAt()
│   │   } catch { jti=null; log debug }
│
├→ ChangePasswordHandler.handle(cmd) @Transactional
│   ├→ user = userRepository.findById(userId).orElseThrow(UserNotFoundException::new)
│   ├→ user.changePassword(Password.createRaw(old), Password.createRaw(new), encoder)
│   │   ├→ encoder.matches(old, password.hash) - verify old
│   │   ├→ if NOT matches → throw InvalidCredentialsException
│   │   ├→ password = Password.createEncoded(encoder.encode(new))
│   │   ├→ incrementTokenVersion()  ← KEY: bump
│   │   └→ registerEvent(new PasswordChangedEvent(userId, username))
│   │
│   ├→ userRepository.save(user) - tokenVersion mới persist
│   │
│   ├→ refreshTokenRepository.revokeAllTokensForUser(userId)
│   │   → SQL: UPDATE refresh_tokens SET revoked=true, revoked_at=now
│   │           WHERE user_id=? AND revoked=false
│   │   (Bulk - không load N row vào memory)
│   │
│   ├→ if jti!=null && exp!=null:
│   │   tokenBlacklist.revoke(jti, userId, exp, "PASSWORD_CHANGED")
│   │   → Đóng race window: trong vài ms giữa bump version và JWT filter read DB,
│   │     blacklist sẽ reject token cũ ngay.
│   │
│   └→ return null
│
├→ Set expired cookies, 204 No Content
└→ Kafka: "auth.user.password_changed.v1"

HỆ QUẢ: Mọi JWT cũ của user (kể cả trên các device khác) đều invalid:
- JwtAuthenticationFilter so JWT.tokenVersion (cũ) với user.tokenVersion (đã bump) → 401
- Refresh token đã revoke → /refresh fail → user phải re-login''')

H(doc, '37. Role assignment flow', level=2)
CODE(doc, '''┌─ POST /api/v1/users/{userId}/roles (ADMIN only)
│  Body: {roleName: "ROLE_MERCHANT"}
│
├→ @PreAuthorize("hasRole('ADMIN')") check
│   → AccessDeniedException nếu không admin → 403
│
├→ AssignRoleHandler.handle(AssignRoleCommand) @Transactional
│   ├→ user = userRepository.findById(userId).orElseThrow
│   ├→ role = roleRepository.findByName(roleName).orElseThrow
│   ├→ user.assignRole(role.getId())
│   │   ├→ if (roles.add(roleId))  // Set.add trả false nếu đã có
│   │   │   ├→ updatedAt = now
│   │   │   └→ registerEvent(new RoleAssignedEvent(userId, roleId))
│   ├→ userRepository.save(user) - publish event
│   └→ return null
│
├→ ResponseEntity 204
└→ Kafka: "auth.user.role_assigned.v1"
   → Service consumer khác cập nhật cache local role của user

LƯU Ý: Role assignment KHÔNG bump tokenVersion. Token cũ của user vẫn valid và
chứa list role cũ. User cần re-login hoặc refresh để thấy role mới.
(Trade-off: nếu bump version mỗi lần đổi role → user bị kick mọi lần admin sửa role)''')

H(doc, '38. gRPC verify token flow', level=2)
CODE(doc, '''Service khác (vd order-service) cần verify JWT:

┌─ gRPC client.verifyToken(VerifyTokenRequest{accessToken: "ey..."})
│  → qua wire (Protobuf, HTTP/2)
│
├→ AuthGrpcService.verifyToken(req, observer)
│   try {
│     claims = tokenProvider.parseAccessToken(req.getAccessToken())
│       → verify signature HS256, verify issuer="hieu.com", check expiry
│     if (tokenBlacklist.isRevoked(claims.tokenId()))
│       → Redis HGET blacklist:{jti}, fallback DB
│       → reply.setValid(false)
│     else
│       → reply.setValid(true)
│              .setUserId(claims.userId())
│              .setUsername(claims.username())
│              .setTokenVersion(claims.tokenVersion())
│              .addAllRoles(claims.roles())
│   } catch (Exception) {
│     reply.setValid(false)  // KHÔNG throw gRPC error
│   }
│   observer.onNext(reply.build())
│   observer.onCompleted()
│
└→ Caller nhận VerifyTokenResponse, branch theo .getValid()

LỢI ÍCH:
- Service consumer chỉ cần parse 1 Boolean - không phân biệt lý do invalid
- gRPC latency ~1-2ms (vs HTTP/JSON ~10-20ms)
- Type-safe qua Protobuf - lỗi field sai bắt ở compile time''')

doc.add_page_break()

# ============================================================
# PHẦN VIII - QUYẾT ĐỊNH THIẾT KẾ & FAQ
# ============================================================
H(doc, 'PHẦN VIII - QUYẾT ĐỊNH THIẾT KẾ & FAQ', level=1)

H(doc, '39. 25 quyết định thiết kế quan trọng', level=2)

decisions = [
    ('Vì sao Hexagonal Architecture (4 tầng)?',
     'Tách nghiệp vụ thuần ra khỏi framework. Đổi DB, đổi message broker, đổi web framework mà nghiệp vụ KHÔNG đổi 1 dòng. Test domain bằng JUnit thuần - không cần startup Spring (~ms thay vì ~seconds).'),

    ('Vì sao tách Domain Model và JpaEntity?',
     'JpaEntity bị "trói" với Hibernate (annotations, lazy loading, setter cho proxy). Domain là POJO sạch với invariants. Tách ra để Hibernate không "đụng" vào quy tắc nghiệp vụ.'),

    ('Vì sao Value Object cho Email, Username, Password, UserId...?',
     'Ép buộc invariant tại điểm tạo. Type-safe ở compile time - hàm register(Username, Email, Password) không thể bị truyền nhầm String/String/String. Password phân biệt raw/encoded ngăn lưu plaintext.'),

    ('Vì sao DomainEvent ở domain/ nhưng IntegrationEvent ở application/?',
     'DomainEvent là quy tắc nghiệp vụ ("đăng ký xong cần báo"). IntegrationEvent là contract giao tiếp với service khác - không phải nghiệp vụ thuần. Tách giúp domain không phụ thuộc Kafka schema.'),

    ('Vì sao publisher event qua port (interface ở domain)?',
     'Dependency Inversion. Domain ra "yêu cầu công việc" (publish event), infrastructure "nhận làm". Mai mốt đổi Spring → Kafka direct → RabbitMQ → chỉ thay adapter, domain không đổi.'),

    ('Vì sao AFTER_COMMIT thay vì publish ngay?',
     'Tránh bug "DB rollback nhưng Kafka đã gửi" → service khác nhận event "ma". Listener fire chỉ khi transaction commit thành công.'),

    ('Vì sao có tokenVersion trong JWT?',
     'JWT stateless không thu hồi được tới khi hết hạn. tokenVersion là "núm vặn" - bump trong DB → mọi JWT cũ tự fail (filter so version). Hiệu lực tức thì khi change password / admin revoke.'),

    ('Vì sao Refresh Token Rotation + Family Revocation?',
     'Phát hiện token theft. Token cũ bị dùng lại → revoke cả family → cả user và attacker đều bị đẩy ra → user login lại trên kênh sạch. Industry best practice (OAuth 2.0 BCP).'),

    ('Vì sao Pessimistic Lock (SELECT FOR UPDATE) khi refresh?',
     'Race condition: 2 request refresh đồng thời cùng token → có thể cả 2 cùng pass = 2 access token. Lock serialize: request 2 thấy token đã revoked → fail.'),

    ('Vì sao Cookie HttpOnly thay vì localStorage?',
     'localStorage đọc được từ JavaScript → XSS đánh cắp token. HttpOnly cookie JS không đọc được. Vẫn fallback Bearer header cho mobile/CLI.'),

    ('Vì sao BCrypt cost 12?',
     'OWASP khuyến nghị tối thiểu 12 cho server modern. ~250ms/hash → user không thấy chậm; attacker brute force trên GPU không khả thi.'),

    ('Vì sao Hybrid Redis + DB cho TokenBlacklist?',
     'Redis nhanh (~1ms, mọi request đều check). DB là source of truth - Redis chết không mất revocation. WarmUp on startup đảm bảo Redis mới không bỏ sót.'),

    ('Vì sao opsForSet cho role-permission cache?',
     'Membership check O(1). Tránh separator-collision bug nếu lưu String "p1,p2,p3". Add/remove single permission cũng dễ.'),

    ('Vì sao TTL 5 phút cho cache?',
     'Self-healing: nếu Kafka consumer offline → invalidation event miss → cache stale tối đa 5 phút. Balance giữa "real-time" và "tolerance" với failure.'),

    ('Vì sao CQRS (Command + Handler) thay vì 1 GodService?',
     '1 handler = 1 use case → đơn nhiệm, dễ test. Controller mỏng. Tránh GodService 5000 dòng. Dễ thêm cross-cutting (logging, metrics, retry) qua decorator.'),

    ('Vì sao record cho Command/Query/DTO?',
     'Immutable tự động, equals/hashCode/toString tự sinh. Boilerplate giảm 90%. Compact constructor cho validation.'),

    ('Vì sao private static final cho hằng số trong handler?',
     'Hằng số class-level, không phụ thuộc instance. 1 ô nhớ duy nhất, không gán lại được. Thread-safe tuyệt đối. Convention SCREAMING_SNAKE_CASE.'),

    ('Vì sao private final cho dependency injection?',
     'Thread-safe (JMM đảm bảo safely published). Fail-fast (constructor injection - thiếu DI → crash ngay startup). Tương thích Lombok @RequiredArgsConstructor.'),

    ('Vì sao Map<String,Object> trong IntegrationEvent payload?',
     'Flexibility - thêm event type mới không cần class mới. Trade-off: mất type safety. Hợp lý vì IntegrationEvent đổi nhanh và consumer thường viết bằng ngôn ngữ khác.'),

    ('Vì sao Issuer Validation trong JWT?',
     'Nhiều service share cùng HS256 secret → có thể giả mạo. requireIssuer("auth-service") chặn token mint bởi service khác.'),

    ('Vì sao RateLimit chạy TRƯỚC JWT filter?',
     'Brute force không nên tốn CPU parse JWT. Throttle ngay → request bị chặn không đụng auth logic → tiết kiệm resource, tăng độ chịu tải.'),

    ('Vì sao gRPC return found=false thay vì NOT_FOUND status?',
     'Tránh gRPC status-code bikeshedding cho lookup. Caller chỉ check 1 Boolean. Errors thực sự (network, internal) vẫn dùng gRPC status.'),

    ('Vì sao SessionIntegrationEvent VÀ UserLifecycleIntegrationEvent cùng field?',
     'Routing tới topic khác → consumer khác (audit-service vs notification-service). Type-safe method overloading: publisher dispatch dựa trên kiểu, không nhầm topic.'),

    ('Vì sao InvalidCredentialsException cho cả "không tồn tại" và "password sai"?',
     'Timing-safe + tránh username enumeration. Attacker không phân biệt được "username chưa register" vs "password wrong" qua response code hay timing.'),

    ('Vì sao Two-phase keyset pagination?',
     'Phase 1 chỉ lấy ID (cheap). Phase 2 LEFT JOIN FETCH chỉ trên IDs có sẵn (bounded). Tránh Cartesian explosion + N+1.'),
]

for i, (q, a) in enumerate(decisions, 1):
    P(doc, f'{i}. {q}', bold=True, size=11)
    P(doc, '→ ' + a, size=10)
    doc.add_paragraph()

H(doc, '40. Cheat sheet & FAQ', level=2)

H(doc, '40.1 Cheatsheet: Modifier nào cho field?', level=3)
TABLE(doc,
      headers=['Tình huống', 'Modifier', 'Ví dụ trong code'],
      rows=[
          ['Hằng số class (prefix, timeout, regex)', 'private static final', 'BLACKLIST_PREFIX, TTL, ALGORITHM'],
          ['Hằng số public', 'public static final', 'AUTH_USER_EVENTS, ACCESS_COOKIE'],
          ['Logger', 'private static final', 'Logger log = LoggerFactory.getLogger(...)'],
          ['Pre-compiled Pattern', 'private static final', 'EMAIL_PATTERN, USERNAME_PATTERN'],
          ['Dependency injection', 'private final', 'redisTemplate, userRepository'],
          ['State thay đổi (single thread)', 'private (no final)', 'retryCount, failureCount'],
          ['State multi-thread atomic', 'private volatile / AtomicXxx', 'AtomicLong counter'],
      ],
      col_widths=[6, 4, 6])

H(doc, '40.2 Cheatsheet: Spring annotation nào cho class?', level=3)
TABLE(doc,
      headers=['Class là gì?', 'Annotation', 'Ví dụ'],
      rows=[
          ['Use case handler, domain service Spring wrapper', '@Service', 'LoginHandler, TokenBlacklistService'],
          ['DB access', '@Repository', 'UserRepositoryImpl, RefreshTokenRepositoryImpl'],
          ['HTTP endpoint', '@RestController', 'AuthController, UserController'],
          ['Bean configuration', '@Configuration', 'SecurityConfig, KafkaConfig'],
          ['Adapter generic (Redis, Kafka, event publisher)', '@Component', 'JwtTokenProvider, RedisRolePermissionCacheAdapter'],
          ['Filter HTTP', '@Component', 'JwtAuthenticationFilter, RateLimitFilter'],
          ['gRPC service', '@GrpcService', 'AuthGrpcService'],
      ],
      col_widths=[6, 3, 7])

H(doc, '40.3 Cheatsheet: Tầng nào làm gì?', level=3)
TABLE(doc,
      headers=['Tầng', 'Trách nhiệm', 'Có Spring annotation?'],
      rows=[
          ['domain/models/', 'Aggregate Root, Value Object, business methods', 'KHÔNG (POJO)'],
          ['domain/services/', 'Cross-aggregate rules (TokenDomainService...)', 'KHÔNG (Spring wrap qua @Bean)'],
          ['domain/repositories/', 'Interface (port)', 'KHÔNG'],
          ['domain/events/', 'DomainEvent base, Publisher interface', 'KHÔNG'],
          ['application/common/', 'Command/Query/Handler interfaces, CursorCodec', 'KHÔNG'],
          ['application/command/', 'Command records', 'KHÔNG'],
          ['application/handler/', 'Orchestrate use case', '@Service'],
          ['application/dto/', 'Read models cho output', 'KHÔNG'],
          ['application/mapper/', 'Domain → DTO', '@Component'],
          ['application/events/', 'IntegrationEvent contract', 'KHÔNG'],
          ['application/port/', 'Outbound port interfaces', 'KHÔNG'],
          ['infrastructure/persistence/impl/', 'Repository implementation', '@Repository'],
          ['infrastructure/security/', 'JWT, BCrypt, UserDetails', '@Component / @Service'],
          ['infrastructure/cache/', 'Redis adapter', '@Component'],
          ['infrastructure/events/', 'Spring event publisher', '@Component'],
          ['infrastructure/messaging/', 'Kafka bridge', '@Component'],
          ['interfaces/rest/', 'REST controllers', '@RestController'],
          ['interfaces/rest/filter/', 'Filter HTTP', '@Component'],
          ['interfaces/grpc/', 'gRPC service', '@GrpcService'],
          ['config/', 'Spring beans + properties binding', '@Configuration'],
          ['exceptions/', 'GlobalExceptionHandler', '@RestControllerAdvice'],
      ],
      col_widths=[6, 6.5, 3.5])

H(doc, '40.4 FAQ', level=3)

faqs = [
    ('Sao có 2 file DomainEvent.java và DomainEventPublisher.java?',
     'DomainEvent là "mẫu giấy thông báo" (abstract class - mọi event subclass kế thừa). '
     'DomainEventPublisher là "người đưa thư" (interface = port mà infra implement). '
     '2 vai trò khác nhau, đặt cùng folder domain/events/ vì cả 2 thuộc về domain.'),

    ('DomainEventPublisher đặt ở domain nhưng class implement ở infrastructure?',
     'Interface đặt ở domain vì là YÊU CẦU của domain. Implementation ở infrastructure vì đó là chi tiết kỹ thuật. '
     'Đây là Dependency Inversion: hướng phụ thuộc đảo ngược, domain không phụ thuộc infra.'),

    ('Vì sao có cả Spring event publisher VÀ Kafka publisher?',
     'Spring publisher (in-process, sync) → KafkaIntegrationEventPublisher lắng nghe bằng '
     '@TransactionalEventListener(AFTER_COMMIT) → map sang Integration Event → gửi Kafka. '
     '2 bước: trong nhà trước, ra ngoài sau.'),

    ('tokenVersion để làm gì?',
     'Chèn vào JWT khi mint. Đổi mật khẩu → bump version trong DB → JwtAuthenticationFilter so '
     'JWT.tokenVersion với user.currentTokenVersion → khác = 401. Tắt mọi JWT cũ tức thì.'),

    ('Vì sao Refresh Token có "family" và "generation"?',
     'Family = ID cố định cho cả 1 phiên login (root + mọi rotation). Generation = thứ tự rotation. '
     'Token bị cướp & dùng lại → server thấy token đã revoke (generation cũ) → revoke cả family → '
     'user thật + attacker đều bị đẩy ra.'),

    ('Vì sao logout cần blacklist mà không chỉ xóa cookie?',
     'JWT stateless: server không lưu state. Xóa cookie chỉ làm browser quên token. NHƯNG nếu attacker '
     'đã copy token (XSS) thì vẫn dùng được. Blacklist Redis + DB chặn token đó ở phía server.'),

    ('Vì sao có cả gRPC và REST?',
     'REST: cho client browser, mobile (public, JSON). gRPC: cho service-to-service nội bộ (nhanh hơn, '
     'type-safe qua proto). order-service kiểm permission user → gRPC, không qua HTTP.'),

    ('Vì sao seal class "public final class User"?',
     'Ngăn extends User và override method phá invariants. Aggregate Root nên là final.'),

    ('AggregateRoot vs Entity khác gì?',
     'Aggregate Root là entity "chính" của 1 nhóm, kiểm soát mọi thao tác. User aggregate gồm User + roles. '
     'Mọi thay đổi đi qua User method. Entity con (Permission là entity của Role aggregate) không truy cập trực tiếp.'),

    ('@Transactional propagation gì?',
     'Mặc định REQUIRED - join transaction hiện có hoặc tạo mới. Khi cần "luôn tạo new" (retry độc lập) dùng REQUIRES_NEW.'),

    ('Vì sao constructor private trong KafkaTopics?',
     'Utility class - chỉ chứa hằng số static. Cấm new → không tạo instance vô nghĩa. JDK Math, Collections cũng vậy.'),

    ('opsForValue() là gì?',
     'Spring Data Redis - thao tác với Redis String type. Trả ValueOperations. Tương tự có opsForHash, opsForSet, opsForList, opsForZSet cho data type khác.'),

    ('record là gì?',
     'Java 14+ syntax cho immutable data class. 1 dòng = 30+ dòng class viết tay (constructor, getter, equals, hashCode, toString tự sinh). Mọi field tự động final.'),

    ('Compact constructor là gì?',
     'Constructor đặc biệt của record - không có parameter list (đã được implicit). Dùng để validate/normalize giá trị trước khi gán vào field.'),

    ('@RequiredArgsConstructor (Lombok) là gì?',
     'Gen constructor nhận tất cả field final + field @NonNull. Tương đương viết tay constructor với super().'),

    ('@Slf4j (Lombok) là gì?',
     'Gen "private static final Logger log = LoggerFactory.getLogger(ClassName.class);". 2 cách (annotation vs viết tay) tương đương 100%.'),

    ('@AuthenticationPrincipal là gì?',
     'Spring Security annotation - inject principal từ SecurityContext vào method parameter. Trong service này inject AuthUserDetails (set bởi JwtAuthenticationFilter).'),

    ('@Value("${prop:default}") là gì?',
     'Inject property từ application.yaml. Cú pháp ${prop:default} với default value nếu không có.'),

    ('@PreAuthorize("hasRole(\'ADMIN\')") chạy khi nào?',
     'Spring Method Security - chặn method gọi nếu authentication không đủ quyền. Throw AccessDeniedException → 403.'),

    ('@JsonInclude(NON_NULL) là gì?',
     'Jackson annotation - serialize JSON skip null fields. UserDTO dùng để response gọn hơn.'),
]

for q, a in faqs:
    P(doc, 'Q: ' + q, bold=True, size=11)
    P(doc, 'A: ' + a, size=10)
    doc.add_paragraph()

# ============================================================
# LỜI KẾT
# ============================================================
doc.add_page_break()
H(doc, 'Lời kết', level=1)

P(doc,
  'auth-service là codebase được thiết kế chỉn chu theo các pattern Java enterprise hiện đại:')
for x in [
    'Hexagonal Architecture - domain tách rời framework, dễ test, dễ thay tech.',
    'CQRS với Command/Query Handler - đơn nhiệm, controller mỏng.',
    'Aggregate Root + Domain Events - quy tắc nghiệp vụ tập trung, audit trail tự nhiên.',
    'Value Objects bất biến - ép invariant tại điểm tạo, type-safe.',
    'Port & Adapter - domain ra yêu cầu, infrastructure nhận làm. Đổi tech không phá nghiệp vụ.',
    'Refresh Token Rotation + Family Revocation - chống token theft.',
    'tokenVersion + Redis Blacklist - JWT stateless nhưng vẫn revoke được tức thì.',
    'Pessimistic Lock + AFTER_COMMIT Listener - race condition + dual-write safe.',
    'Cookie HttpOnly + Bearer fallback - chống XSS, vẫn support non-browser client.',
    'Rate Limit trước Auth Filter - chống brute force, tiết kiệm CPU.',
    'gRPC cho service-to-service - nhanh hơn HTTP, type-safe qua proto.',
    'Hybrid Cache (Redis + DB) - nhanh nhưng vẫn durable.',
]:
    BULLET(doc, x)

P(doc,
  '\nMỗi quyết định thiết kế đều có lý do cụ thể, không phải "trang trí". '
  'Đọc kỹ tài liệu này + thực hành tận tay code là cách nhanh nhất để nắm vững các pattern. '
  'Khi gặp một codebase Spring microservice khác, bạn sẽ nhận ra ngay 80% pattern quen thuộc.\n')

P(doc, 'Hết.', italic=True)

# Final save
doc.save('/Users/admin/HieuTo/ecommerce/auth-service-doc-v2.docx')
print("=" * 60)
print("HOÀN THÀNH: auth-service-doc-v2.docx")
print("=" * 60)
