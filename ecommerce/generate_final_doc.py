"""
Build FINAL comprehensive backend specification document for HIEU E-Commerce.
Output: ecommerce/HIEU-Backend-Documentation-FINAL.docx
Run: python3 generate_final_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===================== HELPERS =====================
def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_)
    tcPr.append(shd)

def page_break(d): d.add_page_break()

def H(d, t, level=1): d.add_heading(t, level=level)

def P(d, t, bold=False, italic=False):
    p = d.add_paragraph()
    r = p.add_run(t); r.font.name = 'Calibri'; r.font.size = Pt(11)
    r.bold = bold; r.italic = italic
    return p

def Code(d, t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(t); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)

def Bul(d, t):
    p = d.add_paragraph(t, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    for r in p.runs: r.font.size = Pt(11)

def Num(d, t):
    p = d.add_paragraph(t, style='List Number')
    p.paragraph_format.left_indent = Cm(0.5)
    for r in p.runs: r.font.size = Pt(11)

def Tbl(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade(c, '2E5A9A'); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10); c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2 == 0: shade(c, 'F2F6FC')

def Cap(d, t):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def Note(d, t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run('Luu y: ' + t); r.italic = True; r.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF4CE')
    pPr.append(shd)

def J(*parts): return '\n'.join(parts)


# ===================== DOCUMENT =====================
doc = Document()
s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)


# ==================== TITLE ====================
t = doc.add_heading('TÀI LIỆU ĐẶC TẢ KỸ THUẬT BACKEND', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

for line, sz, bold, color in [
    ('HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ HIEU', 18, True, None),
    ('KIẾN TRÚC MICROSERVICES — BẢN TỔNG HỢP ĐẦY ĐỦ', 14, True, RGBColor(0x2E,0x5A,0x9A)),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.bold = bold; r.font.size = Pt(sz)
    if color: r.font.color.rgb = color

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HIEU E-Commerce Backend Platform'); r.italic = True; r.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Phiên bản: 3.0 (FINAL)\nCông nghệ: Spring Boot 4.0.5 | Java 25 | gRPC | Kafka | PostgreSQL 16 | Redis 7 | Elasticsearch 8\nNgày cập nhật: 17/05/2026')
r.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('© 2026 HIEU E-Commerce Team — Internal documentation v3.0')
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x88,0x88,0x88)

page_break(doc)


# ==================== TOC ====================
H(doc, 'MỤC LỤC', level=1)
TOC = [
    ('LỊCH SỬ PHIÊN BẢN', True),
    ('BẢNG CHÚ THÍCH THUẬT NGỮ VÀ VIẾT TẮT', True),
    ('', False),
    ('CHƯƠNG 1: TỔNG QUAN DỰ ÁN', True),
    ('CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG', True),
    ('CHƯƠNG 3: ĐẶC TẢ YÊU CẦU', True),
    ('CHƯƠNG 4: USE CASE', True),
    ('CHƯƠNG 5: THIẾT KẾ CƠ SỞ DỮ LIỆU', True),
    ('CHƯƠNG 6: THIẾT KẾ API', True),
    ('CHƯƠNG 7: LUỒNG XỬ LÝ NGHIỆP VỤ', True),
    ('CHƯƠNG 8: GIAO TIẾP GIỮA CÁC SERVICE', True),
    ('CHƯƠNG 9: BẢO MẬT', True),
    ('CHƯƠNG 10: KIỂM THỬ', True),
    ('CHƯƠNG 11: GIÁM SÁT VÀ VẬN HÀNH', True),
    ('CHƯƠNG 12: TRIỂN KHAI', True),
    ('CHƯƠNG 13: LỘ TRÌNH PHÁT TRIỂN', True),
    ('PHỤ LỤC (A. Cổng dịch vụ — B. Mã lỗi — C. Seed — D. Cấu hình — E. Glossary nâng cao)', True),
]
for label, bold in TOC:
    p = doc.add_paragraph(label); p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.bold = bold; r.font.size = Pt(11 if bold else 10)
page_break(doc)


# ==================== VERSION HISTORY ====================
H(doc, 'LỊCH SỬ PHIÊN BẢN', level=1)
P(doc, 'Bảng 0.1: Lịch sử phiên bản tài liệu', italic=True)
Tbl(doc,
    ['Phiên bản', 'Ngày', 'Tác giả', 'Mô tả thay đổi'],
    [
        ['0.1', '12/03/2026', 'Backend Team', 'Khởi tạo với 5 service cơ bản.'],
        ['0.5', '08/04/2026', 'Backend Team', 'Bổ sung inventory, shipping, voucher, notification.'],
        ['1.0', '13/05/2026', 'Backend Team', 'Hoàn thiện 15 service; use case, saga, state machine.'],
        ['2.0', '14/05/2026', 'Backend Team', 'Tái cấu trúc 13 chương + glossary + schema chi tiết.'],
        ['3.0 (FINAL)', '17/05/2026', 'Backend Team',
         'Bản tổng hợp đầy đủ: 18 services + common-lib, CQRS chi tiết, '
         'Domain Events ↔ Integration Events, Outbox Pattern, MapStruct, error codes, '
         'đầy đủ cấu hình mẫu cho 11 PostgreSQL, Redis, Kafka, Elasticsearch, MailHog.'],
    ])
page_break(doc)


# ==================== GLOSSARY ====================
H(doc, 'BẢNG CHÚ THÍCH THUẬT NGỮ VÀ VIẾT TẮT', level=1)
P(doc, 'Tài liệu sử dụng các thuật ngữ và từ viết tắt sau đây.')

H(doc, '1. Thuật ngữ kỹ thuật', level=2)
P(doc, 'Bảng 0.2: Thuật ngữ kỹ thuật', italic=True)
Tbl(doc, ['STT','Thuật ngữ','Tên đầy đủ','Giải thích'], [
    ['1','API','Application Programming Interface','Giao diện cho hệ thống giao tiếp.'],
    ['2','REST','Representational State Transfer','Kiến trúc API HTTP + JSON.'],
    ['3','gRPC','Google Remote Procedure Call','RPC HTTP/2 + Protobuf, latency thấp.'],
    ['4','JWT','JSON Web Token','Token tự đóng gói, ký HMAC/RSA, RFC 7519.'],
    ['5','BCrypt','-','Hash mật khẩu có salt, cost 12.'],
    ['6','HMAC','Hash-based Message Authentication','HS256 = HMAC-SHA-256.'],
    ['7','TTL','Time To Live','Thời gian sống của bản ghi.'],
    ['8','SSE','Server-Sent Events','Push một chiều server→client trên HTTP.'],
    ['9','DLQ','Dead Letter Queue','Queue chứa message fail.'],
    ['10','RBAC','Role-Based Access Control','User→Role→Permission.'],
    ['11','jti','JWT ID','Định danh duy nhất token, dùng blacklist.'],
    ['12','tokenVersion','-','Cho phép thu hồi toàn bộ token bằng tăng version.'],
    ['13','CORS','Cross-Origin Resource Sharing','Kiểm soát truy cập cross-origin.'],
])

H(doc, '2. Thuật ngữ kiến trúc', level=2)
P(doc, 'Bảng 0.3: Thuật ngữ kiến trúc', italic=True)
Tbl(doc, ['STT','Thuật ngữ','Giải thích'], [
    ['1','Microservices','Hệ thống chia thành service nhỏ độc lập.'],
    ['2','DDD','Thiết kế tập trung vào domain model.'],
    ['3','Hexagonal','Tách domain logic khỏi infrastructure (ports & adapters).'],
    ['4','Aggregate','Cụm entity có ranh giới giao dịch chung.'],
    ['5','Value Object','Đối tượng bất biến (UserId, Email).'],
    ['6','Domain Event','Sự kiện trong domain, publish nội bộ.'],
    ['7','Integration Event','Sự kiện phẳng versioned, publish qua Kafka.'],
    ['8','Repository','Pattern truy xuất dữ liệu (port).'],
    ['9','Saga','Chuỗi giao dịch local có compensation.'],
    ['10','Choreography Saga','Saga không orchestrator – mỗi service tự subscribe.'],
    ['11','CQRS','Tách luồng ghi (Command) và đọc (Query).'],
    ['12','Outbox','Lưu event cùng transaction nghiệp vụ, publisher đẩy sau.'],
    ['13','Idempotency','Cho cùng kết quả khi gọi nhiều lần.'],
    ['14','Service Discovery','Tự đăng ký/tìm kiếm service (Eureka).'],
    ['15','API Gateway','Điểm vào duy nhất, routing.'],
    ['16','Circuit Breaker','Ngắt mạch khi downstream lỗi.'],
])

H(doc, '3. Thuật ngữ nghiệp vụ', level=2)
P(doc, 'Bảng 0.4: Thuật ngữ nghiệp vụ', italic=True)
Tbl(doc, ['STT','Thuật ngữ','Giải thích'], [
    ['1','SPU','Standard Product Unit (Áo HIEU phiên bản giới hạn).'],
    ['2','SKU','Stock Keeping Unit – biến thể của SPU.'],
    ['3','Slug','URL-friendly identifier.'],
    ['4','Variant','SKU cụ thể với thuộc tính riêng.'],
    ['5','Attribute','Thuộc tính sản phẩm (màu, size).'],
    ['6','Flash Sale','Khuyến mãi mạnh ngắn hạn, quota giới hạn.'],
    ['7','Voucher','Phiếu giảm giá theo điều kiện.'],
    ['8','Wishlist','Sản phẩm yêu thích.'],
    ['9','Return Request','Yêu cầu đổi/trả cần admin duyệt.'],
])

H(doc, '4. Thuật ngữ hạ tầng', level=2)
P(doc, 'Bảng 0.5: Thuật ngữ hạ tầng', italic=True)
Tbl(doc, ['STT','Thuật ngữ','Giải thích'], [
    ['1','Docker','Container hóa ứng dụng.'],
    ['2','K8s','Container orchestration.'],
    ['3','Kafka KRaft','Streaming sự kiện phân tán (no Zookeeper).'],
    ['4','Redis 7','In-memory cho cache, blacklist, idempotency.'],
    ['5','Elasticsearch','Full-text search + log index.'],
    ['6','Eureka','Service registry Netflix.'],
    ['7','Resilience4J','Circuit breaker, retry.'],
    ['8','MapStruct','Code-gen DTO mapping.'],
    ['9','Lombok','Annotation giảm boilerplate.'],
    ['10','MailHog','Local mail catcher (dev).'],
    ['11','Filebeat','Log shipper.'],
    ['12','Kibana','Log viewer UI.'],
    ['13','Flyway','Schema version control.'],
    ['14','JaCoCo','Test coverage tool.'],
])
page_break(doc)


# ==================== CHAPTER 1 ====================
H(doc, 'CHƯƠNG 1: TỔNG QUAN DỰ ÁN', level=1)

H(doc, '1.1. Giới thiệu', level=2)
P(doc, 'Tài liệu này đặc tả chi tiết backend của hệ thống thương mại điện tử HIEU – '
       'nền tảng bán hàng thời trang luxury xây dựng theo kiến trúc microservices. '
       'Hệ thống gồm 15 service nghiệp vụ + 1 API Gateway + 1 Eureka Server + 1 common-lib, '
       'mỗi service phụ trách một nghiệp vụ nhỏ nhưng phối hợp qua REST, gRPC và Kafka.')

for b in [
    'Microservices: 15 service nghiệp vụ độc lập, 11 PostgreSQL instance.',
    'DDD 4 tầng: interfaces, application, domain, infrastructure.',
    'Hexagonal: Domain định nghĩa port, infrastructure cung cấp adapter.',
    'CQRS: Auth-service có 9 Command + 5 Query handlers.',
    'Choreography Saga: Đặt hàng không có orchestrator.',
    'Idempotency: X-Idempotency-Key UUIDv4 cho mọi command quan trọng.',
    'Outbox Pattern: At-least-once delivery cho Kafka events.',
    'Service Discovery: Netflix Eureka.',
    'Event-driven: Kafka KRaft, ~23 topic.',
    'Database per Service: 11 PostgreSQL riêng.',
    'Cloud-Native: Docker, K8s-ready, Actuator endpoints.',
]:
    Bul(doc, b)

H(doc, '1.2. Mục tiêu dự án', level=2)
for b in [
    'Nền tảng e-commerce hoàn chỉnh cho thời trang luxury.',
    'Microservices đảm bảo độc lập triển khai, scale, cô lập sự cố.',
    'JWT (HS256) – mỗi service tự verify trong filter chain.',
    'REST (FE↔Gateway), gRPC (inter-service sync), Kafka (async).',
    'Toàn vẹn dữ liệu qua saga + outbox + idempotency.',
    'Tích hợp SePay VietQR webhook + GHN/GHTK tracking.',
    'Giám sát Actuator + Filebeat + Elasticsearch + Kibana.',
    'Test pyramid: unit domain, integration controller-service-repo, e2e TestContainers.',
    'Roadmap: AI Recommendation, multi-warehouse, Kubernetes.',
]:
    Bul(doc, b)

H(doc, '1.3. Phạm vi dự án', level=2)
P(doc, 'Bảng 1.1: Danh sách service trong dự án', italic=True)
Tbl(doc, ['STT','Service','Vai trò','Trạng thái'], [
    ['1','eureka-server','Service registry.','Hoàn thành'],
    ['2','api-gateway','Điểm vào – CORS, routing, rate limit.','Hoàn thành'],
    ['3','auth-service','Xác thực, RBAC, JWT, blacklist.','Hoàn thành'],
    ['4','user-profile-service','Hồ sơ + địa chỉ.','Hoàn thành'],
    ['5','catalog-service','Danh mục, sản phẩm, biến thể (REST + gRPC).','Hoàn thành'],
    ['6','inventory-service','Tồn kho, reservation, movements.','Hoàn thành'],
    ['7','cart-service','Giỏ hàng (validate qua gRPC catalog).','Hoàn thành'],
    ['8','order-service','Đặt hàng (saga), state machine, outbox.','Hoàn thành'],
    ['9','payment-service','SePay webhook + refund.','Hoàn thành'],
    ['10','shipping-service','GHN/GHTK tracking.','Hoàn thành'],
    ['11','voucher-service','Voucher + redemption.','Hoàn thành'],
    ['12','flash-sale-service','Atomic Lua quota.','Hoàn thành'],
    ['13','notification-service','In-app + email + SSE.','Hoàn thành'],
    ['14','review-service','Đánh giá sản phẩm.','Hoàn thành'],
    ['15','whishlist-service','Wishlist.','Hoàn thành'],
    ['16','search-service','Elasticsearch.','Hoàn thành'],
    ['17','analytics-service','Dashboard analytics.','Hoàn thành'],
    ['18','recommendation-service','Gợi ý sản phẩm (bản đơn giản).','Hoàn thành'],
    ['19','common-lib','Thư viện chung.','Hoàn thành'],
])

H(doc, '1.4. Đối tượng sử dụng', level=2)
P(doc, 'Bảng 1.2: Các đối tượng sử dụng', italic=True)
Tbl(doc, ['STT','Đối tượng','Tương tác chính'], [
    ['1','Backend Developer','Kiến trúc, schema DB, ma trận giao tiếp.'],
    ['2','Frontend Developer','API contract, error codes, auth flow.'],
    ['3','DevOps / SRE','Docker, K8s, biến môi trường, health check.'],
    ['4','QA / Tester','Use case, test pyramid, dữ liệu seed.'],
    ['5','Solution Architect','Pattern, roadmap.'],
    ['6','Product Owner','Use case, FR/NFR.'],
])

H(doc, '1.5. Cấu trúc tài liệu', level=2)
P(doc, 'Bảng 1.3: Cấu trúc nội dung', italic=True)
Tbl(doc, ['Chương','Tên','Nội dung'], [
    ['1','Tổng quan dự án','Giới thiệu, mục tiêu, phạm vi.'],
    ['2','Kiến trúc','DDD, Hexagonal, CQRS, topology.'],
    ['3','Đặc tả yêu cầu','FR + NFR + ràng buộc.'],
    ['4','Use Case','Actor + đặc tả UC.'],
    ['5','CSDL','ERD, bảng, chỉ mục, migration.'],
    ['6','API','REST endpoints theo service.'],
    ['7','Luồng xử lý','Sequence diagrams.'],
    ['8','Giao tiếp','REST/gRPC/Kafka.'],
    ['9','Bảo mật','JWT, RBAC, revocation.'],
    ['10','Kiểm thử','Unit, integration, coverage.'],
    ['11','Giám sát','Logging, metrics, recovery.'],
    ['12','Triển khai','Docker, K8s, CI/CD.'],
    ['13','Lộ trình','Hiện trạng + roadmap.'],
])
page_break(doc)
print("Done Ch1")


# ==================== CHAPTER 2 ====================
H(doc, 'CHƯƠNG 2: KIẾN TRÚC HỆ THỐNG', level=1)

H(doc, '2.1. Kiến trúc tổng thể & Sơ đồ C4', level=2)
P(doc, 'Hệ thống sử dụng kiến trúc Microservices với API Gateway làm điểm vào duy nhất. Mỗi service hoạt động độc lập, có cơ sở dữ liệu riêng và giao tiếp qua REST, gRPC hoặc Kafka.')

P(doc, 'Các thành phần cốt lõi:')
for b in [
    'API Gateway (port 8080): Spring Cloud Gateway – CORS, routing, trace ID, rate limit.',
    'Eureka Server (port 8761): Service registry – heartbeat 30s.',
    'Auth Service (port 9091, gRPC 9191): JWT, rotation, RBAC, blacklist.',
    'Catalog Service (port 9093, gRPC 9193): REST + gRPC.',
    '13 business services: cart(9094), order(9095), payment(8086), shipping(8087), '
    'inventory(9098), voucher(8094), flash-sale(8089), notification(8090), '
    'user-profile(9099), search(8092), review, whishlist, analytics(8095), recommendation.',
    'Hạ tầng: 11 PostgreSQL 16 (5433-5443), Redis 7 (6379), Kafka KRaft (9092), '
    'Elasticsearch 8 (9200), Kibana (5601), MailHog (1025/8025), Filebeat.',
]:
    Bul(doc, b)

P(doc, 'Sơ đồ topology dạng ASCII:')
Code(doc, J(
    '                  +------------------+',
    '                  |  Client (Next.js)|',
    '                  +--------+---------+',
    '                           |',
    '                           v',
    '                  +------------------+',
    '                  |   API Gateway    |  :8080',
    '                  | (Spring Cloud GW)|  - CORS allowlist',
    '                  |                  |  - Path-prefix routing',
    '                  |                  |  - Rate limit (Redis)',
    '                  +-+-+-+-+-+-+-+-+-++',
    '                    | | | | | | | | |',
    '        +-----------+ | | | | | | | +------+',
    '        v             v v v v v v v        v',
    '  +----------+   +----------+...+----------+  +---------+',
    '  |  Auth    |   | Catalog  |   | Analytics|  | Eureka  |',
    '  |  :9091   |   |  :9093   |   |  :8095   |  |  :8761  |',
    '  |  gRPC    |   |  gRPC    |   |          |  |         |',
    '  |  :9191   |   |  :9193   |   |          |  |         |',
    '  +----+-----+   +----+-----+   +----+-----+  +---------+',
    '       |              |              |',
    '       v              v              v',
    '  +---------+    +----------+   +----------+',
    '  | auth_db |    |catalog_db|   | ES logs  |',
    '  | Postgres|    | Postgres |   | Elastic  |',
    '  +---------+    +----------+   +----------+',
    '',
    '  Shared Infrastructure:',
    '  +---------+ +---------+ +-------------+ +----------+ +---------+',
    '  |  Redis  | |  Kafka  | |Elasticsearch| |  Kibana  | | MailHog |',
    '  |  :6379  | |  :9092  | |   :9200     | |   :5601  | |  :8025  |',
    '  +---------+ +---------+ +-------------+ +----------+ +---------+',
))
Cap(doc, 'Hình 2.1: Sơ đồ topology hệ thống HIEU')

H(doc, '2.2. Bốn tầng kiến trúc DDD + Hexagonal', level=2)
P(doc, 'Mỗi service được thiết kế theo DDD kết hợp Hexagonal với 4 tầng. Phụ thuộc đi MỘT CHIỀU: tầng ngoài → tầng trong, không bao giờ ngược lại. Domain hoàn toàn không phụ thuộc framework.')
P(doc, 'Bảng 2.1: Trách nhiệm 4 tầng DDD', italic=True)
Tbl(doc, ['Tầng','Vai trò','Phụ thuộc','Ví dụ'], [
    ['interfaces','REST controller, gRPC server, Kafka consumer.','application, domain (DTO)','AuthController, AuthGrpcService'],
    ['application','Command/Query handler, DTO, transaction.','domain','LoginHandler, GetUserByIdHandler'],
    ['domain (pure)','Aggregate, Entity, VO, Domain Event, port.','Không phụ thuộc','User aggregate, RefreshToken, UserId VO'],
    ['infrastructure','Implement port: JPA, JWT, Kafka, Redis.','domain','UserRepositoryImpl, JwtTokenProvider'],
])

H(doc, '2.3. CQRS', level=2)
P(doc, 'Mỗi use case ghi là 1 Command + 1 CommandHandler; mỗi use case đọc là 1 Query + 1 QueryHandler.')
Code(doc, J(
    'public interface Command<R> {}',
    'public interface Query<R> {}',
    '',
    '@FunctionalInterface',
    'public interface CommandHandler<C extends Command<R>, R> {',
    '    R handle(C command);',
    '}',
    '',
    'public record LoginCommand(String usernameOrEmail, String rawPassword)',
    '        implements Command<AuthResponseDTO> {}',
    '',
    '@Service @RequiredArgsConstructor',
    'public class LoginHandler implements CommandHandler<LoginCommand, AuthResponseDTO> {',
    '    private final UserRepository userRepository;',
    '    @Override @Transactional',
    '    public AuthResponseDTO handle(LoginCommand command) { /* ... */ }',
    '}',
))

P(doc, 'Bảng 2.2: Danh sách Command/Query của auth-service (9 + 5)', italic=True)
Tbl(doc, ['Loại','Tên','Return','Mô tả'], [
    ['Command','RegisterUserCommand','AuthResponseDTO','Đăng ký + phát JWT.'],
    ['Command','LoginCommand','AuthResponseDTO','Đăng nhập.'],
    ['Command','LogoutCommand','Void','Logout + blacklist + revoke refresh.'],
    ['Command','RefreshTokenCommand','AuthResponseDTO','Rotate refresh.'],
    ['Command','ChangePasswordCommand','Void','Đổi mật khẩu + bump tokenVersion.'],
    ['Command','UpdateEmailCommand','UserDTO','Đổi email.'],
    ['Command','ChangeAccountStatusCommand','Void','LOCK/UNLOCK/ENABLE/DISABLE.'],
    ['Command','AssignRoleCommand','Void','Gán role (admin).'],
    ['Command','UnassignRoleCommand','Void','Bỏ role (admin).'],
    ['Query','GetUserByIdQuery','UserDTO','Lấy user theo UUID.'],
    ['Query','GetUserByUsernameQuery','UserDTO','Lấy user theo username.'],
    ['Query','ListUsersQuery','Page<UserDTO>','Danh sách user (cursor).'],
    ['Query','CheckPermissionQuery','Boolean','Check permission.'],
    ['Query','CheckRoleQuery','Boolean','Check role.'],
])

H(doc, '2.4. Domain Events ↔ Integration Events', level=2)
P(doc, 'Hệ thống tách 2 loại event:')
Bul(doc, 'Domain Event (in-process): UserCreatedEvent, EmailChangedEvent, TokenCreatedEvent... – sống ngắn, trong cùng JVM, abstract class.')
Bul(doc, 'Integration Event (cross-service qua Kafka): UserLifecycleIntegrationEvent, SessionIntegrationEvent – có schemaVersion, eventType chuỗi versioned (auth.user.created.v1), payload phẳng Map.')

Tbl(doc, ['Tiêu chí','Domain Event','Integration Event'], [
    ['Vị trí','domain/events/','application/events/'],
    ['Loại Java','abstract class','interface + records'],
    ['eventType','class.getSimpleName()','Chuỗi "auth.user.created.v1"'],
    ['Versioning','Không','Có (.v1)'],
    ['Phạm vi','1 JVM','Cross-service'],
    ['Publish','Spring ApplicationEventPublisher','KafkaTemplate'],
    ['Schema','Strongly-typed','Flat map JSON'],
    ['Breaking','Refactor OK','Phải bump .v2'],
])

H(doc, '2.5. Cấu trúc thư mục chuẩn', level=2)
Code(doc, J(
    'auth-service/',
    'pom.xml',
    'src/main/java/com/hieu/auth_service/',
    '  AuthServiceApplication.java',
    '  config/                         # Bean config, Security, Kafka, gRPC',
    '  interfaces/                     # Adapters inbound',
    '    rest/, grpc/, kafka/',
    '  application/                    # Use case + DTO',
    '    common/                       # Command/Query/Handler interfaces',
    '    command/, query/, handler/',
    '    dto/, mapper/                 # MapStruct',
    '    port/                         # RolePermissionCachePort, TokenBlacklistPort',
    '    events/                       # IntegrationEvent, KafkaTopics',
    '  domain/                         # Pure domain',
    '    models/                       # User, RefreshToken, Role, Permission',
    '    events/                       # DomainEvent + Publisher port',
    '    repositories/                 # Repository interfaces',
    '    services/                     # TokenDomainService, PasswordEncoderPort',
    '    shared/AggregateRoot.java',
    '  infrastructure/',
    '    persistence/jpa/              # JPA + Spring Data',
    '    security/                     # JwtTokenProvider, BCrypt, TokenBlacklist',
    '    messaging/                    # KafkaIntegrationEventPublisher, Mapper',
    '    events/                       # SpringDomainEventPublisher',
    '    cache/',
    '  exceptions/                     # @RestControllerAdvice',
    'src/main/resources/',
    '  application.yaml',
    '  db/migration/                   # Flyway',
    '  proto/                          # gRPC .proto',
))

H(doc, '2.6. Topology (15 services)', level=2)
P(doc, 'Bảng 2.3: Danh sách service, cổng và DB', italic=True)
Tbl(doc, ['STT','Service','HTTP','gRPC','DB Port','Mô tả'], [
    ['1','eureka-server','8761','-','-','Service registry'],
    ['2','api-gateway','8080','-','-','Gateway'],
    ['3','auth-service','9091','9191','5433','Auth, JWT, RBAC'],
    ['4','user-profile-service','9099','-','5439','Hồ sơ + địa chỉ'],
    ['5','catalog-service','9093','9193','5434','Sản phẩm'],
    ['6','inventory-service','9098','-','5440','Tồn kho'],
    ['7','cart-service','9094','9194','5435','Giỏ hàng'],
    ['8','order-service','9095','-','5436','Saga, outbox'],
    ['9','payment-service','8086','-','5437','SePay'],
    ['10','shipping-service','8087','-','5438','GHN/GHTK'],
    ['11','voucher-service','8094','-','5441','Voucher'],
    ['12','flash-sale-service','8089','-','5442','Atomic Lua'],
    ['13','notification-service','8090','-','5443','In-app + email + SSE'],
    ['14','review-service','-','-','-','Đánh giá'],
    ['15','whishlist-service','-','-','-','Wishlist'],
    ['16','search-service','8092','-','-','Elasticsearch'],
    ['17','analytics-service','8095','-','-','Dashboard'],
    ['18','recommendation-service','-','-','-','Gợi ý'],
    ['19','common-lib','-','-','-','Thư viện share'],
    ['20','Redis','-','-','6379','Cache + blacklist'],
    ['21','Kafka KRaft','-','-','9092','23 topic'],
    ['22','Elasticsearch','-','-','9200','Search + log'],
    ['23','Kibana','-','-','5601','Log UI'],
    ['24','MailHog','-','-','1025/8025','Mail dev'],
])

H(doc, '2.7. Stack công nghệ', level=2)
P(doc, 'Bảng 2.4: Stack công nghệ cốt lõi', italic=True)
Tbl(doc, ['Nhóm','Công nghệ','Phiên bản','Mục đích'], [
    ['Framework','Spring Boot','4.0.5','Framework chính'],
    ['Cloud','Spring Cloud','2025.1.1','Microservices infra'],
    ['Ngôn ngữ','Java','25 (LTS)','Ngôn ngữ'],
    ['Build','Maven','3.9.x','Build tool'],
    ['ORM','Spring Data JPA + Hibernate','7.x','DB access'],
    ['Security','Spring Security','6.x','Auth + RBAC'],
    ['JWT','JJWT','0.12.6','HS256'],
    ['Gateway','Spring Cloud Gateway','-','Routing'],
    ['Discovery','Netflix Eureka','-','Registry'],
    ['gRPC','Spring gRPC + gRPC Java','1.77.1','Inter-service'],
    ['Protobuf','Protocol Buffers','4.33.4','gRPC message'],
    ['Messaging','Kafka (KRaft)','3.7','Event streaming'],
    ['API Docs','SpringDoc OpenAPI','3.0.2','Swagger UI'],
    ['Mapping','MapStruct','1.6.3','DTO mapping'],
    ['Utility','Lombok','1.18.x','Boilerplate'],
    ['Logging','Logback + Logstash','8.x','JSON log'],
    ['Metrics','Micrometer','1.14.x','Prometheus format'],
    ['Database','PostgreSQL','16','CSDL chính'],
    ['Cache','Redis','7','Cache + blacklist'],
    ['Search','Elasticsearch','8.13','Full-text + log'],
    ['Log UI','Kibana','8.13','Log explorer'],
    ['Mail','MailHog','latest','Dev mail'],
    ['Migration','Flyway','10.x','Schema versioning'],
    ['Rate Limit','Bucket4j','8.10.1','Token bucket'],
    ['Container','Docker + Compose','-','Đóng gói'],
    ['Resilience','Resilience4J','-','Circuit breaker'],
    ['Testing','JUnit 5 + Mockito + Testcontainers','-','Unit + integration'],
    ['Coverage','JaCoCo','0.8.12','Coverage'],
    ['Payment','SePay (VN)','-','VietQR + webhook'],
    ['Shipping','GHN, GHTK','-','Tracking'],
])
page_break(doc)
print("Done Ch2")


# ==================== CHAPTER 3 ====================
H(doc, 'CHƯƠNG 3: ĐẶC TẢ YÊU CẦU', level=1)

H(doc, '3.1. Yêu cầu chức năng theo service', level=2)

P(doc, 'Bảng 3.1: FR – Auth Service', italic=True)
Tbl(doc, ['Mã','Tên','Mô tả','Ưu tiên'], [
    ['FR-01','Đăng ký','Username/email/password/họ tên; check trùng, BCrypt cost 12.','Cao'],
    ['FR-02','Đăng nhập','Username hoặc email + password; trả access 15p + refresh 7d qua cookie.','Cao'],
    ['FR-03','Làm mới token','Rotation: revoke cũ + issue mới; reuse detection → revoke family.','Cao'],
    ['FR-04','Xác thực 4-layer','Signature → exp → blacklist → tokenVersion match.','Cao'],
    ['FR-05','Đăng xuất','Revoke refresh + blacklist access.','Cao'],
    ['FR-06','Đổi mật khẩu','Verify old → hash new → bump tokenVersion → revoke all.','Cao'],
    ['FR-07','Đổi email','Validate format + chưa tồn tại.','Trung bình'],
    ['FR-08','Thu hồi toàn bộ token','Admin/self: tăng tokenVersion.','Cao'],
    ['FR-09','Quản lý account status','LOCK/UNLOCK/ENABLE/DISABLE.','Cao'],
    ['FR-10','RBAC – Gán/bỏ role','Admin assign/unassign.','Cao'],
    ['FR-11','Quản lý role/permission','CRUD + cấp/thu hồi.','Cao'],
    ['FR-12','Lấy thông tin user','GetUserById/Username; List (cursor).','Trung bình'],
    ['FR-13','Kiểm tra quyền/role','CheckPermission, CheckRole.','Cao'],
    ['FR-14','Khởi tạo seed','Auto tạo admin + role + permission.','Cao'],
])

P(doc, 'Bảng 3.2: FR – Catalog', italic=True)
Tbl(doc, ['Mã','Tên','Mô tả','Ưu tiên'], [
    ['FR-15','Quản lý danh mục','CRUD phân cấp (parent-child).','Cao'],
    ['FR-16','Quản lý sản phẩm','CRUD product, slug, brand, images.','Cao'],
    ['FR-17','Quản lý variant','CRUD variant: giá, sale price, cân nặng.','Cao'],
    ['FR-18','Quản lý attribute','CRUD attribute + attr_vals.','Cao'],
    ['FR-19','Tìm kiếm sản phẩm','Theo từ khóa, slug, category.','Cao'],
    ['FR-20','Batch API','Lấy nhiều product/variant theo IDs.','Trung bình'],
    ['FR-21','gRPC inter-service','GetProduct, GetVariantBySku, CheckStock.','Cao'],
])

P(doc, 'Bảng 3.3: FR – Cart, Order, Inventory', italic=True)
Tbl(doc, ['Mã','Service','Tên','Mô tả'], [
    ['FR-22','Cart','Xem giỏ','Lấy giỏ hiện tại.'],
    ['FR-23','Cart','Thêm/sửa/xóa item','Validate gRPC catalog; idempotency.'],
    ['FR-24','Cart','Clear cart','Sau order.placed.'],
    ['FR-25','Order','Tạo đơn','Saga + voucher + reserve + create payment.'],
    ['FR-26','Order','Xem đơn','Chi tiết + lịch sử.'],
    ['FR-27','Order','Cập nhật trạng thái','State machine.'],
    ['FR-28','Order','Hủy đơn','PENDING/PAYMENT_PENDING.'],
    ['FR-29','Order','Đổi/trả','ReturnRequest flow.'],
    ['FR-30','Inventory','Quản lý stock','CRUD cho từng SKU.'],
    ['FR-31','Inventory','Reserve/Release','Giữ + release khi hủy/timeout.'],
    ['FR-32','Inventory','Stock movements','Audit append-only.'],
    ['FR-33','Inventory','Cảnh báo','Publish low-stock event.'],
])

P(doc, 'Bảng 3.4: FR – Payment, Shipping, User Profile', italic=True)
Tbl(doc, ['Mã','Service','Tên','Mô tả'], [
    ['FR-34','Payment','Khởi tạo','Payment PENDING + payUrl + qrCodeUrl.'],
    ['FR-35','Payment','Webhook SePay','Verify HMAC → match orderNumber.'],
    ['FR-36','Payment','Refund','Admin trigger.'],
    ['FR-37','Shipping','Tạo shipment','Sau payment.completed.'],
    ['FR-38','Shipping','Tracking','Lưu tracking_events.'],
    ['FR-39','Shipping','Cập nhật','PENDING → IN_TRANSIT → DELIVERED.'],
    ['FR-40','User Profile','Hồ sơ','Update personal info, avatar.'],
    ['FR-41','User Profile','Địa chỉ','CRUD addresses; default.'],
])

P(doc, 'Bảng 3.5: FR – Voucher, Flash Sale, Notification', italic=True)
Tbl(doc, ['Mã','Service','Tên','Mô tả'], [
    ['FR-42','Voucher','Tạo','Admin: code, type, value, conditions.'],
    ['FR-43','Voucher','Apply','Validate khi tạo đơn.'],
    ['FR-44','Voucher','Compensation','Release khi cancel.'],
    ['FR-45','Flash Sale','Tạo campaign','DRAFT → ACTIVE.'],
    ['FR-46','Flash Sale','Tham gia','Atomic check Redis Lua.'],
    ['FR-47','Flash Sale','Chống oversell','Decrement nguyên tử.'],
    ['FR-48','Notification','Email','Transactional.'],
    ['FR-49','Notification','Push','SSE channel.'],
    ['FR-50','Notification','Template','Mustache.'],
])

P(doc, 'Bảng 3.6: FR – Search, Analytics, Review, Wishlist', italic=True)
Tbl(doc, ['Mã','Service','Tên','Mô tả'], [
    ['FR-51','Search','Full-text','Highlight + typo suggest.'],
    ['FR-52','Search','Autocomplete','Edge ngram.'],
    ['FR-53','Search','Faceted','Filter category, price.'],
    ['FR-54','Search','Sync','Subscribe Kafka.'],
    ['FR-55','Analytics','Event','Pageview, addToCart, purchase.'],
    ['FR-56','Analytics','Báo cáo','Doanh thu, top SP.'],
    ['FR-57','Review','Tạo','Chỉ user đã mua.'],
    ['FR-58','Review','Duyệt','APPROVED/REJECTED.'],
    ['FR-59','Wishlist','Thêm/xóa','Đánh dấu yêu thích.'],
    ['FR-60','Wishlist','Xem','Pagination.'],
])

H(doc, '3.2. Yêu cầu phi chức năng', level=2)
P(doc, 'Bảng 3.7: NFR', italic=True)
Tbl(doc, ['Mã','Nhóm','Yêu cầu','Chi tiết'], [
    ['NFR-01','Hiệu năng','API p95 < 500ms','Trong tải bình thường.'],
    ['NFR-02','Hiệu năng','gRPC p95 < 100ms','Inter-service.'],
    ['NFR-03','Mở rộng','Horizontal scale','Container orchestration.'],
    ['NFR-04','Bảo mật','JWT + RBAC + 4-layer','BCrypt 12, HS256, HTTPS.'],
    ['NFR-05','Sẵn sàng','Uptime ≥ 99.5%','Circuit breaker, retry.'],
    ['NFR-06','Bảo trì','Modular','DDD + Hexagonal + CQRS.'],
    ['NFR-07','Giám sát','Logging + metrics','JSON log + Filebeat.'],
    ['NFR-08','Tương thích','RESTful API chuẩn','OpenAPI 3.0.'],
    ['NFR-09','Chịu lỗi','Resilient','Circuit breaker, timeout 3s.'],
    ['NFR-10','Kiểm thử','Coverage > 60%','Domain ≥ 80%.'],
    ['NFR-11','Toàn vẹn','Saga + outbox + idempotency','At-least-once.'],
    ['NFR-12','Vận hành','Secret management','K8s Secret / Vault.'],
    ['NFR-13','Recovery','RTO ≤ 4h, RPO ≤ 1h','Backup S3.'],
    ['NFR-14','Tuân thủ','GDPR','Right to be forgotten.'],
])

H(doc, '3.3. Ràng buộc thiết kế', level=2)
for b in [
    'Database per Service: KHÔNG foreign key cross-service; tham chiếu logical UUID/Long.',
    'Domain layer KHÔNG import framework; phụ thuộc qua port.',
    'Mọi entity dùng Value Object thay vì String/Long trần.',
    'Mọi mutation method phải raise Domain Event.',
    'Mọi command quan trọng PHẢI có X-Idempotency-Key.',
    'JWT KHÔNG chứa permissions; chỉ chứa roles.',
    'Mọi Kafka event có schemaVersion + eventType versioned.',
    'Listener PHẢI dùng @TransactionalEventListener(AFTER_COMMIT).',
    'Test pyramid: unit 70% / integration 25% / e2e 5%.',
]:
    Bul(doc, b)
page_break(doc)
print("Done Ch3")


# ==================== CHAPTER 4 ====================
H(doc, 'CHƯƠNG 4: USE CASE', level=1)

H(doc, '4.1. Danh sách Actor', level=2)
P(doc, 'Bảng 4.1: Actor', italic=True)
Tbl(doc, ['STT','Actor','Mô tả','Tương tác chính'], [
    ['1','Guest','Khách chưa đăng nhập','Duyệt SP, tìm kiếm, đăng ký.'],
    ['2','Customer','Khách đã đăng nhập','Mua hàng, giỏ, đặt, đánh giá.'],
    ['3','Staff','Nhân viên','Xem đơn, duyệt đổi/trả.'],
    ['4','Admin','Quản trị','Toàn quyền: user, role, voucher.'],
    ['5','System','Service nội bộ','gRPC + Kafka.'],
    ['6','External','SePay, GHN, GHTK','Webhook callback.'],
])

H(doc, '4.2. Tổng hợp Use Case', level=2)
P(doc, 'Bảng 4.2: Tổng hợp Use Case', italic=True)
Tbl(doc, ['Mã UC','Tên','Actor','Module'], [
    ['UC-01','Đăng ký','Guest','Auth'],
    ['UC-02','Đăng nhập','Guest → Customer/Admin','Auth'],
    ['UC-03','Làm mới token','Customer, Admin','Auth'],
    ['UC-04','Đăng xuất','Customer, Admin','Auth'],
    ['UC-05','Đổi mật khẩu','Customer, Admin','Auth'],
    ['UC-06','Thu hồi toàn bộ token','Admin, Customer (self)','Auth'],
    ['UC-07','Hồ sơ + địa chỉ','Customer','User Profile'],
    ['UC-08','Quản lý danh mục','Admin, Staff','Catalog'],
    ['UC-09','Quản lý sản phẩm','Admin, Staff','Catalog'],
    ['UC-10','Tìm kiếm','Guest, Customer','Catalog/Search'],
    ['UC-11','Giỏ hàng','Customer','Cart'],
    ['UC-12','Đặt hàng','Customer','Order (saga)'],
    ['UC-13','Thanh toán SePay','Customer','Payment'],
    ['UC-14','Theo dõi vận chuyển','Customer, Admin','Shipping'],
    ['UC-15','Yêu cầu đổi/trả','Customer','Order'],
    ['UC-16','Duyệt đổi/trả','Admin, Staff','Order'],
    ['UC-17','Refund','Admin','Payment'],
    ['UC-18','Quản lý tồn kho','Admin, System','Inventory'],
    ['UC-19','Đánh giá','Customer','Review'],
    ['UC-20','Duyệt đánh giá','Admin','Review'],
    ['UC-21','Sử dụng voucher','Customer','Voucher'],
    ['UC-22','Tạo voucher','Admin','Voucher'],
    ['UC-23','Wishlist','Customer','Wishlist'],
    ['UC-24','Flash Sale','Customer','Flash Sale'],
    ['UC-25','Tạo Flash Sale','Admin','Flash Sale'],
    ['UC-26','Notification','Customer','Notification'],
    ['UC-27','Analytics','Admin','Analytics'],
    ['UC-28','Webhook SePay','External','Payment'],
    ['UC-29','Webhook GHN','External','Shipping'],
])

H(doc, '4.3. Chi tiết Use Case tiêu biểu', level=2)

P(doc, 'UC-02: Đăng nhập', bold=True)
P(doc, 'Bảng 4.3: Đặc tả UC-02', italic=True)
Tbl(doc, ['Thuộc tính','Chi tiết'], [
    ['Mã','UC-02'],
    ['Tên','Đăng nhập'],
    ['Actor','Guest → Customer/Admin'],
    ['Mục tiêu','Xác thực và phát hành cặp JWT.'],
    ['Tiền điều kiện','Có tài khoản (enabled, nonLocked).'],
    ['Luồng chính', J(
        '1. POST /api/v1/auth/login với usernameOrEmail + rawPassword.',
        '2. Gateway forward Auth Service.',
        '3. LoginHandler lookup user.',
        '4. Validate accountStatus.',
        '5. BCrypt.matches(password, hash).',
        '6. user.recordLogin() → raise UserLoggedInEvent.',
        '7. userRepository.save(user).',
        '8. issueAccessToken (JWT 15p).',
        '9. issueForUser (RefreshToken family 0).',
        '10. save refresh → raise TokenCreatedEvent.',
        '11. Trả AuthResponseDTO + Set-Cookie.',
    )],
    ['Luồng thay thế', J(
        'A1: User không tồn tại → 401.',
        'A2: Password sai → 401.',
        'A3: Locked/disabled → 403.',
    )],
    ['Hậu điều kiện','Cookies; refresh lưu DB; Kafka events.'],
    ['Ghi chú','JWT không chứa permissions.'],
])

P(doc, 'UC-12: Đặt hàng từ giỏ', bold=True)
P(doc, 'Bảng 4.4: Đặc tả UC-12 – Choreography Saga', italic=True)
Tbl(doc, ['Thuộc tính','Chi tiết'], [
    ['Mã','UC-12'],
    ['Actor','Customer'],
    ['Tiền điều kiện','Đã đăng nhập, có giỏ.'],
    ['Luồng chính', J(
        '1. POST /api/v1/orders/from-cart; X-Idempotency-Key BẮT BUỘC.',
        '2. Check Redis cart:idem:{key}.',
        '3. gRPC GetCart.',
        '4. (Nếu voucher) validate.',
        '5. gRPC inventory.reserve → reservationId.',
        '6. REST createPayment → Payment PENDING + payUrl.',
        '7. TX: INSERT Order + outbox.order.placed.',
        '8. Trả 200 + OrderDTO + payUrl.',
        '9. Scheduler publish order.placed → Kafka.',
        '10. Cart consume → clear.',
        '11. Notification → email.',
        '12. SePay webhook → payment.completed.',
        '13. Order consume → PAID → confirm inventory + create shipment.',
    )],
    ['Luồng thay thế', J(
        'A1: Giỏ rỗng → 400.',
        'A2: Hết hàng → 409.',
        'A3: Voucher invalid → 400.',
        'A4: Payment fail → rollback.',
        'A5: Timeout 15p → auto cancel.',
    )],
    ['Hậu điều kiện','Order PENDING; reserved; payment PENDING.'],
    ['Ghi chú','Choreography – không orchestrator. State machine enforce ở domain.'],
])

P(doc, 'UC-15: Yêu cầu đổi/trả', bold=True)
P(doc, 'Bảng 4.5: Đặc tả UC-15', italic=True)
Tbl(doc, ['Thuộc tính','Chi tiết'], [
    ['Mã','UC-15'],
    ['Actor','Customer'],
    ['Tiền điều kiện','Order DELIVERED + within window (7 ngày).'],
    ['Luồng chính', J(
        '1. POST /api/v1/orders/{id}/return-requests.',
        '2. Validate: belong to user + DELIVERED + window.',
        '3. INSERT ReturnRequest PENDING.',
        '4. Publish order.return-requested.',
        '5. Notification → admin + customer.',
        '6. Admin approve → APPROVED.',
        '7. Customer ship đồ về.',
        '8. Admin complete → COMPLETED.',
        '9. Publish order.returned.',
        '10. Payment consume → create Refund PENDING.',
        '11. Admin trigger refund SePay → REFUNDED.',
    )],
    ['Luồng thay thế', J(
        'A1: Chưa DELIVERED → 400.',
        'A2: Quá hạn → 400.',
        'A3: Admin reject → REJECTED.',
    )],
    ['Hậu điều kiện','Refund; restock; order=RETURNED.'],
])
page_break(doc)
print("Done Ch4")


# ==================== CHAPTER 5 ====================
H(doc, 'CHƯƠNG 5: THIẾT KẾ CƠ SỞ DỮ LIỆU', level=1)

H(doc, '5.1. Tổng quan Database-per-Service', level=2)
P(doc, 'Mỗi service nghiệp vụ có PostgreSQL riêng, đảm bảo loose coupling. Không có foreign key cross-service. Tham chiếu là logical UUID/Long.')
P(doc, 'Bảng 5.1: Danh sách cơ sở dữ liệu', italic=True)
Tbl(doc, ['STT','Database','Service','Port','Loại'], [
    ['1','authdb','auth-service','5433','PostgreSQL 16'],
    ['2','catalogdb','catalog-service','5434','PostgreSQL 16'],
    ['3','cartdb','cart-service','5435','PostgreSQL 16'],
    ['4','orderdb','order-service','5436','PostgreSQL 16'],
    ['5','paymentdb','payment-service','5437','PostgreSQL 16'],
    ['6','shippingdb','shipping-service','5438','PostgreSQL 16'],
    ['7','profiledb','user-profile-service','5439','PostgreSQL 16'],
    ['8','inventorydb','inventory-service','5440','PostgreSQL 16'],
    ['9','voucherdb','voucher-service','5441','PostgreSQL 16'],
    ['10','flashsaledb','flash-sale-service','5442','PostgreSQL 16'],
    ['11','notificationdb','notification-service','5443','PostgreSQL 16'],
    ['12','(ES)','search/analytics','9200','Elasticsearch 8.13'],
    ['13','(Redis)','tất cả','6379','Redis 7'],
])
Note(doc, 'Mỗi PostgreSQL có user/password riêng. Dev có thể dùng H2 in-memory. Production luôn PostgreSQL 16.')

H(doc, '5.2. Auth Service Database', level=2)
P(doc, 'Database authdb gồm 7 bảng: users, roles, permissions, user_roles, role_permissions, refresh_tokens, token_revocations.')

P(doc, 'Bảng 5.2: Schema users', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Ràng buộc','Mô tả'], [
    ['1','id','UUID','PK','UUID v4.'],
    ['2','username','VARCHAR(50)','UNIQUE','3-50 ký tự.'],
    ['3','email','VARCHAR(100)','UNIQUE','Lowercase + regex.'],
    ['4','password','VARCHAR(255)','NOT NULL','BCrypt cost 12.'],
    ['5','first_name','VARCHAR(50)','NULL','Tên.'],
    ['6','last_name','VARCHAR(50)','NULL','Họ.'],
    ['7','enabled','BOOLEAN','DEFAULT true','Active flag.'],
    ['8','account_non_expired','BOOLEAN','DEFAULT true',''],
    ['9','account_non_locked','BOOLEAN','DEFAULT true',''],
    ['10','credentials_non_expired','BOOLEAN','DEFAULT true',''],
    ['11','token_version','INT','DEFAULT 1','Bump khi đổi pass/revoke all.'],
    ['12','last_login','TIMESTAMPTZ','NULL',''],
    ['13','created_at, updated_at','TIMESTAMPTZ','NOT NULL','Audit.'],
])

P(doc, 'Bảng 5.3: Schema refresh_tokens', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','UUID PK','Token ID.'],
    ['2','value','VARCHAR(64) UNIQUE','Token value.'],
    ['3','user_id','UUID','Owner.'],
    ['4','family','VARCHAR(64)','Family ID.'],
    ['5','generation','INT','Rotation count.'],
    ['6','expiry','TIMESTAMPTZ','createdAt + 7d.'],
    ['7','revoked','BOOLEAN','Revoke flag.'],
    ['8','reason','VARCHAR(32)','NORMAL/REUSE/LOGOUT.'],
    ['9','revoked_at','TIMESTAMPTZ','Revoke time.'],
    ['10','replaced_by_id','UUID','Successor.'],
    ['11','created_at','TIMESTAMPTZ','Audit.'],
])

P(doc, 'Bảng 5.4: Schema token_revocations (blacklist)', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','VARCHAR(64) PK','JWT jti.'],
    ['2','user_id','UUID','Owner.'],
    ['3','expires_at','TIMESTAMPTZ','Token expiry.'],
    ['4','reason','VARCHAR(32)','LOGOUT/ADMIN_REVOKE/PASSWORD_CHANGED.'],
    ['5','revoked_at','TIMESTAMPTZ','Revoke time.'],
])

H(doc, '5.3. Catalog Service Database', level=2)
P(doc, 'catalogdb gồm 6 bảng: categories, products, variants, attrs, attr_vals, variant_attrs.')

P(doc, 'Bảng 5.5: Schema products (SPU)', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','BIGSERIAL PK','Khóa chính.'],
    ['2','name','VARCHAR(200)','Tên SP.'],
    ['3','slug','VARCHAR(200) UNIQUE','URL-friendly.'],
    ['4','description','TEXT','Mô tả.'],
    ['5','cat_id','BIGINT','FK categories.'],
    ['6','brand','VARCHAR(100)','Brand.'],
    ['7','thumb','VARCHAR(500)','Thumbnail URL.'],
    ['8','images','JSONB','Ảnh.'],
    ['9','status','VARCHAR(20)','ACTIVE/INACTIVE/DRAFT.'],
    ['10','version','BIGINT','Optimistic lock.'],
])

P(doc, 'Bảng 5.6: Schema variants (SKU)', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','BIGSERIAL PK',''],
    ['2','product_id','BIGINT','FK products.'],
    ['3','sku','VARCHAR(100) UNIQUE','Mã SKU.'],
    ['4','price','DECIMAL(12,2)','Giá bán.'],
    ['5','cost','DECIMAL(12,2)','Giá vốn.'],
    ['6','sale_price','DECIMAL(12,2)','Sale price.'],
    ['7','qty','INT','Stock cache.'],
    ['8','weight','DECIMAL(10,3)','kg.'],
    ['9','status','VARCHAR(20)','ACTIVE/INACTIVE.'],
])

H(doc, '5.4. Order & Payment Database', level=2)
P(doc, 'orderdb gồm 6 bảng: orders, order_items, order_status_history, return_requests, outbox_events, order_idempotency.')

P(doc, 'Bảng 5.7: Schema orders', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','UUID PK',''],
    ['2','order_number','VARCHAR(32) UNIQUE','HIE2026051700001.'],
    ['3','user_id','UUID','Customer.'],
    ['4','subtotal','DECIMAL(12,2)',''],
    ['5','discount','DECIMAL(12,2)','Voucher.'],
    ['6','shipping_fee','DECIMAL(12,2)',''],
    ['7','total','DECIMAL(12,2)',''],
    ['8','voucher_code','VARCHAR(32)',''],
    ['9','status','VARCHAR(32)','State machine.'],
    ['10','shipping_address','JSONB','Snapshot.'],
    ['11','reservation_id','UUID','From inventory.'],
])

P(doc, 'Bảng 5.8: Schema outbox_events', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','UUID PK',''],
    ['2','aggregate_type','VARCHAR(50)','Order/ReturnRequest.'],
    ['3','aggregate_id','UUID',''],
    ['4','event_type','VARCHAR(64)','order.placed.v1.'],
    ['5','payload','JSONB','Versioned body.'],
    ['6','status','VARCHAR(16)','PENDING/PUBLISHED/FAILED.'],
    ['7','attempts','INT','Retry count.'],
    ['8','created_at, published_at','TIMESTAMPTZ',''],
])

P(doc, 'Scheduler mỗi 1s: SELECT FROM outbox_events WHERE status=PENDING LIMIT 100; publish vào Kafka. Sau N retry → DLQ.', italic=True)

P(doc, 'Bảng 5.9: Schema payments', italic=True)
Tbl(doc, ['STT','Cột','Kiểu','Mô tả'], [
    ['1','id','UUID PK',''],
    ['2','order_id','UUID','Logical FK.'],
    ['3','order_number','VARCHAR(32)','SePay match.'],
    ['4','amount','DECIMAL(12,2)',''],
    ['5','method','VARCHAR(32)','SEPAY_QR/COD.'],
    ['6','status','VARCHAR(16)','PENDING/COMPLETED/FAILED/REFUNDED.'],
    ['7','pay_url, qr_code_url','TEXT',''],
    ['8','transaction_id','VARCHAR(64)','From SePay.'],
    ['9','paid_at','TIMESTAMPTZ',''],
])

H(doc, '5.5. Inventory Database', level=2)
P(doc, 'inventorydb tách 3 bảng theo concern:')
Bul(doc, 'inventories: cân bằng hiện tại (1 row mỗi variant).')
Bul(doc, 'stock_reservations: reservation tạm (TTL).')
Bul(doc, 'stock_movements: audit append-only.')

P(doc, 'Bảng 5.10: Schema inventory', italic=True)
Tbl(doc, ['Bảng','Cột chính','Mục đích'], [
    ['inventories','variant_id PK, available INT, reserved INT','Cân bằng hiện tại.'],
    ['stock_reservations','id PK, variant_id, order_id, qty, status, expires_at','Giữ hàng TTL 15p.'],
    ['stock_movements','id PK, variant_id, qty_delta, type, ref_id','Audit – chỉ INSERT.'],
])

H(doc, '5.6. Các database còn lại', level=2)
P(doc, 'Bảng 5.11: Schema các service nhỏ', italic=True)
Tbl(doc, ['Service','Bảng chính','Mô tả'], [
    ['cart','carts, cart_items','Giỏ + item.'],
    ['voucher','vouchers, voucher_usage_records','Voucher + usage tracking.'],
    ['flash-sale','flash_sales, flash_sale_items','Campaign + items + quota.'],
    ['notification','notifications, notification_templates','In-app + template.'],
    ['shipping','shipments, tracking_events','Shipment + tracking.'],
    ['profile','profiles, addresses','Hồ sơ + địa chỉ.'],
    ['review','reviews, review_reports','Review + báo cáo.'],
    ['wishlist','wishlists, wishlist_items','Wishlist N-N.'],
])

H(doc, '5.7. Chỉ mục quan trọng', level=2)
P(doc, 'Bảng 5.12: Indexes quan trọng', italic=True)
Tbl(doc, ['Service','Bảng','Index','Cột','Mục đích'], [
    ['Auth','users','uk_username','username','Unique'],
    ['Auth','users','uk_email','email','Unique'],
    ['Auth','refresh_tokens','idx_rt_user','user_id','Revoke all'],
    ['Auth','refresh_tokens','idx_rt_family','family','Reuse detect'],
    ['Auth','token_revocations','idx_exp','expires_at','Cleanup'],
    ['Catalog','products','idx_prod_slug','slug','URL lookup'],
    ['Catalog','products','idx_prod_cat','cat_id','Filter'],
    ['Catalog','variants','uk_sku','sku','Unique'],
    ['Order','orders','idx_orders_user','user_id','User history'],
    ['Order','orders','uk_order_number','order_number','SePay match'],
    ['Order','outbox_events','idx_outbox_status','status, created_at','Scheduler'],
    ['Cart','carts','uk_cart_user','user_id','1-1'],
    ['Inventory','stock_reservations','idx_resv_expiry','expires_at','Auto-release'],
    ['Review','reviews','idx_review_product','product_id','SP reviews'],
    ['Voucher','voucher_usage','uk_voucher_user','voucher_id, user_id','Unique/user'],
])

H(doc, '5.8. Migration với Flyway', level=2)
Bul(doc, 'V1__init_<service>_schema.sql – schema ban đầu.')
Bul(doc, 'V2__add_<feature>.sql – thay đổi mới.')
Bul(doc, 'R__<name>.sql – repeatable (seed static).')
Code(doc, J(
    'spring:',
    '  flyway:',
    '    enabled: true',
    '    locations: classpath:db/migration',
    '    baseline-on-migrate: true',
    '  jpa:',
    '    hibernate:',
    '      ddl-auto: validate',
))
Note(doc, 'KHÔNG sửa migration đã merge – luôn tạo file mới.')
page_break(doc)
print("Done Ch5")


# ==================== CHAPTER 6 ====================
H(doc, 'CHƯƠNG 6: THIẾT KẾ API', level=1)

H(doc, '6.1. Quy ước chung', level=2)
for b in [
    'Base URL: http://localhost:8080 (qua API Gateway). Production: https://api.hieu.vn.',
    'Content-Type: application/json (UTF-8).',
    'Auth: HttpOnly cookie ACCESS+REFRESH cho browser; Bearer cho gRPC/Postman.',
    'Response: { success, message, data: T }.',
    'Error: { code: "AUTH-401", status, message, timestamp, traceId }.',
    'Pagination: cursor-based ?cursor=&size=20 → {items, nextCursor, pageSize, totalElements}.',
    'Idempotency: X-Idempotency-Key (UUIDv4) BẮT BUỘC cho POST /orders, /payments, /refunds.',
    'Versioning: dự kiến /api/v1/...',
    'Correlation: X-Request-Id propagate qua MDC.',
    'CORS: chỉ allow origin từ allowedOriginPatterns.',
]:
    Bul(doc, b)

H(doc, '6.2. Auth Service API', level=2)
P(doc, 'Public:', bold=True)
Tbl(doc, ['Method','Endpoint','Mô tả'], [
    ['POST','/api/v1/auth/register','Đăng ký + Set-Cookie.'],
    ['POST','/api/v1/auth/login','Đăng nhập + Set-Cookie.'],
    ['POST','/api/v1/auth/refresh','Rotation.'],
    ['GET','/api/v1/auth/validate','Validate JWT (legacy).'],
])
P(doc, 'Protected:', bold=True)
Tbl(doc, ['Method','Endpoint','Mô tả','Quyền'], [
    ['POST','/api/v1/auth/logout','Logout','User+'],
    ['POST','/api/v1/auth/change-password','Đổi mật khẩu','User'],
    ['POST','/api/v1/auth/revoke-all','Revoke all','Admin/self'],
    ['POST','/api/v1/auth/users/{id}/lock','LOCK','Admin'],
    ['POST','/api/v1/auth/users/{id}/unlock','UNLOCK','Admin'],
    ['POST','/api/v1/auth/users/{id}/disable','DISABLE','Admin'],
    ['POST','/api/v1/auth/users/{id}/enable','ENABLE','Admin'],
    ['POST','/api/v1/auth/users/{id}/roles/{role}','Gán role','Admin'],
    ['DELETE','/api/v1/auth/users/{id}/roles/{role}','Bỏ role','Admin'],
    ['GET','/api/v1/auth/users/{id}','User by ID','Admin/self'],
    ['GET','/api/v1/auth/users','List users','Admin'],
    ['GET','/api/v1/auth/check-permission','Check permission','Admin'],
    ['GET','/api/v1/auth/check-role','Check role','Admin'],
])

H(doc, '6.3. User Profile API', level=2)
Tbl(doc, ['Method','Endpoint','Mô tả'], [
    ['GET','/api/v1/profiles/me','Profile.'],
    ['PUT','/api/v1/profiles/me','Update.'],
    ['GET','/api/v1/profiles/me/addresses','Addresses.'],
    ['POST','/api/v1/profiles/me/addresses','Thêm.'],
    ['PUT','/api/v1/profiles/me/addresses/{id}','Update.'],
    ['DELETE','/api/v1/profiles/me/addresses/{id}','Xóa.'],
    ['POST','/api/v1/profiles/me/addresses/{id}/default','Set default.'],
])

H(doc, '6.4. Catalog API', level=2)
Tbl(doc, ['Method','Endpoint','Mô tả','Quyền'], [
    ['POST','/api/v1/categories','Tạo','Admin'],
    ['PUT','/api/v1/categories/{id}','Update','Admin'],
    ['GET','/api/v1/categories','List','Public'],
    ['GET','/api/v1/categories/{id}','Detail','Public'],
    ['POST','/api/v1/products','Tạo SP','Admin/Staff'],
    ['PUT','/api/v1/products/{id}','Update','Admin/Staff'],
    ['GET','/api/v1/products','List','Public'],
    ['GET','/api/v1/products/{id}','Detail','Public'],
    ['GET','/api/v1/products/slug/{slug}','By slug','Public'],
    ['GET','/api/v1/products/search?q=','Search','Public'],
    ['POST','/api/v1/products/batch','Batch IDs','Public'],
    ['POST','/api/v1/variants','Tạo variant','Admin/Staff'],
    ['POST','/api/v1/attrs','Tạo attribute','Admin'],
])

H(doc, '6.5 - 6.14. API các service còn lại', level=2)
Tbl(doc, ['Service','Method','Endpoint','Mô tả'], [
    ['Inventory','GET','/api/v1/inventory/{variantId}','Xem stock'],
    ['Inventory','POST','/api/v1/inventory/reserve','Giữ hàng (System)'],
    ['Inventory','POST','/api/v1/inventory/release','Release'],
    ['Inventory','POST','/api/v1/inventory/confirm','Confirm'],
    ['Cart','GET','/api/v1/cart','Xem giỏ'],
    ['Cart','POST','/api/v1/cart/items','Thêm item (idem)'],
    ['Cart','PUT','/api/v1/cart/items/{id}','Update qty'],
    ['Cart','DELETE','/api/v1/cart/items/{id}','Xóa item'],
    ['Order','POST','/api/v1/orders/from-cart','Tạo đơn (idem)'],
    ['Order','GET','/api/v1/orders/{id}','Chi tiết'],
    ['Order','GET','/api/v1/orders','User orders'],
    ['Order','GET','/api/v1/admin/orders','All orders (Admin)'],
    ['Order','POST','/api/v1/orders/{id}/cancel','Hủy'],
    ['Order','POST','/api/v1/orders/{id}/return-requests','Yêu cầu trả'],
    ['Order','POST','/api/v1/return-requests/{id}/approve','Duyệt (Admin)'],
    ['Order','POST','/api/v1/return-requests/{id}/complete','Hoàn tất'],
    ['Payment','GET','/api/v1/payments/{id}','Chi tiết'],
    ['Payment','GET','/api/v1/payments/by-order/{orderId}','Theo order'],
    ['Payment','POST','/api/v1/payments/webhook/sepay','SePay webhook (HMAC)'],
    ['Payment','POST','/api/v1/payments/{id}/refund','Refund (Admin)'],
    ['Shipping','POST','/api/v1/shipments','Tạo (System)'],
    ['Shipping','GET','/api/v1/shipments/{id}','Chi tiết + tracking'],
    ['Shipping','POST','/api/v1/shipments/webhook/ghn','GHN webhook'],
    ['Voucher','POST','/api/v1/vouchers','Tạo (Admin)'],
    ['Voucher','GET','/api/v1/vouchers','Active list'],
    ['Voucher','POST','/api/v1/vouchers/validate','Validate (System)'],
    ['Voucher','POST','/api/v1/vouchers/redeem','Apply (System)'],
    ['Voucher','POST','/api/v1/vouchers/release','Compensation'],
    ['Flash Sale','POST','/api/v1/flash-sales','Tạo (Admin)'],
    ['Flash Sale','GET','/api/v1/flash-sales/active','Active'],
    ['Flash Sale','GET','/api/v1/flash-sales/{id}/availability','Quota còn'],
    ['Flash Sale','POST','/api/v1/flash-sales/{id}/participate','Tham gia (Lua)'],
    ['Notification','GET','/api/v1/notifications','List user'],
    ['Notification','POST','/api/v1/notifications/{id}/read','Mark read'],
    ['Notification','GET','/api/v1/notifications/sse','SSE push'],
    ['Review','POST','/api/v1/reviews','Tạo (đã mua)'],
    ['Review','GET','/api/v1/reviews/product/{id}','By product'],
    ['Review','POST','/api/v1/reviews/{id}/approve','Duyệt (Admin)'],
    ['Wishlist','POST','/api/v1/wishlist','Thêm'],
    ['Wishlist','GET','/api/v1/wishlist','Xem list'],
    ['Search','GET','/api/v1/search?q=&category=','Filter search'],
    ['Search','GET','/api/v1/search/autocomplete','Gợi ý'],
])
page_break(doc)
print("Done Ch6")


# ==================== CHAPTER 7 ====================
H(doc, 'CHƯƠNG 7: LUỒNG XỬ LÝ NGHIỆP VỤ', level=1)

H(doc, '7.1. Luồng đăng ký', level=2)
Code(doc, J(
    'Client          API Gateway          Auth Service          Database          Kafka',
    '  |--POST /register->|                     |                    |                |',
    '  |                  |--Forward---------->|                    |                |',
    '  |                  |                     |--Check username/email->|             |',
    '  |                  |                     |--BCrypt(rawPwd, 12) |                |',
    '  |                  |                     |--User.register()    |                |',
    '  |                  |                     |   raise UserCreatedEvent             |',
    '  |                  |                     |--save User -------->|                |',
    '  |                  |                     |--assign ROLE_CUSTOMER                |',
    '  |                  |                     |--issueAccessToken (15p)              |',
    '  |                  |                     |--issue RefreshToken (7d, gen 0)      |',
    '  |                  |                     |--save refresh ----->|                |',
    '  |                  |                     |--publish auth.user.created.v1 ------>|',
    '  |                  |                     |--publish auth.token.created.v1 ----->|',
    '  |                  |<--AuthResponseDTO --|                    |                |',
    '  |<--Set-Cookie ----|                     |                    |                |',
))
Cap(doc, 'Hình 7.1: Đăng ký')

H(doc, '7.2. Luồng đăng nhập', level=2)
P(doc, 'Phát hành 2 token: ACCESS (15p) + REFRESH (7d). HS256 với JWT_SECRET. Mỗi service tự verify.')
Code(doc, J(
    'Client          API Gateway          Auth Service          Database          Kafka',
    '  |--POST /login --->|                     |                    |                |',
    '  |                  |--Forward---------->|                    |                |',
    '  |                  |                     |--lookup user                         |',
    '  |                  |                     |--ensureAuthenticatable()             |',
    '  |                  |                     |--BCrypt.matches                      |',
    '  |                  |                     |--user.recordLogin()                  |',
    '  |                  |                     |   raise UserLoggedInEvent            |',
    '  |                  |                     |--save user -------->|                |',
    '  |                  |                     |--issueAccessToken (JWT)              |',
    '  |                  |                     |   { jti, userId, email, roles,       |',
    '  |                  |                     |     tokenVersion, iat, exp }         |',
    '  |                  |                     |--issueForUser (RefreshToken root)    |',
    '  |                  |                     |--save refresh ----->|                |',
    '  |                  |                     |--publish events----------------------->|',
    '  |                  |<--AuthResponseDTO --|                    |                |',
    '  |<--Set-Cookie ----|                     |                    |                |',
))
Cap(doc, 'Hình 7.2: Đăng nhập (JWT không chứa permissions)')

H(doc, '7.3. Request qua Gateway', level=2)
Code(doc, J(
    'Client         API Gateway       JWT Filter        Downstream Service',
    '  |--GET /api/v1/xxx ->|                |                     |',
    '  | (Cookie ACCESS) |--Route match-->|                     |',
    '  |                 |--Add trace ID--|                     |',
    '  |                 |                |--Verify HS256       |',
    '  |                 |                |--Check exp          |',
    '  |                 |                |--Extract claims     |',
    '  |                 |<--Add headers--|                     |',
    '  |                 |  X-User-Id, X-Username, X-Request-Id |',
    '  |                 |--Forward ------------------------- >|',
    '  |                 |                                      |--local JWT verify',
    '  |                 |                                      |--check blacklist (Redis)',
    '  |                 |                                      |--check tokenVersion',
    '  |                 |                                      |--@PreAuthorize',
    '  |                 |                                      |--process logic',
    '  |                 |<--Response--------------------------|',
    '  |<--200 + data ---|                                      |',
))
Cap(doc, 'Hình 7.3: Request qua Gateway')

H(doc, '7.4. addItem trong Cart', level=2)
P(doc, 'Cart-service round-trip tới catalog qua gRPC mỗi addItem/updateItem để xác thực status + stock thực tế.')
Code(doc, J(
    'Client          Cart Service          Redis           Catalog (gRPC)         Cart DB',
    '  |--POST addItem ---->|                                                              |',
    '  |  X-Idem-Key: k     |--GET cart:idem:k                                             |',
    '  |     (if hit) <-----|<--Cached---------                                            |',
    '  |                    |--gRPC GetProduct(productId) -------->|                       |',
    '  |                    |<--Product (status, variants) -------|                       |',
    '  |                    |--Validate status = ACTIVE                                    |',
    '  |                    |--Find variant + Validate stock >= qty                        |',
    '  |                    |--@Retryable upsert cart_item ---------------------------->|',
    '  |                    |<--Updated -------------------------------------------------|',
    '  |                    |--SET cart:idem:k (TTL 24h)----------->|                       |',
    '  |<--200 CartDTO -----|                                                              |',
))
Cap(doc, 'Hình 7.4: addItem cart-service')

H(doc, '7.5. Saga đặt hàng (choreography)', level=2)
P(doc, 'Saga distributed: order, voucher, inventory, payment. Order-service local TX, phát event qua outbox, downstream tự subscribe.')
P(doc, 'Compensation:')
Bul(doc, 'Reserve fail → release qua order.cancelled.')
Bul(doc, 'Voucher applied → release qua API khi cancelled.')
Bul(doc, 'Payment PENDING → auto cancel sau N phút nếu không pay.')
Bul(doc, 'Shipment chỉ tạo sau payment.completed.')
Code(doc, J(
    'Client    Order Svc      Voucher    Inventory    Payment    Cart   Kafka    Shipping  Notif',
    '  |-POST-->|              |            |            |          |       |          |          |',
    '  |   from-|--gRPC GetCart-----------------------------------|       |          |          |',
    '  |   cart |<--CartDTO -------------------------------------|         |          |          |',
    '  |        |--REST validate->          |            |          |       |          |          |',
    '  |        |<--discount---|             |            |          |       |          |          |',
    '  |        |--gRPC reserve(items)----->|            |          |       |          |          |',
    '  |        |<--reservationId-----------|            |          |       |          |          |',
    '  |        |--REST createPayment-------------------->|         |       |          |          |',
    '  |        |<--payUrl + qrCode---------------------|           |       |          |          |',
    '  |        |--TX: INSERT order + outbox.order.placed ---|     |       |          |          |',
    '  |<-200 + payUrl --|     |            |            |          |       |          |          |',
    '  |        |--scheduler poll outbox + publish ----------------------->|         |          |',
    '  |        |              |            |            |          |<--clear cart---|          |',
    '  |        |              |            |            |          |       |--placed-->-->email',
    '  |(later) |              |            |            |--webhook|       |          |          |',
    '  |        |              |            |            |--COMPLETED+publish payment.completed-->|',
    '  |        |<--consume payment.completed -----------------------|     |          |          |',
    '  |        |--inventory.confirm------->|            |          |       |          |          |',
    '  |        |--create Shipment----------------------|----------|------|->        |          |',
    '  |        |--publish order.confirmed------------------------------|--->        |--->push   |',
))
Cap(doc, 'Hình 7.5: Saga đặt hàng')

H(doc, '7.6. Luồng thanh toán SePay', level=2)
P(doc, 'User chọn SEPAY_QR → payment-service tạo Payment PENDING + payUrl + qrCodeUrl. User quét VietQR + chuyển khoản → SePay gửi webhook → verify HMAC + match orderNumber → mark COMPLETED + publish Kafka.')
Code(doc, J(
    'User       FE          Payment Svc          SePay Gateway        Bank        Order Svc',
    ' |--checkout->|              |                       |                |              |',
    ' |          |--/from-cart-->                         |                |              |',
    ' |          |<--payUrl + qr -|                       |                |              |',
    ' |--scan QR -->-->-->-->-->-->-->-->-->-->-->-->-->|--bank notif--->|              |',
    ' |                                                  |<--credited ---|              |',
    ' |                          |<--webhook --|                          |              |',
    ' |                          |--verify HMAC                            |              |',
    ' |                          |--match orderNumber                      |              |',
    ' |                          |--mark COMPLETED                         |              |',
    ' |                          |--publish payment.completed ----------------->          |',
    ' |                                                                                  |--inventory.confirm',
    ' |                                                                                  |--create shipment',
))
Cap(doc, 'Hình 7.6: Thanh toán SePay')

H(doc, '7.7. Luồng đổi/trả', level=2)
Code(doc, J(
    'Customer    Order Svc        Kafka        Notification        Admin         Payment Svc',
    '  |--POST /return-requests------->|                |                |                |',
    '  |            |--validate cond.                  |                |                |',
    '  |            |--INSERT ReturnRequest PENDING                     |                |',
    '  |            |--publish order.return-requested ->|                |                |',
    '  |            |                  |                |--email admin + customer        |',
    '  |<-200-------|                  |                |                |                |',
    '  |            |<--POST /approve --------------------------------|     |                |',
    '  |            |--status=APPROVED                                |    |                |',
    '  |            |--publish order.return-approved -------------->|    |                |',
    '  |                                              (ships back)                         |',
    '  |            |<--POST /complete ---------------------------------|                |',
    '  |            |--status=COMPLETED                                                   |',
    '  |            |--publish order.returned ---------------------------------->         |',
    '  |                                                                                  |--Refund PENDING',
    '  |            |<--POST /process-refund (admin) -----------------|                  |',
    '  |                                                                                  |--SePay refund API',
    '  |                                                                                  |--mark REFUNDED',
))
Cap(doc, 'Hình 7.7: Đổi/trả')

H(doc, '7.8. Flash Sale (atomic Lua)', level=2)
Code(doc, J(
    'Admin       FE           Flash-Sale Svc       Redis           Order Svc      Kafka',
    '  |--create->|                  |                   |                |              |',
    '  |<--id-----|                  |                   |                |              |',
    '  |              (at startTime) |                   |                |              |',
    '  |                             |--scheduler activate                |              |',
    '  |                             |--SET fs:{id}:qty:{vId}=quota ----->|              |',
    '  |          |--GET /availability                   |                |              |',
    '  |          |<--remaining ------|                  |                |              |',
    '  |          |--POST /participate->                 |                |              |',
    '  |                             |--EVAL Lua atomic:                  |              |',
    '  |                             |   local q = redis.call(GET,k)      |              |',
    '  |                             |   if tonumber(q)>0 then            |              |',
    '  |                             |     redis.call(DECR,k)             |              |',
    '  |                             |     return OK                      |              |',
    '  |                             |   else return EXHAUSTED            |              |',
    '  |                             |<--ok or EXHAUSTED                  |              |',
    '  |          |<--flashPrice + ttl                                    |              |',
    '  |          |--checkout (within 15min) --------------------------> |              |',
    '  |                                                                  |--use flashPrice',
    '  |                                                                  |--place order',
))
Cap(doc, 'Hình 7.8: Flash Sale')

H(doc, '7.9. State machine Order', level=2)
Code(doc, J(
    '                       +---------+',
    '                       | PENDING |',
    '                       +----+----+',
    '                            | reserve OK',
    '                            v',
    '                +---------------------+',
    '                | INVENTORY_RESERVED  |',
    '                +----------+----------+',
    '                           | create payment',
    '                           v',
    '                  +-------------------+   payment.failed',
    '                  |  PAYMENT_PENDING  |--------+',
    '                  +-------+-----------+        |',
    '                          | payment.completed   |',
    '                          v                    v',
    '                       +------+          +-----------+',
    '                       | PAID |          | CANCELLED |',
    '                       +--+---+          +-----------+',
    '                          | confirmed         ^',
    '                          v                   | cancel',
    '                     +-----------+            |',
    '                     | CONFIRMED |            |',
    '                     +-----+-----+            |',
    '                           | create shipment  |',
    '                           v                  |',
    '                      +---------+             |',
    '                      | SHIPPED |             |',
    '                      +----+----+             |',
    '                           | delivered        |',
    '                           v                  |',
    '                     +-----------+            |',
    '                     | DELIVERED |---+        |',
    '                     +-----------+   |        |',
    '                           |  return v        |',
    '                           v   +---------------+',
    '                  +---------------+ RETURN_REQUESTED',
    '                  |   APPROVED    |',
    '                  +-------+-------+',
    '                          v',
    '                     +----------+',
    '                     | RETURNED |',
    '                     +----------+',
))
Cap(doc, 'Hình 7.9: State machine Order')

H(doc, '7.10. 4-layer token validation', level=2)
Code(doc, J(
    'Client             Service          Auth Service        Redis      Database',
    '  |--Request + JWT -->|                  |                |            |',
    '  |                   |--1. Verify HS256 sig                          |',
    '  |                   |--2. Check exp                                  |',
    '  |                   |--3. Check blacklist (Redis EXISTS) ---------> |',
    '  |                   |    (fallback DB if Redis down) -----------------> |',
    '  |                   |<--EXISTS / NOT-EXIST -------------------------|',
    '  |                   |--4. Check tokenVersion match user                 |',
    '  |                   |    (Caffeine → gRPC → DB)                         |',
    '  |                   |<--match / mismatch -----------------------------|',
    '  |  any layer fail   |                                                  |',
    '  |<--401 ------------|                                                  |',
    '  |<--200 process ----|                                                  |',
))
Cap(doc, 'Hình 7.10: 4-layer validation (defense in depth)')
page_break(doc)
print("Done Ch7")


# ==================== CHAPTER 8 ====================
H(doc, 'CHƯƠNG 8: GIAO TIẾP GIỮA CÁC SERVICE', level=1)

H(doc, '8.1. Ba mô hình giao tiếp', level=2)
P(doc, 'Bảng 8.1: REST vs gRPC vs Kafka', italic=True)
Tbl(doc, ['Tiêu chí','REST','gRPC','Kafka'], [
    ['Phương thức','HTTP/1.1 + JSON','HTTP/2 + Protobuf','TCP + binary'],
    ['Đồng bộ','Sync','Sync','Async'],
    ['Latency','~50ms','~10-20ms','~100ms'],
    ['Throughput','Trung bình','Cao','Rất cao'],
    ['Use case','FE↔Gateway','Inter-service sync','Event streaming'],
    ['Schema','OpenAPI','Protobuf','JSON + schemaVersion'],
])

H(doc, '8.2. Ma trận giao tiếp', level=2)
P(doc, 'Bảng 8.2: Ma trận giao tiếp', italic=True)
Tbl(doc, ['Từ','Đến','Giao thức','Mục đích'], [
    ['Client','API Gateway','REST','Mọi request.'],
    ['Gateway','All services','REST','Routing.'],
    ['Cart','Catalog','gRPC','GetProduct, CheckStock.'],
    ['Order','Cart','gRPC','GetCart.'],
    ['Order','Inventory','gRPC','Reserve/Release/Confirm.'],
    ['Order','Voucher','REST','Validate.'],
    ['Order','Payment','REST','Khởi tạo Payment.'],
    ['Order','Catalog','gRPC','Lookup snapshot.'],
    ['Auth','All','gRPC','ValidateToken.'],
    ['Payment','Kafka','Producer','payment.completed/failed.'],
    ['Order','Kafka','Outbox','order.placed/cancelled/confirmed/returned.'],
    ['Cart','Kafka','Consumer','order.placed → clear.'],
    ['Notification','Kafka','Consumer','All events → email/SSE.'],
    ['Search','Kafka','Consumer','catalog.product.* → reindex.'],
    ['Analytics','Kafka','Consumer','All events → log.'],
    ['Shipping','Kafka','Consumer','order.confirmed → create shipment.'],
    ['All','Redis','TCP','Cache, blacklist, idem.'],
    ['All','Eureka','REST','Registration.'],
    ['SePay','Payment','HTTP webhook','Notify.'],
    ['GHN/GHTK','Shipping','HTTP webhook','Tracking.'],
])

H(doc, '8.3. Danh mục Kafka Topics', level=2)
P(doc, 'Convention: <service>.<aggregate>.<event>.v<n>')
P(doc, 'Bảng 8.3: Kafka topics (~23)', italic=True)
Tbl(doc, ['Topic','Producer','Consumer chính'], [
    ['auth.user.events.v1','auth','analytics, notification'],
    ['auth.session.events.v1','auth','analytics, audit'],
    ['catalog.product.created.v1','catalog','search, analytics'],
    ['catalog.product.updated.v1','catalog','search, cart'],
    ['catalog.variant.stock-changed.v1','inventory','catalog, notification'],
    ['order.placed.v1','order (outbox)','cart, notification, analytics'],
    ['order.payment-pending.v1','order','notification'],
    ['order.confirmed.v1','order','shipping, notification'],
    ['order.cancelled.v1','order','inventory, voucher, notification'],
    ['order.shipped.v1','order','notification'],
    ['order.delivered.v1','order','notification, review'],
    ['order.return-requested.v1','order','notification'],
    ['order.return-approved.v1','order','notification'],
    ['order.returned.v1','order','payment, inventory'],
    ['payment.completed.v1','payment','order'],
    ['payment.failed.v1','payment','order'],
    ['payment.refunded.v1','payment','order, notification'],
    ['shipping.created.v1','shipping','order, notification'],
    ['shipping.in-transit.v1','shipping','notification'],
    ['shipping.delivered.v1','shipping','order, notification'],
    ['voucher.created.v1','voucher','analytics'],
    ['flashsale.activated.v1','flash-sale','notification'],
    ['analytics.event.v1','all','analytics'],
])

H(doc, '8.4. gRPC Service Definitions', level=2)
P(doc, 'Auth Service gRPC (port 9191):', bold=True)
Code(doc, J(
    'service AuthService {',
    '  rpc ValidateToken (TokenRequest) returns (TokenValidation);',
    '  rpc GetUserById (UserIdRequest) returns (UserSummary);',
    '  rpc IncrementTokenVersion (UserIdRequest) returns (Empty);',
    '  rpc CheckPermission (PermissionRequest) returns (BoolResponse);',
    '}',
    'message TokenValidation { bool valid; string userId; repeated string roles; }',
    'message UserSummary { string id; string username; string email; bool enabled; repeated string roles; }',
))

P(doc, 'Catalog Service gRPC (port 9193):', bold=True)
Code(doc, J(
    'service CatalogService {',
    '  rpc GetProduct (GetProductRequest) returns (Product);',
    '  rpc GetVariantBySku (SkuRequest) returns (Variant);',
    '  rpc CheckStock (StockRequest) returns (StockResponse);',
    '  rpc BatchGetProducts (BatchProductRequest) returns (ProductList);',
    '}',
    'message Product { int64 id; string name; string slug; string status; repeated Variant variants; }',
    'message Variant { int64 id; string sku; double price; double salePrice; int32 quantity; string status; }',
))

P(doc, 'Cart Service gRPC (port 9194):', bold=True)
Code(doc, J(
    'service CartService { rpc GetCart (UserIdRequest) returns (Cart); }',
    'message Cart { string userId; repeated CartItem items; double total; }',
    'message CartItem { int64 productId; int64 variantId; string sku; int32 quantity; double unitPrice; }',
))

H(doc, '8.5. Outbox Pattern', level=2)
P(doc, 'Order-service áp dụng Outbox Pattern để at-least-once delivery. Event INSERT cùng transaction nghiệp vụ. Scheduler quét outbox publish lên Kafka.')
Code(doc, J(
    '@Scheduled(fixedDelay = 1000)',
    '@Transactional',
    'public void publishOutbox() {',
    '    List<OutboxEvent> events = outboxRepo.findByStatusOrderByCreatedAt(PENDING, PAGE_100);',
    '    for (OutboxEvent e : events) {',
    '        try {',
    '            kafkaTemplate.send(e.getEventType(), e.getAggregateId(), e.getPayload()).get();',
    '            e.markPublished();',
    '        } catch (Exception ex) {',
    '            e.incrementAttempts();',
    '            if (e.getAttempts() > 10) e.moveToDLQ();',
    '        }',
    '        outboxRepo.save(e);',
    '    }',
    '}',
))

H(doc, '8.6. Domain Events ↔ Integration Events', level=2)
P(doc, 'Mapper là chỗ DUY NHẤT 2 thế giới gặp nhau:')
Code(doc, J(
    'public Routed map(DomainEvent event) {',
    '    return switch (event) {',
    '        case UserCreatedEvent e -> lifecycle(',
    '            e, "auth.user.created.v1",',
    '            Map.of("username", e.username(), "email", e.email()));',
    '        case EmailChangedEvent e -> lifecycle(',
    '            e, "auth.user.email_changed.v1",',
    '            Map.of("oldEmail", e.oldEmail(), "newEmail", e.newEmail()));',
    '        case TokenCreatedEvent e -> session(',
    '            e, "auth.token.created.v1",',
    '            fields("userId", e.userId(), "family", e.family()));',
    '        default -> null;  // event không leak ra ngoài',
    '    };',
    '}',
))
page_break(doc)
print("Done Ch8")


# ==================== CHAPTER 9 ====================
H(doc, 'CHƯƠNG 9: BẢO MẬT', level=1)

H(doc, '9.1. Tổng quan bảo mật nhiều lớp', level=2)
P(doc, 'Defense in depth – nhiều lớp bảo vệ độc lập.')
P(doc, 'Bảng 9.1: Các lớp bảo mật', italic=True)
Tbl(doc, ['Lớp','Vị trí','Trách nhiệm'], [
    ['1','API Gateway','CORS, rate limit, trace ID, JWT pass-through.'],
    ['2','Service Filter','JWT verify, blacklist check, SecurityContext.'],
    ['3','Method Security','@PreAuthorize role/permission check.'],
    ['4','Domain Validation','Business rule.'],
    ['5','Database','Unique constraint, BCrypt hash.'],
    ['6','Redis','Token blacklist + idempotency.'],
    ['7','Webhook','HMAC verify + IP allowlist.'],
])

H(doc, '9.2. JWT', level=2)
P(doc, 'Bảng 9.2: Cấu hình JWT', italic=True)
Tbl(doc, ['Thuộc tính','Giá trị','Mô tả'], [
    ['Algorithm','HS256','HMAC-SHA-256, JWT_SECRET ≥ 32 ký tự.'],
    ['Access TTL','900000ms (15p)','Ngắn để giảm cửa sổ tấn công.'],
    ['Refresh TTL','604800000ms (7d)','Dài; chống lộ bằng rotation.'],
    ['Header','Authorization: Bearer','Tùy chọn (gRPC).'],
    ['Cookie','HttpOnly; Secure; SameSite=Strict','Browser default.'],
    ['Storage client','Cookie / RAM','KHÔNG localStorage.'],
])

Code(doc, J(
    '{',
    '  "sub": "john_doe",',
    '  "userId": "550e8400-e29b-41d4-a716-446655440000",',
    '  "email": "john@hieu.vn",',
    '  "roles": ["ROLE_CUSTOMER"],',
    '  "jti": "abc-123-def",',
    '  "tokenVersion": 1,',
    '  "iat": 1747200000,',
    '  "exp": 1747200900,',
    '  "iss": "hieu-auth-service"',
    '}',
))
Note(doc, 'JWT KHÔNG chứa permissions để giữ kích thước nhỏ. Resolve runtime qua Redis cache hoặc gRPC.')

H(doc, '9.3. RBAC', level=2)
P(doc, 'Mô hình chuẩn: User N-N Role, Role N-N Permission.')
Code(doc, J(
    'User (1) ---< (N) UserRole (N) >--- (1) Role',
    'Role (1) ---< (N) RolePermission (N) >--- (1) Permission',
    '',
    'Ví dụ:',
    '  admin    -> ROLE_ADMIN    -> [USER_*, PRODUCT_*, ORDER_*, VOUCHER_*]',
    '  staff    -> ROLE_STAFF    -> [ORDER_READ_ALL, RETURN_APPROVE]',
    '  customer -> ROLE_CUSTOMER -> [PRODUCT_READ, CART_*, ORDER_OWN_*]',
))

P(doc, 'Bảng 9.3: Roles tiêu biểu', italic=True)
Tbl(doc, ['Role','Permissions','Use case'], [
    ['ROLE_ADMIN','USER_*, ROLE_*, PRODUCT_*, ORDER_*, REFUND_PROCESS','Toàn quyền'],
    ['ROLE_STAFF','ORDER_READ_ALL, RETURN_APPROVE, INVENTORY_ADJUST','Vận hành'],
    ['ROLE_CUSTOMER','PRODUCT_READ, CART_*, ORDER_OWN_*, WISHLIST_*','Khách hàng'],
])

H(doc, '9.4. Refresh Token Rotation & Blacklist', level=2)
P(doc, 'Bảng 9.4: Bốn lớp kiểm tra token', italic=True)
Tbl(doc, ['Lớp','Kiểm tra','Nguồn','Lỗi fail'], [
    ['1','Signature HS256','JWT_SECRET','Invalid signature (401)'],
    ['2','Chưa hết hạn','exp claim','Token expired (401)'],
    ['3','Không blacklist','Redis (fallback DB)','Token revoked (401)'],
    ['4','tokenVersion match','users.token_version','Version mismatch (401)'],
])
P(doc, 'Cơ chế revoke:')
Bul(doc, 'Logout: jti vào Redis blacklist (TTL=remaining).')
Bul(doc, 'Đổi pass / Revoke all: bump tokenVersion.')
Bul(doc, 'Refresh rotation: cũ revoked + mới phát hành.')
Bul(doc, 'Reuse detection: revoked dùng lại → revoke family.')
Bul(doc, 'Admin revoke: endpoint /revoke-all/{userId}.')

H(doc, '9.5. Token Revocation chi tiết', level=2)
Code(doc, J(
    'public Authentication validate(String token) {',
    '    Claims claims = parseToken(token);  // L1+L2: sig + exp',
    '    if (tokenBlacklist.isRevoked(claims.getId())) throw new TokenRevokedException();',
    '    int tokenVer = claims.get("tokenVersion", Integer.class);',
    '    int currentVer = userVersionCache.get(userId);',
    '    if (tokenVer != currentVer) throw new TokenVersionMismatchException();',
    '    return buildAuthentication(claims);',
    '}',
))
Note(doc, 'Redis blacklist check ~0.5ms. Caffeine cache ~1µs. Tổng overhead < 2ms cho cached case.')

H(doc, '9.6. Idempotency & Anti-replay', level=2)
P(doc, 'X-Idempotency-Key UUIDv4; service lưu (key→result) Redis TTL 24h.')
P(doc, 'Bảng 9.5: Idempotency theo endpoint', italic=True)
Tbl(doc, ['Endpoint','Key','TTL','Behavior'], [
    ['POST /api/v1/orders/from-cart','X-Idempotency-Key','24h','Trả OrderDTO cũ.'],
    ['POST /api/v1/payments','X-Idempotency-Key','24h','Trả Payment cũ.'],
    ['POST /api/v1/payments/{id}/refund','X-Idempotency-Key','24h','Trả Refund cũ.'],
    ['POST /api/v1/cart/items','X-Idempotency-Key','24h','Trả Cart cũ.'],
    ['SePay webhook','transactionId','7d','Bỏ qua duplicate.'],
])

H(doc, '9.7. Rate Limiting & CORS', level=2)
P(doc, 'Bảng 9.6: Rate limit', italic=True)
Tbl(doc, ['Endpoint','Limit','Window','Implementation'], [
    ['/api/v1/auth/login','5/IP','15p','Bucket4j + Redis'],
    ['/api/v1/auth/register','3/IP','60p','Bucket4j'],
    ['/api/v1/auth/refresh','20/user','5p','Bucket4j'],
    ['/api/v1/orders/from-cart','10/user','60p','Bucket4j'],
    ['/api/v1/* (mọi endpoint)','100/IP','60s','Redis counter'],
    ['User authenticated','1000/user','60s','Redis counter'],
])

P(doc, 'CORS config:')
Code(doc, J(
    "globalcors:",
    "  cors-configurations:",
    "    '[/**]':",
    "      allowedOriginPatterns:",
    "        - https://hieu.vn",
    "        - https://admin.hieu.vn",
    "        - http://localhost:3000  # dev only",
    "      allowedMethods: [GET, POST, PATCH, PUT, DELETE, OPTIONS]",
    "      allowCredentials: true",
    "      maxAge: 3600",
))

H(doc, '9.8. Bảo mật vận hành', level=2)
for b in [
    'JWT_SECRET ≥ 32 ký tự; rotate định kỳ.',
    'BCrypt cost 12.',
    'PostgreSQL: user riêng/service, password qua env, pg_hba md5.',
    'Kafka prod: SASL/SCRAM + ACL.',
    'Redis prod: requirepass + TLS.',
    'Webhook SePay: verify HMAC + IP allowlist.',
    'Secret prod: K8s Secret hoặc Vault.',
    'Audit log: admin action vào auth.user.events.v1.',
    'OWASP Dependency Check + Snyk CI/CD.',
    'Pen-test 6 tháng/lần.',
]:
    Bul(doc, b)
page_break(doc)
print("Done Ch9")


# ==================== CHAPTER 10 ====================
H(doc, 'CHƯƠNG 10: KIỂM THỬ', level=1)

H(doc, '10.1. Test Pyramid', level=2)
Code(doc, J(
    '                  /\\',
    '                 / E2E\\         5%  (TestContainers)',
    '                /------\\',
    '               / Integ. \\       25% (H2 + @SpringBootTest)',
    '              /----------\\',
    '             /   Unit     \\     70% (Mockito, no Spring)',
    '            /--------------\\',
))
Cap(doc, 'Hình 10.1: Test pyramid')

P(doc, 'Bảng 10.1: Framework test', italic=True)
Tbl(doc, ['Framework','Phiên bản','Mục đích'], [
    ['JUnit 5','5.10.x','Framework chính.'],
    ['Mockito','5.x','Mock dependency.'],
    ['AssertJ','3.24.x','Fluent assertion.'],
    ['Spring Boot Test','4.0.5','@WebMvcTest, @SpringBootTest.'],
    ['spring-security-test','6.x','@WithMockUser.'],
    ['Testcontainers','1.20.4','PostgreSQL/Redis/Kafka thật.'],
    ['WireMock','3.x','Stub external API.'],
    ['JaCoCo','0.8.12','Coverage report.'],
])

H(doc, '10.2. Unit Test', level=2)
P(doc, 'Bảng 10.2: Phân bổ unit test', italic=True)
Tbl(doc, ['Service','Test class','# test','Phạm vi'], [
    ['auth-service','UserTest','15+','Register, changePassword, transitionStatus.'],
    ['auth-service','RefreshTokenTest','10+','Create, rotate, revoke, family.'],
    ['auth-service','JwtTokenProviderTest','8+','Issue, parse, signature, claims.'],
    ['auth-service','LoginHandlerTest','6+','Success, wrong password, locked.'],
    ['auth-service','TokenBlacklistServiceTest','5+','Revoke, isRevoked, fallback.'],
    ['catalog-service','ProductTest','12+','Slug, status, variants.'],
    ['order-service','OrderTest','20+','State machine, totals, return.'],
    ['order-service','OrderApplicationServiceTest','10+','Create from cart, voucher.'],
    ['inventory-service','InventoryTest','12+','Reserve, release, audit.'],
    ['payment-service','PaymentTest','10+','Create, webhook HMAC.'],
    ['flash-sale-service','FlashSaleTest','10+','Atomic Lua, race condition.'],
])

H(doc, '10.3. Integration Test', level=2)
P(doc, 'Integration test kiểm tra:')
Bul(doc, 'Controller → Service → Repo với H2 hoặc Postgres TestContainer.')
Bul(doc, 'Spring Security filter chain.')
Bul(doc, 'gRPC endpoint với grpc-test inproc.')
Bul(doc, 'Kafka producer/consumer với EmbeddedKafka.')
Bul(doc, 'Outbox publisher scheduler.')
Bul(doc, 'Webhook contract với WireMock.')

Code(doc, J(
    '@SpringBootTest',
    '@Testcontainers',
    'class AuthControllerIntegrationTest {',
    '    @Container static PostgreSQLContainer<?> postgres = new PostgreSQLContainer<>("postgres:16-alpine");',
    '    @Container static GenericContainer<?> redis = new GenericContainer<>("redis:7-alpine").withExposedPorts(6379);',
    '    @Test void registerAndLogin_endToEnd() { /* ... */ }',
    '}',
))

H(doc, '10.4. Coverage', level=2)
P(doc, 'Bảng 10.3: Coverage mục tiêu', italic=True)
Tbl(doc, ['Tầng','Coverage','Lý do'], [
    ['Domain','80%','Business logic quan trọng nhất.'],
    ['Application (handlers)','70%','Use case orchestration.'],
    ['Infrastructure','50%','Phần lớn là wrapper.'],
    ['Interfaces','60%','Integration test cover.'],
    ['Overall','60%','Mục tiêu tổng thể.'],
])
P(doc, 'Tool: JaCoCo. Report: target/site/jacoco/index.html. Lệnh: mvn test.')
page_break(doc)
print("Done Ch10")


# ==================== CHAPTER 11 ====================
H(doc, 'CHƯƠNG 11: GIÁM SÁT VÀ VẬN HÀNH', level=1)

H(doc, '11.1. Boot stack local', level=2)
Code(doc, J(
    '# Bước 1: hạ tầng',
    'cd /Users/admin/HieuTo/projectjn/ecommerce',
    'docker compose up -d',
    '',
    '# Bước 2: services Spring Boot (Java 25 + Maven 3.9)',
    './boot-bg.sh        # boot 15 service nền',
    '',
    '# Bước 3: seed dữ liệu',
    './seed.sh           # users + categories + permissions',
    'python3 seed-hieu-luxe.py --wipe   # 25 sản phẩm HIEU',
    '',
    '# Bước 4: kiểm tra',
    'curl http://localhost:8080/api/v1/products | jq .totalElements',
    'curl http://localhost:9091/actuator/health',
    'open http://localhost:5601    # Kibana',
    'open http://localhost:8025    # MailHog',
))

P(doc, 'Thứ tự boot:')
Bul(doc, 'eureka-server boot đầu tiên (chờ 25s).')
Bul(doc, 'auth-service tiếp theo.')
Bul(doc, 'Các business service song song.')
Bul(doc, 'api-gateway boot cuối.')

H(doc, '11.2. Migrations & Seeding', level=2)
Bul(doc, 'Flyway tự chạy khi service start.')
Bul(doc, 'Convention: V{n}__init_{service}.sql.')
Bul(doc, 'KHÔNG sửa migration đã merge.')
Bul(doc, 'seed.sh idempotent (409 swallow).')

H(doc, '11.3. Structured Logging & Correlation ID', level=2)
P(doc, 'Logback + Logstash encoder → JSON; Filebeat → Elasticsearch → Kibana.')
Code(doc, J(
    '<configuration>',
    '  <appender name="STDOUT" class="ch.qos.logback.core.ConsoleAppender">',
    '    <encoder class="net.logstash.logback.encoder.LogstashEncoder">',
    '      <includeMdcKeyName>traceId</includeMdcKeyName>',
    '      <includeMdcKeyName>userId</includeMdcKeyName>',
    '    </encoder>',
    '  </appender>',
    '</configuration>',
))
P(doc, 'X-Request-Id propagate qua MDC.')

P(doc, 'Bảng 11.1: Log levels', italic=True)
Tbl(doc, ['Level','Dùng khi','Ví dụ'], [
    ['ERROR','Lỗi nghiêm trọng','Exception unhandled.'],
    ['WARN','Bất thường','Redis down, retry exhausted.'],
    ['INFO','Sự kiện quan trọng','Login, order placed.'],
    ['DEBUG','Chi tiết (dev)','Cache hit/miss.'],
    ['TRACE','Chi tiết nhất','SQL params.'],
])

H(doc, '11.4. Metrics & Actuator', level=2)
P(doc, 'Bảng 11.2: Metrics quan trọng', italic=True)
Tbl(doc, ['Metric','Loại','Mô tả'], [
    ['http.server.requests','Timer','HTTP request time (p95/p99).'],
    ['jvm.memory.used','Gauge','JVM memory.'],
    ['jdbc.connections.active','Gauge','HikariCP active.'],
    ['resilience4j.circuitbreaker.state','Gauge','CLOSED/OPEN/HALF_OPEN.'],
    ['jwt.validation.success','Counter','Token valid.'],
    ['jwt.validation.failure','Counter','Token invalid.'],
    ['cache.gets','Counter','Caffeine cache.'],
    ['kafka.consumer.records-lag-max','Gauge','Consumer lag.'],
    ['order.placed','Counter','Business metric.'],
])

P(doc, 'Bảng 11.3: Actuator endpoints', italic=True)
Tbl(doc, ['Endpoint','Mô tả'], [
    ['/actuator/health','UP/DOWN từng component.'],
    ['/actuator/health/liveness','K8s liveness probe.'],
    ['/actuator/health/readiness','K8s readiness probe.'],
    ['/actuator/info','Version, build, git.'],
    ['/actuator/metrics','Danh sách metrics.'],
    ['/actuator/prometheus','Prometheus format.'],
    ['/actuator/loggers','Đổi log level runtime.'],
    ['/actuator/threaddump','Debug deadlock.'],
])

H(doc, '11.5. Dashboard Grafana (dự kiến)', level=2)
Bul(doc, 'microservices-overview.json: uptime, request rate, error rate, p95.')
Bul(doc, 'JVM: heap, GC, threads.')
Bul(doc, 'Database: HikariCP, query time, slow query.')
Bul(doc, 'Security: login success/fail, blacklist hits.')
Bul(doc, 'Business: orders/min, revenue, top products.')
Bul(doc, 'Kafka: producer rate, consumer lag.')

H(doc, '11.6. Health Check & Liveness/Readiness', level=2)
Code(doc, J(
    'spec:',
    '  containers:',
    '    - name: auth-service',
    '      livenessProbe:',
    '        httpGet: { path: /actuator/health/liveness, port: 9091 }',
    '        initialDelaySeconds: 60',
    '        periodSeconds: 30',
    '      readinessProbe:',
    '        httpGet: { path: /actuator/health/readiness, port: 9091 }',
    '        initialDelaySeconds: 30',
    '        periodSeconds: 10',
))
P(doc, 'Liveness DOWN → restart pod. Readiness DOWN → ngưng route traffic.')

H(doc, '11.7. Khôi phục sau sự cố', level=2)
for b in [
    'Backup Postgres pg_dump hàng đêm; retention 7d + 4w + 6m (3-2-1 rule).',
    'Kafka retention 7 ngày – replay được.',
    'Idempotency record 24h.',
    'Outbox: scheduler retry mỗi 10s, sau 10 lần move DLQ.',
    'Saga compensation: order consume payment.failed → release inventory.',
    'Database-per-service: hỏng 1 DB chỉ gãy 1 service.',
    'DR: RTO ≤ 4h, RPO ≤ 1h; backup S3 + Postgres replica.',
    'Circuit breaker tự ngắt downstream lỗi.',
    'Graceful shutdown: timeout 30s xử lý in-flight.',
]:
    Bul(doc, b)
page_break(doc)
print("Done Ch11")


# ==================== CHAPTER 12 ====================
H(doc, 'CHƯƠNG 12: TRIỂN KHAI', level=1)

H(doc, '12.1. Docker Compose', level=2)
P(doc, 'Bảng 12.1: Docker Compose files', italic=True)
Tbl(doc, ['File','Mục đích','Services'], [
    ['docker-compose.yml','Hạ tầng chính','11 PostgreSQL, Redis, Kafka, ES, Kibana, MailHog, Filebeat.'],
    ['docker-compose-services.yml (dự kiến)','Spring Boot','15 microservices Docker image.'],
    ['docker-compose-monitoring.yml (dự kiến)','Monitoring','Prometheus, Grafana, exporters.'],
])

H(doc, '12.2. Biến môi trường', level=2)
P(doc, 'Bảng 12.2: Biến môi trường chính', italic=True)
Tbl(doc, ['Biến','Default','Mô tả'], [
    ['SPRING_PROFILES_ACTIVE','dev','dev/test/prod.'],
    ['JWT_SECRET','(256-bit)','HS256 secret ≥ 32 ký tự.'],
    ['JWT_EXPIRATION','900000','Access TTL (ms).'],
    ['JWT_REFRESH_EXPIRATION','604800000','Refresh TTL.'],
    ['EUREKA_SERVER_URL','http://localhost:8761/eureka/','Registry URL.'],
    ['SPRING_DATASOURCE_URL','jdbc:postgresql://localhost:5433/authdb','Service DB URL.'],
    ['SPRING_DATASOURCE_USERNAME','authuser','DB user.'],
    ['SPRING_DATASOURCE_PASSWORD','authpass','DB password.'],
    ['SPRING_REDIS_HOST','localhost','Redis.'],
    ['KAFKA_BOOTSTRAP_SERVERS','localhost:9092','Kafka.'],
    ['SEPAY_WEBHOOK_SECRET','(HMAC)','Verify webhook.'],
    ['MAIL_HOST','localhost','MailHog dev.'],
    ['ELASTICSEARCH_HOST','localhost','ES.'],
])
P(doc, 'Production: K8s Secret hoặc Vault. Dev: .env.example commit, .env gitignore.')

H(doc, '12.3. Các lệnh khởi động', level=2)
Tbl(doc, ['Lệnh','Mô tả'], [
    ['docker compose up -d','Hạ tầng.'],
    ['docker compose ps','List + health.'],
    ['docker compose logs -f auth-service','Tail log.'],
    ['docker compose down -v','Reset.'],
    ['./mvnw clean install -DskipTests','Build.'],
    ['./mvnw spring-boot:run -pl auth-service','Chạy 1 service.'],
    ['./boot-bg.sh','Boot toàn bộ background.'],
    ['./mvnw test -pl auth-service','Test 1 module.'],
    ['./mvnw test','Test toàn dự án.'],
    ['./seed.sh','Seed data.'],
])

H(doc, '12.4. Dự kiến: Kubernetes', level=2)
for b in [
    'Deployment với replicaCount (default 2).',
    'Service (ClusterIP) inter-service.',
    'Ingress NGINX thay Gateway.',
    'ConfigMap cho non-sensitive config.',
    'Secret cho JWT_SECRET, DB password.',
    'HPA scale theo CPU/memory/custom metric.',
    'PVC cho database storage (gp3 SSD).',
    'Helm Charts per environment.',
    'Istio mTLS (tùy chọn).',
    'Cert-Manager + Let\'s Encrypt.',
]:
    Bul(doc, b)

H(doc, '12.5. Dự kiến: CI/CD Pipeline', level=2)
P(doc, 'GitHub Actions hoặc Jenkins:')
for n in [
    'Code Push: branch feature/* hoặc main.',
    'Build: Maven package + Docker image.',
    'Unit Test: mvn test; fail pipeline nếu fail.',
    'Code Quality: SonarQube.',
    'Security Scan: OWASP Dependency Check + Snyk.',
    'Docker Build & Push: tag :git-sha + :latest.',
    'Deploy Staging: K8s staging cluster.',
    'Integration Test: e2e với data seed.',
    'Manual Approval: reviewer approve.',
    'Deploy Production: rolling canary (5% → 25% → 100%).',
    'Monitor: theo dõi error rate; auto-rollback nếu > 5%.',
]:
    Num(doc, n)
page_break(doc)
print("Done Ch12")


# ==================== CHAPTER 13 ====================
H(doc, 'CHƯƠNG 13: LỘ TRÌNH PHÁT TRIỂN', level=1)

H(doc, '13.1. Hiện trạng triển khai', level=2)
P(doc, 'Tính đến phiên bản 3.0, đã hoàn thành phần code cơ bản với DDD + Hexagonal + CQRS + Choreography Saga:')
for b in [
    'JWT cookie + RBAC + refresh rotation + blacklist 4-layer.',
    'Catalog Service (REST + gRPC) – sản phẩm, danh mục, biến thể, thuộc tính.',
    'E-commerce core: cart → order → payment → shipping → inventory với saga + outbox.',
    'Voucher, flash sale (atomic Lua), notification (in-app + email + SSE).',
    'Search (Elasticsearch), analytics, log explorer.',
    'Đổi/trả: ReturnRequest + refund.',
    'Tích hợp SePay VietQR webhook.',
    'Unit test domain + security critical path.',
    'Docker Compose hạ tầng.',
    'Filebeat → ES → Kibana log stack.',
    'Domain Events ↔ Integration Events với mapper.',
    'MapStruct DTO mapping (auth-service đã migrate).',
]:
    Bul(doc, b)

H(doc, '13.2. Công nghệ dự kiến bổ sung', level=2)
P(doc, 'Bảng 13.1: Roadmap công nghệ', italic=True)
Tbl(doc, ['STT','Công nghệ','Mục đích','Ghi chú'], [
    ['1','Prometheus + Grafana','Metrics dashboard','Setup dashboard chuẩn.'],
    ['2','OpenTelemetry + Jaeger','Distributed tracing','Request xuyên services.'],
    ['3','Python + FastAPI + ML','AI Recommendation','Collaborative + Content-Based.'],
    ['4','Kubernetes','Container orchestration','Auto scale, rolling update.'],
    ['5','Helm','K8s package','Deployment template.'],
    ['6','GitHub Actions','CI/CD','Build, test, deploy automation.'],
    ['7','SonarQube','Code quality','Static analysis.'],
    ['8','HashiCorp Vault','Secret','Secret cho K8s.'],
    ['9','NGINX Ingress','L7 LB','Thay Gateway.'],
    ['10','AWS RDS','Managed Postgres','HA + read replica.'],
    ['11','Confluent Kafka','Managed Kafka','Schema registry.'],
    ['12','Schema Registry','Kafka schema','Tránh breaking change.'],
    ['13','Istio','Service mesh','mTLS, canary.'],
    ['14','ArgoCD','GitOps','Sync K8s manifests.'],
    ['15','Sentry / Datadog','APM','Production monitoring.'],
])

H(doc, '13.3. Lộ trình theo giai đoạn', level=2)
P(doc, 'Bảng 13.2: Lộ trình', italic=True)
Tbl(doc, ['Giai đoạn','Nội dung','Công nghệ','Trạng thái'], [
    ['Phase 1 (Q1/2026)','Core services (Auth, Catalog, Gateway, Eureka)','Spring Boot, JWT, gRPC','Hoàn thành'],
    ['Phase 2 (Q2/2026)','15 business services, Saga, Outbox, Idempotency','Spring Boot, Kafka','Hoàn thành'],
    ['Phase 3 (Q2/2026)','Refactor security (4-layer, rotation, hybrid blacklist)','JWT, Redis','Hoàn thành'],
    ['Phase 4 (Q2/2026)','Infrastructure & monitoring cơ bản','Docker, Filebeat, ES, Kibana','Hoàn thành'],
    ['Phase 5 (Q3/2026)','Testing & quality + MapStruct migration','JUnit, Testcontainers, MapStruct','Đang triển khai'],
    ['Phase 6 (Q3/2026)','Production monitoring','Prometheus, Grafana, OpenTelemetry, Jaeger','Dự kiến'],
    ['Phase 7 (Q4/2026)','AI features','Python + FastAPI + ML','Dự kiến'],
    ['Phase 8 (Q4/2026)','DevOps','K8s, Helm, ArgoCD, GitHub Actions, Vault','Dự kiến'],
    ['Phase 9 (Q1/2027)','Multi-warehouse + multi-region','Sharding, geo Kafka','Dự kiến'],
    ['Phase 10 (Q1/2027)','Observability','Sentry, Datadog APM','Dự kiến'],
])
page_break(doc)
print("Done Ch13")


# ==================== APPENDIX ====================
H(doc, 'PHỤ LỤC', level=1)

H(doc, 'A. Bảng cổng dịch vụ', level=2)
P(doc, 'Bảng A.1: Tổng hợp cổng và database', italic=True)
Tbl(doc, ['STT','Thành phần','Cổng','Mục đích'], [
    ['1','API Gateway','8080','HTTP entry'],
    ['2','Eureka Server','8761','Registry + dashboard'],
    ['3','Auth (HTTP)','9091','REST API'],
    ['4','Auth (gRPC)','9191','Inter-service'],
    ['5','User Profile','9099','Profile + addresses'],
    ['6','Catalog (HTTP)','9093','REST API'],
    ['7','Catalog (gRPC)','9193','Product lookup'],
    ['8','Inventory','9098','Stock'],
    ['9','Cart (HTTP)','9094','REST'],
    ['10','Cart (gRPC)','9194','GetCart'],
    ['11','Order','9095','Saga'],
    ['12','Payment','8086','Payment + webhook'],
    ['13','Shipping','8087','Shipment'],
    ['14','Voucher','8094','Voucher'],
    ['15','Flash Sale','8089','Atomic quota'],
    ['16','Notification','8090','In-app + email + SSE'],
    ['17','Search','8092','Elasticsearch'],
    ['18','Analytics','8095','Dashboard'],
    ['19','Review','-','Review'],
    ['20','Wishlist','-','Wishlist'],
    ['21','Recommendation','-','v1 đơn giản'],
    ['22','PostgreSQL auth','5433','authdb'],
    ['23','PostgreSQL catalog','5434','catalogdb'],
    ['24','PostgreSQL cart','5435','cartdb'],
    ['25','PostgreSQL order','5436','orderdb'],
    ['26','PostgreSQL payment','5437','paymentdb'],
    ['27','PostgreSQL shipping','5438','shippingdb'],
    ['28','PostgreSQL profile','5439','profiledb'],
    ['29','PostgreSQL inventory','5440','inventorydb'],
    ['30','PostgreSQL voucher','5441','voucherdb'],
    ['31','PostgreSQL flash-sale','5442','flashsaledb'],
    ['32','PostgreSQL notification','5443','notificationdb'],
    ['33','Redis','6379','Cache + blacklist'],
    ['34','Kafka','9092','Event streaming'],
    ['35','Elasticsearch','9200','Search + log'],
    ['36','Kibana','5601','Log UI'],
    ['37','MailHog SMTP','1025','Dev SMTP'],
    ['38','MailHog UI','8025','Mail viewer'],
])

H(doc, 'B. Tham chiếu mã lỗi', level=2)
P(doc, 'Bảng B.1: Mã lỗi chuẩn', italic=True)
Tbl(doc, ['Code','Status','Mô tả'], [
    ['AUTH-400-MALFORMED','400','Request body malformed.'],
    ['AUTH-400-VALIDATION','400','Field validation fail.'],
    ['AUTH-401-INVALID-CREDENTIALS','401','Username/password sai.'],
    ['AUTH-401-TOKEN-EXPIRED','401','Access token hết hạn.'],
    ['AUTH-401-TOKEN-REVOKED','401','Token bị blacklist.'],
    ['AUTH-401-TOKEN-VERSION-MISMATCH','401','tokenVersion không khớp.'],
    ['AUTH-401-TOKEN-REUSE-DETECTED','401','Refresh token reused → revoke family.'],
    ['AUTH-403-ACCOUNT-LOCKED','403','Account bị khóa.'],
    ['AUTH-403-INSUFFICIENT-PERMISSION','403','Thiếu role/permission.'],
    ['AUTH-404-USER-NOT-FOUND','404','User không tồn tại.'],
    ['AUTH-409-USERNAME-EXISTS','409','Username đã tồn tại.'],
    ['AUTH-409-EMAIL-EXISTS','409','Email đã tồn tại.'],
    ['ORDER-400-EMPTY-CART','400','Giỏ rỗng.'],
    ['ORDER-409-INSUFFICIENT-STOCK','409','Hết hàng.'],
    ['ORDER-400-INVALID-VOUCHER','400','Voucher không hợp lệ.'],
    ['ORDER-422-INVALID-STATE-TRANSITION','422','Chuyển state sai.'],
    ['PAYMENT-400-INVALID-AMOUNT','400','Amount không khớp.'],
    ['PAYMENT-401-INVALID-WEBHOOK-SIGNATURE','401','HMAC sai.'],
    ['INVENTORY-409-CONCURRENT-UPDATE','409','Optimistic lock conflict.'],
    ['FLASHSALE-409-QUOTA-EXHAUSTED','409','Hết quota.'],
    ['COMMON-429-RATE-LIMITED','429','Quá tốc độ.'],
    ['COMMON-500-INTERNAL-ERROR','500','Server error.'],
    ['COMMON-503-SERVICE-UNAVAILABLE','503','Circuit open.'],
])

H(doc, 'C. Tài khoản và dữ liệu seed', level=2)
P(doc, 'Bảng C.1: Tài khoản mặc định', italic=True)
Tbl(doc, ['Username','Email','Password','Role'], [
    ['admin','admin@hieu.vn','Admin@123','ROLE_ADMIN'],
    ['staff01','staff01@hieu.vn','Staff@123','ROLE_STAFF'],
    ['customer01','customer01@hieu.vn','Customer@123','ROLE_CUSTOMER'],
    ['customer02','customer02@hieu.vn','Customer@123','ROLE_CUSTOMER'],
])

P(doc, 'Bảng C.2: Dữ liệu seed', italic=True)
Tbl(doc, ['Loại','SL','Mô tả'], [
    ['Categories','~10','Áo, quần, giày, phụ kiện.'],
    ['Products','25','HIEU Luxe.'],
    ['Variants','~75','3 variant/product.'],
    ['Attributes','~5','Màu, size, chất liệu.'],
    ['Vouchers','~5','WELCOME10, FREESHIP, FLASH50K.'],
    ['Flash Sales','1','Demo active.'],
    ['Roles','4','ADMIN, STAFF, CUSTOMER, GUEST.'],
    ['Permissions','~30','USER_*, PRODUCT_*, ORDER_*.'],
])

H(doc, 'D. Cấu hình mẫu', level=2)
P(doc, 'D.1. application.yaml (auth-service):', bold=True)
Code(doc, J(
    'spring:',
    '  application:',
    '    name: auth-service',
    '  datasource:',
    '    url: ${SPRING_DATASOURCE_URL:jdbc:postgresql://localhost:5433/authdb}',
    '    username: ${SPRING_DATASOURCE_USERNAME:authuser}',
    '    password: ${SPRING_DATASOURCE_PASSWORD:authpass}',
    '  jpa:',
    '    hibernate:',
    '      ddl-auto: validate',
    '  flyway:',
    '    locations: classpath:db/migration',
    '  data:',
    '    redis:',
    '      host: ${SPRING_REDIS_HOST:localhost}',
    '      port: ${SPRING_REDIS_PORT:6379}',
    '  kafka:',
    '    bootstrap-servers: ${KAFKA_BOOTSTRAP_SERVERS:localhost:9092}',
    '  grpc:',
    '    server:',
    '      port: 9191',
    'eureka:',
    '  client:',
    '    service-url:',
    '      defaultZone: ${EUREKA_SERVER_URL:http://localhost:8761/eureka/}',
    'server:',
    '  port: 9091',
    '  shutdown: graceful',
    'jwt:',
    '  secret: ${JWT_SECRET}',
    '  expiration: ${JWT_EXPIRATION:900000}',
    '  refresh-expiration: ${JWT_REFRESH_EXPIRATION:604800000}',
    'management:',
    '  endpoints:',
    '    web:',
    '      exposure:',
    '        include: health,info,metrics,prometheus,loggers',
    '  endpoint:',
    '    health:',
    '      probes:',
    '        enabled: true',
))

P(doc, 'D.2. API Gateway routes:', bold=True)
Code(doc, J(
    'spring:',
    '  cloud:',
    '    gateway:',
    '      routes:',
    '        - id: auth-public',
    '          uri: lb://AUTH-SERVICE',
    '          predicates:',
    '            - Path=/api/v1/auth/register,/api/v1/auth/login,/api/v1/auth/refresh',
    '        - id: auth-protected',
    '          uri: lb://AUTH-SERVICE',
    '          predicates: [Path=/api/v1/auth/**]',
    '          filters: [JwtCookieAuthFilter]',
    '        - id: catalog',
    '          uri: lb://CATALOG-SERVICE',
    '          predicates: [Path=/api/v1/products/**,/api/v1/categories/**]',
    '        - id: cart',
    '          uri: lb://CART-SERVICE',
    '          predicates: [Path=/api/v1/cart/**]',
    '        - id: order',
    '          uri: lb://ORDER-SERVICE',
    '          predicates: [Path=/api/v1/orders/**,/api/v1/return-requests/**]',
    '        - id: payment',
    '          uri: lb://PAYMENT-SERVICE',
    '          predicates: [Path=/api/v1/payments/**]',
))

P(doc, 'D.3. Docker Compose hạ tầng:', bold=True)
Code(doc, J(
    'services:',
    '  postgres-auth:',
    '    image: postgres:16-alpine',
    '    environment:',
    '      POSTGRES_DB: authdb',
    '      POSTGRES_USER: authuser',
    '      POSTGRES_PASSWORD: authpass',
    '    ports: ["5433:5432"]',
    '    volumes: [pgdata-auth:/var/lib/postgresql/data]',
    '  redis:',
    '    image: redis:7-alpine',
    '    ports: ["6379:6379"]',
    '  kafka:',
    '    image: bitnami/kafka:3.7',
    '    environment:',
    '      KAFKA_CFG_NODE_ID: 0',
    '      KAFKA_CFG_PROCESS_ROLES: controller,broker',
    '    ports: ["9092:9092"]',
    '  elasticsearch:',
    '    image: docker.elastic.co/elasticsearch/elasticsearch:8.13.0',
    '    environment:',
    '      discovery.type: single-node',
    '    ports: ["9200:9200"]',
    '  kibana:',
    '    image: docker.elastic.co/kibana/kibana:8.13.0',
    '    ports: ["5601:5601"]',
    '  mailhog:',
    '    image: mailhog/mailhog',
    '    ports: ["1025:1025","8025:8025"]',
))

H(doc, 'E. Glossary nâng cao – Patterns & Decisions', level=2)
P(doc, 'Bảng E.1: Pattern + lý do chọn', italic=True)
Tbl(doc, ['Pattern','Mục đích','Tại sao chọn'], [
    ['Aggregate (DDD)','Invariant trong 1 đơn vị giao dịch.','Tránh inconsistent state.'],
    ['Repository (port)','Trừu tượng data access.','Domain không biết JPA.'],
    ['Domain Event','Decouple aggregate ↔ subscriber.','Loose coupling.'],
    ['Outbox Pattern','At-least-once delivery.','Event không mất khi Kafka down.'],
    ['CQRS','Tách read/write.','1 use case = 1 file Handler.'],
    ['Hexagonal','Tách domain khỏi framework.','Domain testable không Spring.'],
    ['Choreography Saga','Distributed transaction.','Không cần orchestrator, scale tốt.'],
    ['Idempotency Key','An toàn khi retry.','Mạng/Kafka duplicate.'],
    ['Refresh Token Rotation','Phát hiện token theft.','Reused → revoke family.'],
    ['Token Blacklist','Revoke trước expiry.','Kill 1 token cụ thể.'],
    ['tokenVersion','Kill all token của user.','1 INT update.'],
    ['Circuit Breaker','Tránh cascading failure.','Ngắt khi downstream lỗi.'],
    ['Caffeine + Redis','2-layer cache.','L1 1µs + L2 0.5ms.'],
    ['Optimistic Lock','Concurrent update.','Tránh lost update.'],
    ['Lua atomic','Race-condition-free counter.','Flash sale quota.'],
    ['Value Object','Type safety + invariant.','UserId không nhầm với String.'],
])

# ==================== FOOTER ====================
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('--- HẾT TÀI LIỆU ---')
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('© 2026 HIEU E-Commerce Backend Team — Documentation v3.0 (FINAL)')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88,0x88,0x88)


# ==================== SAVE ====================
out = '/Users/admin/HieuTo/projectjn/ecommerce/HIEU-Backend-Documentation-FINAL.docx'
doc.save(out)
import os
print(f"\n✓ Done. File size: {os.path.getsize(out)/1024:.1f} KB")
print(f"✓ Saved to: {out}")
